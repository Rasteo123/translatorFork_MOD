import sys
import os
import json
import asyncio
import shutil
import re
import array
import queue
import time
import threading
import traceback
import platform
import subprocess
import io
import html
import logging
import hashlib
import zipfile
import concurrent.futures
from datetime import datetime, timedelta

try:
    from zoneinfo import ZoneInfo
except ImportError:
    ZoneInfo = None

# PyQt6 Imports
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QFileDialog, QVBoxLayout,
    QHBoxLayout, QWidget, QTextEdit, QPushButton, QLabel,
    QProgressBar, QMessageBox, QInputDialog, QSplitter,
    QListWidget, QListWidgetItem, QToolBar, QSlider,
    QSizePolicy, QCheckBox, QMenu, QComboBox, QSpinBox,
    QDialog, QDialogButtonBox, QScrollArea, QTabWidget, QPlainTextEdit
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QObject, QTimer
from PyQt6.QtGui import QTextCursor, QTextCharFormat, QColor, QAction, QTextBlockFormat, QDragEnterEvent, QDropEvent, QIcon

# Libraries
try:
    import ebooklib
    from ebooklib import epub
except ImportError:
    ebooklib = None
    epub = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import nltk
except ImportError:
    nltk = None

try:
    import pyaudio
except ImportError:
    pyaudio = None

try:
    from pydub import AudioSegment
except ImportError:
    AudioSegment = None

try:
    import edge_tts  # pip install edge-tts
except ImportError:
    edge_tts = None

# Google SDK
try:
    from google import genai
    from google.genai import types as genai_types
except ImportError:
    genai = None
    genai_types = None

try:
    from loguru import logger as _loguru_logger
except ImportError:
    _loguru_logger = None

try:
    from gemini_translator.api import config as reader_api_config
except Exception:
    reader_api_config = None

try:
    from gemini_translator.core.worker_helpers.rpm_limiter import RPMLimiter
except Exception:
    RPMLimiter = None

try:
    from gemini_translator.utils.helpers import TokenCounter
except Exception:
    TokenCounter = None

try:
    from gemini_translator.utils.settings import SettingsManager
except Exception:
    SettingsManager = None

if platform.system() == "Windows":
    import subprocess
    # Патч: заставляем все процессы запускаться без окна консоли
    _orig_popen = subprocess.Popen
    def _hidden_popen(*args, **kwargs):
        if 'creationflags' not in kwargs:
            kwargs['creationflags'] = subprocess.CREATE_NO_WINDOW
        return _orig_popen(*args, **kwargs)
    subprocess.Popen = _hidden_popen





# --- КОНФИГУРАЦИЯ ---
MODEL_ID = "gemini-3.1-flash-live-preview" # Legacy default for Live API
AUDIO_RATE = 24000
AUDIO_CHANNELS = 1
READER_LIVE_WORKER_START_STAGGER_SECONDS = 1.5
READER_PARALLEL_WORKER_START_STAGGER_SECONDS = 1.0
READER_FLASH_WORKER_START_STAGGER_SECONDS = 1.0
WINDOWS_CREATE_NO_WINDOW = subprocess.CREATE_NO_WINDOW if platform.system() == "Windows" else 0
READER_SETTINGS_KEY = "gemini_reader_settings"
READER_RPD_STATE_KEY = "gemini_reader_rpd_state"
READER_INVALID_KEYS_KEY = "gemini_reader_invalid_keys"
READER_PROJECT_SETTINGS_FILENAME = "project_settings.json"
LEGACY_SETTINGS_FILE = "settings.json"
READER_BOOKS_DIRNAME = "gemini_reader_books"
LIVE_PARALLEL_DIRNAME = "_live_parallel"
MAX_PART_SIZE = 1.45 * 1024 * 1024 * 1024 # ~1.5 ГБ порог для склейки
SCRIPT_SUFFIX = ".tts.txt"
VIDEO_COVER_BASENAME = "video_cover"
VIDEO_COVER_EXTENSIONS = (".jpg", ".jpeg", ".png", ".webp", ".bmp")
READER_SUPPORTED_BOOK_EXTENSIONS = (".epub", ".zip", ".txt", ".md", ".html", ".htm")
READER_BOOK_FILE_FILTER = "Books (*.epub *.zip *.txt *.md *.html *.htm)"
VIDEO_FRAME_SIZE = "1920:1080"
CHAPTER_TRIM_SILENCE_THRESHOLD_DB = -45
CHAPTER_TRIM_SILENCE_STEP_MS = 10
CHAPTER_TRIM_KEEP_MS = 70
READER_LOG_FLUSH_INTERVAL_MS = 120
READER_PROGRESS_FLUSH_INTERVAL_MS = 100
READER_PROGRESS_EMIT_INTERVAL_SEC = 0.12
READER_LOG_MAX_BLOCKS = 2000
READER_WORKER_SLEEP_STEP_SEC = 0.5
READER_LIVE_CONNECT_TIMEOUT_SEC = 30
READER_LIVE_SEND_TIMEOUT_SEC = 30
READER_LIVE_CLOSE_TIMEOUT_SEC = 10
READER_LIVE_FIRST_CHUNK_TIMEOUT_SEC = 30
READER_LIVE_NEXT_CHUNK_TIMEOUT_SEC = 15
READER_GENERATE_CONTENT_TIMEOUT_SEC = 300
READER_EDGE_TTS_TOTAL_TIMEOUT_SEC = 90
READER_EDGE_TTS_CHUNK_TIMEOUT_SEC = 20
READER_AUDIO_EXPORT_TIMEOUT_SEC = 1800
READER_FFPROBE_TIMEOUT_SEC = 60
READER_FFMPEG_CONCAT_TIMEOUT_SEC = 3600
READER_FFMPEG_VIDEO_TIMEOUT_SEC = 7200

ENGINE_MODES = {
    "Live API": "live",
    "Flash TTS": "flash_tts",
}

LIVE_AUDIO_MODELS = {
    "Gemini 3.1 Flash Live Preview": "gemini-3.1-flash-live-preview",
    "Gemini 2.5 Flash Native Audio Preview": "gemini-2.5-flash-native-audio-preview-12-2025",
}

LIVE_SEGMENT_OPTIONS = {
    "Предложения": "sentences",
    "Параграфы": "paragraphs",
}

FLASH_TTS_MODELS = {
    "Gemini 3.1 Flash TTS Preview": "gemini-3.1-flash-tts-preview",
    "Gemini 2.5 Flash TTS Preview": "gemini-2.5-flash-preview-tts",
    "Gemini 2.5 Pro TTS Preview": "gemini-2.5-pro-preview-tts",
}

VOICE_MODE_OPTIONS = {
    "Один голос": "single",
    "Два голоса (Narrator + Dialogue)": "duo",
    "Автор + Муж./Жен. роли": "author_gender",
}

PIPELINE_MODE_OPTIONS = {
    "Авто: AI сценарий + TTS": "auto",
    "По шагам: сохранённый сценарий": "staged",
    "Без AI сценария": "raw",
}

PREPROCESS_PROFILE_OPTIONS = {
    "Бережно": "Keep tags sparse. Only add them where emotion or sound is explicit in the source text.",
    "Выразительно": "Add tasteful performance cues and a few non-verbal sounds when clearly supported by the scene.",
}

TTS_SPEAKER_NARRATOR = "Narrator"
TTS_SPEAKER_DIALOGUE = "Dialogue"
LIVE_ROLE_AUTHOR = "Author"
LIVE_ROLE_MALE = "Male"
LIVE_ROLE_FEMALE = "Female"
ROLE_SPEAKER_PATTERN = re.compile(
    rf"^\s*({TTS_SPEAKER_NARRATOR}|{TTS_SPEAKER_DIALOGUE})\s*:\s*(.+)$",
    re.IGNORECASE,
)
AUTHOR_GENDER_ROLE_PATTERN = re.compile(
    rf"^\s*({LIVE_ROLE_AUTHOR}|{LIVE_ROLE_MALE}|{LIVE_ROLE_FEMALE})\s*:\s*(.+)$",
    re.IGNORECASE,
)
LIVE_META_TAG_PATTERN = re.compile(r"^\[\s*([^\]\n]{1,40})\s*\]\s*")
ROLE_QUOTE_PATTERN = re.compile(r'«([^»]+)»|"([^"]+)"')
SCRIPT_WORD_PATTERN = re.compile(r"[A-Za-zА-Яа-яЁё0-9-]+")
FEMALE_DIALOGUE_HINTS = re.compile(
    r"\b(сказала|спросила|ответила|прошептала|крикнула|произнесла|буркнула|добавила|"
    r"пробормотала|прошипела|вздохнула|усмехнулась|улыбнулась|девушка|женщина|сестра|"
    r"мать|героиня|она)\b",
    re.IGNORECASE,
)
MALE_DIALOGUE_HINTS = re.compile(
    r"\b(сказал|спросил|ответил|прошептал|крикнул|произнес|буркнул|добавил|"
    r"пробормотал|прошипел|вздохнул|усмехнулся|улыбнулся|парень|мужчина|брат|"
    r"отец|герой|он)\b",
    re.IGNORECASE,
)

AUTHOR_GENDER_ATTRIBUTION_WORDS = (
    r"сказал(?:а)?|спросил(?:а)?|ответил(?:а)?|добавил(?:а)?|"
    r"произнес(?:ла)?|прошептал(?:а)?|крикнул(?:а)?|буркнул(?:а)?|"
    r"пробормотал(?:а)?|процедил(?:а)?|усмехнул(?:ся|ась)|улыбнул(?:ся|ась)|"
    r"вздохнул(?:а)?|подумал(?:а)?|заметил(?:а)?|отметил(?:а)?|бросил(?:а)|"
    r"раздал(?:ся|ась)|донес(?:ся|лась)|послышал(?:ся|ась)|прозвучал(?:а)?|"
    r"отозвал(?:ся|ась)|разнес(?:ся|лась)|прокатил(?:ся|ась)"
)
AUTHOR_GENDER_AUTHORIAL_TAIL_PATTERN = re.compile(
    rf"(?:"
    rf"(?:[»\"”]|[.!?…])\s*,?\s*(?:[—-]\s*)?"
    rf"|,\s*[—-]\s*"
    rf")[^\n]{{0,140}}\b"
    rf"(?:{AUTHOR_GENDER_ATTRIBUTION_WORDS}|голос|бас|баритон|шепот|крик|рык)\b",
    re.IGNORECASE,
)
AUTHOR_GENDER_AUTHORIAL_START_PATTERN = re.compile(
    rf"^\s*(?:[—-]\s*)?(?:{AUTHOR_GENDER_ATTRIBUTION_WORDS})\b",
    re.IGNORECASE,
)
AUTHOR_GENDER_AUTHORIAL_KEYWORDS = (
    "сказал", "сказала", "спросил", "спросила", "ответил", "ответила", "добавил", "добавила",
    "произнес", "произнесла", "прошептал", "прошептала", "крикнул", "крикнула", "буркнул",
    "буркнула", "пробормотал", "пробормотала", "процедил", "процедила", "усмехнулся",
    "усмехнулась", "улыбнулся", "улыбнулась", "вздохнул", "вздохнула", "подумал", "подумала",
    "заметил", "заметила", "отметил", "отметила", "бросил", "бросила", "раздался", "раздалась",
    "донесся", "донеслась", "послышался", "послышалась", "прозвучал", "прозвучала", "отозвался",
    "отозвалась", "разнесся", "разнеслась", "прокатился", "прокатилась", "голос", "бас",
    "баритон", "шепот", "крик", "рык",
)

TTS_SPEED_PROMPTS = {
    "Very Slow": "Read very slowly, with clear breathing room between phrases.",
    "Slow": "Read slowly and steadily, with calm pacing.",
    "Normal": "Read naturally, like a professional audiobook narrator.",
    "Fast": "Read with energetic but still intelligible pacing.",
    "Very Fast": "Read very quickly but keep the diction understandable.",
}

PREPROCESS_MODEL_FALLBACKS = {
    "Gemini 2.5 Flash": "gemini-2.5-flash",
    "Gemini 3.1 Flash-Lite": "gemini-3.1-flash-lite-preview",
    "Gemini 3.0 Flash Preview": "gemini-3-flash-preview",
}

TTS_AUDIO_TAG_HINT = (
    "[whispers] [laughs] [sighs] [gasp] [shouting] "
    "[serious] [excited] [tired] [very slow] [very fast]"
)

LIVE_TAG_STYLE_HINTS = {
    "serious": "Интонация серьёзная, собранная и сдержанная.",
    "excited": "Интонация взволнованная и приподнятая, но без переигрывания.",
    "tired": "Интонация усталая, с лёгкой тяжестью и выдохом в голосе.",
    "whispers": "Подача тихая, почти шёпотом, но с разборчивой дикцией.",
    "shouting": "Подача напряжённая и громкая, как крик, но слова остаются чёткими.",
    "laughs": "Передай улыбку или смешок только тембром и интонацией, не добавляя новые слова.",
    "sighs": "Передай усталый выдох только интонацией, не добавляя новые слова.",
    "gasp": "Передай резкий вдох или испуг только интонацией, не добавляя новые слова.",
    "very slow": "Темп заметно медленнее обычного.",
    "slow": "Темп слегка замедленный.",
    "fast": "Темп слегка ускоренный, но дикция остаётся чёткой.",
    "very fast": "Темп заметно быстрее обычного, но дикция остаётся чёткой.",
}

DEFAULT_PREPROCESS_DIRECTIVE = (
    "Prefer compact, production-safe markup. Avoid overacting, dense tag spam and cinematic additions "
    "that are not grounded in the original text. In role-labeled scripts, keep narration, attribution and "
    "scene description out of character lines."
)

DEFAULT_TTS_DIRECTIVE = (
    "Keep diction clear and stable. Prioritize audiobook intelligibility over theatrical exaggeration."
)

LIVE_API_DEFAULT_RPM = 5
FLASH_PREPROCESS_DEFAULT_RPM = 5
FLASH_TTS_DEFAULT_RPM = 3
FLASH_TTS_CHAPTER_INTERVAL_SECONDS = 60
DEFAULT_TPM_LIMIT = 0
RATE_LIMIT_BACKOFF_SECONDS = 65
LIVE_PARAGRAPH_MAX_CHARS = 2200
LIVE_AUTHOR_GENDER_BLOCK_MAX_CHARS = 1100
LIVE_AUTHOR_GENDER_BLOCK_MAX_LINES = 6
LIVE_REQUEST_TRIM_KEEP_MS = 20
LIVE_REQUEST_JOIN_GAP_MS = 40
LIVE_REQUEST_ROLE_SWITCH_GAP_MS = 75
LIVE_REQUEST_SHORT_CUE_GAP_MS = 240
LIVE_REQUEST_COLON_GAP_MS = 260

RPD_LIMIT_FALLBACKS = {
    "gemini-3.1-flash-tts-preview": 10,
    "gemini-2.5-flash-preview-tts": 10,
    "gemini-2.5-pro-preview-tts": 10,
}

VOICES_MAP = {
    "Puck": "М", "Charon": "М", "Kore": "Ж", "Fenrir": "М", "Aoede": "Ж",
    "Zephyr": "Ж", "Leda": "Ж", "Orus": "М", "Callirrhoe": "Ж", "Autonoe": "Ж",
    "Enceladus": "М", "Iapetus": "М", "Umbriel": "М", "Algieba": "М", "Despina": "Ж",
    "Erinome": "Ж", "Algenib": "М", "Rasalgethi": "М", "Laomedeia": "Ж", "Achernar": "Ж",
    "Alnilam": "М", "Schedar": "М", "Gacrux": "Ж", "Pulcherrima": "Ж", "Achird": "М",
    "Zubenelgenubi": "М", "Vindemiatrix": "Ж", "Sadachbia": "М", "Sadaltager": "М", "Sulafat": "Ж"
}

EDGE_VOICE = "ru-RU-SvetlanaNeural" # Голос для подмены при цензуре

SPEED_PROMPTS = {
    "Very Slow": "Читай очень медленно и размеренно. Делай отчетливые паузы между фразами.",
    "Slow": "Читай медленно, в спокойном и расслабленном темпе.",
    "Normal": "Читай в естественном, нормальном темпе, как при обычном разговоре.",
    "Fast": "Читай в быстром, энергичном темпе.",
    "Very Fast": "Читай очень быстро, максимально динамично."
}

# --- СИСТЕМА ЛОГИРОВАНИЯ В GUI ---
class LogSignal(QObject):
    new_log = pyqtSignal(str, str) # msg, level

log_fifo = LogSignal()

def custom_log_handler(message):
    """Глобальный обработчик для перехвата сообщений из loguru и отправки в GUI"""
    try:
        record = message.record
        msg = record["message"]
        
        # Фильтрация мусора и длинных системных ошибок
        if any(x in msg for x in ["Traceback", "stack trace", "ProactorEventLoop"]):
            return
            
        # Упрощение ошибок сети для читаемости в логе
        if "websockets.exceptions" in msg or "1008" in msg:
            msg = "Сеть: Ошибка протокола (возможно цензура фрагмента). Воркер пробует обход."
        elif "1011" in msg or "timeout" in msg.lower():
            msg = "Сеть: Сервер Gemini временно недоступен или занят."

        level = record["level"].name
        
        # Отправка сигнала в GUI (если объект log_fifo создан)
        if 'log_fifo' in globals():
            log_fifo.new_log.emit(msg, level)
            
    except Exception as e:
        # Если что-то пошло не так в самом обработчике, выводим в стандартную консоль
        print(f"Error in custom_log_handler: {e}")

class _GuiLoggingHandler(logging.Handler):
    def emit(self, record):
        try:
            msg = record.getMessage()
            if any(x in msg for x in ["Traceback", "stack trace", "ProactorEventLoop"]):
                return
            if "websockets.exceptions" in msg or "1008" in msg:
                msg = "Сеть: Ошибка протокола (возможно цензура фрагмента). Воркер пробует обход."
            elif "1011" in msg or "timeout" in msg.lower():
                msg = "Сеть: Сервер Gemini временно недоступен или занят."
            log_fifo.new_log.emit(msg, record.levelname)
        except Exception:
            pass


def _run_subprocess(args, **kwargs):
    if WINDOWS_CREATE_NO_WINDOW:
        kwargs.setdefault("creationflags", WINDOWS_CREATE_NO_WINDOW)
    return subprocess.run(args, **kwargs)


async def _to_thread_with_timeout(label, timeout_seconds, func, *args, **kwargs):
    loop = asyncio.get_running_loop()
    executor = concurrent.futures.ThreadPoolExecutor(max_workers=1, thread_name_prefix="reader-blocking")
    future = loop.run_in_executor(executor, lambda: func(*args, **kwargs))
    try:
        return await asyncio.wait_for(future, timeout=timeout_seconds)
    except asyncio.TimeoutError as exc:
        future.cancel()
        raise RuntimeError(f"{label} timed out after {int(timeout_seconds)} seconds.") from exc
    finally:
        executor.shutdown(wait=False, cancel_futures=True)


async def _open_live_session_with_timeout(client, model_id, config):
    session_cm = client.aio.live.connect(model=model_id, config=config)
    try:
        session = await asyncio.wait_for(
            session_cm.__aenter__(),
            timeout=READER_LIVE_CONNECT_TIMEOUT_SEC,
        )
    except asyncio.TimeoutError as exc:
        raise RuntimeError(
            f"Live API connect timed out after {READER_LIVE_CONNECT_TIMEOUT_SEC} seconds."
        ) from exc
    return session_cm, session


async def _close_live_session_quietly(session_cm):
    if session_cm is None:
        return
    try:
        await asyncio.wait_for(
            session_cm.__aexit__(None, None, None),
            timeout=READER_LIVE_CLOSE_TIMEOUT_SEC,
        )
    except Exception:
        pass


def _make_genai_client(api_key, timeout_seconds=READER_GENERATE_CONTENT_TIMEOUT_SEC):
    if genai_types is not None and hasattr(genai_types, "HttpOptions"):
        return genai.Client(
            api_key=api_key,
            http_options=genai_types.HttpOptions(timeout=int(timeout_seconds * 1000)),
        )
    return genai.Client(api_key=api_key)


def _runtime_base_dir():
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def _resolve_tool_path(tool_name):
    direct_path = shutil.which(tool_name)
    if direct_path:
        return direct_path

    candidates = [tool_name]
    if platform.system() == "Windows" and not tool_name.lower().endswith(".exe"):
        candidates.insert(0, f"{tool_name}.exe")

    search_roots = [_runtime_base_dir(), os.getcwd()]
    for root in search_roots:
        for candidate in candidates:
            full_path = os.path.join(root, candidate)
            if os.path.isfile(full_path):
                return full_path
    return None


def _get_app_settings_manager():
    app = QApplication.instance()
    if app is None or not hasattr(app, "get_settings_manager"):
        return None
    try:
        return app.get_settings_manager()
    except Exception:
        return None


def _reader_books_dir(settings_manager=None):
    if settings_manager is not None and getattr(settings_manager, "config_dir", None):
        return os.path.join(settings_manager.config_dir, READER_BOOKS_DIRNAME)
    return "books"


def _load_legacy_settings():
    if not os.path.exists(LEGACY_SETTINGS_FILE):
        return {}
    try:
        with open(LEGACY_SETTINGS_FILE, "r", encoding="utf-8") as file_obj:
            return json.load(file_obj)
    except Exception:
        return {}


def _build_preprocess_models_map():
    models_map = {}
    if reader_api_config is not None:
        try:
            provider_cfg = reader_api_config.api_providers().get("gemini", {})
            for display_name, model_cfg in provider_cfg.get("models", {}).items():
                model_id = str(model_cfg.get("id") or "").strip()
                if not model_id or "tts" in model_id.lower():
                    continue
                models_map[display_name] = model_id
        except Exception:
            models_map = {}

    if models_map:
        return models_map
    return dict(PREPROCESS_MODEL_FALLBACKS)


def _default_preprocess_model_label(models_map):
    for preferred in ("Gemini 2.5 Flash", "Gemini 3.1 Flash-Lite", "Gemini 3.0 Flash Preview"):
        if preferred in models_map:
            return preferred
    return next(iter(models_map), "")


def _extract_response_text(response):
    direct_text = getattr(response, "text", None)
    if isinstance(direct_text, str) and direct_text.strip():
        return direct_text.strip()

    parts = []
    for candidate in getattr(response, "candidates", []) or []:
        content = getattr(candidate, "content", None)
        for part in getattr(content, "parts", []) or []:
            text_value = getattr(part, "text", None)
            if text_value:
                parts.append(text_value)
    return "\n".join(parts).strip()


def _iter_inline_audio_parts(response):
    for candidate in getattr(response, "candidates", []) or []:
        content = getattr(candidate, "content", None)
        for part in getattr(content, "parts", []) or []:
            inline_data = getattr(part, "inline_data", None)
            chunk = getattr(inline_data, "data", None) if inline_data is not None else None
            mime_type = str(getattr(inline_data, "mime_type", "") or "").strip().lower()
            if chunk:
                yield bytes(chunk), mime_type


def _guess_audio_format_from_mime(mime_type):
    mime_type = (mime_type or "").lower()
    if "mpeg" in mime_type or "mp3" in mime_type:
        return "mp3"
    if "wav" in mime_type:
        return "wav"
    if "ogg" in mime_type:
        return "ogg"
    if "flac" in mime_type:
        return "flac"
    if "aac" in mime_type:
        return "aac"
    if "mp4" in mime_type or "m4a" in mime_type:
        return "mp4"
    return None


def _decode_generated_audio_chunk(chunk, mime_type=""):
    mime_type = (mime_type or "").lower()
    if not chunk:
        return b""

    if any(marker in mime_type for marker in ("audio/l16", "audio/pcm", "audio/raw")):
        return bytes(chunk)

    if AudioSegment is None:
        if mime_type:
            raise RuntimeError(
                f"Модель вернула audio chunk в формате {mime_type}, но pydub недоступен для декодирования."
            )
        return bytes(chunk)

    audio_format = _guess_audio_format_from_mime(mime_type)
    if audio_format is None and not mime_type:
        return bytes(chunk)
    if audio_format is None:
        raise RuntimeError(f"Неподдерживаемый mime type аудио: {mime_type}")

    fp = io.BytesIO(chunk)
    segment = AudioSegment.from_file(fp, format=audio_format)
    segment = segment.set_frame_rate(AUDIO_RATE).set_channels(AUDIO_CHANNELS).set_sample_width(2)
    return segment.raw_data


def _extract_audio_bytes(response):
    data = bytearray()
    for chunk, mime_type in _iter_inline_audio_parts(response):
        data.extend(_decode_generated_audio_chunk(chunk, mime_type))
    return bytes(data)


def _build_preprocess_prompt(raw_text, voice_mode, profile_prompt, extra_directive=""):
    profile_line = profile_prompt or PREPROCESS_PROFILE_OPTIONS["Бережно"]
    extra_line = (extra_directive or "").strip() or DEFAULT_PREPROCESS_DIRECTIVE
    common_rules = (
        "You are preparing a Russian audiobook script for Gemini Flash TTS.\n"
        "Rules:\n"
        "- Preserve plot facts, meaning, names and chronology.\n"
        "- Keep the script in Russian.\n"
        "- Do not summarize, shorten, paraphrase or explain the text.\n"
        "- Do not invent any words, reactions, acknowledgements, connective phrases or narration that are absent from the source.\n"
        "- Use the Russian letter Ё/ё wherever it is orthographically appropriate and unambiguous; do not replace Ё/ё with Е/е in words such as всё, ещё, её, идёт, шёл, нём.\n"
        f"- You may use sparse English audio tags such as {TTS_AUDIO_TAG_HINT}.\n"
        "- Tags must be short, tasteful and only where justified by the source scene.\n"
        "- Return only the final script text, with no markdown and no comments.\n"
        f"- Direction profile: {profile_line}\n"
        f"- Additional director note: {extra_line}\n"
    )

    if voice_mode == "duo":
        return (
            f"{common_rules}\n"
            f"Transform the source into a strict two-speaker script using exactly these speakers: "
            f"{TTS_SPEAKER_NARRATOR} and {TTS_SPEAKER_DIALOGUE}.\n"
            f"Rules for speakers:\n"
            f"- Prefix every spoken line with `{TTS_SPEAKER_NARRATOR}:` or `{TTS_SPEAKER_DIALOGUE}:`.\n"
            f"- Narrative prose, author remarks, scene description and exposition go to `{TTS_SPEAKER_NARRATOR}`.\n"
            f"- Direct speech and quoted dialogue go to `{TTS_SPEAKER_DIALOGUE}`.\n"
            f"- If a paragraph mixes narration and dialogue, split it into several speaker lines.\n"
            f"- Do not invent extra speakers beyond those two names.\n\n"
            f"SOURCE TEXT:\n{raw_text}"
        )

    if voice_mode == "author_gender":
        return (
            f"{common_rules}\n"
            f"Transform the source into a strict labeled script using exactly these labels: "
            f"{LIVE_ROLE_AUTHOR}, {LIVE_ROLE_MALE}, {LIVE_ROLE_FEMALE}.\n"
            "Rules:\n"
            f"- Prefix every output line with `{LIVE_ROLE_AUTHOR}:`, `{LIVE_ROLE_MALE}:`, or `{LIVE_ROLE_FEMALE}:`.\n"
            f"- Use `{LIVE_ROLE_AUTHOR}:` for narration, scene description, speech attributions, inner thoughts, and any dialogue whose speaker gender is not explicitly clear in the source.\n"
            f"- Use `{LIVE_ROLE_MALE}:` only for direct speech when the source clearly identifies the speaker as male.\n"
            f"- Use `{LIVE_ROLE_FEMALE}:` only for direct speech when the source clearly identifies the speaker as female.\n"
            "- Never guess. If the source does not explicitly support male or female attribution, keep that line under Author.\n"
            f"- `{LIVE_ROLE_MALE}:` and `{LIVE_ROLE_FEMALE}:` may contain only the words the character actually says aloud.\n"
            "- Never leave author text inside character lines: no speech attributions, no action beats, no narration, no scene description, no voice description.\n"
            "- If a source line mixes dialogue and author text, split it into multiple labeled lines.\n"
            "- Tails such as `— сказал он`, `— спросила девушка`, `раздался густой бас`, `послышался голос`, `усмехнулся мужчина` must go to Author, not to Male/Female.\n"
            "- If you are not fully certain that the speaker is male or female from the source itself, keep the whole line under Author.\n"
            "- Preserve the wording of the source as closely as possible.\n"
            "- Do not invent standalone cues like `[Да]`, `[Нет]`, `[Хм]`, unless that exact spoken word is present in the source at that point.\n"
            "- Do not add extra labels, speaker names, markdown or explanations.\n"
            "Examples:\n"
            f"- WRONG: {LIVE_ROLE_MALE}: — Парень, посторонись немного, — раздался густой бас.\n"
            f"  RIGHT: {LIVE_ROLE_MALE}: — Парень, посторонись немного.\n"
            f"  RIGHT: {LIVE_ROLE_AUTHOR}: Раздался густой бас.\n"
            f"- WRONG: {LIVE_ROLE_FEMALE}: «Ты идёшь?» — спросила девушка.\n"
            f"  RIGHT: {LIVE_ROLE_FEMALE}: «Ты идёшь?»\n"
            f"  RIGHT: {LIVE_ROLE_AUTHOR}: — спросила девушка.\n\n"
            f"SOURCE TEXT:\n{raw_text}"
        )

    return (
        f"{common_rules}\n"
        "Keep the result as a single-speaker performance script.\n"
        "Do not add speaker labels.\n\n"
        f"SOURCE TEXT:\n{raw_text}"
    )


def _build_tts_generation_prompt(script_text, voice_mode, speed_key, extra_directive=""):
    speed_prompt = TTS_SPEED_PROMPTS.get(speed_key, TTS_SPEED_PROMPTS["Normal"])
    director_note = (extra_directive or "").strip() or DEFAULT_TTS_DIRECTIVE
    if voice_mode == "duo":
        return (
            "Perform the following Russian two-speaker audiobook script exactly as written.\n"
            f"Use the speaker names `{TTS_SPEAKER_NARRATOR}` and `{TTS_SPEAKER_DIALOGUE}` exactly as provided.\n"
            "Respect inline audio tags and emotional cues.\n"
            f"{speed_prompt}\n"
            f"{director_note}\n"
            "Preserve the wording of the script.\n\n"
            f"{script_text}"
        )

    return (
        "Perform the following Russian audiobook script exactly as written.\n"
        "Respect inline audio tags and emotional cues.\n"
        f"{speed_prompt}\n"
        f"{director_note}\n"
        "Preserve the wording of the script.\n\n"
        f"{script_text}"
    )


def _build_single_voice_speech_config(voice_name):
    return genai_types.SpeechConfig(
        language_code="ru-RU",
        voice_config=genai_types.VoiceConfig(
            prebuilt_voice_config=genai_types.PrebuiltVoiceConfig(
                voice_name=voice_name,
            )
        ),
    )


def _build_duo_voice_speech_config(primary_voice, secondary_voice):
    return genai_types.SpeechConfig(
        language_code="ru-RU",
        multi_speaker_voice_config=genai_types.MultiSpeakerVoiceConfig(
            speaker_voice_configs=[
                genai_types.SpeakerVoiceConfig(
                    speaker=TTS_SPEAKER_NARRATOR,
                    voice_config=genai_types.VoiceConfig(
                        prebuilt_voice_config=genai_types.PrebuiltVoiceConfig(
                            voice_name=primary_voice,
                        )
                    ),
                ),
                genai_types.SpeakerVoiceConfig(
                    speaker=TTS_SPEAKER_DIALOGUE,
                    voice_config=genai_types.VoiceConfig(
                        prebuilt_voice_config=genai_types.PrebuiltVoiceConfig(
                            voice_name=secondary_voice,
                        )
                    ),
                ),
            ]
        ),
    )


def _normalize_live_tag_name(tag_text):
    return re.sub(r"\s+", " ", (tag_text or "").strip().lower())


def _extract_leading_live_tags(text):
    remaining = (text or "").strip()
    tags = []
    while remaining:
        match = LIVE_META_TAG_PATTERN.match(remaining)
        if not match:
            break
        tag_name = _normalize_live_tag_name(match.group(1))
        if tag_name not in LIVE_TAG_STYLE_HINTS and not tag_name.startswith("pause"):
            break
        tags.append(tag_name)
        remaining = remaining[match.end():].lstrip()
    return tags, remaining


def _unwrap_live_spoken_cue(text):
    stripped = (text or "").strip()
    match = re.fullmatch(r"\[\s*([^\[\]\n]{1,24})\s*\]", stripped)
    if not match:
        return stripped
    inner_text = match.group(1).strip()
    if re.search(r"[A-Za-z]{2,}", inner_text):
        return stripped
    if re.search(r"[А-Яа-яЁё0-9]", inner_text):
        return inner_text
    return stripped


def _parse_live_pause_tag_ms(tag_name):
    normalized = _normalize_live_tag_name(tag_name)
    if normalized in {"pause", "beat", "short pause", "small pause"}:
        return LIVE_REQUEST_SHORT_CUE_GAP_MS
    if normalized in {"long pause", "dramatic pause"}:
        return LIVE_REQUEST_COLON_GAP_MS + 80
    match = re.fullmatch(r"pause\s*(\d{2,4})\s*ms", normalized)
    if match:
        return max(40, min(int(match.group(1)), 1200))
    return 0


def _build_live_segment_style(tags):
    style_parts = []
    for tag_name in tags or []:
        if tag_name in LIVE_TAG_STYLE_HINTS:
            style_parts.append(LIVE_TAG_STYLE_HINTS[tag_name])
    return " ".join(style_parts).strip()


def _is_short_live_cue(text):
    normalized = re.sub(r"\s+", " ", (text or "").strip())
    compact = normalized.strip(" .!?…:;,-")
    if not compact or len(compact) > 12:
        return False
    words = [word for word in re.findall(r"[А-Яа-яЁёA-Za-z0-9-]+", compact) if word]
    return 0 < len(words) <= 2


def _prepare_live_segment_text(text):
    tags, remaining = _extract_leading_live_tags(text)
    spoken_text = _unwrap_live_spoken_cue(remaining)
    spoken_text = re.sub(r"[ \t]+", " ", spoken_text).strip()
    pause_after_ms = 0
    for tag_name in tags:
        pause_after_ms = max(pause_after_ms, _parse_live_pause_tag_ms(tag_name))
    if _is_short_live_cue(spoken_text):
        if spoken_text and not re.search(r"[.!?…:]$", spoken_text):
            spoken_text = f"{spoken_text}."
        pause_after_ms = max(pause_after_ms, LIVE_REQUEST_SHORT_CUE_GAP_MS)
    if spoken_text.endswith(":"):
        pause_after_ms = max(pause_after_ms, LIVE_REQUEST_COLON_GAP_MS)
    return {
        "text": spoken_text,
        "style_hint": _build_live_segment_style(tags),
        "pause_after_ms": pause_after_ms,
    }


def _build_live_single_voice_request_config(
    voice_name,
    role_hint="default",
    style_prompt="",
    speed_instr=None,
    segment_directive="",
):
    speed_instr = speed_instr or SPEED_PROMPTS.get("Normal", "")
    gender = VOICES_MAP.get(voice_name, "Ж")
    dictor_type = "диктор-мужчина" if gender == "М" else "диктор-женщина"

    if role_hint == "author":
        role_instr = (
            "Озвучь только авторский текст, ремарки и описание сцены как рассказчик аудиокниги. "
            "Не добавляй реплики персонажей."
        )
    elif role_hint == "male_dialogue":
        role_instr = (
            "Озвучь только прямую речь мужского персонажа. "
            "Не добавляй авторский текст, ремарки или чужие реплики."
        )
    elif role_hint == "female_dialogue":
        role_instr = (
            "Озвучь только прямую речь женского персонажа. "
            "Не добавляй авторский текст, ремарки или чужие реплики."
        )
    else:
        role_instr = "Озвучь предоставленный текст СЛОВО В СЛОВО."

    segment_instr = ""
    if segment_directive:
        segment_instr = (
            f"Сценическая пометка для текущего фрагмента: {segment_directive} "
            "Реализуй её только тембром, интонацией и темпом. "
            "Не проговаривай названия тегов и не добавляй новые слова, междометия или звуки."
        )

    return genai_types.LiveConnectConfig(
        response_modalities=["AUDIO"],
        speech_config=_build_single_voice_speech_config(voice_name),
        system_instruction=(
            f"Ты профессиональный {dictor_type}. {style_prompt} "
            f"{speed_instr} "
            f"{role_instr} "
            f"{segment_instr} "
            "Читай только входной текст СЛОВО В СЛОВО. "
            "Не перефразируй, не продолжай сцену, не добавляй вступления, не переставляй фразы и не переключайся на другую роль. "
            "Обязательно дочитывай текст до самой последней точки и выдай полный аудиоответ. "
            "Язык: Русский."
        ),
    )


def _split_author_gender_script_segments(script_text, segment_mode):
    script_text = (script_text or "").strip()
    if not script_text:
        return []

    if segment_mode == "paragraphs":
        blocks = [block.strip() for block in re.split(r"\n\s*\n|\r\n\s*\r\n", script_text) if block.strip()]
        if len(blocks) > 1:
            return blocks

        labeled_lines = [line.strip() for line in script_text.splitlines() if line.strip()]
        if labeled_lines and all(AUTHOR_GENDER_ROLE_PATTERN.match(line) for line in labeled_lines):
            packed = []
            current_lines = []
            current_chars = 0
            for line in labeled_lines:
                if current_lines and (
                    len(current_lines) >= LIVE_AUTHOR_GENDER_BLOCK_MAX_LINES
                    or current_chars + len(line) + 1 > LIVE_AUTHOR_GENDER_BLOCK_MAX_CHARS
                ):
                    packed.append("\n".join(current_lines))
                    current_lines = []
                    current_chars = 0
                current_lines.append(line)
                current_chars += len(line) + 1
            if current_lines:
                packed.append("\n".join(current_lines))
            if packed:
                return packed
        return blocks or [script_text]

    return [line.strip() for line in script_text.splitlines() if line.strip()]


def _join_live_request_segments(segments, segment_mode, voice_mode="single"):
    parts = [segment for segment in (segments or []) if segment]
    if not parts:
        return ""
    if voice_mode == "author_gender":
        separator = "\n\n" if segment_mode == "paragraphs" else "\n"
        return separator.join(parts).strip()
    if segment_mode == "paragraphs":
        return "\n\n".join(parts).strip()
    return " ".join(parts).strip()


def _append_role_segment(segments, speaker, text):
    normalized = re.sub(r"\s+", " ", (text or "").strip())
    if not normalized:
        return

    is_punctuation_tail = bool(re.fullmatch(r"[.,!?…:;)\]»]+", normalized))
    if segments and is_punctuation_tail:
        prev_speaker, prev_text = segments[-1]
        segments[-1] = (prev_speaker, f"{prev_text.rstrip()}{normalized}")
        return

    if segments and segments[-1][0] == speaker:
        prev_speaker, prev_text = segments[-1]
        segments[-1] = (prev_speaker, f"{prev_text.rstrip()} {normalized}")
        return

    segments.append((speaker, normalized))


def _split_role_paragraph(paragraph):
    paragraph = re.sub(r"\s+", " ", (paragraph or "").strip())
    if not paragraph:
        return []

    existing_label_match = ROLE_SPEAKER_PATTERN.match(paragraph)
    if existing_label_match:
        speaker = existing_label_match.group(1).strip()
        speaker = TTS_SPEAKER_DIALOGUE if speaker.lower() == TTS_SPEAKER_DIALOGUE.lower() else TTS_SPEAKER_NARRATOR
        return [(speaker, existing_label_match.group(2).strip())]

    if re.match(r"^\s*[—–-]\s*", paragraph):
        core = re.sub(r"^\s*[—–-]\s*", "", paragraph).strip()
        raw_parts = [part.strip() for part in re.split(r"\s+[—–-]\s+", core) if part.strip()]
        if len(raw_parts) >= 2:
            segments = []
            for idx, part in enumerate(raw_parts):
                speaker = TTS_SPEAKER_DIALOGUE if idx % 2 == 0 else TTS_SPEAKER_NARRATOR
                _append_role_segment(segments, speaker, part)
            if segments:
                return segments
        return [(TTS_SPEAKER_DIALOGUE, core)]

    quote_matches = list(ROLE_QUOTE_PATTERN.finditer(paragraph))
    if quote_matches:
        segments = []
        cursor = 0
        for match in quote_matches:
            before_text = paragraph[cursor:match.start()]
            quote_text = (match.group(1) or match.group(2) or "").strip()
            _append_role_segment(segments, TTS_SPEAKER_NARRATOR, before_text)
            _append_role_segment(segments, TTS_SPEAKER_DIALOGUE, quote_text)
            cursor = match.end()
        _append_role_segment(segments, TTS_SPEAKER_NARRATOR, paragraph[cursor:])
        if segments:
            return segments

    return [(TTS_SPEAKER_NARRATOR, paragraph)]


def _build_live_role_script(raw_text):
    paragraphs = [
        part.strip()
        for part in re.split(r"\n\s*\n|\r\n\s*\r\n", raw_text or "")
        if part.strip()
    ]
    if not paragraphs and (raw_text or "").strip():
        paragraphs = [(raw_text or "").strip()]

    segments = []
    for paragraph in paragraphs:
        for speaker, text in _split_role_paragraph(paragraph):
            _append_role_segment(segments, speaker, text)

    if not segments:
        return ""

    return "\n".join(f"{speaker}: {text}" for speaker, text in segments if text.strip())


def _infer_dialogue_gender(text):
    sample = ((text or "").strip().lower()).replace("ё", "е")
    if not sample:
        return None
    female_score = len(FEMALE_DIALOGUE_HINTS.findall(sample))
    male_score = len(MALE_DIALOGUE_HINTS.findall(sample))
    if female_score > male_score and female_score > 0:
        return "female"
    if male_score > female_score and male_score > 0:
        return "male"
    return None


def _parse_author_gender_script(script_text):
    entries = []
    for raw_line in (script_text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        match = AUTHOR_GENDER_ROLE_PATTERN.match(line)
        if not match:
            if not entries:
                return []
            entries[-1]["text"] = f"{entries[-1]['text'].rstrip()}\n{line}"
            continue

        role_label = match.group(1).strip().lower()
        text = match.group(2).strip()
        if not text:
            continue

        if role_label == LIVE_ROLE_MALE.lower():
            role_hint = "male_dialogue"
        elif role_label == LIVE_ROLE_FEMALE.lower():
            role_hint = "female_dialogue"
        else:
            role_hint = "author"

        entries.append({"role_hint": role_hint, "text": text})

    return entries


def _has_author_gender_authorial_tail(text):
    normalized = re.sub(r"\s+", " ", (text or "").lower().replace("ё", "е")).strip()
    if not normalized:
        return False
    if AUTHOR_GENDER_AUTHORIAL_START_PATTERN.search(normalized):
        return True

    trimmed = normalized.lstrip("—- ").strip()
    first_token = trimmed.split(" ", 1)[0].strip(",.!?…:;\"'«»()[]") if trimmed else ""
    if first_token in AUTHOR_GENDER_AUTHORIAL_KEYWORDS:
        return True

    dash_parts = [
        part.strip(" ,.!?…:;\"'«»()[]")
        for part in re.split(r"\s*[—-]\s*", normalized)
        if part.strip(" ,.!?…:;\"'«»()[]")
    ]
    if len(dash_parts) >= 2:
        for tail in dash_parts[1:]:
            if any(keyword in tail for keyword in AUTHOR_GENDER_AUTHORIAL_KEYWORDS):
                return True

    for splitter in ("»", "\"", "”"):
        if splitter in normalized:
            tail = normalized.split(splitter, 1)[1].strip(" ,.!?…:;\"'«»()[]")
            if tail and any(keyword in tail for keyword in AUTHOR_GENDER_AUTHORIAL_KEYWORDS):
                return True
    return False


def _normalize_script_word(token):
    return (token or "").strip().lower().replace("ё", "е")


def _extract_source_word_set(source_text):
    words = set()
    for token in SCRIPT_WORD_PATTERN.findall(source_text or ""):
        normalized = _normalize_script_word(token)
        if len(normalized) < 3:
            continue
        words.add(normalized)
    return words


def _strip_author_gender_line_markup(text):
    prepared = _prepare_live_segment_text(text)
    stripped = prepared.get("text", "").strip()
    return re.sub(r"\[[^\]\n]{1,40}\]", " ", stripped).strip()


def _collect_author_gender_source_drift_issues(script_text, source_text):
    if not (script_text or "").strip() or not (source_text or "").strip():
        return []

    issues = []
    source_words = _extract_source_word_set(source_text)
    if not source_words:
        return issues

    for line_no, raw_line in enumerate((script_text or "").splitlines(), start=1):
        line = raw_line.strip()
        if not line:
            continue
        match = AUTHOR_GENDER_ROLE_PATTERN.match(line)
        text = match.group(2).strip() if match else line
        plain_text = _strip_author_gender_line_markup(text)
        if not plain_text:
            continue

        line_words = []
        unexpected_words = []
        for token in SCRIPT_WORD_PATTERN.findall(plain_text):
            normalized = _normalize_script_word(token)
            if len(normalized) < 3:
                continue
            if re.fullmatch(r"[a-z0-9-]+", normalized):
                continue
            line_words.append(normalized)
            if normalized not in source_words:
                unexpected_words.append(normalized)

        if len(unexpected_words) >= 2 and len(unexpected_words) >= max(2, len(line_words) // 3):
            unique_preview = ", ".join(sorted(dict.fromkeys(unexpected_words))[:6])
            issues.append(
                {
                    "line_no": line_no,
                    "kind": "source_drift",
                    "line": line,
                    "reason": f"line adds words not grounded in source: {unique_preview}",
                }
            )
    return issues


def _find_author_gender_script_issues(script_text, source_text=None):
    issues = []
    for line_no, raw_line in enumerate((script_text or "").splitlines(), start=1):
        line = raw_line.strip()
        if not line:
            continue
        match = AUTHOR_GENDER_ROLE_PATTERN.match(line)
        if not match:
            issues.append(
                {
                    "line_no": line_no,
                    "kind": "unlabeled",
                    "line": line,
                    "reason": "line does not start with Author/Male/Female",
                }
            )
            continue

        role_label = match.group(1).strip()
        text = match.group(2).strip()
        if role_label.lower() not in {LIVE_ROLE_MALE.lower(), LIVE_ROLE_FEMALE.lower()}:
            continue
        if not text:
            issues.append(
                {
                    "line_no": line_no,
                    "kind": "empty",
                    "line": line,
                    "reason": "character line is empty",
                }
            )
            continue
        if _has_author_gender_authorial_tail(text):
            issues.append(
                {
                    "line_no": line_no,
                    "kind": "authorial_tail",
                    "line": line,
                    "reason": "character line contains dialogue mixed with author attribution or narration",
                }
            )
    if source_text:
        issues.extend(_collect_author_gender_source_drift_issues(script_text, source_text))
    return issues


def _build_author_gender_repair_prompt(raw_text, draft_script, issues, profile_prompt, extra_directive=""):
    profile_line = profile_prompt or PREPROCESS_PROFILE_OPTIONS["Бережно"]
    extra_line = (extra_directive or "").strip() or DEFAULT_PREPROCESS_DIRECTIVE
    issue_lines = "\n".join(
        f"- line {item['line_no']}: {item['reason']} :: {item['line']}"
        for item in (issues or [])[:8]
    ) or "- The draft does not follow the strict Author/Male/Female format."
    return (
        "You are repairing a Russian audiobook role script.\n"
        "Rewrite the draft into a strict labeled script using exactly these labels: "
        f"{LIVE_ROLE_AUTHOR}, {LIVE_ROLE_MALE}, {LIVE_ROLE_FEMALE}.\n"
        "Hard rules:\n"
        "- Preserve plot facts, meaning, names and chronology.\n"
        "- Keep the script in Russian.\n"
        "- Do not summarize, shorten, paraphrase or explain the text.\n"
        "- Do not invent any new words, filler phrases, reactions or connective text that are absent from the source.\n"
        "- Do not invent standalone cues like `[Да]`, `[Нет]`, `[Хм]` unless that exact spoken word is present in the source at that point.\n"
        "- Use the Russian letter Ё/ё wherever it is orthographically appropriate and unambiguous; do not replace Ё/ё with Е/е in words such as всё, ещё, её, идёт, шёл, нём.\n"
        f"- Direction profile: {profile_line}\n"
        f"- Additional director note: {extra_line}\n"
        f"- `{LIVE_ROLE_MALE}:` and `{LIVE_ROLE_FEMALE}:` may contain only words actually spoken aloud by that character.\n"
        f"- `{LIVE_ROLE_AUTHOR}:` must contain narration, attribution, inner thoughts, scene description, voice description and all ambiguous dialogue.\n"
        "- Split mixed lines into separate labeled lines.\n"
        "- Never leave tails like `— сказал он`, `— спросила девушка`, `раздался густой бас`, `послышался голос`, `усмехнулся мужчина` inside Male/Female lines.\n"
        "- If gender is not explicit in the source, move the whole line to Author.\n"
        "- Return only the corrected final script, with no markdown and no comments.\n"
        "Examples:\n"
        f"- WRONG: {LIVE_ROLE_MALE}: — Парень, посторонись немного, — раздался густой бас.\n"
        f"  RIGHT: {LIVE_ROLE_MALE}: — Парень, посторонись немного.\n"
        f"  RIGHT: {LIVE_ROLE_AUTHOR}: Раздался густой бас.\n"
        f"- WRONG: {LIVE_ROLE_FEMALE}: «Ты идёшь?» — спросила девушка.\n"
        f"  RIGHT: {LIVE_ROLE_FEMALE}: «Ты идёшь?»\n"
        f"  RIGHT: {LIVE_ROLE_AUTHOR}: — спросила девушка.\n\n"
        f"PROBLEMS TO FIX:\n{issue_lines}\n\n"
        f"SOURCE TEXT:\n{raw_text}\n\n"
        f"DRAFT SCRIPT:\n{draft_script}"
    )


def _append_gendered_segment(segments, kind, text, gender=None):
    normalized = re.sub(r"\s+", " ", (text or "").strip())
    if not normalized:
        return

    is_punctuation_tail = bool(re.fullmatch(r"[.,!?…:;)\]»]+", normalized))
    if segments and is_punctuation_tail:
        segments[-1]["text"] = f"{segments[-1]['text'].rstrip()}{normalized}"
        return

    if segments and segments[-1]["kind"] == kind and segments[-1].get("gender") == gender:
        segments[-1]["text"] = f"{segments[-1]['text'].rstrip()} {normalized}"
        return

    segments.append({"kind": kind, "gender": gender, "text": normalized})


def _split_gendered_paragraph(paragraph, previous_dialogue_gender=None):
    paragraph = re.sub(r"\s+", " ", (paragraph or "").strip())
    if not paragraph:
        return [], previous_dialogue_gender

    existing_label_match = ROLE_SPEAKER_PATTERN.match(paragraph)
    if existing_label_match:
        speaker = existing_label_match.group(1).strip()
        text = existing_label_match.group(2).strip()
        if speaker.lower() == TTS_SPEAKER_DIALOGUE.lower():
            return ([{"kind": "dialogue", "gender": previous_dialogue_gender, "text": text}], previous_dialogue_gender)
        return ([{"kind": "narrator", "gender": None, "text": text}], previous_dialogue_gender)

    if re.match(r"^\s*[—–-]\s*", paragraph):
        core = re.sub(r"^\s*[—–-]\s*", "", paragraph).strip()
        raw_parts = [part.strip() for part in re.split(r"\s+[—–-]\s+", core) if part.strip()]
        if len(raw_parts) >= 2:
            segments = []
            running_gender = previous_dialogue_gender
            for idx, part in enumerate(raw_parts):
                if idx % 2 == 0:
                    local_gender = running_gender
                    if idx > 0:
                        local_gender = _infer_dialogue_gender(raw_parts[idx - 1]) or local_gender
                    if idx + 1 < len(raw_parts):
                        local_gender = _infer_dialogue_gender(raw_parts[idx + 1]) or local_gender
                    if local_gender:
                        running_gender = local_gender
                    _append_gendered_segment(segments, "dialogue", part, local_gender)
                else:
                    narrator_gender = _infer_dialogue_gender(part)
                    if narrator_gender:
                        running_gender = narrator_gender
                    _append_gendered_segment(segments, "narrator", part)
            return segments, running_gender
        return ([{"kind": "dialogue", "gender": previous_dialogue_gender, "text": core}], previous_dialogue_gender)

    quote_matches = list(ROLE_QUOTE_PATTERN.finditer(paragraph))
    if quote_matches:
        segments = []
        cursor = 0
        running_gender = previous_dialogue_gender
        for idx, match in enumerate(quote_matches):
            before = paragraph[cursor:match.start()]
            _append_gendered_segment(segments, "narrator", before)

            next_start = quote_matches[idx + 1].start() if idx + 1 < len(quote_matches) else len(paragraph)
            after_context = paragraph[match.end():next_start]
            quote_text = (match.group(1) or match.group(2) or "").strip()
            local_gender = _infer_dialogue_gender(before) or _infer_dialogue_gender(after_context) or running_gender
            if local_gender:
                running_gender = local_gender
            _append_gendered_segment(segments, "dialogue", quote_text, local_gender)
            cursor = match.end()
        _append_gendered_segment(segments, "narrator", paragraph[cursor:])
        return segments, running_gender

    return ([{"kind": "narrator", "gender": None, "text": paragraph}], previous_dialogue_gender)


def _build_author_gender_live_plan(raw_text, author_voice, male_voice, female_voice):
    labeled_entries = _parse_author_gender_script(raw_text)
    if labeled_entries and not _find_author_gender_script_issues(raw_text):
        plan = []
        for entry in labeled_entries:
            role_hint = entry["role_hint"]
            voice_name = author_voice
            if role_hint == "male_dialogue":
                voice_name = male_voice
            elif role_hint == "female_dialogue":
                voice_name = female_voice

            prepared_entry = _prepare_live_segment_text(entry.get("text", "").strip())
            text_value = prepared_entry.get("text", "").strip()
            if not text_value:
                continue

            plan.append(
                {
                    "voice": voice_name,
                    "role_hint": role_hint,
                    "text": text_value,
                    "segment_directive": prepared_entry.get("style_hint", ""),
                    "pause_after_ms": int(prepared_entry.get("pause_after_ms", 0) or 0),
                }
            )
        if plan:
            return plan
    return []


def _clean_tts_markup(text):
    cleaned = re.sub(rf"^\s*(?:{TTS_SPEAKER_NARRATOR}|{TTS_SPEAKER_DIALOGUE})\s*:\s*", "", text, flags=re.MULTILINE)
    cleaned = re.sub(r"\[[^\]\n]{1,40}\]", " ", cleaned)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _split_long_tts_line(prefix, content, max_chars):
    content = content.strip()
    if not content:
        return []
    if len(prefix) + len(content) + 2 <= max_chars:
        return [f"{prefix}: {content}"]

    sentences = _sentence_tokenize(content)
    if not sentences:
        sentences = [content]

    chunks = []
    current = []
    current_len = len(prefix) + 2
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        addition = len(sentence) + (1 if current else 0)
        if current and current_len + addition > max_chars:
            chunks.append(f"{prefix}: {' '.join(current)}")
            current = [sentence]
            current_len = len(prefix) + 2 + len(sentence)
        else:
            current.append(sentence)
            current_len += addition
    if current:
        chunks.append(f"{prefix}: {' '.join(current)}")
    return chunks


def _split_tts_script(script_text, voice_mode, max_chars):
    normalized_text = (script_text or "").strip()
    if not normalized_text:
        return []

    max_chars = max(900, int(max_chars))
    if voice_mode == "duo":
        raw_lines = [line.strip() for line in normalized_text.splitlines() if line.strip()]
        prepared_lines = []
        for line in raw_lines:
            if ":" in line:
                prefix, content = line.split(":", 1)
                prefix = prefix.strip()
                content = content.strip()
                if prefix in {TTS_SPEAKER_NARRATOR, TTS_SPEAKER_DIALOGUE}:
                    prepared_lines.extend(_split_long_tts_line(prefix, content, max_chars))
                    continue
            prepared_lines.extend(_split_long_tts_line(TTS_SPEAKER_NARRATOR, line, max_chars))

        chunks = []
        current_lines = []
        current_len = 0
        for line in prepared_lines:
            candidate_len = current_len + len(line) + (1 if current_lines else 0)
            if current_lines and candidate_len > max_chars:
                chunks.append("\n".join(current_lines))
                current_lines = [line]
                current_len = len(line)
            else:
                current_lines.append(line)
                current_len = candidate_len
        if current_lines:
            chunks.append("\n".join(current_lines))
        return chunks

    paragraphs = [part.strip() for part in re.split(r"\n\s*\n", normalized_text) if part.strip()]
    if not paragraphs:
        paragraphs = [normalized_text]

    chunks = []
    current_parts = []
    current_len = 0
    for paragraph in paragraphs:
        if len(paragraph) > max_chars:
            split_parts = []
            current_sentences = []
            current_sentence_len = 0
            for sentence in _sentence_tokenize(paragraph):
                sentence = sentence.strip()
                if not sentence:
                    continue
                addition = len(sentence) + (1 if current_sentences else 0)
                if current_sentences and current_sentence_len + addition > max_chars:
                    split_parts.append(" ".join(current_sentences))
                    current_sentences = [sentence]
                    current_sentence_len = len(sentence)
                else:
                    current_sentences.append(sentence)
                    current_sentence_len += addition
            if current_sentences:
                split_parts.append(" ".join(current_sentences))
        else:
            split_parts = [paragraph]

        for part in split_parts:
            candidate_len = current_len + len(part) + (2 if current_parts else 0)
            if current_parts and candidate_len > max_chars:
                chunks.append("\n\n".join(current_parts))
                current_parts = [part]
                current_len = len(part)
            else:
                current_parts.append(part)
                current_len = candidate_len

    if current_parts:
        chunks.append("\n\n".join(current_parts))
    return chunks


def _script_matches_voice_mode(script_text, voice_mode):
    script_text = (script_text or "").strip()
    if not script_text:
        return False

    has_duo_labels = bool(
        re.search(
            rf"^\s*(?:{TTS_SPEAKER_NARRATOR}|{TTS_SPEAKER_DIALOGUE})\s*:",
            script_text,
            flags=re.MULTILINE,
        )
    )
    has_author_gender_labels = bool(
        re.search(
            rf"^\s*(?:{LIVE_ROLE_AUTHOR}|{LIVE_ROLE_MALE}|{LIVE_ROLE_FEMALE})\s*:",
            script_text,
            flags=re.MULTILINE,
        )
    )
    if voice_mode == "duo":
        return has_duo_labels
    if voice_mode == "author_gender":
        return bool(
            has_author_gender_labels
            and _parse_author_gender_script(script_text)
            and not _find_author_gender_script_issues(script_text)
        )
    return not has_duo_labels and not has_author_gender_labels


class InvalidApiKeyError(RuntimeError):
    pass


class ReaderWorkerStopped(RuntimeError):
    pass


class RateLimitBudgetError(RuntimeError):
    def __init__(self, message, model_id=None):
        super().__init__(message)
        self.model_id = model_id


class ProjectRateLimitReachedError(RuntimeError):
    def __init__(self, message, model_id=None):
        super().__init__(message)
        self.model_id = model_id


def _is_invalid_api_key_error(exc):
    text = str(exc or "")
    return "API_KEY_INVALID" in text or "API key not valid" in text


def _mask_api_key(api_key):
    api_key = (api_key or "").strip()
    if not api_key:
        return "<пустой>"
    if len(api_key) <= 10:
        return f"{api_key[:3]}...{api_key[-2:]}"
    return f"{api_key[:6]}...{api_key[-4:]}"


def _format_local_timestamp(timestamp):
    if not timestamp:
        return ""
    try:
        return datetime.fromtimestamp(float(timestamp)).astimezone().strftime("%H:%M (%d.%m)")
    except Exception:
        return ""


def _ffmpeg_concat_path(path):
    normalized = os.path.abspath(path).replace("\\", "/")
    return normalized.replace("'", "'\\''")


def _format_chapter_scope_slug(chapter_indices):
    chapter_numbers = sorted({idx + 1 for idx in (chapter_indices or []) if isinstance(idx, int) and idx >= 0})
    if not chapter_numbers:
        return ""

    ranges = []
    start = chapter_numbers[0]
    prev = chapter_numbers[0]
    for number in chapter_numbers[1:]:
        if number == prev + 1:
            prev = number
            continue
        ranges.append(f"{start}-{prev}" if start != prev else str(start))
        start = prev = number
    ranges.append(f"{start}-{prev}" if start != prev else str(start))
    return "Ch" + "_".join(ranges)


def _detect_leading_silence_ms(
    segment,
    silence_threshold=CHAPTER_TRIM_SILENCE_THRESHOLD_DB,
    chunk_size=CHAPTER_TRIM_SILENCE_STEP_MS,
):
    if segment is None:
        return 0
    trim_ms = 0
    segment_len = len(segment)
    while trim_ms < segment_len:
        chunk = segment[trim_ms:trim_ms + chunk_size]
        if len(chunk) <= 0:
            break
        if chunk.dBFS != float("-inf") and chunk.dBFS > silence_threshold:
            break
        if chunk.dBFS == float("-inf") or chunk.dBFS <= silence_threshold:
            trim_ms += chunk_size
            continue
        break
    return min(trim_ms, segment_len)


def _trim_audio_segment_boundaries(
    segment,
    silence_threshold=CHAPTER_TRIM_SILENCE_THRESHOLD_DB,
    chunk_size=CHAPTER_TRIM_SILENCE_STEP_MS,
    keep_ms=CHAPTER_TRIM_KEEP_MS,
):
    if segment is None or len(segment) <= 0:
        return segment

    leading_ms = _detect_leading_silence_ms(segment, silence_threshold=silence_threshold, chunk_size=chunk_size)
    trailing_ms = _detect_leading_silence_ms(
        segment.reverse(),
        silence_threshold=silence_threshold,
        chunk_size=chunk_size,
    )
    if leading_ms <= 0 and trailing_ms <= 0:
        return segment

    start_ms = max(0, leading_ms - keep_ms)
    end_ms = len(segment) - max(0, trailing_ms - keep_ms)
    if end_ms <= start_ms:
        return segment
    return segment[start_ms:end_ms]


def _load_trimmed_mp3_segment(path):
    if AudioSegment is None:
        raise RuntimeError("Для обрезки пауз нужен pydub.")
    segment = AudioSegment.from_file(path, format="mp3")
    return _trim_audio_segment_boundaries(segment)


def _export_trimmed_mp3_file(source_path, output_path):
    segment = _load_trimmed_mp3_segment(source_path)
    tmp_path = f"{output_path}.tmp"
    try:
        segment.export(tmp_path, format="mp3")
        os.replace(tmp_path, output_path)
    except Exception as exc:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass
        raise RuntimeError(f"Не удалось сохранить MP3 после обрезки пауз: {exc}") from exc


def _export_pcm_to_mp3(raw_audio, output_path):
    if AudioSegment is None:
        raise RuntimeError("Невозможно сохранить MP3: не найден pydub.")
    if not raw_audio:
        raise RuntimeError("Невозможно сохранить пустой аудиоблок.")

    tmp_path = f"{output_path}.tmp"
    try:
        snap = AudioSegment(data=raw_audio, sample_width=2, frame_rate=AUDIO_RATE, channels=AUDIO_CHANNELS)
        snap = _trim_audio_segment_boundaries(snap)
        snap.export(tmp_path, format="mp3")
        os.replace(tmp_path, output_path)
    except Exception as exc:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass
        raise RuntimeError(f"Не удалось сохранить MP3: {exc}") from exc


def _load_mp3_as_raw_pcm(path):
    if AudioSegment is None:
        raise RuntimeError("Невозможно загрузить частичный MP3: не найден pydub.")
    segment = AudioSegment.from_file(path, format="mp3")
    segment = segment.set_frame_rate(AUDIO_RATE).set_channels(AUDIO_CHANNELS).set_sample_width(2)
    return segment.raw_data


def _trim_raw_pcm_boundaries(raw_audio, keep_ms=LIVE_REQUEST_TRIM_KEEP_MS):
    if not raw_audio or AudioSegment is None:
        return raw_audio
    try:
        snap = AudioSegment(data=raw_audio, sample_width=2, frame_rate=AUDIO_RATE, channels=AUDIO_CHANNELS)
        snap = _trim_audio_segment_boundaries(snap, keep_ms=keep_ms)
        return snap.raw_data
    except Exception:
        return raw_audio


def _pcm_silence(duration_ms):
    duration_ms = int(duration_ms or 0)
    if duration_ms <= 0:
        return b""
    samples = max(1, int(AUDIO_RATE * duration_ms / 1000.0))
    return b"\x00\x00" * samples * AUDIO_CHANNELS


def _combine_mp3_sequence(input_paths, output_path):
    normalized_paths = [os.path.abspath(path) for path in (input_paths or []) if path and os.path.exists(path)]
    if not normalized_paths:
        raise RuntimeError("Нет MP3-файлов для сборки.")

    tmp_output = f"{output_path}.tmp"
    ffmpeg_path = _resolve_tool_path("ffmpeg")
    concat_list_path = f"{output_path}.concat.txt"

    if ffmpeg_path:
        try:
            with open(concat_list_path, "w", encoding="utf-8") as list_file:
                for file_path in normalized_paths:
                    list_file.write(f"file '{_ffmpeg_concat_path(file_path)}'\n")
            combine_cmd = [
                ffmpeg_path,
                "-y",
                "-f", "concat",
                "-safe", "0",
                "-i", concat_list_path,
                "-c", "copy",
                tmp_output,
            ]
            result = _run_subprocess(
                combine_cmd,
                capture_output=True,
                text=True,
                timeout=READER_FFMPEG_CONCAT_TIMEOUT_SEC,
            )
            if result.returncode == 0 and os.path.exists(tmp_output):
                os.replace(tmp_output, output_path)
                return
            raise RuntimeError(result.stderr.strip() or "ffmpeg не смог собрать итоговый MP3.")
        except Exception:
            if os.path.exists(tmp_output):
                try:
                    os.remove(tmp_output)
                except Exception:
                    pass
        finally:
            if os.path.exists(concat_list_path):
                try:
                    os.remove(concat_list_path)
                except Exception:
                    pass

    if AudioSegment is None:
        raise RuntimeError("Для сборки MP3 нужны ffmpeg или pydub.")

    combined = None
    try:
        for file_path in normalized_paths:
            segment = AudioSegment.from_file(file_path, format="mp3")
            combined = segment if combined is None else (combined + segment)
        if combined is None:
            raise RuntimeError("Нет валидных MP3-файлов для сборки.")
        combined.export(tmp_output, format="mp3")
        os.replace(tmp_output, output_path)
    except Exception as exc:
        if os.path.exists(tmp_output):
            try:
                os.remove(tmp_output)
            except Exception:
                pass
        raise RuntimeError(f"Не удалось собрать MP3: {exc}") from exc


def _is_rate_limited_error(exc):
    text = str(exc or "")
    text_low = text.lower()
    return (
        "resource_exhausted" in text_low
        or "quota" in text_low
        or "429" in text_low
        or "rate limit" in text_low
    )


def _lookup_model_limit(model_id, field_name, fallback=None):
    if not model_id or reader_api_config is None:
        if field_name == "rpd":
            return RPD_LIMIT_FALLBACKS.get(model_id, fallback)
        return fallback
    try:
        provider_cfg = reader_api_config.api_providers().get("gemini", {})
        for model_cfg in provider_cfg.get("models", {}).values():
            if str(model_cfg.get("id") or "").strip() == model_id:
                value = model_cfg.get(field_name, fallback)
                if value is None and field_name == "rpd":
                    return RPD_LIMIT_FALLBACKS.get(model_id, fallback)
                return fallback if value is None else value
    except Exception:
        if field_name == "rpd":
            return RPD_LIMIT_FALLBACKS.get(model_id, fallback)
        return fallback
    if field_name == "rpd":
        return RPD_LIMIT_FALLBACKS.get(model_id, fallback)
    return fallback


def _make_rpm_limiter(model_id, fallback_rpm):
    rpm_value = _lookup_model_limit(model_id, "rpm", fallback_rpm)
    try:
        rpm_value = int(rpm_value)
    except Exception:
        rpm_value = fallback_rpm
    if RPMLimiter is None:
        return None
    return RPMLimiter(max(1, rpm_value))


class TPMLimiter:
    def __init__(self, tpm_limit=0):
        try:
            self.tpm_limit = max(0, int(tpm_limit or 0))
        except Exception:
            self.tpm_limit = 0
        self._events = []
        self._lock = threading.Lock()

    def _prune(self, now_ts):
        cutoff = now_ts - 60.0
        self._events = [(ts, tokens) for ts, tokens in self._events if ts > cutoff]

    def get_required_delay(self, token_count):
        if self.tpm_limit <= 0:
            return 0.0
        token_count = max(1, int(token_count or 1))
        with self._lock:
            now_ts = time.time()
            self._prune(now_ts)
            used_tokens = sum(tokens for _, tokens in self._events)
            if used_tokens + token_count <= self.tpm_limit:
                return 0.0

            running_total = used_tokens + token_count
            for event_ts, event_tokens in self._events:
                running_total -= event_tokens
                if running_total <= self.tpm_limit:
                    return max(0.25, 60.0 - (now_ts - event_ts))
            return 60.0

    def register(self, token_count):
        if self.tpm_limit <= 0:
            return
        token_count = max(1, int(token_count or 1))
        with self._lock:
            now_ts = time.time()
            self._prune(now_ts)
            self._events.append((now_ts, token_count))


def _gemini_reset_policy():
    if reader_api_config is not None:
        try:
            provider_cfg = reader_api_config.api_providers().get("gemini", {})
            policy = provider_cfg.get("reset_policy", {})
            if policy:
                return dict(policy)
        except Exception:
            pass
    return {
        "type": "daily",
        "timezone": "America/Los_Angeles",
        "reset_hour": 0,
        "reset_minute": 1,
    }


def _policy_timezone(policy):
    tz_name = (policy or {}).get("timezone") or "America/Los_Angeles"
    if ZoneInfo is not None:
        try:
            return ZoneInfo(tz_name)
        except Exception:
            return None
    return None


def _current_policy_bucket(policy):
    policy = policy or _gemini_reset_policy()
    tz = _policy_timezone(policy)
    now_dt = datetime.now(tz) if tz is not None else datetime.utcnow()

    if policy.get("type") == "daily":
        boundary = now_dt.replace(
            hour=policy.get("reset_hour", 0),
            minute=policy.get("reset_minute", 1),
            second=0,
            microsecond=0,
        )
        if now_dt < boundary:
            boundary -= timedelta(days=1)
        return boundary.strftime("%Y-%m-%dT%H:%M")

    return now_dt.strftime("%Y-%m-%d")


def _next_policy_reset_text(policy):
    policy = policy or _gemini_reset_policy()
    tz = _policy_timezone(policy)
    now_dt = datetime.now(tz) if tz is not None else datetime.utcnow()

    if policy.get("type") == "daily":
        next_reset = now_dt.replace(
            hour=policy.get("reset_hour", 0),
            minute=policy.get("reset_minute", 1),
            second=0,
            microsecond=0,
        )
        if next_reset <= now_dt:
            next_reset += timedelta(days=1)
        return next_reset.strftime("%H:%M (%d.%m)")

    return "следующее окно"


class ProjectDailyRequestLimiter:
    def __init__(self, settings_manager=None):
        self.settings_manager = settings_manager
        self.policy = _gemini_reset_policy()
        self.lock = threading.Lock()
        self.state = self._load_state()

    def _load_state(self):
        bucket = _current_policy_bucket(self.policy)
        default_state = {"bucket": bucket, "counts": {}, "counts_by_key": {}}
        if self.settings_manager is not None:
            try:
                stored = self.settings_manager.load_settings().get(READER_RPD_STATE_KEY, {})
                if isinstance(stored, dict):
                    default_state.update(stored)
            except Exception:
                pass
        else:
            legacy_data = _load_legacy_settings()
            stored = legacy_data.get(READER_RPD_STATE_KEY, {})
            if isinstance(stored, dict):
                default_state.update(stored)
        if not isinstance(default_state.get("counts"), dict):
            default_state["counts"] = {}
        if not isinstance(default_state.get("counts_by_key"), dict):
            default_state["counts_by_key"] = {}
        return default_state

    def _save_state_locked(self):
        payload = {
            "bucket": self.state.get("bucket"),
            "counts": dict(self.state.get("counts", {})),
            "counts_by_key": {
                model_id: dict(key_counts)
                for model_id, key_counts in self.state.get("counts_by_key", {}).items()
                if isinstance(key_counts, dict)
            },
        }
        if self.settings_manager is not None:
            try:
                self.settings_manager.save_ui_state({READER_RPD_STATE_KEY: payload})
                return
            except Exception:
                pass

        legacy_data = _load_legacy_settings()
        legacy_data[READER_RPD_STATE_KEY] = payload
        with open(LEGACY_SETTINGS_FILE, "w", encoding="utf-8") as file_obj:
            json.dump(legacy_data, file_obj, ensure_ascii=False, indent=2)

    def _rollover_locked(self):
        current_bucket = _current_policy_bucket(self.policy)
        if self.state.get("bucket") != current_bucket:
            self.state = {"bucket": current_bucket, "counts": {}, "counts_by_key": {}}
            self._save_state_locked()

    def _local_count_key(self, api_key):
        return (api_key or "").strip() or "__project__"

    def _get_local_count_locked(self, model_id, api_key=None):
        counts_by_key = self.state.setdefault("counts_by_key", {})
        key_counts = counts_by_key.get(model_id)
        if isinstance(key_counts, dict):
            if api_key:
                return int(key_counts.get(self._local_count_key(api_key), 0) or 0)
            return sum(int(value or 0) for value in key_counts.values())
        return int(self.state.get("counts", {}).get(model_id, 0) or 0)

    def _set_local_count_locked(self, model_id, api_key, value):
        counts_by_key = self.state.setdefault("counts_by_key", {})
        key_counts = counts_by_key.setdefault(model_id, {})
        key_counts[self._local_count_key(api_key)] = max(0, int(value or 0))

    def _settings_count(self, api_key, model_id):
        if self.settings_manager is None or not api_key:
            return None
        try:
            key_info = self.settings_manager.get_key_info(api_key)
            if not key_info:
                return None
            return int(self.settings_manager.get_request_count(key_info, model_id) or 0)
        except Exception:
            return None

    def try_acquire(self, model_id, daily_limit, amount=1, api_key=None):
        if not model_id or not daily_limit or daily_limit <= 0:
            return True, 0, daily_limit, _next_policy_reset_text(self.policy)

        amount = max(1, int(amount or 1))
        with self.lock:
            self._rollover_locked()
            settings_count = self._settings_count(api_key, model_id)
            current_count = (
                settings_count
                if settings_count is not None
                else self._get_local_count_locked(model_id, api_key)
            )
            if current_count + amount > daily_limit:
                return False, current_count, daily_limit, _next_policy_reset_text(self.policy)

            if self.settings_manager is not None and api_key:
                updated = False
                for _ in range(amount):
                    updated = self.settings_manager.increment_request_count(api_key, model_id) or updated
                if updated:
                    return True, current_count + amount, daily_limit, _next_policy_reset_text(self.policy)

            self._set_local_count_locked(model_id, api_key, current_count + amount)
            self._save_state_locked()
            return True, current_count + amount, daily_limit, _next_policy_reset_text(self.policy)

    def release(self, model_id, amount=1, api_key=None):
        if not model_id:
            return False

        amount = max(1, int(amount or 1))
        with self.lock:
            self._rollover_locked()
            if self.settings_manager is not None and api_key:
                changed = False
                for _ in range(amount):
                    changed = self.settings_manager.decrement_request_count(api_key, model_id) or changed
                if changed:
                    return True

            current_count = self._get_local_count_locked(model_id, api_key)
            self._set_local_count_locked(model_id, api_key, current_count - amount)
            self._save_state_locked()
            return True

    def get_count(self, model_id, api_key=None):
        with self.lock:
            self._rollover_locked()
            settings_count = self._settings_count(api_key, model_id)
            if settings_count is not None:
                return settings_count
            return self._get_local_count_locked(model_id, api_key)


def _sentence_tokenize(text):
    fallback_sentences = [
        part.strip()
        for part in re.split(r"(?<=[.!?])\s+|\n+", text)
        if part.strip()
    ]
    if nltk is None:
        return fallback_sentences

    try:
        return nltk.sent_tokenize(text)
    except LookupError:
        try:
            nltk.download("punkt", quiet=True)
            return nltk.sent_tokenize(text)
        except Exception:
            return fallback_sentences
    except Exception:
        return fallback_sentences


def _normalize_reader_text(text):
    cleaned = re.sub(r'[^\x20-\x7E\u0400-\u04FF\s\.,!\?\-:;\"\'«»—]', '', text or "")
    cleaned = re.sub(r"\s+", " ", cleaned)
    return cleaned.strip()


def _merge_reader_paragraphs(raw_paragraphs, min_alnum=15):
    paragraphs = []
    pending = ""

    for raw_paragraph in raw_paragraphs or []:
        paragraph = _normalize_reader_text(raw_paragraph)
        if not paragraph:
            continue

        if pending:
            candidate = f"{pending}\n\n{paragraph}"
        else:
            candidate = paragraph

        alnum_count = sum(ch.isalnum() for ch in candidate)
        if alnum_count >= min_alnum:
            paragraphs.append(candidate)
            pending = ""
        else:
            pending = candidate

    if pending:
        if paragraphs:
            paragraphs[-1] = f"{paragraphs[-1]}\n\n{pending}"
        else:
            paragraphs.append(pending)

    return paragraphs


def _split_live_paragraph(paragraph, max_chars=LIVE_PARAGRAPH_MAX_CHARS):
    paragraph = (paragraph or "").strip()
    if not paragraph or len(paragraph) <= max_chars:
        return [paragraph] if paragraph else []

    sentences = _sentence_tokenize(paragraph)
    if len(sentences) <= 1:
        return [paragraph]

    chunks = []
    current = []
    current_len = 0
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        addition = len(sentence) + (1 if current else 0)
        if current and current_len + addition > max_chars:
            chunks.append(" ".join(current))
            current = [sentence]
            current_len = len(sentence)
        else:
            current.append(sentence)
            current_len += addition
    if current:
        chunks.append(" ".join(current))
    return chunks or [paragraph]


if platform.system() == "Windows" and "_orig_popen" in globals():
    subprocess.Popen = _orig_popen

if _loguru_logger is not None:
    logger = _loguru_logger
    logger.remove()
    logger.add(custom_log_handler, level="INFO")
else:
    logger = logging.getLogger("gemini_reader")
    if not getattr(logger, "_gemini_reader_configured", False):
        logger.setLevel(logging.INFO)
        logger.handlers.clear()
        logger.addHandler(_GuiLoggingHandler())
        logger.propagate = False
        logger._gemini_reader_configured = True

# --- ПОДГОТОВКА ТЕКСТА ---
def _is_supported_reader_book_path(path):
    return isinstance(path, str) and path.lower().endswith(READER_SUPPORTED_BOOK_EXTENSIONS)


def _reader_fallback_title(raw_title, fallback):
    title = _normalize_reader_text(raw_title or "")
    return title or fallback


def _reader_natural_sort_key(value):
    return [int(token) if token.isdigit() else token.lower() for token in re.split(r"(\d+)", value)]


def _reader_paragraphs_to_html(paragraphs):
    html_parts = []
    for paragraph in paragraphs:
        normalized = _normalize_reader_text(paragraph)
        if normalized:
            html_parts.append(f"<p>{html.escape(normalized)}</p>")
    return "".join(html_parts)


def _reader_plain_text_to_html(text):
    paragraphs = []
    current_lines = []
    for raw_line in (text or "").splitlines():
        line = _normalize_reader_text(raw_line)
        if line:
            current_lines.append(line)
        elif current_lines:
            paragraphs.append(" ".join(current_lines))
            current_lines = []
    if current_lines:
        paragraphs.append(" ".join(current_lines))
    return _reader_paragraphs_to_html(paragraphs)


def _reader_parse_epub_chapters(filepath):
    if epub is None or ebooklib is None:
        raise RuntimeError("Для импорта EPUB требуется пакет ebooklib.")

    chapters = []
    book = epub.read_epub(filepath)
    for item_ref in book.spine:
        item = book.get_item_with_id(item_ref[0])
        if item and (item.get_type() == ebooklib.ITEM_DOCUMENT or item.get_name().lower().endswith((".xhtml", ".html", ".htm"))):
            raw_title = ""
            if BeautifulSoup is not None:
                soup = BeautifulSoup(item.get_content(), "html.parser")
                heading = soup.find(["h1", "h2", "h3"])
                if heading:
                    raw_title = heading.get_text(" ", strip=True)
            chapter = Chapter(
                _reader_fallback_title(raw_title, f"Chapter {len(chapters) + 1}"),
                item.get_content(),
            )
            if chapter.flat_sentences:
                chapters.append(chapter)
    return chapters


def _reader_parse_zip_docx_chapters(filepath):
    if Document is None:
        raise RuntimeError("Для импорта ZIP(DOCX) требуется пакет python-docx.")

    chapters = []
    with zipfile.ZipFile(filepath, "r") as archive:
        docx_files = sorted(
            [
                name
                for name in archive.namelist()
                if name.lower().endswith(".docx") and not os.path.basename(name).startswith("~")
            ],
            key=_reader_natural_sort_key,
        )
        for index, name in enumerate(docx_files, start=1):
            doc = Document(io.BytesIO(archive.read(name)))
            paragraphs = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
            if paragraphs:
                last_paragraph = paragraphs[-1].lower()
                if any(marker in last_paragraph for marker in ("rulate", "boosty", "http", "patreon", "t.me")):
                    paragraphs.pop()
            content = _reader_paragraphs_to_html(paragraphs)
            if content:
                title = os.path.splitext(os.path.basename(name))[0].replace("_", " ")
                chapters.append(Chapter(_reader_fallback_title(title, f"Chapter {index}"), content))
    return chapters


def _reader_parse_text_chapters(filepath):
    with open(filepath, "r", encoding="utf-8", errors="ignore") as file_obj:
        text = file_obj.read()

    chunks = re.split(r"^\s*#\s*\[(.*?)\]", text, flags=re.MULTILINE)
    chapters = []
    if len(chunks) <= 1:
        content = _reader_plain_text_to_html(text)
        if content:
            fallback_title = os.path.splitext(os.path.basename(filepath))[0]
            return [Chapter(_reader_fallback_title(fallback_title, "Chapter 1"), content)]
        return []

    for index in range(1, len(chunks), 2):
        title = chunks[index].strip()
        content = chunks[index + 1].strip() if index + 1 < len(chunks) else ""
        content_html = _reader_plain_text_to_html(content)
        if content_html:
            chapters.append(Chapter(_reader_fallback_title(title, f"Chapter {len(chapters) + 1}"), content_html))
    return chapters


def _reader_parse_html_chapters(filepath):
    if BeautifulSoup is None:
        raise RuntimeError("Для импорта HTML требуется пакет beautifulsoup4.")

    with open(filepath, "r", encoding="utf-8", errors="ignore") as file_obj:
        markup = file_obj.read()

    soup = BeautifulSoup(markup, "html.parser")
    headings = soup.find_all(["h1", "h2"])
    chapters = []

    if not headings:
        content = _reader_paragraphs_to_html(
            [tag.get_text(" ", strip=True) for tag in soup.find_all(["p", "div", "li"]) if tag.get_text(" ", strip=True)]
        )
        if content:
            raw_title = soup.title.get_text(" ", strip=True) if soup.title else os.path.splitext(os.path.basename(filepath))[0]
            chapters.append(Chapter(_reader_fallback_title(raw_title, "Chapter 1"), content))
        return chapters

    for index, heading in enumerate(headings, start=1):
        raw_title = heading.get_text(" ", strip=True)
        content_lines = []
        for sibling in heading.find_next_siblings():
            if sibling.name in ("h1", "h2"):
                break
            text = sibling.get_text(" ", strip=True)
            if text:
                content_lines.append(text)
        content = _reader_paragraphs_to_html(content_lines)
        if content:
            chapters.append(Chapter(_reader_fallback_title(raw_title, f"Chapter {index}"), content))
    return chapters


def _reader_load_supported_chapters(filepath):
    lower_path = filepath.lower()
    if lower_path.endswith(".epub"):
        return _reader_parse_epub_chapters(filepath)
    if lower_path.endswith(".zip"):
        return _reader_parse_zip_docx_chapters(filepath)
    if lower_path.endswith((".txt", ".md")):
        return _reader_parse_text_chapters(filepath)
    if lower_path.endswith((".html", ".htm")):
        return _reader_parse_html_chapters(filepath)
    raise RuntimeError(f"Неподдерживаемый формат книги: {os.path.basename(filepath)}")


class Chapter:
    def __init__(self, title, content):
        self.title = title
        self.flat_sentences = []
        self.paragraphs = []
        self.raw_text = ""
        self._parse_content(content)

    def _parse_content(self, content):
        if BeautifulSoup is None:
            raise RuntimeError("Для чтения EPUB требуется пакет beautifulsoup4.")
        soup = BeautifulSoup(content, 'html.parser')
        for s in soup(["script", "style", "title"]):
            s.extract()

        lines =[]
        for el in soup.find_all(['p', 'h1', 'h2', 'h3', 'div', 'li']):
            t = el.get_text().strip()
            if t:
                lines.append(t)

        normalized_lines = [_normalize_reader_text(line) for line in lines]
        normalized_lines = [line for line in normalized_lines if line]

        self.paragraphs = _merge_reader_paragraphs(normalized_lines)
        self.raw_text = "\n".join(normalized_lines)
        full_text = self.raw_text

        raw_sentences = _sentence_tokenize(full_text)
            
        merged_sentences =[]
        current_sentence = ""
        
        for s in raw_sentences:
            s = s.strip()
            if not s:
                continue
                
            # Добавляем к текущему буферу
            if current_sentence:
                current_sentence += " " + s
            else:
                current_sentence = s
                
            # Считаем только реальные буквы/цифры (без пробелов, тире, кавычек и точек)
            alnum_count = sum(c.isalnum() for c in current_sentence)
            
            # Если набралось 10 или больше символов — фиксируем как готовое предложение
            if alnum_count >= 10:
                merged_sentences.append(current_sentence)
                current_sentence = ""
                
        # Если в самом конце остался короткий "хвост" (меньше 10 символов)
        if current_sentence:
            if merged_sentences:
                # Приклеиваем его к последнему нормальному предложению
                merged_sentences[-1] += " " + current_sentence
            else:
                # Если во всей главе вообще меньше 10 символов
                merged_sentences.append(current_sentence)
                
        self.flat_sentences = merged_sentences


class BookManager:
    def __init__(self, filepath=None, base_dir=None):
        self.base_dir = base_dir or "books"
        if not os.path.exists(self.base_dir):
            os.makedirs(self.base_dir)
        self.chapters =[]
        self.title = "Unknown"
        self.book_dir = ""
        self.lock = threading.Lock()
        if filepath:
            self._import_book(filepath)

    def _import_book_epub_legacy(self, filepath):
        if epub is None or ebooklib is None:
            raise RuntimeError("Для импорта EPUB требуется пакет ebooklib.")
        filename = os.path.basename(filepath)
        self.title = "".join([c for c in os.path.splitext(filename)[0] if c.isalnum() or c in (' ', '-', '_')]).strip()
        self.book_dir = os.path.join(self.base_dir, self.title)
        
        if not os.path.exists(self.book_dir):
            os.makedirs(self.book_dir)
        
        try:
            shutil.copy(filepath, os.path.join(self.book_dir, filename))
        except Exception as e:
            logger.error(f"Ошибка при копировании файла: {e}")

        try:
            book = epub.read_epub(filepath)
            for item_ref in book.spine:
                item = book.get_item_with_id(item_ref[0])
                # ИЗМЕНЕНИЕ: Улучшенная поддержка файлов .html и .htm, если они не помечены как DOCUMENT
                if item and (item.get_type() == ebooklib.ITEM_DOCUMENT or item.get_name().lower().endswith(('.xhtml', '.html', '.htm'))):
                    chap = Chapter(f"Chapter {len(self.chapters) + 1}", item.get_content())
                    if chap.flat_sentences:
                        self.chapters.append(chap)
        except Exception as e:
            logger.error(f"Epub Error: {e}")

    def _import_book(self, filepath):
        filename = os.path.basename(filepath)
        self.title = "".join([c for c in os.path.splitext(filename)[0] if c.isalnum() or c in (' ', '-', '_')]).strip()
        self.book_dir = os.path.join(self.base_dir, self.title)

        if not os.path.exists(self.book_dir):
            os.makedirs(self.book_dir)

        try:
            shutil.copy(filepath, os.path.join(self.book_dir, filename))
        except Exception as e:
            logger.error(f"РћС€РёР±РєР° РїСЂРё РєРѕРїРёСЂРѕРІР°РЅРёРё С„Р°Р№Р»Р°: {e}")

        try:
            self.chapters = _reader_load_supported_chapters(filepath)
        except Exception as e:
            logger.error(f"Book import error: {e}")
            raise

    def get_paths(self):
        return os.path.join(self.book_dir, "progress_v19.json")

    def save_progress(self, c, s):
        with self.lock:
            try:
                with open(self.get_paths(), 'w') as f:
                    json.dump({"chapter": c, "sentence": s}, f)
            except Exception as e:
                print(f"Save progress error: {e}")

    def load_progress(self):
        with self.lock:
            try:
                with open(self.get_paths(), 'r') as f:
                    d = json.load(f)
                    return d.get("chapter", 0), d.get("sentence", 0)
            except:
                return 0, 0

    def is_chapter_done(self, c_idx):
        return os.path.exists(os.path.join(self.book_dir, f"Ch{c_idx + 1}.done"))

    def mark_chapter_done(self, c_idx):
        path = os.path.join(self.book_dir, f"Ch{c_idx + 1}.done")
        with open(path, 'w') as f:
            f.write("done")
        self.clear_tts_progress(c_idx)
            
        # Удаляем маркер переозвучки, если он был, чтобы статус вернулся на "Готово"
        revoice_path = os.path.join(self.book_dir, f"Ch{c_idx + 1}.revoice")
        if os.path.exists(revoice_path):
            os.remove(revoice_path)


    def is_chapter_skipped(self, c_idx):
        return os.path.exists(os.path.join(self.book_dir, f"Ch{c_idx + 1}.skip"))

    def toggle_chapter_skip(self, c_idx):
        path = os.path.join(self.book_dir, f"Ch{c_idx + 1}.skip")
        if os.path.exists(path):
            os.remove(path)
        else:
            with open(path, 'w') as f:
                f.write("skip")




    def get_mp3_path(self, c):
        return os.path.join(self.book_dir, f"Ch{c + 1}.mp3")

    def get_tts_progress_path(self, c_idx):
        return os.path.join(self.book_dir, f"Ch{c_idx + 1}.tts_progress.json")

    def load_tts_progress(self, c_idx):
        path = self.get_tts_progress_path(c_idx)
        if not os.path.exists(path):
            return {}
        try:
            with open(path, "r", encoding="utf-8") as file_obj:
                data = json.load(file_obj)
            return data if isinstance(data, dict) else {}
        except Exception:
            return {}

    def save_tts_progress(self, c_idx, progress):
        path = self.get_tts_progress_path(c_idx)
        with open(path, "w", encoding="utf-8") as file_obj:
            json.dump(progress or {}, file_obj, ensure_ascii=False, indent=2)
        return path

    def clear_tts_progress(self, c_idx):
        path = self.get_tts_progress_path(c_idx)
        if os.path.exists(path):
            try:
                os.remove(path)
            except Exception:
                pass

    def get_tts_script_path(self, c_idx):
        return os.path.join(self.book_dir, f"Ch{c_idx + 1}{SCRIPT_SUFFIX}")

    def load_tts_script(self, c_idx):
        path = self.get_tts_script_path(c_idx)
        if not os.path.exists(path):
            return ""
        try:
            with open(path, "r", encoding="utf-8") as file_obj:
                return file_obj.read()
        except Exception:
            return ""

    def save_tts_script(self, c_idx, script_text):
        path = self.get_tts_script_path(c_idx)
        with open(path, "w", encoding="utf-8") as file_obj:
            file_obj.write(script_text or "")
        self.clear_tts_progress(c_idx)
        return path

    def has_tts_script(self, c_idx):
        return os.path.exists(self.get_tts_script_path(c_idx))

    @staticmethod
    def _chapter_index_from_marker(name, suffix):
        if not name.startswith("Ch") or not name.endswith(suffix):
            return None
        raw_index = name[2:-len(suffix)]
        if not raw_index.isdigit():
            return None
        index = int(raw_index) - 1
        return index if index >= 0 else None

    def chapter_status_snapshot(self):
        snapshot = {"done": set(), "skipped": set(), "scripts": set()}
        if not self.book_dir or not os.path.isdir(self.book_dir):
            return snapshot

        markers = (
            (".done", "done"),
            (".skip", "skipped"),
            (SCRIPT_SUFFIX, "scripts"),
        )
        try:
            for entry in os.scandir(self.book_dir):
                if not entry.is_file():
                    continue
                for suffix, bucket in markers:
                    index = self._chapter_index_from_marker(entry.name, suffix)
                    if index is not None:
                        snapshot[bucket].add(index)
                        break
        except OSError:
            pass
        return snapshot

    def get_video_cover_path(self):
        for extension in VIDEO_COVER_EXTENSIONS:
            candidate = os.path.join(self.book_dir, f"{VIDEO_COVER_BASENAME}{extension}")
            if os.path.exists(candidate):
                return candidate
        return ""

    def save_video_cover_image(self, source_path):
        if not self.book_dir:
            raise RuntimeError("Книга ещё не загружена.")
        if not source_path or not os.path.isfile(source_path):
            raise FileNotFoundError("Файл картинки не найден.")

        extension = os.path.splitext(source_path)[1].lower()
        if not extension:
            extension = ".png"
        destination_path = os.path.join(self.book_dir, f"{VIDEO_COVER_BASENAME}{extension}")

        for existing_extension in VIDEO_COVER_EXTENSIONS:
            existing_path = os.path.join(self.book_dir, f"{VIDEO_COVER_BASENAME}{existing_extension}")
            if os.path.exists(existing_path) and os.path.abspath(existing_path) != os.path.abspath(destination_path):
                try:
                    os.remove(existing_path)
                except Exception:
                    pass

        if os.path.abspath(source_path) != os.path.abspath(destination_path):
            shutil.copy2(source_path, destination_path)
        return destination_path


# --- УМНЫЙ КОМБАЙН (ПО 1.5 ГБ) ---
class AudioCombinerWorker(QThread):
    finished_signal = pyqtSignal(str)
    progress_signal = pyqtSignal(str)

    def __init__(
        self,
        book_manager,
        voice_name="Unknown",
        ffmpeg_path=None,
        ffprobe_path=None,
        video_image_path=None,
        chapter_indices=None,
    ):
        super().__init__()
        self.bm = book_manager
        self.voice_name = voice_name
        self.ffmpeg_path = ffmpeg_path or _resolve_tool_path("ffmpeg")
        self.ffprobe_path = ffprobe_path or _resolve_tool_path("ffprobe")
        self.video_image_path = video_image_path or ""
        self.chapter_indices = sorted(
            {idx for idx in (chapter_indices or []) if isinstance(idx, int) and idx >= 0}
        )

    def run(self):
        try:
            self.progress_signal.emit("Подготовка файлов...")
            allowed_files = None
            if self.chapter_indices:
                allowed_files = {f"Ch{idx + 1}.mp3" for idx in self.chapter_indices}
            files = sorted(
                [
                    file_name
                    for file_name in os.listdir(self.bm.book_dir)
                    if file_name.startswith("Ch")
                    and file_name.endswith(".mp3")
                    and "Part" not in file_name
                    and (allowed_files is None or file_name in allowed_files)
                ],
                key=lambda name: int(re.search(r'\d+', name).group()),
            )

            if not files:
                if self.chapter_indices:
                    self.finished_signal.emit(
                        "Ошибка: среди выбранных глав нет готовых MP3 для экспорта."
                    )
                else:
                    self.finished_signal.emit("Ошибка: нет файлов Ch*.mp3")
                return

            parts = []
            current_part_files = []
            current_size = 0
            created_videos = 0
            scope_suffix = _format_chapter_scope_slug(self.chapter_indices)
            base_title = self.bm.title if not scope_suffix else f"{self.bm.title}_{scope_suffix}"

            for file_name in files:
                file_path = os.path.join(self.bm.book_dir, file_name)
                file_size = os.path.getsize(file_path)
                if current_size + file_size > MAX_PART_SIZE and current_part_files:
                    parts.append(current_part_files)
                    current_part_files = []
                    current_size = 0
                current_part_files.append(file_name)
                current_size += file_size

            if current_part_files:
                parts.append(current_part_files)

            for idx, part_files in enumerate(parts):
                part_name = f"{base_title}_Part_{idx + 1}_{self.voice_name}.mp3"
                part_output_path = os.path.join(self.bm.book_dir, part_name)
                trim_dir_path = os.path.join(self.bm.book_dir, "_combine_trim_cache", f"part_{idx}")
                if len(part_files) == 1:
                    file_name = part_files[0]
                    file_path = os.path.join(self.bm.book_dir, file_name)
                    if not os.path.exists(file_path):
                        raise RuntimeError(f"Не найден файл главы для экспорта: {file_name}")
                    if os.path.getsize(file_path) <= 0:
                        raise RuntimeError(f"Файл главы пустой или повреждён: {file_name}")
                    try:
                        duration_result = _run_subprocess(
                            [
                                self.ffprobe_path,
                                "-v", "error",
                                "-show_entries", "format=duration",
                                "-of", "default=noprint_wrappers=1:nokey=1",
                                file_path,
                            ],
                            capture_output=True,
                            text=True,
                            timeout=READER_FFPROBE_TIMEOUT_SEC,
                        )
                        if duration_result.returncode != 0:
                            raise RuntimeError(duration_result.stderr.strip() or "ffprobe не смог прочитать MP3")
                    except Exception as exc:
                        raise RuntimeError(
                            f"Глава {file_name} содержит битый или неполный MP3. "
                            f"Нужно переозвучить главу перед экспортом. Детали: {exc}"
                        ) from exc

                    self.progress_signal.emit(f"Подготовка части {idx + 1} из {len(parts)}...")
                    if AudioSegment is not None:
                        _export_trimmed_mp3_file(file_path, part_output_path)
                    else:
                        shutil.copy2(file_path, part_output_path)
                else:
                    list_txt_path = os.path.join(self.bm.book_dir, f"concat_list_{idx}.txt")
                    meta_txt_path = os.path.join(self.bm.book_dir, f"metadata_{idx}.txt")

                    meta_lines = [";FFMETADATA1", f"title={self.bm.title} - Часть {idx + 1} ({self.voice_name})", ""]
                    current_time_ms = 0
                    shutil.rmtree(trim_dir_path, ignore_errors=True)
                    os.makedirs(trim_dir_path, exist_ok=True)

                    with open(list_txt_path, 'w', encoding='utf-8') as list_file:
                        for file_name in part_files:
                            file_path = os.path.join(self.bm.book_dir, file_name)

                            if not os.path.exists(file_path):
                                raise RuntimeError(f"Не найден файл главы для склейки: {file_name}")
                            if os.path.getsize(file_path) <= 0:
                                raise RuntimeError(f"Файл главы пустой или повреждён: {file_name}")

                            prepared_file_path = file_path
                            if AudioSegment is not None:
                                prepared_file_path = os.path.join(trim_dir_path, file_name)
                                _export_trimmed_mp3_file(file_path, prepared_file_path)

                            list_file.write(f"file '{_ffmpeg_concat_path(prepared_file_path)}'\n")

                            duration_cmd = [
                                self.ffprobe_path,
                                "-v", "error",
                                "-show_entries", "format=duration",
                                "-of", "default=noprint_wrappers=1:nokey=1",
                                prepared_file_path,
                            ]
                            duration_sec = 0.0
                            try:
                                duration_result = _run_subprocess(
                                    duration_cmd,
                                    capture_output=True,
                                    text=True,
                                    timeout=READER_FFPROBE_TIMEOUT_SEC,
                                )
                                if duration_result.returncode != 0:
                                    raise RuntimeError(duration_result.stderr.strip() or "ffprobe не смог прочитать MP3")
                                duration_sec = float(duration_result.stdout.strip())
                            except Exception as exc:
                                logger.error(f"Не удалось прочитать главу {file_name}: {exc}")
                                raise RuntimeError(
                                    f"Глава {file_name} содержит битый или неполный MP3. "
                                    f"Нужно переозвучить главу перед склейкой. Детали: {exc}"
                                ) from exc

                            duration_ms = int(duration_sec * 1000)
                            chapter_match = re.search(r'Ch(\d+)\.mp3', file_name)
                            chapter_number = chapter_match.group(1) if chapter_match else "?"

                            meta_lines.append("[CHAPTER]")
                            meta_lines.append("TIMEBASE=1/1000")
                            meta_lines.append(f"START={current_time_ms}")
                            meta_lines.append(f"END={current_time_ms + duration_ms}")
                            meta_lines.append(f"title=Глава {chapter_number}")
                            meta_lines.append("")
                            current_time_ms += duration_ms

                    with open(meta_txt_path, 'w', encoding='utf-8') as meta_file:
                        meta_file.write("\n".join(meta_lines))

                    self.progress_signal.emit(f"Склейка части {idx + 1} из {len(parts)}...")
                    combine_cmd = [
                        self.ffmpeg_path,
                        "-y",
                        "-f", "concat",
                        "-safe", "0",
                        "-i", list_txt_path,
                        "-i", meta_txt_path,
                        "-map_metadata", "1",
                        "-c", "copy",
                        part_output_path,
                    ]
                    combine_process = _run_subprocess(
                        combine_cmd,
                        capture_output=True,
                        text=True,
                        timeout=READER_FFMPEG_CONCAT_TIMEOUT_SEC,
                    )

                    if os.path.exists(list_txt_path):
                        os.remove(list_txt_path)
                    if os.path.exists(meta_txt_path):
                        os.remove(meta_txt_path)
                    if os.path.isdir(trim_dir_path):
                        shutil.rmtree(trim_dir_path, ignore_errors=True)

                    if combine_process.returncode != 0:
                        logger.error(f"FFMPEG Error Part {idx + 1}: {combine_process.stderr}")
                        raise RuntimeError(
                            f"Не удалось собрать часть {idx + 1}: "
                            f"{combine_process.stderr.strip() or 'ошибка ffmpeg'}"
                        )

                if self.video_image_path:
                    video_name = os.path.splitext(part_name)[0] + ".mp4"
                    video_output_path = os.path.join(self.bm.book_dir, video_name)
                    self.progress_signal.emit(f"Создание видео {idx + 1} из {len(parts)}...")
                    video_cmd = [
                        self.ffmpeg_path,
                        "-y",
                        "-loop", "1",
                        "-framerate", "1",
                        "-i", self.video_image_path,
                        "-i", part_output_path,
                        "-map", "0:v:0",
                        "-map", "1:a:0",
                        "-c:v", "libx264",
                        "-preset", "medium",
                        "-tune", "stillimage",
                        "-vf",
                        (
                            f"scale={VIDEO_FRAME_SIZE}:force_original_aspect_ratio=decrease,"
                            f"pad={VIDEO_FRAME_SIZE}:(ow-iw)/2:(oh-ih)/2,format=yuv420p"
                        ),
                        "-c:a", "aac",
                        "-b:a", "192k",
                        "-movflags", "+faststart",
                        "-shortest",
                        video_output_path,
                    ]
                    video_process = _run_subprocess(
                        video_cmd,
                        capture_output=True,
                        text=True,
                        timeout=READER_FFMPEG_VIDEO_TIMEOUT_SEC,
                    )
                    if video_process.returncode != 0:
                        logger.error(f"FFMPEG Video Error Part {idx + 1}: {video_process.stderr}")
                        raise RuntimeError(
                            f"Не удалось создать видео для части {idx + 1}: "
                            f"{video_process.stderr.strip() or 'ошибка ffmpeg'}"
                        )
                    created_videos += 1

            if created_videos:
                self.finished_signal.emit(f"Готово! Создано частей: {len(parts)}, видео: {created_videos}")
            else:
                self.finished_signal.emit(f"Готово! Создано частей: {len(parts)}")

        except Exception as e:
            logger.exception(f"Непредвиденная ошибка при объединении аудио: {e}")
            self.finished_signal.emit(f"Ошибка: {str(e)}")

class AudioPlayer(QThread):
    def __init__(self, audio_queue, vol=80):
        super().__init__()
        if pyaudio is None:
            raise RuntimeError("Для live-воспроизведения требуется PyAudio.")
        self.audio_queue = audio_queue
        self.is_running = True
        self.vol = float(vol) / 100.0
        self.p = pyaudio.PyAudio()
        self.stream = self.p.open(
            format=pyaudio.paInt16,
            channels=AUDIO_CHANNELS,
            rate=AUDIO_RATE,
            output=True
        )

    def set_volume(self, v):
        self.vol = float(v) / 100.0

    def run(self):
        while self.is_running:
            try:
                try:
                    item = self.audio_queue.get(timeout=0.1)
                except queue.Empty:
                    continue
                if item is None:
                    break
                data, c, s, ui = item
                if data:
                    arr = array.array('h', data)
                    if self.vol < 1.0:
                        for i in range(len(arr)):
                            arr[i] = max(min(int(arr[i] * self.vol), 32767), -32768)
                    self.stream.write(arr.tobytes())
                self.audio_queue.task_done()
            except Exception as e:
                continue

    def stop(self):
        self.is_running = False
        try:
            self.stream.stop_stream()
            self.stream.close()
            self.p.terminate()
        except:
            pass


# --- ВОРКЕР (С ПОДМЕНОЙ EDGE TTS) ---
class GeminiWorker(QThread):
    change_chapter_signal = pyqtSignal(int)
    worker_progress = pyqtSignal(int, int, int, int) 
    status_signal = pyqtSignal(str)
    finished_signal = pyqtSignal(int)
    error_signal = pyqtSignal(int, str)
    chapter_done_ui_signal = pyqtSignal(int)
    invalid_key_signal = pyqtSignal(int, str, str, int)
    quota_key_signal = pyqtSignal(int, str, str, str, int)
    project_quota_signal = pyqtSignal(int, str, str, int)

    def __init__(
        self,
        worker_id,
        api_key,
        bm,
        audio_queue,
        model_id,
        voice,
        style_prompt,
        speed,
        record,
        fast,
        chunk,
        segment_mode,
        manager_chapter_queue,
        daily_request_limiter=None,
        voice_mode="single",
        secondary_voice=None,
        tertiary_voice=None,
        allow_edge_fallback=True,
    ):
        super().__init__()
        self.worker_id = worker_id
        self.start_stagger_index = worker_id
        self.api_key = api_key
        self.bm = bm
        self.audio_queue = audio_queue
        self.model_id = model_id  # СОХРАНЯЕМ ВЫБРАННУЮ МОДЕЛЬ
        self.voice = voice
        self.style_prompt = style_prompt
        self.speed = speed
        self.record = record
        self.fast = fast
        self.chunk = chunk
        self.segment_mode = segment_mode or "sentences"
        self.voice_mode = voice_mode or "single"
        self.secondary_voice = secondary_voice or voice
        self.tertiary_voice = tertiary_voice or self.secondary_voice or voice
        self.allow_edge_fallback = bool(allow_edge_fallback)
        self.manager_chapter_queue = manager_chapter_queue
        self.c_idx = -1
        self.s_idx = 0
        self._is_running = True
        self._is_paused = False
        self.audio_chunks = []
        self.buffer_lock = threading.Lock()
        self.token_counter = TokenCounter() if TokenCounter is not None else None
        self.request_rpm_limiter = _make_rpm_limiter(self.model_id, LIVE_API_DEFAULT_RPM)
        self.request_tpm_limiter = TPMLimiter(_lookup_model_limit(self.model_id, "tpm", DEFAULT_TPM_LIMIT))
        self.daily_request_limiter = daily_request_limiter
        self.request_rpd_limit = _lookup_model_limit(self.model_id, "rpd", 0)
        self._last_progress_emit_payload = None
        self._last_progress_emit_at = 0.0
        self._finished_emitted = False

    def _emit_finished(self):
        if self._finished_emitted:
            return
        self._finished_emitted = True
        self.finished_signal.emit(self.worker_id)

    def _ensure_running(self):
        if not self._is_running:
            raise ReaderWorkerStopped()

    async def _sleep_interruptibly(self, seconds):
        deadline = time.monotonic() + max(0.0, float(seconds or 0.0))
        while self._is_running:
            remaining = deadline - time.monotonic()
            if remaining <= 0:
                return True
            await asyncio.sleep(min(remaining, READER_WORKER_SLEEP_STEP_SEC))
        return False

    def update_chunk_size(self, v): self.chunk = v

    def _emit_worker_progress(self, chapter_index, step_index, total_steps, force=False):
        payload = (self.worker_id, chapter_index, step_index, total_steps)
        if total_steps <= 0:
            force = True
        now = time.monotonic()
        if not force and self._last_progress_emit_payload is not None:
            prev_worker_id, prev_chapter_index, prev_step_index, prev_total_steps = self._last_progress_emit_payload
            if (
                prev_worker_id == self.worker_id
                and prev_chapter_index == chapter_index
                and prev_total_steps == total_steps
            ):
                prev_pct = int((min(prev_step_index, total_steps) / total_steps) * 100) if total_steps > 0 else 0
                current_pct = int((min(step_index, total_steps) / total_steps) * 100) if total_steps > 0 else 0
                if current_pct == prev_pct and (now - self._last_progress_emit_at) < READER_PROGRESS_EMIT_INTERVAL_SEC:
                    return
        self._last_progress_emit_payload = payload
        self._last_progress_emit_at = now
        self.worker_progress.emit(*payload)

    def _chapter_segments(self, chapter_index):
        chapter = self.bm.chapters[chapter_index]
        if self.voice_mode == "author_gender":
            saved_script = self.bm.load_tts_script(chapter_index) if self.bm is not None else ""
            if _script_matches_voice_mode(saved_script, "author_gender"):
                script_segments = _split_author_gender_script_segments(saved_script, self.segment_mode)
                if script_segments:
                    return script_segments
            return []
        if self.segment_mode == "paragraphs":
            paragraphs = []
            for part in getattr(chapter, "paragraphs", []) or []:
                paragraphs.extend(_split_live_paragraph(part))
            paragraphs = [part.strip() for part in paragraphs if (part or "").strip()]
            if paragraphs:
                return paragraphs
            raw_text = (chapter.raw_text or "").strip()
            return [raw_text] if raw_text else []
        return [part.strip() for part in chapter.flat_sentences if (part or "").strip()]

    def _join_segments_for_request(self, segments):
        return _join_live_request_segments(
            segments,
            self.segment_mode,
            self.voice_mode,
        )

    def _live_voice_descriptor(self):
        if self.voice_mode == "author_gender":
            return f"author={self.voice}, male={self.secondary_voice}, female={self.tertiary_voice}"
        if self.voice_mode == "duo":
            return f"{self.voice}/{self.secondary_voice}"
        return self.voice

    def _build_live_single_config(self, voice_name, role_hint="default", segment_directive=""):
        return _build_live_single_voice_request_config(
            voice_name,
            role_hint=role_hint,
            style_prompt=self.style_prompt,
            speed_instr=SPEED_PROMPTS.get(self.speed, SPEED_PROMPTS["Normal"]),
            segment_directive=segment_directive,
        )

    def _build_live_request_payload(self, text):
        prepared = (text or "").strip()
        if not prepared:
            return None

        if self.voice_mode == "author_gender":
            plan = _build_author_gender_live_plan(
                prepared,
                self.voice,
                self.secondary_voice,
                self.tertiary_voice,
            )
            if not plan:
                return None
            return {
                "mode": "author_gender",
                "plan": [
                    {
                        **item,
                        "config": self._build_live_single_config(
                            item["voice"],
                            item["role_hint"],
                            item.get("segment_directive", ""),
                        ),
                    }
                    for item in plan
                    if item.get("text", "").strip()
                ],
            }

        if self.voice_mode == "duo":
            return {
                "mode": "duo",
                "text": _build_live_role_script(prepared) or prepared,
                "config": self._build_live_connect_config()[0],
            }

        return {
            "mode": "single",
            "text": prepared,
            "config": self._build_live_single_config(self.voice, "default"),
        }

    def _estimate_tokens(self, text):
        if self.token_counter is not None:
            try:
                return max(1, int(self.token_counter.estimate_tokens(text)))
            except Exception:
                pass
        return max(1, int(len(text or "") / 3.0))

    def _rpm_required_delay(self, limiter=None):
        limiter = limiter or self.request_rpm_limiter
        if limiter is None or not hasattr(limiter, "interval") or not hasattr(limiter, "last_request_time"):
            return 0.0
        try:
            with limiter.lock:
                now_ts = time.time()
                elapsed = now_ts - limiter.last_request_time
                if elapsed >= limiter.interval:
                    return 0.0
                return max(0.0, limiter.interval - elapsed)
        except Exception:
            return 0.0

    async def _wait_for_request_budget(
        self,
        payload_text,
        rpm_limiter=None,
        tpm_limiter=None,
        request_label="API",
        daily_request_limiter=None,
        model_id=None,
        rpd_limit=0,
    ):
        rpm_limiter = rpm_limiter or self.request_rpm_limiter
        tpm_limiter = tpm_limiter or self.request_tpm_limiter
        daily_request_limiter = daily_request_limiter or self.daily_request_limiter
        model_id = model_id or self.model_id
        token_cost = self._estimate_tokens(payload_text)
        budget_acquired = False

        while self._is_running:
            rpm_delay = self._rpm_required_delay(rpm_limiter)
            tpm_delay = tpm_limiter.get_required_delay(token_cost) if tpm_limiter is not None else 0.0
            if rpm_delay <= 0 and tpm_delay <= 0:
                if daily_request_limiter is not None and model_id and rpd_limit and rpd_limit > 0:
                    acquired, current_count, limit_value, reset_text = daily_request_limiter.try_acquire(
                        model_id, rpd_limit, amount=1, api_key=self.api_key
                    )
                    if not acquired:
                        key_suffix = f" для ключа {_mask_api_key(self.api_key)}" if self.api_key else ""
                        message = (
                            f"Дневной лимит {limit_value} запросов для модели {model_id}{key_suffix} исчерпан. "
                            f"Следующий сброс около {reset_text}."
                        )
                        if self.api_key:
                            raise RateLimitBudgetError(message, model_id=model_id)
                        raise ProjectRateLimitReachedError(message, model_id=model_id)
                    budget_acquired = True
                if rpm_limiter is not None:
                    rpm_limiter.can_proceed()
                if tpm_limiter is not None:
                    tpm_limiter.register(token_cost)
                return budget_acquired

            sleep_candidates = [delay for delay in (rpm_delay, tpm_delay) if delay and delay > 0]
            sleep_for = min(sleep_candidates) if sleep_candidates else 0.5
            if not await self._sleep_interruptibly(max(0.25, min(sleep_for, 5.0))):
                raise ReaderWorkerStopped()
        raise ReaderWorkerStopped()

        return False

    def _release_request_budget(self, model_id=None, budget_acquired=False, amount=1):
        if not budget_acquired or self.daily_request_limiter is None:
            return
        try:
            self.daily_request_limiter.release(model_id or self.model_id, amount=amount, api_key=self.api_key)
        except Exception:
            pass

    async def _handle_rate_limit(self, exc, attempt, model_id=None, rpm_limiter=None):
        delay_seconds = RATE_LIMIT_BACKOFF_SECONDS + max(0, attempt - 1) * 20
        limiter = rpm_limiter or self.request_rpm_limiter
        if limiter is not None and hasattr(limiter, "decrease_rpm"):
            try:
                limiter.decrease_rpm(percentage=25)
                limiter.update_last_request_time(delay_seconds)
            except Exception:
                pass

        model_label = model_id or self.model_id or "<unknown>"
        logger.warning(
            f"[W{self.worker_id}] {model_label}: получен RESOURCE_EXHAUSTED/429. "
            f"Backoff {delay_seconds} сек и снижение RPM для ключа {_mask_api_key(self.api_key)}."
        )
        if not await self._sleep_interruptibly(delay_seconds):
            raise ReaderWorkerStopped()

    def _requeue_current_chapter(self):
        if self.c_idx == -1 or self.manager_chapter_queue is None:
            return
        try:
            if not self.bm.is_chapter_done(self.c_idx) and not self.bm.is_chapter_skipped(self.c_idx):
                self.manager_chapter_queue.put(self.c_idx)
        except Exception:
            pass

    def _abort_for_invalid_key(self, error_text):
        chapter_index = self.c_idx
        self._requeue_current_chapter()
        masked_key = _mask_api_key(self.api_key)
        logger.error(
            f"[W{self.worker_id}] Ключ {masked_key} отклонён API Google как невалидный. "
            "Воркер остановлен, глава возвращена в очередь."
        )
        self.invalid_key_signal.emit(self.worker_id, self.api_key, error_text, chapter_index)
        self.c_idx = -1
        self._is_running = False
        self._emit_finished()

    def _abort_for_quota_key(self, error_text, model_id):
        chapter_index = self.c_idx
        self._requeue_current_chapter()
        masked_key = _mask_api_key(self.api_key)
        logger.error(
            f"[W{self.worker_id}] Ключ {masked_key} упёрся в квоту/лимит модели {model_id}. "
            "Воркер остановлен, глава возвращена в очередь."
        )
        self.quota_key_signal.emit(self.worker_id, self.api_key, model_id or "", error_text, chapter_index)
        self.c_idx = -1
        self._is_running = False
        self._emit_finished()

    def _abort_for_project_quota(self, error_text, model_id):
        chapter_index = self.c_idx
        self._requeue_current_chapter()
        logger.error(
            f"[W{self.worker_id}] Дневной лимит проекта по модели {model_id} исчерпан. "
            "Глава возвращена в очередь, воркер остановлен."
        )
        self.project_quota_signal.emit(self.worker_id, model_id or "", error_text, chapter_index)
        self.c_idx = -1
        self._is_running = False
        self._emit_finished()

    def run(self):
        try:
            asyncio.run(self.main_loop())
        except ReaderWorkerStopped:
            pass
        except Exception as e:
            logger.exception(f"[W{self.worker_id}] Worker crashed: {e}")
            self.error_signal.emit(self.worker_id, f"CRASH: {str(e)}")
        finally:
            self._is_running = False
            self._emit_finished()

    async def _wait_before_worker_start(self, stagger_seconds):
        try:
            worker_index = max(0, int(getattr(self, "start_stagger_index", self.worker_id)))
        except (TypeError, ValueError):
            worker_index = 0
        delay_seconds = max(0.0, worker_index * float(stagger_seconds or 0.0))
        while self._is_running and delay_seconds > 0:
            sleep_step = min(delay_seconds, 1.0)
            await asyncio.sleep(sleep_step)
            delay_seconds -= sleep_step

    async def get_edge_tts_fallback(self, text):
        """Озвучка забаненного текста через Edge TTS с умным выбором пола"""
        import random
        if not self.allow_edge_fallback:
            logger.warning(f"[W{self.worker_id}] Edge TTS fallback отключён в настройках reader.")
            return None
        if edge_tts is None or AudioSegment is None:
            logger.warning(f"[W{self.worker_id}] Edge TTS fallback недоступен: нет edge-tts или pydub.")
            return None
        # Разгружаем сервера: небольшая случайная пауза перед запросом
        if not await self._sleep_interruptibly(random.uniform(0.5, 2.0)):
            return None
        
        # --- УМНЫЙ ВЫБОР ГОЛОСА EDGE TTS ПО ПОЛУ ---
        gender = VOICES_MAP.get(self.voice, "Ж")
        edge_voice = "ru-RU-DmitryNeural" if gender == "М" else "ru-RU-SvetlanaNeural"
        
        try:
            # 1. Мягкая очистка (заменяем тире и многоточия, которые ломают серверы Microsoft)
            clean_text = text.replace("…", ".").replace("—", "").replace("«", "").replace("»", "")
            clean_text = re.sub(r'[\n\r\t]', ' ', clean_text).strip()
            
            if not clean_text:
                return None

            # Обязательно добавляем точку в конце, чтобы генератор плавно завершил звук
            clean_text += " ."

            data_out = b""
            
            # Делаем 3 попытки, если Microsoft банит по лимитам запросов
            for attempt in range(3):
                data_out = b""
                try:
                    communicate = edge_tts.Communicate(clean_text, edge_voice)
                    stream_iterator = communicate.stream().__aiter__()
                    deadline = time.monotonic() + READER_EDGE_TTS_TOTAL_TIMEOUT_SEC
                    while self._is_running:
                        remaining = deadline - time.monotonic()
                        if remaining <= 0:
                            raise TimeoutError("Edge TTS stream total timeout.")
                        try:
                            chunk = await asyncio.wait_for(
                                stream_iterator.__anext__(),
                                timeout=min(READER_EDGE_TTS_CHUNK_TIMEOUT_SEC, remaining),
                            )
                        except StopAsyncIteration:
                            break
                        if chunk["type"] == "audio":
                            data_out += chunk["data"]
                            
                    if len(data_out) > 100:
                        break # Успех! Выходим из цикла попыток
                except Exception as stream_err:
                    logger.warning(f"[W{self.worker_id}] Ошибка потока Edge TTS (попытка {attempt+1}): {stream_err}")
                
                # Если ответ пустой, ждем и делаем более агрессивную очистку
                if len(data_out) <= 100:
                    logger.warning(f"[W{self.worker_id}] Edge TTS вернул пустой звук. Повтор через {attempt+2} сек...")
                    if not await self._sleep_interruptibly(attempt + 2):
                        return None
                    # 2. Агрессивная очистка: оставляем только буквы, цифры и базовую пунктуацию
                    clean_text = re.sub(r'[^\w\sа-яА-ЯёЁ.,?!]', ' ', text).strip() + " ."

            # Если после всех попыток данных нет
            if not data_out or len(data_out) < 100:
                logger.error(f"[W{self.worker_id}] Edge TTS не выдал звук. Возможно, жесткий фильтр Microsoft.")
                return None

            def _decode_silently():
                import subprocess
                startupinfo = None
                # Скрываем всплывающие окна консоли на Windows
                if platform.system() == "Windows":
                    startupinfo = subprocess.STARTUPINFO()
                    startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW
                    startupinfo.wShowWindow = 0 # SW_HIDE

                fp = io.BytesIO(data_out)
                try:
                    # Конвертируем ответ в нужный нам формат
                    segment = AudioSegment.from_file(fp, format="mp3")
                    segment = segment.set_frame_rate(AUDIO_RATE).set_channels(AUDIO_CHANNELS).set_sample_width(2)
                    return segment.raw_data
                except Exception as e:
                    logger.error(f"[W{self.worker_id}] Ошибка pydub при декодировании Edge TTS: {e}")
                    return None

            return await _to_thread_with_timeout(
                "Edge TTS decode",
                READER_AUDIO_EXPORT_TIMEOUT_SEC,
                _decode_silently,
            )

        except Exception as e:
            logger.error(f"[W{self.worker_id}] Критическая ошибка Edge TTS: {e}")
            return None

    def _build_live_connect_config(self):
        speed_instr = SPEED_PROMPTS.get(self.speed, SPEED_PROMPTS["Normal"])
        if self.voice_mode == "duo":
            sys_instr = (
                f"Ты профессиональная студия озвучки аудиокниг. {self.style_prompt} "
                f"{speed_instr} "
                "Озвучивай входной текст СЛОВО В СЛОВО. "
                f"Во входе используются только роли {TTS_SPEAKER_NARRATOR} и {TTS_SPEAKER_DIALOGUE}. "
                f"{TTS_SPEAKER_NARRATOR} — авторский текст, ремарки и описание сцены. "
                f"{TTS_SPEAKER_DIALOGUE} — прямая речь и цитаты. "
                "Если встретишь сценические теги в квадратных скобках, воспринимай их как пометки к подаче и не произноси их вслух. "
                "Не меняй слова, не добавляй новых ролей, новых фраз или междометий и обязательно дочитывай текст до самой последней точки. "
                "Выдай полный аудиоответ. Язык: Русский."
            )

            config = genai_types.LiveConnectConfig(
                response_modalities=["AUDIO"],
                speech_config=_build_duo_voice_speech_config(self.voice, self.secondary_voice),
                system_instruction=sys_instr
            )
            return config, f"{self.voice}/{self.secondary_voice}"

        return self._build_live_single_config(self.voice, "default"), self.voice

    async def _send_live_text_turn(self, session, text_to_send):
        text_value = (text_to_send or "").strip()
        if not text_value:
            return
        self._ensure_running()
        try:
            await asyncio.wait_for(
                session.send_client_content(turns=text_value, turn_complete=True),
                timeout=READER_LIVE_SEND_TIMEOUT_SEC,
            )
        except Exception:
            await asyncio.wait_for(
                session.send_realtime_input(text=text_value),
                timeout=READER_LIVE_SEND_TIMEOUT_SEC,
            )

    async def _collect_live_request_raw_audio(self, client, config, text_to_send, on_chunk=None):
        received_chunks = []
        budget_acquired = False
        session_cm = None
        try:
            budget_acquired = await self._wait_for_request_budget(
                text_to_send,
                request_label="Live API",
                daily_request_limiter=self.daily_request_limiter,
                model_id=self.model_id,
                rpd_limit=self.request_rpd_limit,
            )
            self._ensure_running()
            session_cm, session = await _open_live_session_with_timeout(client, self.model_id, config)
            await self._send_live_text_turn(session, text_to_send)

            receive_iterator = session.receive().__aiter__()
            current_timeout = READER_LIVE_FIRST_CHUNK_TIMEOUT_SEC

            while self._is_running:
                try:
                    response = await asyncio.wait_for(receive_iterator.__anext__(), timeout=current_timeout)
                    current_timeout = READER_LIVE_NEXT_CHUNK_TIMEOUT_SEC

                    if response.server_content:
                        if response.server_content.model_turn:
                            for part in response.server_content.model_turn.parts:
                                if part.inline_data:
                                    data = bytes(part.inline_data.data)
                                    received_chunks.append(data)
                                    if on_chunk is not None:
                                        on_chunk(data)
                        if getattr(response.server_content, "turn_complete", False):
                            break
                except asyncio.TimeoutError:
                    break
                except StopAsyncIteration:
                    break
        except Exception:
            self._release_request_budget(self.model_id, budget_acquired)
            raise
        finally:
            await _close_live_session_quietly(session_cm)

        return b"".join(received_chunks)

    async def _request_live_audio_bytes(self, client, config, text_to_send):
        raw_audio = await self._collect_live_request_raw_audio(client, config, text_to_send)
        return _trim_raw_pcm_boundaries(raw_audio)

    def _commit_live_audio_bytes(self, audio_bytes):
        if not audio_bytes:
            return False
        if self.record:
            with self.buffer_lock:
                self.audio_chunks.append(audio_bytes)
        if self.audio_queue and not self.fast:
            self.audio_queue.put((audio_bytes, self.c_idx, self.s_idx, False))
        return True

    async def _collect_live_payload_audio(self, client, payload):
        if not payload:
            return b""
        if payload.get("mode") == "author_gender":
            combined = bytearray()
            plan = [item for item in payload.get("plan", []) if item.get("text", "").strip()]
            for idx, item in enumerate(plan):
                chunk_bytes = await self._request_live_audio_bytes(client, item["config"], item["text"])
                if not chunk_bytes:
                    return b""
                combined.extend(chunk_bytes)
                if idx + 1 < len(plan):
                    next_item = plan[idx + 1]
                    gap_ms = LIVE_REQUEST_ROLE_SWITCH_GAP_MS if next_item.get("role_hint") != item.get("role_hint") else LIVE_REQUEST_JOIN_GAP_MS
                    gap_ms = max(gap_ms, int(item.get("pause_after_ms", 0) or 0))
                    combined.extend(_pcm_silence(gap_ms))
            return bytes(combined)
        return await self._request_live_audio_bytes(client, payload["config"], payload["text"])

    async def main_loop(self):
        
        # Обновленный клиент (новый SDK обрабатывает Live API без костылей v1alpha)
        if genai is None or genai_types is None:
            raise RuntimeError("Для озвучки Gemini Reader требуется пакет google-genai.")
        await self._wait_before_worker_start(READER_LIVE_WORKER_START_STAGGER_SECONDS)
        if not self._is_running:
            return
        client = _make_genai_client(self.api_key)
        voice_descriptor = self._live_voice_descriptor()

        # Используем выбранную в UI модель
        logger.info(f"Воркер {self.worker_id} запущен (Модель: {self.model_id}, Голоса: {voice_descriptor}).")

        total_sent = 0
        single_sentence_mode_remaining = 0
        fail_count = 0 
        gemini_retry_count = 0 # Счетчик для повторных попыток пробиться к Gemini (одиночные)
        batch_retry_count = 0  # НОВЫЙ СЧЕТЧИК: для повторных попыток целого батча

        while self._is_running:
            if self.c_idx == -1:
                try:
                    self.c_idx = self.manager_chapter_queue.get_nowait()
                    self.s_idx = 0 
                    single_sentence_mode_remaining = 0
                    fail_count = 0
                    gemini_retry_count = 0
                    batch_retry_count = 0
                    total_sent = len(self._chapter_segments(self.c_idx))
                    if self.voice_mode == "author_gender" and total_sent == 0:
                        self.error_signal.emit(
                            self.worker_id,
                            f"Глава {self.c_idx + 1}: нет корректного AI-сценария Author/Male/Female. Сначала подготовьте AI-сценарий.",
                        )
                        self.c_idx = -1
                        continue
                    segment_label = "абз." if self.segment_mode == "paragraphs" else "предл."
                    logger.info(f"Воркер {self.worker_id} взял Главу {self.c_idx + 1} ({total_sent} {segment_label})")
                except queue.Empty:
                    self._emit_finished()
                    break
            else:
                total_sent = len(self._chapter_segments(self.c_idx))
                if self.voice_mode == "author_gender" and total_sent == 0:
                    self.error_signal.emit(
                        self.worker_id,
                        f"Глава {self.c_idx + 1}: нет корректного AI-сценария Author/Male/Female. Сначала подготовьте AI-сценарий.",
                    )
                    self.c_idx = -1
                    continue

            if self.bm.is_chapter_done(self.c_idx):
                self.chapter_done_ui_signal.emit(self.c_idx)
                self.c_idx = -1
                continue

            if self.s_idx >= total_sent:
                logger.info(f"Воркер {self.worker_id} сохраняет и завершает Главу {self.c_idx + 1}")
                await self.save_file(final=True)
                self.bm.mark_chapter_done(self.c_idx)
                self.chapter_done_ui_signal.emit(self.c_idx)
                if self.worker_id == 0: 
                    self.change_chapter_signal.emit(self.c_idx)
                self.c_idx = -1
                continue

            if total_sent > 0:
                self._emit_worker_progress(self.c_idx, self.s_idx, total_sent)

            current_chunk_size = 1 if single_sentence_mode_remaining > 0 else self.chunk
            
            segments = self._chapter_segments(self.c_idx)
            text_parts = []
            actual_count = 0
            end_range = min(self.s_idx + current_chunk_size, total_sent)
            
            for i in range(self.s_idx, end_range):
                text_parts.append(segments[i])
                actual_count += 1
            
            request_payload = self._build_live_request_payload(self._join_segments_for_request(text_parts))
            payload_preview = ""
            if request_payload:
                if request_payload.get("mode") == "author_gender":
                    payload_preview = " | ".join(
                        item.get("text", "")[:20] for item in request_payload.get("plan", [])[:2]
                    )
                else:
                    payload_preview = request_payload.get("text", "")[:30]
            if not request_payload:
                self.s_idx += 1
                if single_sentence_mode_remaining > 0:
                    single_sentence_mode_remaining -= 1
                continue

            data_received = False
            request_budget_acquired = False
            try:
                if request_payload.get("mode") == "author_gender":
                    audio_bytes = await self._collect_live_payload_audio(client, request_payload)
                    data_received = self._commit_live_audio_bytes(audio_bytes)
                else:
                    text_to_send = request_payload["text"]
                    config = request_payload["config"]
                    request_audio = bytearray()

                    def _on_live_chunk(data):
                        request_audio.extend(data)
                        if self.audio_queue and not self.fast:
                            self.audio_queue.put((data, self.c_idx, self.s_idx, False))

                    raw_audio = await self._collect_live_request_raw_audio(
                        client,
                        config,
                        text_to_send,
                        on_chunk=_on_live_chunk,
                    )
                    data_received = bool(raw_audio)
                    if self.record and request_audio:
                        trimmed_audio = _trim_raw_pcm_boundaries(bytes(request_audio))
                        with self.buffer_lock:
                            self.audio_chunks.append(trimmed_audio)
            except Exception as e:
                if isinstance(e, ReaderWorkerStopped):
                    break
                if isinstance(e, ProjectRateLimitReachedError):
                    self._abort_for_project_quota(str(e), e.model_id or self.model_id)
                    break
                if isinstance(e, RateLimitBudgetError):
                    self._abort_for_quota_key(str(e), e.model_id or self.model_id)
                    break
                if _is_invalid_api_key_error(e):
                    self._abort_for_invalid_key(str(e))
                    break
                if _is_rate_limited_error(e):
                    self._abort_for_quota_key(str(e), self.model_id)
                    break
                # Ошибки сети или внезапные разрывы логируем, чтобы не было "тихих" провалов
                logger.debug(f"[W{self.worker_id}] Внутренняя ошибка сессии Gemini: {e}")

            if not data_received and self._is_running:
                if current_chunk_size > 1:
                    # ДАЕМ БАТЧУ 3 ПОПЫТКИ ПЕРЕД ДРОБЛЕНИЕМ НА ОДИНОЧНЫЕ ПРЕДЛОЖЕНИЯ
                    if batch_retry_count < 3:
                        batch_retry_count += 1
                        logger.warning(f"[W{self.worker_id}] Ошибка API Gemini (батч). Попытка {batch_retry_count}/3 для батча: '{payload_preview}...'")
                        if not await self._sleep_interruptibly(2):
                            break
                        continue # Возвращаемся в начало и пробуем этот же батч целиком
                    else:
                        logger.warning(f"[W{self.worker_id}] Батч не прошел после 3 попыток! Дробим {actual_count} предл. по одному...")
                        single_sentence_mode_remaining = actual_count
                        fail_count = 0 
                        gemini_retry_count = 0 
                        batch_retry_count = 0
                        continue
                else:
                    # ДАЕМ ОДИНОЧНОМУ ПРЕДЛОЖЕНИЮ 3 ПОПЫТКИ ПЕРЕД EDGE TTS
                    if gemini_retry_count < 3:
                        gemini_retry_count += 1
                        logger.warning(f"[W{self.worker_id}] Ошибка API Gemini (одиночное). Попытка {gemini_retry_count}/3 для: '{payload_preview}...'")
                        if not await self._sleep_interruptibly(2):
                            break
                        continue
                    else:
                        fallback_source_text = request_payload.get("text", "") if request_payload else ""
                        if request_payload and request_payload.get("mode") == "author_gender":
                            fallback_source_text = "\n\n".join(
                                item.get("text", "") for item in request_payload.get("plan", [])
                            )
                        logger.warning(f"[W{self.worker_id}] Gemini сдался после 3 попыток. Озвучка Edge TTS: '{payload_preview}...'")
                        fallback_data = await self.get_edge_tts_fallback(fallback_source_text)
                        if fallback_data:
                            data_received = self._commit_live_audio_bytes(fallback_data)

            # УСПЕХ ИЛИ ПРОПУСК
            if data_received:
                fail_count = 0 
                gemini_retry_count = 0 
                batch_retry_count = 0 # Сбрасываем все счетчики ошибок при успехе
                self.s_idx = min(self.s_idx + actual_count, total_sent)
                
                if single_sentence_mode_remaining > 0:
                    single_sentence_mode_remaining -= actual_count
                    
                self._emit_worker_progress(self.c_idx, self.s_idx, total_sent, force=self.s_idx >= total_sent)
                
                if self.worker_id == 0: 
                    self.bm.save_progress(self.c_idx, self.s_idx)
                
                if self.record and self.s_idx % 5 == 0: 
                    await self.save_file()
            else:
                # Сюда программа дойдет только если даже Edge TTS не смог сгенерировать звук
                fail_count += 1
                if fail_count >= 3:
                    logger.error(f"[W{self.worker_id}] Пропуск предложения после неудач Edge TTS: '{payload_preview}...'")
                    self.s_idx = min(self.s_idx + actual_count, total_sent)
                    fail_count = 0
                    gemini_retry_count = 0
                    batch_retry_count = 0
                    if single_sentence_mode_remaining > 0:
                        single_sentence_mode_remaining -= actual_count
                else:
                    logger.error(f"[W{self.worker_id}] Ошибка Edge TTS. Попытка {fail_count}/3. Пауза 5 сек...")
                    if not await self._sleep_interruptibly(5):
                        break




    async def save_file(self, final=False):
        path = self.bm.get_mp3_path(self.c_idx)
        with self.buffer_lock:
            if not self.audio_chunks: return
            combined_data = b"".join(self.audio_chunks)
        
        def _exp():
            try:
                _export_pcm_to_mp3(combined_data, path)
            except Exception as exc:
                raise RuntimeError(f"[W{self.worker_id}] Не удалось сохранить MP3 главы {self.c_idx + 1}: {exc}") from exc
        await _to_thread_with_timeout(
            "MP3 export",
            READER_AUDIO_EXPORT_TIMEOUT_SEC,
            _exp,
        )
        if final:
            with self.buffer_lock:
                self.audio_chunks = []

    def pause(self): self._is_paused = not self._is_paused
    def stop(self): self._is_running = False


class GeminiParallelChapterWorker(GeminiWorker):
    def __init__(
        self,
        worker_id,
        api_key,
        bm,
        model_id,
        voice,
        secondary_voice,
        tertiary_voice,
        speed,
        task_queue,
        parallel_state,
        daily_request_limiter=None,
        voice_mode="single",
        allow_edge_fallback=True,
    ):
        super().__init__(
            worker_id,
            api_key,
            bm,
            None,
            model_id,
            voice,
            "Ты диктор.",
            speed,
            True,
            True,
            1,
            parallel_state.get("segment_mode", "sentences"),
            task_queue,
            daily_request_limiter=daily_request_limiter,
            voice_mode=voice_mode,
            secondary_voice=secondary_voice,
            tertiary_voice=tertiary_voice,
            allow_edge_fallback=allow_edge_fallback,
        )
        self.parallel_task_queue = task_queue
        self.parallel_state = parallel_state
        self.current_task = None

    def _requeue_current_task(self):
        if self.current_task is None or self.parallel_task_queue is None:
            return
        if self.parallel_state.get("cancelled"):
            return
        try:
            self.parallel_task_queue.put(self.current_task)
        except Exception:
            pass

    def _abort_for_invalid_key(self, error_text):
        chapter_index = self.current_task.get("chapter_index", -1) if self.current_task else -1
        self._requeue_current_task()
        masked_key = _mask_api_key(self.api_key)
        logger.error(
            f"[W{self.worker_id}] Ключ {masked_key} отклонён API Google во время параллельной озвучки главы. "
            "Текущий блок возвращён в очередь."
        )
        self.invalid_key_signal.emit(self.worker_id, self.api_key, error_text, chapter_index)
        self.current_task = None
        self.c_idx = -1
        self._is_running = False
        self._emit_finished()

    def _abort_for_quota_key(self, error_text, model_id):
        chapter_index = self.current_task.get("chapter_index", -1) if self.current_task else -1
        self._requeue_current_task()
        masked_key = _mask_api_key(self.api_key)
        logger.error(
            f"[W{self.worker_id}] Ключ {masked_key} упёрся в лимит {model_id} во время параллельной озвучки. "
            "Текущий блок возвращён в очередь."
        )
        self.quota_key_signal.emit(self.worker_id, self.api_key, model_id or "", error_text, chapter_index)
        self.current_task = None
        self.c_idx = -1
        self._is_running = False
        self._emit_finished()

    def _abort_for_project_quota(self, error_text, model_id):
        chapter_index = self.current_task.get("chapter_index", -1) if self.current_task else -1
        self._requeue_current_task()
        logger.error(
            f"[W{self.worker_id}] Дневной лимит проекта по модели {model_id} исчерпан во время параллельной озвучки. "
            "Текущий блок возвращён в очередь."
        )
        self.project_quota_signal.emit(self.worker_id, model_id or "", error_text, chapter_index)
        self.current_task = None
        self.c_idx = -1
        self._is_running = False
        self._emit_finished()

    def _register_completed_task(self, task_index):
        with self.parallel_state["lock"]:
            completed_tasks = self.parallel_state.setdefault("completed_tasks", set())
            completed_tasks.add(task_index)
            completed_count = len(completed_tasks)
            self.parallel_state["completed_count"] = completed_count
        return completed_count

    async def _process_task(self, client, task):
        payload = self._build_live_request_payload(task.get("text", "").strip())
        fallback_source_text = task.get("text", "").strip()
        if not payload:
            raise RuntimeError("Пустой блок главы для озвучки.")

        gemini_retry_count = 0
        fail_count = 0

        while self._is_running:
            audio_bytes = b""
            try:
                audio_bytes = await self._collect_live_payload_audio(client, payload)
            except Exception as exc:
                if isinstance(exc, RateLimitBudgetError):
                    self._abort_for_quota_key(str(exc), exc.model_id or self.model_id)
                    return False
                if isinstance(exc, ProjectRateLimitReachedError):
                    self._abort_for_project_quota(str(exc), exc.model_id or self.model_id)
                    return False
                if isinstance(exc, RateLimitBudgetError):
                    self._abort_for_quota_key(str(exc), exc.model_id or self.model_id)
                    return False
                if _is_invalid_api_key_error(exc):
                    self._abort_for_invalid_key(str(exc))
                    return False
                if _is_rate_limited_error(exc):
                    self._abort_for_quota_key(str(exc), self.model_id)
                    return False
                logger.debug(f"[W{self.worker_id}] Ошибка Live API для блока главы: {exc}")

            if not audio_bytes and self._is_running:
                if gemini_retry_count < 3:
                    gemini_retry_count += 1
                    logger.warning(
                        f"[W{self.worker_id}] Пустой ответ Gemini для блока {task['task_index'] + 1}/{task['total_tasks']}. "
                        f"Повтор {gemini_retry_count}/3."
                    )
                    if not await self._sleep_interruptibly(2):
                        return False
                    continue

                fallback_data = await self.get_edge_tts_fallback(fallback_source_text)
                if fallback_data:
                    audio_bytes = fallback_data
                else:
                    fail_count += 1
                    if fail_count >= 3:
                        raise RuntimeError(
                            f"Не удалось озвучить блок {task['task_index'] + 1}/{task['total_tasks']} даже через Edge TTS."
                        )
                    if not await self._sleep_interruptibly(5):
                        return False
                    continue

            if audio_bytes:
                await _to_thread_with_timeout(
                    "parallel MP3 export",
                    READER_AUDIO_EXPORT_TIMEOUT_SEC,
                    _export_pcm_to_mp3,
                    audio_bytes,
                    task["output_path"],
                )
                completed_count = self._register_completed_task(task["task_index"])
                self._emit_worker_progress(task["chapter_index"], completed_count, task["total_tasks"])
                return True

        return False

    async def main_loop(self):
        if genai is None or genai_types is None:
            raise RuntimeError("Для озвучки Gemini Reader требуется пакет google-genai.")

        await self._wait_before_worker_start(READER_PARALLEL_WORKER_START_STAGGER_SECONDS)
        if not self._is_running:
            return
        client = _make_genai_client(self.api_key)
        voice_descriptor = self._live_voice_descriptor()
        logger.info(
            f"Параллельный Live воркер {self.worker_id} запущен (Модель: {self.model_id}, Голоса: {voice_descriptor})."
        )

        while self._is_running:
            try:
                task = self.parallel_task_queue.get_nowait()
            except queue.Empty:
                self._emit_finished()
                break

            self.current_task = task
            self.c_idx = task["chapter_index"]
            self.s_idx = task["task_index"]

            try:
                completed = await self._process_task(client, task)
                if not completed and not self._is_running:
                    break
            except Exception as exc:
                logger.error(
                    f"[W{self.worker_id}] Ошибка блока {task['task_index'] + 1}/{task['total_tasks']} главы "
                    f"{task['chapter_index'] + 1}: {exc}"
                )
                self.error_signal.emit(
                    self.worker_id,
                    f"Глава {task['chapter_index'] + 1}, блок {task['task_index'] + 1}: {exc}"
                )
                self._requeue_current_task()
                self.current_task = None
                self.c_idx = -1
                self._is_running = False
                self._emit_finished()
                break

            self.current_task = None
            self.c_idx = -1


class FlashTtsWorker(GeminiWorker):
    script_ready_signal = pyqtSignal(int)

    def __init__(
        self,
        worker_id,
        api_key,
        bm,
        audio_queue,
        tts_model_id,
        primary_voice,
        secondary_voice,
        speed,
        record,
        fast,
        chunk,
        manager_chapter_queue,
        preprocess_model_id,
        preprocess_profile,
        voice_mode,
        run_mode,
        preprocess_directive,
        tts_directive,
        daily_request_limiter,
        allow_edge_fallback=True,
    ):
        super().__init__(
            worker_id,
            api_key,
            bm,
            audio_queue,
            tts_model_id,
            primary_voice,
            "Professional audiobook performance.",
            speed,
            record,
            fast,
            chunk,
            "sentences",
            manager_chapter_queue,
            daily_request_limiter=daily_request_limiter,
            allow_edge_fallback=allow_edge_fallback,
        )
        self.secondary_voice = secondary_voice or primary_voice
        self.preprocess_model_id = preprocess_model_id
        self.preprocess_profile = preprocess_profile or PREPROCESS_PROFILE_OPTIONS["Бережно"]
        self.voice_mode = voice_mode
        self.run_mode = run_mode
        self.preprocess_directive = preprocess_directive or DEFAULT_PREPROCESS_DIRECTIVE
        self.tts_directive = tts_directive or DEFAULT_TTS_DIRECTIVE
        self._last_chapter_started_at = 0.0
        self.preprocess_rpm_limiter = _make_rpm_limiter(self.preprocess_model_id, FLASH_PREPROCESS_DEFAULT_RPM)
        self.preprocess_tpm_limiter = TPMLimiter(_lookup_model_limit(self.preprocess_model_id, "tpm", DEFAULT_TPM_LIMIT))
        self.tts_rpm_limiter = _make_rpm_limiter(self.model_id, FLASH_TTS_DEFAULT_RPM)
        self.tts_tpm_limiter = TPMLimiter(_lookup_model_limit(self.model_id, "tpm", DEFAULT_TPM_LIMIT))
        self.preprocess_rpd_limit = _lookup_model_limit(self.preprocess_model_id, "rpd", 0)
        self.tts_rpd_limit = _lookup_model_limit(self.model_id, "rpd", 0)

    def _tts_chunk_limit(self):
        base = max(1, int(self.chunk))
        if self.voice_mode == "duo":
            return max(1200, base * 700)
        return max(1600, base * 900)

    def _build_speech_config(self):
        if self.voice_mode == "duo":
            return _build_duo_voice_speech_config(self.voice, self.secondary_voice)
        return _build_single_voice_speech_config(self.voice)

    async def _wait_for_flash_chapter_slot(self, chapter_index):
        if self.run_mode == "prepare":
            return
        wait_seconds = 0.0
        if self._last_chapter_started_at > 0:
            elapsed = time.time() - self._last_chapter_started_at
            wait_seconds = max(0.0, FLASH_TTS_CHAPTER_INTERVAL_SECONDS - elapsed)
        if wait_seconds > 0:
            logger.info(
                f"[W{self.worker_id}] Ключ {_mask_api_key(self.api_key)} ждёт {wait_seconds:.1f} сек "
                f"перед стартом главы {chapter_index + 1} (лимит: 1 глава/мин на ключ)."
            )
        while self._is_running and wait_seconds > 0:
            sleep_step = min(wait_seconds, 5.0)
            if not await self._sleep_interruptibly(sleep_step):
                return
            wait_seconds -= sleep_step
        if self._is_running:
            self._last_chapter_started_at = time.time()

    async def _call_text_generation(self, client, model_id, prompt_text):
        last_error = None
        for attempt in range(1, 4):
            budget_acquired = False
            try:
                budget_acquired = await self._wait_for_request_budget(
                    prompt_text,
                    rpm_limiter=self.preprocess_rpm_limiter,
                    tpm_limiter=self.preprocess_tpm_limiter,
                    request_label=f"text-model {model_id}",
                    daily_request_limiter=self.daily_request_limiter,
                    model_id=model_id,
                    rpd_limit=self.preprocess_rpd_limit,
                )
                response = await _to_thread_with_timeout(
                    f"text-model {model_id}",
                    READER_GENERATE_CONTENT_TIMEOUT_SEC,
                    client.models.generate_content,
                    model=model_id,
                    contents=prompt_text,
                    config=genai_types.GenerateContentConfig(temperature=0.1),
                )
                generated_text = _extract_response_text(response)
                if generated_text:
                    return generated_text
                raise RuntimeError("Пустой текстовый ответ модели.")
            except Exception as exc:
                self._release_request_budget(model_id, budget_acquired)
                if isinstance(exc, ReaderWorkerStopped):
                    raise
                if isinstance(exc, (RateLimitBudgetError, ProjectRateLimitReachedError)):
                    raise
                if _is_invalid_api_key_error(exc):
                    raise InvalidApiKeyError(str(exc)) from exc
                if _is_rate_limited_error(exc):
                    raise RateLimitBudgetError(
                        f"Ключ {_mask_api_key(self.api_key)} получил RESOURCE_EXHAUSTED/429 для модели {model_id}: {exc}",
                        model_id=model_id,
                    ) from exc
                last_error = exc
                logger.warning(
                    f"[W{self.worker_id}] Ошибка text-model {model_id} (попытка {attempt}/3): {exc}"
                )
                if not await self._sleep_interruptibly(1.5 + attempt):
                    raise ReaderWorkerStopped()
        if last_error is not None and _is_rate_limited_error(last_error):
            raise RateLimitBudgetError(
                f"Не удалось получить текст от модели {model_id}: {last_error}",
                model_id=model_id,
            )
        raise RuntimeError(f"Не удалось получить текст от модели {model_id}: {last_error}")

    async def _normalize_author_gender_script(self, text_client, raw_text, prepared_script):
        candidate = (prepared_script or "").strip()
        for repair_attempt in range(3):
            issues = _find_author_gender_script_issues(candidate, raw_text)
            if _parse_author_gender_script(candidate) and not issues:
                return candidate

            repair_prompt = _build_author_gender_repair_prompt(
                raw_text,
                candidate,
                issues,
                profile_prompt=self.preprocess_profile,
                extra_directive=self.preprocess_directive,
            )
            logger.warning(
                f"[W{self.worker_id}] AI-сценарий Author/Male/Female требует repair-pass "
                f"(попытка {repair_attempt + 1}/3, проблем: {max(1, len(issues))})."
            )
            candidate = await self._call_text_generation(
                text_client,
                self.preprocess_model_id,
                repair_prompt,
            )
            candidate = (candidate or "").strip()

        issues = _find_author_gender_script_issues(candidate, raw_text)
        if issues:
            issue_preview = "; ".join(
                f"line {item['line_no']}: {item['reason']}"
                for item in issues[:3]
            )
            raise RuntimeError(
                "Не удалось получить строгий AI-сценарий Author/Male/Female без смешения реплик и авторского текста"
                + (f" ({issue_preview})." if issue_preview else ".")
            )
        raise RuntimeError("Не удалось получить строгий AI-сценарий Author/Male/Female.")

    async def _prepare_script_for_chapter(self, text_client, chapter_index, force=False):
        existing_script = self.bm.load_tts_script(chapter_index)
        if existing_script and not force and _script_matches_voice_mode(existing_script, self.voice_mode):
            return existing_script

        raw_text = self.bm.chapters[chapter_index].raw_text.strip()
        if not raw_text:
            raise RuntimeError("Глава не содержит текста для озвучки.")

        preprocess_prompt = _build_preprocess_prompt(
            raw_text,
            voice_mode=self.voice_mode,
            profile_prompt=self.preprocess_profile,
            extra_directive=self.preprocess_directive,
        )
        prepared_script = await self._call_text_generation(
            text_client,
            self.preprocess_model_id,
            preprocess_prompt,
        )
        if self.voice_mode == "author_gender":
            prepared_script = await self._normalize_author_gender_script(
                text_client,
                raw_text,
                prepared_script,
            )
        self.bm.save_tts_script(chapter_index, prepared_script)
        self.script_ready_signal.emit(chapter_index)
        logger.info(f"[W{self.worker_id}] Сценарий для главы {chapter_index + 1} подготовлен.")
        return prepared_script

    async def _resolve_script_for_run(self, text_client, chapter_index):
        if self.run_mode == "prepare":
            return await self._prepare_script_for_chapter(text_client, chapter_index, force=True)

        if self.run_mode == "auto":
            return await self._prepare_script_for_chapter(text_client, chapter_index, force=False)

        if self.run_mode == "staged":
            prepared_script = self.bm.load_tts_script(chapter_index)
            if prepared_script.strip() and _script_matches_voice_mode(prepared_script, self.voice_mode):
                return prepared_script
            if prepared_script.strip():
                raise RuntimeError(
                    f"Сохранённый сценарий главы {chapter_index + 1} не соответствует режиму "
                    f"{'двух голосов' if self.voice_mode == 'duo' else 'одного голоса'}."
                )
            raise RuntimeError(
                f"Для главы {chapter_index + 1} нет сохранённого TTS-сценария. "
                "Сначала выполните шаг AI-подготовки."
            )

        if self.run_mode == "raw":
            if self.voice_mode == "duo":
                raise RuntimeError("Двухголосый режим требует AI-сценарий. Выберите режим 'Авто' или 'По шагам'.")
            return self.bm.chapters[chapter_index].raw_text.strip()

        raise RuntimeError(f"Неизвестный режим Flash TTS: {self.run_mode}")

    async def _synthesize_chunk(self, tts_client, script_chunk):
        prompt_text = _build_tts_generation_prompt(
            script_chunk,
            self.voice_mode,
            self.speed,
            extra_directive=self.tts_directive,
        )
        config = genai_types.GenerateContentConfig(
            response_modalities=["AUDIO"],
            speech_config=self._build_speech_config(),
        )

        last_error = None
        for attempt in range(1, 4):
            budget_acquired = False
            try:
                budget_acquired = await self._wait_for_request_budget(
                    prompt_text,
                    rpm_limiter=self.tts_rpm_limiter,
                    tpm_limiter=self.tts_tpm_limiter,
                    request_label=f"flash-tts {self.model_id}",
                    daily_request_limiter=self.daily_request_limiter,
                    model_id=self.model_id,
                    rpd_limit=self.tts_rpd_limit,
                )
                response = await _to_thread_with_timeout(
                    f"flash-tts {self.model_id}",
                    READER_GENERATE_CONTENT_TIMEOUT_SEC,
                    tts_client.models.generate_content,
                    model=self.model_id,
                    contents=prompt_text,
                    config=config,
                )
                audio_bytes = _extract_audio_bytes(response)
                if audio_bytes:
                    return audio_bytes
                raise RuntimeError("TTS модель вернула пустой audio payload.")
            except Exception as exc:
                self._release_request_budget(self.model_id, budget_acquired)
                if isinstance(exc, ReaderWorkerStopped):
                    raise
                if isinstance(exc, (RateLimitBudgetError, ProjectRateLimitReachedError)):
                    raise
                if _is_invalid_api_key_error(exc):
                    raise InvalidApiKeyError(str(exc)) from exc
                if _is_rate_limited_error(exc):
                    raise RateLimitBudgetError(
                        f"Ключ {_mask_api_key(self.api_key)} получил RESOURCE_EXHAUSTED/429 для модели {self.model_id}: {exc}",
                        model_id=self.model_id,
                    ) from exc
                last_error = exc
                logger.warning(
                    f"[W{self.worker_id}] Ошибка Flash TTS (попытка {attempt}/3): {exc}"
                )
                if not await self._sleep_interruptibly(2 + attempt):
                    raise ReaderWorkerStopped()

        fallback_text = _clean_tts_markup(script_chunk)
        fallback_audio = await self.get_edge_tts_fallback(fallback_text)
        if fallback_audio:
            logger.warning(
                f"[W{self.worker_id}] Flash TTS не ответил. Использован аварийный Edge fallback "
                f"для {'двухголосого' if self.voice_mode == 'duo' else 'одноголосого'} блока."
            )
            return fallback_audio

        if last_error is not None and _is_rate_limited_error(last_error):
            raise RateLimitBudgetError(
                f"Не удалось синтезировать блок через Flash TTS ({self.model_id}): {last_error}",
                model_id=self.model_id,
            )
        raise RuntimeError(f"Не удалось синтезировать блок через Flash TTS: {last_error}")

    def _flash_tts_progress_signature(self, script_text, script_chunks):
        payload = {
            "script": script_text or "",
            "total_chunks": len(script_chunks or []),
            "voice_mode": self.voice_mode,
            "tts_model_id": self.model_id,
            "primary_voice": self.voice,
            "secondary_voice": self.secondary_voice,
            "speed": self.speed,
            "chunk_limit": self._tts_chunk_limit(),
            "tts_directive": self.tts_directive,
        }
        raw_value = json.dumps(payload, ensure_ascii=False, sort_keys=True)
        return hashlib.sha256(raw_value.encode("utf-8")).hexdigest()

    async def _load_flash_tts_resume(self, chapter_index, script_text, script_chunks):
        if not self.record:
            return 0
        progress = self.bm.load_tts_progress(chapter_index)
        if not progress:
            return 0

        signature = self._flash_tts_progress_signature(script_text, script_chunks)
        total_chunks = len(script_chunks)
        completed_chunks = int(progress.get("completed_chunks", 0) or 0)
        mp3_path = self.bm.get_mp3_path(chapter_index)

        if (
            progress.get("signature") != signature
            or int(progress.get("total_chunks", 0) or 0) != total_chunks
            or completed_chunks <= 0
            or not os.path.exists(mp3_path)
        ):
            self.bm.clear_tts_progress(chapter_index)
            return 0

        completed_chunks = min(completed_chunks, total_chunks)
        try:
            raw_audio = await asyncio.to_thread(_load_mp3_as_raw_pcm, mp3_path)
        except Exception as exc:
            logger.warning(f"[W{self.worker_id}] Не удалось загрузить частичный MP3 главы {chapter_index + 1}: {exc}")
            self.bm.clear_tts_progress(chapter_index)
            return 0
        if not raw_audio:
            self.bm.clear_tts_progress(chapter_index)
            return 0

        self.audio_chunks = [raw_audio]
        logger.info(
            f"[W{self.worker_id}] Возобновление главы {chapter_index + 1}: "
            f"готово {completed_chunks}/{total_chunks} TTS-блок(ов)."
        )
        self.worker_progress.emit(self.worker_id, chapter_index, completed_chunks, total_chunks)
        return completed_chunks

    def _save_flash_tts_progress(self, chapter_index, script_text, script_chunks, completed_chunks):
        if not self.record:
            return
        self.bm.save_tts_progress(
            chapter_index,
            {
                "signature": self._flash_tts_progress_signature(script_text, script_chunks),
                "completed_chunks": int(completed_chunks),
                "total_chunks": len(script_chunks),
                "updated_at": time.time(),
            },
        )

    async def _synthesize_chapter(self, tts_client, chapter_index, script_text):
        self.audio_chunks = []
        script_chunks = _split_tts_script(script_text, self.voice_mode, self._tts_chunk_limit())
        if not script_chunks:
            raise RuntimeError("Подготовленный сценарий пустой после разбиения на блоки.")

        total_chunks = len(script_chunks)
        logger.info(
            f"[W{self.worker_id}] Глава {chapter_index + 1}: {total_chunks} TTS-блок(ов), "
            f"режим={self.voice_mode}, модель={self.model_id}."
        )

        completed_chunks = await self._load_flash_tts_resume(chapter_index, script_text, script_chunks)
        if completed_chunks >= total_chunks:
            self.bm.mark_chapter_done(chapter_index)
            self.chapter_done_ui_signal.emit(chapter_index)
            return True

        for chunk_index, script_chunk in enumerate(script_chunks[completed_chunks:], start=completed_chunks + 1):
            if not self._is_running:
                return False
            audio_bytes = await self._synthesize_chunk(tts_client, script_chunk)
            if self.record:
                with self.buffer_lock:
                    self.audio_chunks.append(audio_bytes)
            if self.audio_queue and not self.fast:
                self.audio_queue.put((audio_bytes, chapter_index, chunk_index - 1, False))
            if self.record:
                await self.save_file()
                self._save_flash_tts_progress(chapter_index, script_text, script_chunks, chunk_index)
            self._emit_worker_progress(chapter_index, chunk_index, total_chunks, force=chunk_index >= total_chunks)

        await self.save_file(final=True)
        self.bm.mark_chapter_done(chapter_index)
        self.chapter_done_ui_signal.emit(chapter_index)
        return True

    async def main_loop(self):
        if genai is None or genai_types is None:
            raise RuntimeError("Для Flash TTS и AI-предобработки требуется пакет google-genai.")

        if self.run_mode in {"auto", "prepare"} and not self.preprocess_model_id:
            raise RuntimeError("Не выбрана модель для AI-предобработки сценария.")

        await self._wait_before_worker_start(READER_FLASH_WORKER_START_STAGGER_SECONDS)
        if not self._is_running:
            return
        text_client = _make_genai_client(self.api_key)
        tts_client = _make_genai_client(self.api_key)
        logger.info(
            f"FlashTTS воркер {self.worker_id} запущен (TTS={self.model_id}, preprocess={self.preprocess_model_id}, "
            f"voice_mode={self.voice_mode}, primary={self.voice}, secondary={self.secondary_voice})."
        )

        while self._is_running:
            if self.c_idx == -1:
                try:
                    self.c_idx = self.manager_chapter_queue.get_nowait()
                except queue.Empty:
                    self._emit_finished()
                    break

            chapter_index = self.c_idx
            if self.run_mode != "prepare" and self.bm.is_chapter_done(chapter_index):
                self.chapter_done_ui_signal.emit(chapter_index)
                self.c_idx = -1
                continue

            try:
                if self.run_mode == "prepare":
                    self._emit_worker_progress(chapter_index, 0, 1, force=True)
                    await self._resolve_script_for_run(text_client, chapter_index)
                    self._emit_worker_progress(chapter_index, 1, 1, force=True)
                    self.c_idx = -1
                    continue

                await self._wait_for_flash_chapter_slot(chapter_index)
                if not self._is_running:
                    break
                script_text = await self._resolve_script_for_run(text_client, chapter_index)
                if not script_text.strip():
                    raise RuntimeError(f"Сценарий главы {chapter_index + 1} пуст.")

                completed = await self._synthesize_chapter(tts_client, chapter_index, script_text)
                if completed and self.worker_id == 0:
                    self.change_chapter_signal.emit(chapter_index)
            except ReaderWorkerStopped:
                break
            except InvalidApiKeyError as exc:
                self._abort_for_invalid_key(str(exc))
                break
            except ProjectRateLimitReachedError as exc:
                self._abort_for_project_quota(str(exc), exc.model_id or self.model_id)
                break
            except RateLimitBudgetError as exc:
                self._abort_for_quota_key(str(exc), exc.model_id or self.model_id)
                break
            except Exception as exc:
                logger.error(f"[W{self.worker_id}] Ошибка обработки главы {chapter_index + 1}: {exc}")
                self.error_signal.emit(self.worker_id, f"Глава {chapter_index + 1}: {exc}")
            finally:
                self.c_idx = -1


# --- UI КОМПОНЕНТЫ ---
# --- UI КОМПОНЕНТЫ ---
class DashboardRow(QWidget):
    def __init__(self, w_id):
        super().__init__()
        self.w_id = w_id
        self._last_progress_signature = None
        
        layout = QHBoxLayout(self)
        layout.setContentsMargins(5, 5, 5, 5)
        
        self.lbl_info = QLabel(f"[W{w_id}] Ожидание...")
        self.lbl_info.setFixedWidth(160)
        self.lbl_info.setStyleSheet("font-weight: bold; font-size: 14px;")
        
        self.pbar = QProgressBar()
        self.pbar.setRange(0, 100)
        self.pbar.setValue(0)
        # ОТКЛЮЧАЕМ текст внутри самой полоски, чтобы небыло "96% 96%"
        self.pbar.setTextVisible(False) 
        self.pbar.setFixedHeight(20)
        
        self.lbl_status = QLabel("--")
        self.lbl_status.setFixedWidth(100)
        self.lbl_status.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
        self.lbl_status.setStyleSheet("font-size: 14px;")
        
        layout.addWidget(self.lbl_info)
        layout.addWidget(self.pbar)
        layout.addWidget(self.lbl_status)

    def update_progress(self, w_id, c_idx, s_idx, total_s):
        """Обновление прогресс-бара с правильным приемом аргументов от сигнала"""
        if total_s <= 0:
            pct_chap = 0
        else:
            # Ограничиваем s_idx, чтобы он не превышал total_s (защита от >100%)
            current_s = min(s_idx, total_s)
            pct_chap = int((current_s / total_s) * 100)

        signature = (c_idx, pct_chap)
        if signature == self._last_progress_signature:
            return
        self._last_progress_signature = signature
        
        self.lbl_info.setText(f"[W{w_id}] Глава {c_idx+1}")
        self.lbl_status.setText(f"{pct_chap}%")
        self.lbl_status.setStyleSheet("font-size: 14px; font-weight: bold;")
        
        # Устанавливаем значение прогресс-бара
        self.pbar.setValue(pct_chap)

    def set_finished(self):
        self.lbl_status.setText("ГОТОВО")
        self.pbar.setValue(100)
        self.hide()

class ApiKeysDialog(QDialog):
    def __init__(self, keys_str, parent=None):
        super().__init__(parent)
        self.setWindowTitle("API Ключи")
        self.resize(400, 300)
        layout = QVBoxLayout(self)
        self.text_edit = QTextEdit()
        self.text_edit.setPlainText(keys_str)
        layout.addWidget(QLabel("Введите ключи (каждый с новой строки):"))
        layout.addWidget(self.text_edit)
        btn = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        btn.accepted.connect(self.accept)
        btn.rejected.connect(self.reject)
        layout.addWidget(btn)
    def get_keys(self):
        return [k.strip() for k in self.text_edit.toPlainText().split('\n') if k.strip()]


class ReaderKeyStatusDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Состояние ключей")
        self.resize(760, 520)

        layout = QVBoxLayout(self)
        self.lbl_summary = QLabel("Ключи: 0")
        self.lbl_summary.setStyleSheet("font-size: 10pt; color: #cfcfcf;")
        layout.addWidget(self.lbl_summary)

        self.list_widget = QListWidget()
        self.list_widget.setSelectionMode(QListWidget.SelectionMode.ExtendedSelection)
        self.list_widget.setToolTip(
            "Список ключей для текущих моделей reader.\n"
            "Если ничего не выделено, действия применяются ко всем ключам из списка."
        )
        layout.addWidget(self.list_widget, 1)

        buttons_row = QHBoxLayout()
        self.btn_refresh = QPushButton("Обновить")
        self.btn_clear_limits = QPushButton("Сбросить лимиты")
        self.btn_clear_invalid = QPushButton("Снять invalid")
        buttons_row.addWidget(self.btn_refresh)
        buttons_row.addWidget(self.btn_clear_limits)
        buttons_row.addWidget(self.btn_clear_invalid)
        buttons_row.addStretch(1)
        layout.addLayout(buttons_row)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        buttons.rejected.connect(self.reject)
        buttons.accepted.connect(self.accept)
        layout.addWidget(buttons)

    def set_snapshot(self, snapshot):
        rows = (snapshot or {}).get("rows", [])
        self.lbl_summary.setText((snapshot or {}).get("summary_text", "Ключи: 0"))
        self.lbl_summary.setToolTip((snapshot or {}).get("summary_tooltip", ""))
        self.list_widget.clear()

        state_colors = {
            "ready": "#90EE90",
            "limited": "#F08080",
            "invalid": "#FFAB91",
            "session": "#FFD54F",
        }

        for row in rows:
            item = QListWidgetItem(row.get("display_text", row.get("key", "")))
            item.setData(Qt.ItemDataRole.UserRole, row.get("key", ""))
            item.setToolTip(row.get("tooltip", ""))
            item.setForeground(QColor(state_colors.get(row.get("state"), "#ECECEC")))
            self.list_widget.addItem(item)

    def selected_keys(self):
        selected = []
        for item in self.list_widget.selectedItems():
            key = item.data(Qt.ItemDataRole.UserRole)
            if key:
                selected.append(key)
        return selected

    def visible_keys(self):
        keys = []
        for i in range(self.list_widget.count()):
            key = self.list_widget.item(i).data(Qt.ItemDataRole.UserRole)
            if key:
                keys.append(key)
        return keys


class PromptTuningDialog(QDialog):
    def __init__(self, preprocess_directive, tts_directive, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Flash TTS: Режиссура промптов")
        self.resize(760, 620)

        layout = QVBoxLayout(self)
        layout.addWidget(QLabel("Дополнительные указания для AI-предобработки сценария:"))

        self.preprocess_edit = QPlainTextEdit()
        self.preprocess_edit.setPlaceholderText(DEFAULT_PREPROCESS_DIRECTIVE)
        self.preprocess_edit.setPlainText(preprocess_directive or "")
        layout.addWidget(self.preprocess_edit, 1)

        layout.addWidget(QLabel("Дополнительные указания для TTS-исполнения:"))

        self.tts_edit = QPlainTextEdit()
        self.tts_edit.setPlaceholderText(DEFAULT_TTS_DIRECTIVE)
        self.tts_edit.setPlainText(tts_directive or "")
        layout.addWidget(self.tts_edit, 1)

        buttons_row = QHBoxLayout()
        btn_reset = QPushButton("Сбросить к стандарту")
        btn_reset.clicked.connect(self._reset_defaults)
        buttons_row.addWidget(btn_reset)
        buttons_row.addStretch(1)
        layout.addLayout(buttons_row)

        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)

    def _reset_defaults(self):
        self.preprocess_edit.setPlainText(DEFAULT_PREPROCESS_DIRECTIVE)
        self.tts_edit.setPlainText(DEFAULT_TTS_DIRECTIVE)

    def get_values(self):
        return (
            self.preprocess_edit.toPlainText().strip(),
            self.tts_edit.toPlainText().strip(),
        )


# --- ТЕСТОВЫЙ ВОСПРОИЗВОДИТЕЛЬ ГОЛОСА ---
class VoiceSampleWorker(QThread):
    finished_signal = pyqtSignal()
    error_signal = pyqtSignal(str)

    def __init__(
        self,
        api_key,
        model_id,
        voice,
        engine_mode="live",
        voice_mode="single",
        secondary_voice=None,
        tertiary_voice=None,
        tts_directive="",
        daily_request_limiter=None,
    ):
        super().__init__()
        self.api_key = api_key
        self.model_id = model_id
        self.voice = voice
        self.engine_mode = engine_mode or "live"
        self.voice_mode = voice_mode or "single"
        self.secondary_voice = secondary_voice or voice
        self.tertiary_voice = tertiary_voice or self.secondary_voice or voice
        self.tts_directive = tts_directive or DEFAULT_TTS_DIRECTIVE
        self.daily_request_limiter = daily_request_limiter
        self.rpd_limit = _lookup_model_limit(self.model_id, "rpd", 0)
        self.test_text = "Проверка голоса. Саша сушит сушки."

    def run(self):
        try:
            asyncio.run(self.get_and_play())
        except Exception as e:
            self.error_signal.emit(str(e))
        finally:
            self.finished_signal.emit()

    async def get_and_play(self):
        if genai is None or genai_types is None:
            raise RuntimeError("Для теста голоса требуется пакет google-genai.")

        if self.daily_request_limiter is not None and self.rpd_limit and self.rpd_limit > 0:
            request_amount = 3 if self.engine_mode == "live" and self.voice_mode == "author_gender" else 1
            acquired, _, limit_value, reset_text = self.daily_request_limiter.try_acquire(
                self.model_id, self.rpd_limit, amount=request_amount, api_key=self.api_key
            )
            if not acquired:
                raise RuntimeError(
                    f"Дневной лимит {limit_value} запросов для модели {self.model_id} исчерпан. "
                    f"Следующий сброс около {reset_text}."
                )

        client = _make_genai_client(self.api_key)

        if self.engine_mode == "flash_tts":
            if self.voice_mode == "author_gender":
                self.voice_mode = "duo"
            if self.voice_mode == "duo":
                test_script = (
                    f"{TTS_SPEAKER_NARRATOR}: [serious] Рассказчик открывает сцену.\n"
                    f"{TTS_SPEAKER_DIALOGUE}: [excited] А я отвечаю другим голосом."
                )
                speech_config = _build_duo_voice_speech_config(self.voice, self.secondary_voice)
            else:
                test_script = "[clear, upbeat] Проверка голоса. Саша сушит сушки."
                speech_config = _build_single_voice_speech_config(self.voice)

            config = genai_types.GenerateContentConfig(
                response_modalities=["AUDIO"],
                speech_config=speech_config,
            )
            response = await _to_thread_with_timeout(
                f"voice sample {self.model_id}",
                READER_GENERATE_CONTENT_TIMEOUT_SEC,
                client.models.generate_content,
                model=self.model_id,
                contents=_build_tts_generation_prompt(
                    test_script,
                    self.voice_mode,
                    "Normal",
                    extra_directive=self.tts_directive,
                ),
                config=config,
            )
            audio_data = _extract_audio_bytes(response)
        else:
            live_test_text = self.test_text
            request_plan = []
            if self.voice_mode == "author_gender":
                request_plan = [
                    (
                        _build_live_single_voice_request_config(self.voice, "author"),
                        "Рассказчик открывает сцену.",
                    ),
                    (
                        _build_live_single_voice_request_config(self.secondary_voice, "male_dialogue"),
                        "Я пришёл первым.",
                    ),
                    (
                        _build_live_single_voice_request_config(self.tertiary_voice, "female_dialogue"),
                        "А я уже всё проверила.",
                    ),
                ]
            elif self.voice_mode == "duo":
                live_test_text = (
                    f"{TTS_SPEAKER_NARRATOR}: Рассказчик открывает сцену.\n"
                    f"{TTS_SPEAKER_DIALOGUE}: А я отвечаю другим голосом."
                )
                request_plan = [
                    (
                        genai_types.LiveConnectConfig(
                            response_modalities=["AUDIO"],
                            speech_config=_build_duo_voice_speech_config(self.voice, self.secondary_voice),
                            system_instruction="Ты профессиональная студия озвучки. Строго соблюдай роли Narrator и Dialogue.",
                        ),
                        live_test_text,
                    )
                ]
            else:
                request_plan = [
                    (
                        _build_live_single_voice_request_config(self.voice, "default"),
                        live_test_text,
                    )
                ]

            audio_data = b""

            for config, text_value in request_plan:
                request_audio = bytearray()
                session_cm = None
                try:
                    session_cm, session = await _open_live_session_with_timeout(client, self.model_id, config)
                    try:
                        await asyncio.wait_for(
                            session.send_client_content(turns=text_value, turn_complete=True),
                            timeout=READER_LIVE_SEND_TIMEOUT_SEC,
                        )
                    except Exception:
                        await asyncio.wait_for(
                            session.send_realtime_input(text=text_value),
                            timeout=READER_LIVE_SEND_TIMEOUT_SEC,
                        )
                    receive_iterator = session.receive().__aiter__()

                    current_timeout = READER_LIVE_FIRST_CHUNK_TIMEOUT_SEC

                    while True:
                        try:
                            response = await asyncio.wait_for(receive_iterator.__anext__(), timeout=current_timeout)
                            current_timeout = READER_LIVE_NEXT_CHUNK_TIMEOUT_SEC

                            if response.server_content:
                                if response.server_content.model_turn:
                                    for part in response.server_content.model_turn.parts:
                                        if part.inline_data:
                                            request_audio.extend(part.inline_data.data)

                                if getattr(response.server_content, "turn_complete", False):
                                    break

                        except asyncio.TimeoutError:
                            break
                        except StopAsyncIteration:
                            break
                finally:
                    await _close_live_session_quietly(session_cm)
                audio_data += _trim_raw_pcm_boundaries(bytes(request_audio))

        if audio_data:
            self.play_audio(audio_data)

    def play_audio(self, data):
        if pyaudio is None:
            raise RuntimeError("Для воспроизведения теста голоса требуется PyAudio.")
        p = pyaudio.PyAudio()
        stream = p.open(
            format=pyaudio.paInt16,
            channels=AUDIO_CHANNELS,
            rate=AUDIO_RATE,
            output=True
        )
        
        arr = array.array('h', data)
        vol = 0.5 # Жестко фиксируем громкость на 50%
        for i in range(len(arr)):
            arr[i] = max(min(int(arr[i] * vol), 32767), -32768)
            
        stream.write(arr.tobytes())
        stream.stop_stream()
        stream.close()
        p.terminate()




# --- ОСНОВНОЕ ОКНО ---
class MainWindow(QMainWindow):
    def __init__(self, settings_manager=None):
        super().__init__()
        self.setWindowTitle("Gemini Reader v21 + Edge Fallback")
        self.resize(1150, 850)
        self.setAcceptDrops(True)
        self.settings_manager = settings_manager or _get_app_settings_manager()
        self.reader_books_dir = _reader_books_dir(self.settings_manager)
        self.daily_request_limiter = ProjectDailyRequestLimiter(self.settings_manager)
        self.preprocess_models_map = _build_preprocess_models_map()
        self.preprocess_directive = DEFAULT_PREPROCESS_DIRECTIVE
        self.tts_directive = DEFAULT_TTS_DIRECTIVE
        self.bm = None
        self.project_settings_manager = None
        self.workers = []
        self.api_keys = []
        self.audio_queue = queue.Queue(maxsize=100)
        self.player = None
        self.combiner = None
        self.tester_worker = None
        self.worker_widgets = {}
        self._return_to_menu_handler = None
        self._returning_to_main_menu = False
        self._active_job_kind = "tts"
        self._active_reader_engine = None
        self._active_flash_run_mode = None
        self._active_manager_queue = None
        self._parallel_live_state = None
        self._run_had_invalid_keys = False
        self._project_quota_message = ""
        self._stop_requested = False
        self._current_chapter_index = None
        self.disabled_api_keys = set()
        self._settings_event_bus = getattr(self.settings_manager, "bus", None) if self.settings_manager is not None else None
        self._loading_settings = False
        self._checked_chapter_indices_state = set()
        self._chapter_check_state_refresh = False
        self._chapter_check_anchor_index = None
        self._chapter_last_press_index = None
        self._chapter_last_press_modifiers = Qt.KeyboardModifier.NoModifier
        self._pending_worker_progress = {}
        self._log_buffer = []
        self._log_flush_timer = QTimer(self)
        self._log_flush_timer.setSingleShot(True)
        self._log_flush_timer.timeout.connect(self._flush_log_buffer)
        self._progress_flush_timer = QTimer(self)
        self._progress_flush_timer.setSingleShot(True)
        self._progress_flush_timer.timeout.connect(self._flush_worker_progress)
        self.init_ui()
        self.load_settings()
        if self._settings_event_bus is not None:
            try:
                self._settings_event_bus.event_posted.connect(self._on_settings_bus_event)
            except Exception:
                pass
        self.apply_runtime_capabilities()
        self._update_key_state_ui()
        log_fifo.new_log.connect(self.add_log_to_ui)

    def init_ui(self):
        # Toolbar
        toolbar = QToolBar()
        self.addToolBar(toolbar)
        self.act_return_to_menu = QAction("← В меню", self)
        self.act_return_to_menu.triggered.connect(self._return_to_menu)
        toolbar.addAction(self.act_return_to_menu)
        toolbar.addSeparator()
        act_key = QAction("🔑 Ключи API", self)
        act_key.triggered.connect(self.set_api_keys)
        toolbar.addAction(act_key)
        
        act_open = QAction("📂 Открыть книгу", self)
        act_open.triggered.connect(self.open_book_dialog)
        toolbar.addAction(act_open)

        self.act_prompt_tuning = QAction("🧠 Flash TTS", self)
        self.act_prompt_tuning.triggered.connect(self.edit_flash_tts_prompts)
        toolbar.addAction(self.act_prompt_tuning)
        
        self.lbl_info = QLabel("Перетащите файл книги")
        toolbar.addSeparator()
        toolbar.addWidget(self.lbl_info)

        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QVBoxLayout(central)

        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # Лево: Список глав
        self.list_chapters = QListWidget()
        self.list_chapters.setFixedWidth(250)
        
        # Включаем множественное выделение (через Shift и Ctrl)
        self.list_chapters.setSelectionMode(QListWidget.SelectionMode.ExtendedSelection)
        self.list_chapters.itemPressed.connect(self._on_chapter_item_pressed)
        self.list_chapters.itemClicked.connect(self.on_chapter_clicked)
        self.list_chapters.itemSelectionChanged.connect(self._on_chapter_selection_changed)
        self.list_chapters.itemChanged.connect(self._on_chapter_item_changed)
        self.list_chapters.setToolTip("Ставьте галочки для batch-обработки. Shift+клик по чекбоксу отмечает диапазон. Ctrl/Shift-выделение строк остаётся для просмотра и контекстных действий.")
        
        # --- ВКЛЮЧАЕМ КОНТЕКСТНОЕ МЕНЮ (ПРАВАЯ КНОПКА МЫШИ) ---
        self.list_chapters.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.list_chapters.customContextMenuRequested.connect(self.show_chapter_context_menu)
        
        splitter.addWidget(self.list_chapters)

        # Право: Табы
        self.tabs = QTabWidget()
        
        # Таб 1: Текст
        self.txt_view = QTextEdit()
        self.txt_view.setReadOnly(True)
        self.tabs.addTab(self.txt_view, "📖 Текст")

        # Таб 2: TTS-сценарий
        self.script_view = QPlainTextEdit()
        self.script_view.setPlaceholderText(
            "Здесь хранится подготовленный TTS-сценарий главы.\n"
            "В режиме 'По шагам' вы можете отредактировать его вручную и сохранить."
        )
        self.tabs.addTab(self.script_view, "🎭 Сценарий")

        # Таб 3: Промпты
        self.prompts_tab = QWidget()
        prompts_layout = QVBoxLayout(self.prompts_tab)
        prompts_layout.setContentsMargins(8, 8, 8, 8)

        self.lbl_prompts_hint = QLabel(
            "Промпты Flash TTS. Здесь видны и редактируются дополнительные указания "
            "для AI-предобработки и TTS-исполнения."
        )
        self.lbl_prompts_hint.setWordWrap(True)
        prompts_layout.addWidget(self.lbl_prompts_hint)

        prompts_layout.addWidget(QLabel("AI-предобработка сценария:"))
        self.preprocess_prompt_view = QPlainTextEdit()
        self.preprocess_prompt_view.setPlaceholderText(DEFAULT_PREPROCESS_DIRECTIVE)
        prompts_layout.addWidget(self.preprocess_prompt_view, 1)

        prompts_layout.addWidget(QLabel("TTS-исполнение:"))
        self.tts_prompt_view = QPlainTextEdit()
        self.tts_prompt_view.setPlaceholderText(DEFAULT_TTS_DIRECTIVE)
        prompts_layout.addWidget(self.tts_prompt_view, 1)

        prompts_buttons = QHBoxLayout()
        self.btn_save_prompts = QPushButton("💾 Сохранить промпты")
        self.btn_save_prompts.clicked.connect(self.save_prompt_settings)
        prompts_buttons.addWidget(self.btn_save_prompts)

        self.btn_reset_prompts = QPushButton("↺ Сбросить")
        self.btn_reset_prompts.clicked.connect(self.reset_prompt_settings)
        prompts_buttons.addWidget(self.btn_reset_prompts)

        self.btn_open_prompt_dialog = QPushButton("🧠 Открыть диалог")
        self.btn_open_prompt_dialog.clicked.connect(self.edit_flash_tts_prompts)
        prompts_buttons.addWidget(self.btn_open_prompt_dialog)
        prompts_buttons.addStretch(1)
        prompts_layout.addLayout(prompts_buttons)

        self.tabs.addTab(self.prompts_tab, "🧠 Промпты")

        # Таб 4: Дашборд
        self.scroll_dash = QScrollArea()
        self.scroll_dash.setWidgetResizable(True)
        self.dash_content = QWidget()
        self.dash_layout = QVBoxLayout(self.dash_content)
        self.dash_layout.setAlignment(Qt.AlignmentFlag.AlignTop)
        self.scroll_dash.setWidget(self.dash_content)
        self.tabs.addTab(self.scroll_dash, "📊 Воркеры")

        # Таб 5: Лог
        self.log_view = QPlainTextEdit()
        self.log_view.setReadOnly(True)
        self.log_view.document().setMaximumBlockCount(READER_LOG_MAX_BLOCKS)
        self.log_view.setStyleSheet("background-color: #1e1e1e; color: #ececec; font-family: Consolas; font-size: 13px;")
        self.tabs.addTab(self.log_view, "📋 Лог")

        splitter.addWidget(self.tabs)
        main_layout.addWidget(splitter)

        # Контролы
        controls_split = QHBoxLayout()
        controls_split.setSpacing(12)
        workflow_controls = QVBoxLayout()
        workflow_controls.setSpacing(8)
        top_controls = QHBoxLayout()
        middle_controls = QHBoxLayout()
        script_controls = QHBoxLayout()
        option_controls = QHBoxLayout()

        self.live_models_map = dict(LIVE_AUDIO_MODELS)
        self.flash_tts_models_map = dict(FLASH_TTS_MODELS)
        self.models_map = {}

        self.combo_engine = QComboBox()
        for label, mode_id in ENGINE_MODES.items():
            self.combo_engine.addItem(label, mode_id)
        self.combo_engine.currentIndexChanged.connect(self._on_engine_changed)
        self.combo_engine.currentIndexChanged.connect(self._update_worker_spinbox_limit)
        self.combo_engine.currentIndexChanged.connect(self._update_key_state_ui)
        self.combo_engine.currentIndexChanged.connect(self.save_settings)
        top_controls.addWidget(QLabel("Движок:"))
        top_controls.addWidget(self.combo_engine)

        self.combo_model = QComboBox()
        self.combo_model.currentIndexChanged.connect(self._update_worker_spinbox_limit)
        self.combo_model.currentIndexChanged.connect(self._update_key_state_ui)
        self.combo_model.currentTextChanged.connect(self.save_settings)
        top_controls.addWidget(QLabel("Модель TTS:"))
        top_controls.addWidget(self.combo_model)

        self.combo_voice_mode = QComboBox()
        for label, mode_id in VOICE_MODE_OPTIONS.items():
            self.combo_voice_mode.addItem(label, mode_id)
        self.combo_voice_mode.currentIndexChanged.connect(self._on_voice_mode_changed)
        self.combo_voice_mode.currentIndexChanged.connect(self.save_settings)
        top_controls.addWidget(QLabel("Голоса:"))
        top_controls.addWidget(self.combo_voice_mode)
        top_controls.addStretch(1)

        self.combo_voices = QComboBox()
        self.combo_voice_secondary = QComboBox()
        self.combo_voice_tertiary = QComboBox()
        for voice_id, gender in VOICES_MAP.items():
            display_text = f"{voice_id} ({gender})"
            self.combo_voices.addItem(display_text, voice_id)
            self.combo_voice_secondary.addItem(display_text, voice_id)
            self.combo_voice_tertiary.addItem(display_text, voice_id)

        self.combo_voices.currentIndexChanged.connect(self.save_settings)
        self.combo_voice_secondary.currentIndexChanged.connect(self.save_settings)
        self.combo_voice_tertiary.currentIndexChanged.connect(self.save_settings)
        self.lbl_voice_primary = QLabel("Voice A:")
        middle_controls.addWidget(self.lbl_voice_primary)
        middle_controls.addWidget(self.combo_voices)
        self.lbl_voice_secondary = QLabel("Voice B:")
        middle_controls.addWidget(self.lbl_voice_secondary)
        middle_controls.addWidget(self.combo_voice_secondary)
        self.lbl_voice_tertiary = QLabel("Voice C:")
        middle_controls.addWidget(self.lbl_voice_tertiary)
        middle_controls.addWidget(self.combo_voice_tertiary)

        self.btn_test_voice = QPushButton("🔊 Плей")
        self.btn_test_voice.setFixedSize(90, 28)
        self.btn_test_voice.clicked.connect(self.test_voice_sample)
        middle_controls.addWidget(self.btn_test_voice)

        self.combo_speed = QComboBox()
        self.combo_speed.addItems(list(SPEED_PROMPTS.keys()))
        self.combo_speed.setCurrentText("Normal")
        self.combo_speed.currentTextChanged.connect(self.save_settings)
        middle_controls.addWidget(QLabel("Скорость:"))
        middle_controls.addWidget(self.combo_speed)

        self.combo_live_segment_mode = QComboBox()
        for label, mode_id in LIVE_SEGMENT_OPTIONS.items():
            self.combo_live_segment_mode.addItem(label, mode_id)
        self.combo_live_segment_mode.currentIndexChanged.connect(self._refresh_live_segment_controls)
        self.combo_live_segment_mode.currentIndexChanged.connect(self.save_settings)
        self.lbl_live_segment_mode = QLabel("Live:")
        middle_controls.addWidget(self.lbl_live_segment_mode)
        middle_controls.addWidget(self.combo_live_segment_mode)

        self.spin_chunk = QSpinBox()
        self.spin_chunk.setRange(1, 50)
        self.spin_chunk.setValue(2)
        self.spin_chunk.valueChanged.connect(self.save_settings)
        self.lbl_chunk = QLabel("Блок:")
        middle_controls.addWidget(self.lbl_chunk)
        middle_controls.addWidget(self.spin_chunk)

        self.spin_workers = QSpinBox()
        self.spin_workers.setRange(1, 1)
        self.spin_workers.setValue(1)
        self.spin_workers.setToolTip(
            "Максимальное число параллельных воркеров. "
            "Фактический запуск дополнительно ограничивается доступными ключами и числом глав."
        )
        self.spin_workers.valueChanged.connect(self.save_settings)
        middle_controls.addWidget(QLabel("Воркеры:"))
        middle_controls.addWidget(self.spin_workers)

        self.btn_play = QPushButton("▶ СТАРТ")
        self.btn_play.setFixedSize(110, 38)
        self.btn_play.clicked.connect(self.toggle_play)
        middle_controls.addWidget(self.btn_play)

        self.btn_stop = QPushButton("⏹ СТОП")
        self.btn_stop.setFixedSize(110, 38)
        self.btn_stop.setEnabled(False)
        self.btn_stop.clicked.connect(self.force_stop)
        middle_controls.addWidget(self.btn_stop)
        middle_controls.addStretch(1)

        self.combo_preprocess_model = QComboBox()
        for display_name, model_id in self.preprocess_models_map.items():
            self.combo_preprocess_model.addItem(display_name, model_id)
        self.combo_preprocess_model.currentIndexChanged.connect(self._update_worker_spinbox_limit)
        self.combo_preprocess_model.currentIndexChanged.connect(self._update_key_state_ui)
        self.combo_preprocess_model.currentIndexChanged.connect(self.save_settings)
        script_controls.addWidget(QLabel("AI сценарий:"))
        script_controls.addWidget(self.combo_preprocess_model)

        self.combo_preprocess_profile = QComboBox()
        for label, profile_prompt in PREPROCESS_PROFILE_OPTIONS.items():
            self.combo_preprocess_profile.addItem(label, profile_prompt)
        self.combo_preprocess_profile.currentIndexChanged.connect(self.save_settings)
        script_controls.addWidget(QLabel("Профиль:"))
        script_controls.addWidget(self.combo_preprocess_profile)

        self.combo_pipeline_mode = QComboBox()
        for label, mode_id in PIPELINE_MODE_OPTIONS.items():
            self.combo_pipeline_mode.addItem(label, mode_id)
        self.combo_pipeline_mode.currentIndexChanged.connect(self._update_worker_spinbox_limit)
        self.combo_pipeline_mode.currentIndexChanged.connect(self._update_key_state_ui)
        self.combo_pipeline_mode.currentIndexChanged.connect(self.save_settings)
        script_controls.addWidget(QLabel("Pipeline:"))
        script_controls.addWidget(self.combo_pipeline_mode)

        self.btn_prepare_script = QPushButton("🪄 AI сценарий")
        self.btn_prepare_script.setFixedSize(125, 30)
        self.btn_prepare_script.clicked.connect(self.prepare_selected_scripts)
        script_controls.addWidget(self.btn_prepare_script)

        self.btn_save_script = QPushButton("💾 Сохранить сценарий")
        self.btn_save_script.setFixedSize(160, 30)
        self.btn_save_script.clicked.connect(self.save_current_script)
        script_controls.addWidget(self.btn_save_script)
        script_controls.addStretch(1)

        self.chk_mp3 = QCheckBox("Запись MP3")
        self.chk_mp3.setChecked(True)
        self.chk_mp3.stateChanged.connect(self.save_settings)
        self.chk_fast = QCheckBox("Только экспорт")
        self.chk_fast.stateChanged.connect(self.save_settings)
        self.chk_edge_fallback = QCheckBox("Edge fallback")
        self.chk_edge_fallback.setChecked(True)
        self.chk_edge_fallback.setToolTip("Если отключено, reader не будет использовать Microsoft Edge TTS как аварийную озвучку.")
        self.chk_edge_fallback.stateChanged.connect(self.save_settings)
        self.chk_selected_only = QCheckBox("Только отмеченные главы")
        self.chk_selected_only.stateChanged.connect(self.save_settings)
        self.chk_selected_only.stateChanged.connect(self._on_chapter_selection_changed)
        self.chk_parallel_single_chapter = QCheckBox("1 глава = много воркеров")
        self.chk_parallel_single_chapter.setToolTip(
            "Только для Live API. Если выбрана одна глава, reader разделит её на блоки и "
            "раздаст их нескольким воркерам с последующей автоматической сборкой."
        )
        self.chk_parallel_single_chapter.stateChanged.connect(self.save_settings)
        option_controls.addWidget(self.chk_mp3)
        option_controls.addWidget(self.chk_fast)
        option_controls.addWidget(self.chk_edge_fallback)
        option_controls.addWidget(self.chk_selected_only)
        option_controls.addWidget(self.chk_parallel_single_chapter)
        option_controls.addStretch(1)

        workflow_controls.addLayout(top_controls)
        workflow_controls.addLayout(middle_controls)
        workflow_controls.addLayout(script_controls)
        workflow_controls.addLayout(option_controls)

        actions_panel_widget = QWidget()
        actions_panel_widget.setMinimumWidth(210)
        actions_panel_widget.setMaximumWidth(230)
        actions_panel = QVBoxLayout(actions_panel_widget)
        actions_panel.setContentsMargins(0, 0, 0, 0)
        actions_panel.setSpacing(6)

        lbl_project_actions = QLabel("Проект")
        actions_panel.addWidget(lbl_project_actions)

        self.lbl_chapter_scope = QLabel("Главы: все")
        self.lbl_chapter_scope.setToolTip("Область обработки для кнопок Старт и AI сценарий.")
        self.lbl_chapter_scope.setWordWrap(True)
        actions_panel.addWidget(self.lbl_chapter_scope)

        self.lbl_key_state = QLabel("Ключи: 0")
        self.lbl_key_state.setToolTip("Сводка по состоянию ключей для текущих моделей reader.")
        self.lbl_key_state.setWordWrap(True)
        actions_panel.addWidget(self.lbl_key_state)

        self.btn_key_status = QPushButton("Статус ключей")
        self.btn_key_status.setFixedHeight(30)
        self.btn_key_status.clicked.connect(self.show_key_status_dialog)
        actions_panel.addWidget(self.btn_key_status)

        self.btn_pick_chapters = QPushButton("Выбрать главы")
        self.btn_pick_chapters.setFixedHeight(30)
        self.btn_pick_chapters.clicked.connect(self.pick_chapters_dialog)
        actions_panel.addWidget(self.btn_pick_chapters)

        self.btn_clear_chapter_selection = QPushButton("Сбросить отметки")
        self.btn_clear_chapter_selection.setFixedHeight(30)
        self.btn_clear_chapter_selection.clicked.connect(self.clear_chapter_selection)
        actions_panel.addWidget(self.btn_clear_chapter_selection)

        self.btn_clean_stuck = QPushButton("🧹 Очистить зависшие")
        self.btn_clean_stuck.setFixedHeight(30)
        self.btn_clean_stuck.setStyleSheet("background-color: #ffe0b2;")
        self.btn_clean_stuck.clicked.connect(self.clean_stuck_files)
        actions_panel.addWidget(self.btn_clean_stuck)

        actions_panel.addSpacing(4)
        lbl_export_actions = QLabel("Экспорт")
        actions_panel.addWidget(lbl_export_actions)

        self.lbl_export_folder = QLabel("Экспорт: откройте книгу")
        self.lbl_export_folder.setToolTip("Папка текущей книги, куда сохраняются MP3 и видео.")
        self.lbl_export_folder.setWordWrap(True)
        actions_panel.addWidget(self.lbl_export_folder)

        self.btn_open_export_folder = QPushButton("📁 Папка экспорта")
        self.btn_open_export_folder.setFixedHeight(30)
        self.btn_open_export_folder.clicked.connect(self.open_export_folder)
        actions_panel.addWidget(self.btn_open_export_folder)

        self.btn_combine = QPushButton("🧩 Склеить MP3")
        self.btn_combine.setFixedHeight(30)
        self.btn_combine.clicked.connect(self.run_combine)
        actions_panel.addWidget(self.btn_combine)

        actions_panel.addSpacing(4)
        lbl_video_actions = QLabel("Видео")
        actions_panel.addWidget(lbl_video_actions)

        self.lbl_video_cover = QLabel("Видео: картинка не выбрана")
        self.lbl_video_cover.setToolTip("Выбранная картинка будет скопирована в папку книги и использована для экспорта видео.")
        self.lbl_video_cover.setWordWrap(True)
        actions_panel.addWidget(self.lbl_video_cover)

        self.btn_select_video_cover = QPushButton("Картинка видео")
        self.btn_select_video_cover.setFixedHeight(30)
        self.btn_select_video_cover.clicked.connect(self.select_video_cover)
        actions_panel.addWidget(self.btn_select_video_cover)

        self.btn_export_video = QPushButton("Экспорт видео")
        self.btn_export_video.setFixedHeight(30)
        self.btn_export_video.clicked.connect(self.run_export_video)
        actions_panel.addWidget(self.btn_export_video)
        actions_panel.addStretch(1)

        controls_split.addLayout(workflow_controls, 1)
        controls_split.addWidget(actions_panel_widget, 0)
        main_layout.addLayout(controls_split)

        self._on_engine_changed()
        self._on_voice_mode_changed()
        self._on_chapter_selection_changed()
        self._update_export_folder_status()
        self._update_video_cover_status()
        self._sync_prompt_views()

    def _running_tasks_exist(self):
        active_workers = any(getattr(worker, "isRunning", lambda: False)() for worker in self.workers)
        combiner_running = bool(self.combiner and self.combiner.isRunning())
        tester_running = bool(self.tester_worker and self.tester_worker.isRunning())
        return active_workers or combiner_running or tester_running

    def _refresh_runtime_controls(self):
        running = self._running_tasks_exist()
        is_flash_tts = self._is_flash_tts_mode()
        has_genai = genai is not None and genai_types is not None
        has_pyaudio = pyaudio is not None

        self.btn_play.setText("⏳ ИДЁТ" if running else "▶ СТАРТ")
        self.btn_play.setEnabled(not running and has_genai)
        self.btn_stop.setEnabled(running)
        self.btn_clean_stuck.setEnabled(not running)
        chapter_scope_enabled = not running and self.bm is not None
        self.chk_selected_only.setEnabled(chapter_scope_enabled)
        self.btn_pick_chapters.setEnabled(chapter_scope_enabled)
        self.btn_clear_chapter_selection.setEnabled(chapter_scope_enabled)
        self.btn_select_video_cover.setEnabled(chapter_scope_enabled)
        has_video_tools = _resolve_tool_path("ffmpeg") is not None and _resolve_tool_path("ffprobe") is not None
        has_audio_tools = AudioSegment is not None and has_video_tools
        self.btn_open_export_folder.setEnabled(bool(self.bm and self.bm.book_dir))
        self.btn_combine.setEnabled(chapter_scope_enabled and has_audio_tools)
        has_video_cover = bool(self.bm and self.bm.get_video_cover_path())
        self.btn_export_video.setEnabled(chapter_scope_enabled and has_video_tools and has_video_cover)

        author_gender_script_mode = not is_flash_tts and self._selected_voice_mode() == "author_gender"
        flash_controls_enabled = not running and is_flash_tts and has_genai
        ai_script_controls_enabled = not running and has_genai and (is_flash_tts or author_gender_script_mode)
        self.btn_prepare_script.setEnabled(ai_script_controls_enabled)
        self.btn_save_script.setEnabled(ai_script_controls_enabled)
        self.combo_preprocess_model.setEnabled(ai_script_controls_enabled)
        self.combo_preprocess_profile.setEnabled(ai_script_controls_enabled)
        self.combo_pipeline_mode.setEnabled(flash_controls_enabled)
        self.combo_voice_mode.setEnabled(not running and has_genai)
        self.act_prompt_tuning.setEnabled((is_flash_tts or author_gender_script_mode) and not running)
        self.preprocess_prompt_view.setReadOnly(running)
        self.tts_prompt_view.setReadOnly(running)
        self.btn_save_prompts.setEnabled(not running)
        self.btn_reset_prompts.setEnabled(not running)
        self.btn_open_prompt_dialog.setEnabled(not running)
        live_controls_enabled = not running and not is_flash_tts
        self.combo_live_segment_mode.setEnabled(live_controls_enabled)
        self.chk_parallel_single_chapter.setEnabled(live_controls_enabled)

        self.btn_test_voice.setEnabled(not running and has_genai and has_pyaudio)
        self._update_prompt_hint()

    def _set_reading_controls_running(self, running):
        self._refresh_runtime_controls()

    def _current_engine_id(self):
        return self.combo_engine.currentData() or "live"

    def _selected_voice_mode(self):
        return self.combo_voice_mode.currentData() or "single"

    def _selected_live_segment_mode(self):
        return self.combo_live_segment_mode.currentData() or "sentences"

    def _selected_pipeline_mode(self):
        return self.combo_pipeline_mode.currentData() or "auto"

    def _is_flash_tts_mode(self):
        return self._current_engine_id() == "flash_tts"

    def _refresh_model_combo(self, models_map):
        previous_label = self.combo_model.currentText()
        self.models_map = dict(models_map)
        self.combo_model.blockSignals(True)
        self.combo_model.clear()
        for display_name, model_id in self.models_map.items():
            self.combo_model.addItem(display_name, model_id)
        if previous_label and previous_label in self.models_map:
            self.combo_model.setCurrentText(previous_label)
        self.combo_model.blockSignals(False)

    def _on_engine_changed(self):
        is_flash_tts = self._is_flash_tts_mode()
        self._refresh_model_combo(self.flash_tts_models_map if is_flash_tts else self.live_models_map)

        if not is_flash_tts:
            self.combo_pipeline_mode.setCurrentIndex(self.combo_pipeline_mode.findData("auto"))
        elif self._selected_voice_mode() == "author_gender":
            duo_idx = self.combo_voice_mode.findData("duo")
            if duo_idx >= 0:
                self.combo_voice_mode.setCurrentIndex(duo_idx)

        self._on_voice_mode_changed()
        self._refresh_live_segment_controls()
        self._refresh_runtime_controls()

    def _refresh_live_segment_controls(self, *_args):
        if not hasattr(self, "combo_live_segment_mode"):
            return
        is_live_mode = not self._is_flash_tts_mode()
        self.lbl_live_segment_mode.setVisible(is_live_mode)
        self.combo_live_segment_mode.setVisible(is_live_mode)
        self.chk_parallel_single_chapter.setVisible(is_live_mode)
        if not is_live_mode:
            self.lbl_chunk.setText("Блок:")
            self.spin_chunk.setToolTip(
                "Базовый размер блока для текущего TTS-режима."
            )
            return
        segment_mode = self._selected_live_segment_mode()
        if segment_mode == "paragraphs":
            self.lbl_chunk.setText("Блок абз.:")
            self.spin_chunk.setToolTip(
                "Сколько параграфов отправлять одним запросом в Live API."
            )
        else:
            self.lbl_chunk.setText("Блок предл.:")
            self.spin_chunk.setToolTip(
                "Сколько предложений отправлять одним запросом в Live API."
            )

    def _worker_models_for_limit(self):
        if self._is_flash_tts_mode():
            model_ids = []
            pipeline_mode = self._selected_pipeline_mode()
            if pipeline_mode in {"auto", "prepare"}:
                model_ids.append(self._selected_preprocess_model_id())
            if pipeline_mode != "prepare":
                model_ids.append(self._selected_model_id())
            return [model_id for model_id in model_ids if model_id]
        return [self._selected_model_id()]

    def _update_worker_spinbox_limit(self, *_args):
        if not hasattr(self, "spin_workers"):
            return
        current_value = self.spin_workers.value()
        available_count = len(self._get_available_api_keys(self._worker_models_for_limit()))
        max_workers = max(1, available_count)
        self.spin_workers.blockSignals(True)
        self.spin_workers.setMaximum(max_workers)
        self.spin_workers.setValue(max(1, min(current_value, max_workers)))
        self.spin_workers.blockSignals(False)

    def _on_voice_mode_changed(self):
        mode = self._selected_voice_mode()
        show_secondary_voice = mode in {"duo", "author_gender"}
        show_tertiary_voice = mode == "author_gender"

        if mode == "single":
            self.lbl_voice_primary.setText("Голос:")
            self.lbl_voice_secondary.setText("Voice B:")
            self.lbl_voice_tertiary.setText("Voice C:")
        elif mode == "duo":
            self.lbl_voice_primary.setText("Narrator:")
            self.lbl_voice_secondary.setText("Dialogue:")
            self.lbl_voice_tertiary.setText("Voice C:")
        else:
            self.lbl_voice_primary.setText("Автор:")
            self.lbl_voice_secondary.setText("Муж.:")
            self.lbl_voice_tertiary.setText("Жен.:")

        self.lbl_voice_secondary.setVisible(show_secondary_voice)
        self.combo_voice_secondary.setVisible(show_secondary_voice)
        self.lbl_voice_secondary.setEnabled(show_secondary_voice)
        self.combo_voice_secondary.setEnabled(show_secondary_voice)

        self.lbl_voice_tertiary.setVisible(show_tertiary_voice)
        self.combo_voice_tertiary.setVisible(show_tertiary_voice)
        self.lbl_voice_tertiary.setEnabled(show_tertiary_voice)
        self.combo_voice_tertiary.setEnabled(show_tertiary_voice)

    def _selected_model_id(self):
        return self.combo_model.currentData() or self.models_map.get(self.combo_model.currentText(), MODEL_ID)

    def _selected_preprocess_model_id(self):
        return self.combo_preprocess_model.currentData() or self.preprocess_models_map.get(
            self.combo_preprocess_model.currentText(), ""
        )

    def _reader_status_model_ids(self):
        if not hasattr(self, "combo_engine"):
            return []
        if self._is_flash_tts_mode():
            pipeline_mode = self._selected_pipeline_mode()
            if pipeline_mode == "prepare":
                return [model_id for model_id in [self._selected_preprocess_model_id()] if model_id]
            return [model_id for model_id in [self._selected_model_id()] if model_id]
        return [model_id for model_id in [self._selected_model_id()] if model_id]

    def _reader_request_count_for_key(self, api_key, model_id):
        if not api_key or not model_id:
            return 0
        if self.settings_manager is not None:
            try:
                key_info = self.settings_manager.get_key_info(api_key)
                if key_info:
                    return int(self.settings_manager.get_request_count(key_info, model_id) or 0)
            except Exception:
                pass
        if self.daily_request_limiter is not None:
            try:
                return int(self.daily_request_limiter.get_count(model_id, api_key=api_key) or 0)
            except Exception:
                pass
        return 0

    def _reader_request_limit_for_model(self, model_id):
        try:
            return int(_lookup_model_limit(model_id, "rpd", 0) or 0)
        except Exception:
            return 0

    def _reader_key_model_quota_exhausted(self, api_key, model_id):
        limit_value = self._reader_request_limit_for_model(model_id)
        if limit_value <= 0:
            return False
        return self._reader_request_count_for_key(api_key, model_id) >= limit_value

    def _load_invalid_key_states(self):
        raw_state = {}
        if self.settings_manager is not None:
            try:
                raw_state = self.settings_manager.load_settings().get(READER_INVALID_KEYS_KEY, {}) or {}
            except Exception:
                raw_state = {}
        else:
            raw_state = _load_legacy_settings().get(READER_INVALID_KEYS_KEY, {}) or {}

        normalized = {}
        if not isinstance(raw_state, dict):
            return normalized

        for api_key, state in raw_state.items():
            if not isinstance(api_key, str) or not api_key.strip() or not isinstance(state, dict):
                continue
            normalized[api_key] = {
                "invalid_at": state.get("invalid_at"),
                "message": str(state.get("message") or "").strip(),
                "source_model": str(state.get("source_model") or "").strip(),
            }
        return normalized

    def _save_invalid_key_states(self, states):
        normalized = {}
        for api_key, state in (states or {}).items():
            if not isinstance(api_key, str) or not api_key.strip() or not isinstance(state, dict):
                continue
            normalized[api_key] = {
                "invalid_at": state.get("invalid_at"),
                "message": str(state.get("message") or "").strip(),
                "source_model": str(state.get("source_model") or "").strip(),
            }

        if self.settings_manager is not None:
            try:
                self.settings_manager.save_ui_state({READER_INVALID_KEYS_KEY: normalized})
                return
            except Exception:
                pass

        legacy_data = _load_legacy_settings()
        legacy_data[READER_INVALID_KEYS_KEY] = normalized
        with open(LEGACY_SETTINGS_FILE, "w", encoding="utf-8") as file_obj:
            json.dump(legacy_data, file_obj, ensure_ascii=False, indent=2)

    def _prune_invalid_key_states(self, valid_keys=None):
        valid_keys = {key for key in (valid_keys or self.api_keys) if (key or "").strip()}
        current_states = self._load_invalid_key_states()
        pruned_states = {key: value for key, value in current_states.items() if key in valid_keys}
        if pruned_states != current_states:
            self._save_invalid_key_states(pruned_states)
        return pruned_states

    def _mark_key_invalid(self, api_key, error_text="", model_id=""):
        if not (api_key or "").strip():
            return
        current_states = self._load_invalid_key_states()
        current_states[api_key] = {
            "invalid_at": time.time(),
            "message": (str(error_text or "").strip() or "API key not valid.")[:500],
            "source_model": str(model_id or "").strip(),
        }
        self._save_invalid_key_states(current_states)

    def _clear_invalid_key_states(self, keys=None):
        target_keys = {key for key in (keys or self.api_keys) if (key or "").strip()}
        if not target_keys:
            return False
        current_states = self._load_invalid_key_states()
        changed = False
        for key in list(current_states.keys()):
            if key in target_keys:
                current_states.pop(key, None)
                self.disabled_api_keys.discard(key)
                changed = True
        if changed:
            self._save_invalid_key_states(current_states)
        return changed

    def _save_api_keys_with_status_preservation(self, keys):
        normalized_keys = [key for key in dict.fromkeys((key or "").strip() for key in keys) if key]
        if self.settings_manager is not None:
            try:
                existing_statuses = self.settings_manager.load_key_statuses() or []
            except Exception:
                existing_statuses = []
            existing_map = {
                item.get("key"): item
                for item in existing_statuses
                if isinstance(item, dict) and item.get("key")
            }
            merged_statuses = []
            for key in normalized_keys:
                if key in existing_map:
                    merged_statuses.append(existing_map[key])
                else:
                    merged_statuses.append({"key": key, "provider": "gemini", "status_by_model": {}})
            self.settings_manager.save_key_statuses(merged_statuses)
        else:
            legacy_data = _load_legacy_settings()
            legacy_data["api_keys"] = normalized_keys
            with open(LEGACY_SETTINGS_FILE, "w", encoding="utf-8") as file_obj:
                json.dump(legacy_data, file_obj, ensure_ascii=False, indent=2)

        self._prune_invalid_key_states(normalized_keys)
        self.api_keys = normalized_keys

    def _reader_request_count_for_key(self, api_key, model_id):
        if not api_key or not model_id:
            return 0
        if self.settings_manager is not None:
            try:
                key_info = self.settings_manager.get_key_info(api_key)
                if key_info:
                    return int(self.settings_manager.get_request_count(key_info, model_id) or 0)
            except Exception:
                pass
        if self.daily_request_limiter is not None:
            return int(self.daily_request_limiter.get_count(model_id, api_key=api_key) or 0)
        return 0

    def _reader_request_limit_for_model(self, model_id):
        if not model_id:
            return 0
        return int(_lookup_model_limit(model_id, "rpd", 0) or 0)

    def _reader_key_model_quota_exhausted(self, api_key, model_id):
        limit = self._reader_request_limit_for_model(model_id)
        if not limit or limit <= 0:
            return False
        return self._reader_request_count_for_key(api_key, model_id) >= limit

    def _build_key_state_snapshot(self, required_model_ids=None):
        model_ids = [model_id for model_id in (required_model_ids or self._reader_status_model_ids()) if model_id]
        invalid_states = self._prune_invalid_key_states()
        rows = []
        counts = {"ready": 0, "limited": 0, "invalid": 0, "session": 0}
        usage_by_model = {
            model_id: {"used": 0, "limit": 0}
            for model_id in model_ids
        }

        for api_key in [key for key in self.api_keys if (key or "").strip()]:
            key_info = self.settings_manager.get_key_info(api_key) if self.settings_manager is not None else None
            request_parts = []
            limited_parts = []
            request_total = 0
            request_limit_total = 0
            for model_id in model_ids:
                request_count = self._reader_request_count_for_key(api_key, model_id)
                request_limit = self._reader_request_limit_for_model(model_id)
                if model_id in usage_by_model:
                    usage_by_model[model_id]["used"] += int(request_count)
                    if request_limit > 0:
                        usage_by_model[model_id]["limit"] += int(request_limit)
                if request_limit > 0:
                    request_limit_total += request_limit
                if request_count:
                    request_total += int(request_count)
                if request_count or request_limit:
                    if request_limit > 0:
                        request_parts.append(f"{model_id}: {request_count}/{request_limit}")
                    else:
                        request_parts.append(f"{model_id}: {request_count}")
                try:
                    settings_limited = bool(
                        key_info
                        and self.settings_manager is not None
                        and self.settings_manager.is_key_limit_active(key_info, model_id)
                    )
                except Exception:
                    settings_limited = False
                count_limited = self._reader_key_model_quota_exhausted(api_key, model_id)
                if settings_limited:
                    if key_info and self.settings_manager is not None:
                        reset_text = self.settings_manager.get_key_reset_time_str(key_info, model_id)
                    else:
                        reset_text = _next_policy_reset_text(_gemini_reset_policy())
                    limited_parts.append(f"{model_id} ({reset_text})")
                elif count_limited:
                    reset_text = f"сброс около {_next_policy_reset_text(_gemini_reset_policy())}"
                    limited_parts.append(f"{model_id} ({reset_text})")

            invalid_state = invalid_states.get(api_key)
            if invalid_state:
                state = "invalid"
                status_text = "невалидный"
            elif limited_parts:
                state = "limited"
                status_text = "лимит"
            elif api_key in self.disabled_api_keys:
                state = "session"
                status_text = "отключён в сессии"
            else:
                state = "ready"
                status_text = "готов"
            counts[state] += 1

            tooltip_lines = [f"Полный ключ: {api_key}"]
            if model_ids:
                tooltip_lines.append(f"Модели: {', '.join(model_ids)}")
            if request_parts:
                tooltip_lines.append("Запросы: " + ", ".join(request_parts))
            if limited_parts:
                tooltip_lines.append("Лимиты: " + "; ".join(limited_parts))
            if invalid_state:
                invalid_at = _format_local_timestamp(invalid_state.get("invalid_at"))
                if invalid_at:
                    tooltip_lines.append(f"Помечен invalid: {invalid_at}")
                if invalid_state.get("source_model"):
                    tooltip_lines.append(f"Модель ошибки: {invalid_state['source_model']}")
                if invalid_state.get("message"):
                    tooltip_lines.append(invalid_state["message"])
            elif api_key in self.disabled_api_keys:
                tooltip_lines.append("Ключ временно отключён в текущем запуске reader.")
            elif not limited_parts:
                tooltip_lines.append("Статус: активен.")

            if request_limit_total > 0:
                short_requests = f" | req {request_total}/{request_limit_total}"
            else:
                short_requests = f" | req {request_total}" if request_total else ""
            display_text = f"{_mask_api_key(api_key)} — {status_text}{short_requests}"
            rows.append(
                {
                    "key": api_key,
                    "state": state,
                    "display_text": display_text,
                    "tooltip": "\n".join(tooltip_lines),
                }
            )

        state_order = {"invalid": 0, "limited": 1, "session": 2, "ready": 3}
        rows.sort(key=lambda row: (state_order.get(row.get("state"), 9), row.get("display_text", "")))
        total = sum(counts.values())
        summary_parts = [f"готовы {counts['ready']}"]
        if counts["limited"]:
            summary_parts.append(f"лимит {counts['limited']}")
        if counts["invalid"]:
            summary_parts.append(f"невалидны {counts['invalid']}")
        if counts["session"]:
            summary_parts.append(f"сессия {counts['session']}")
        summary_text = f"Ключи: {total}" if not total else f"Ключи: {total} ({' / '.join(summary_parts)})"
        usage_parts = []
        for model_id, usage in usage_by_model.items():
            if usage["limit"] > 0:
                usage_parts.append(f"{model_id}: {usage['used']}/{usage['limit']}")
            elif usage["used"]:
                usage_parts.append(f"{model_id}: {usage['used']}")
        if usage_parts:
            summary_text += " | Запросы: " + "; ".join(usage_parts)
        summary_tooltip = "Актуальные модели: " + (", ".join(model_ids) if model_ids else "не выбраны")
        if usage_parts:
            summary_tooltip += "\nПотраченные запросы: " + "; ".join(usage_parts)
        return {
            "rows": rows,
            "counts": counts,
            "total": total,
            "model_ids": model_ids,
            "summary_text": summary_text,
            "summary_tooltip": summary_tooltip,
        }

    def _update_key_state_ui(self, *_args):
        if not hasattr(self, "lbl_key_state"):
            return
        snapshot = self._build_key_state_snapshot()
        self.lbl_key_state.setText(snapshot["summary_text"])
        self.lbl_key_state.setToolTip(snapshot["summary_tooltip"])
        if hasattr(self, "btn_key_status"):
            self.btn_key_status.setEnabled(bool(self.api_keys))
            self.btn_key_status.setToolTip(snapshot["summary_tooltip"])

    def _on_settings_bus_event(self, event):
        event_name = (event or {}).get("event")
        if event_name not in {"key_statuses_updated", "request_count_updated"}:
            return
        self._update_worker_spinbox_limit()
        self._update_key_state_ui()

    def _clear_key_limit_states(self, keys=None):
        if self._running_tasks_exist():
            QMessageBox.information(self, "Ключи", "Сброс статусов ключей доступен только когда фоновые задачи остановлены.")
            return False
        if self.settings_manager is None:
            return False

        model_ids = self._reader_status_model_ids()
        target_keys = [key for key in (keys or self.api_keys) if (key or "").strip()]
        if not model_ids or not target_keys:
            return False

        changed = False
        for api_key in target_keys:
            for model_id in model_ids:
                try:
                    if self.settings_manager.clear_key_exhaustion_status(api_key, model_id):
                        changed = True
                except Exception:
                    continue
            self.disabled_api_keys.discard(api_key)

        if changed:
            self._update_worker_spinbox_limit()
            self._update_key_state_ui()
        return changed

    def show_key_status_dialog(self):
        dlg = ReaderKeyStatusDialog(self)

        def refresh_dialog():
            dlg.set_snapshot(self._build_key_state_snapshot())

        def selected_or_all_keys():
            keys = dlg.selected_keys()
            return keys or dlg.visible_keys()

        def clear_limits():
            if self._clear_key_limit_states(selected_or_all_keys()):
                self.statusBar().showMessage("Статусы лимитов ключей обновлены.")
            refresh_dialog()

        def clear_invalid():
            if self._running_tasks_exist():
                QMessageBox.information(self, "Ключи", "Снимать invalid-пометки можно только после остановки текущих задач.")
                return
            if self._clear_invalid_key_states(selected_or_all_keys()):
                self._update_worker_spinbox_limit()
                self._update_key_state_ui()
                self.statusBar().showMessage("Invalid-пометки ключей сняты.")
            refresh_dialog()

        dlg.btn_refresh.clicked.connect(refresh_dialog)
        dlg.btn_clear_limits.clicked.connect(clear_limits)
        dlg.btn_clear_invalid.clicked.connect(clear_invalid)
        refresh_dialog()
        dlg.exec()

    def _get_available_api_keys(self, required_model_ids=None):
        model_ids = [model_id for model_id in (required_model_ids or []) if model_id]
        invalid_states = self._load_invalid_key_states()
        available_keys = []
        for api_key in self.api_keys:
            if not (api_key or "").strip():
                continue
            if api_key in self.disabled_api_keys:
                continue
            if api_key in invalid_states:
                continue
            if model_ids and any(self._reader_key_model_quota_exhausted(api_key, model_id) for model_id in model_ids):
                continue
            if self.settings_manager is None or not model_ids:
                available_keys.append(api_key)
                continue

            key_info = self.settings_manager.get_key_info(api_key)
            if not key_info:
                available_keys.append(api_key)
                continue

            is_blocked = any(
                self.settings_manager.is_key_limit_active(key_info, model_id)
                or self._reader_key_model_quota_exhausted(api_key, model_id)
                for model_id in model_ids
            )
            if not is_blocked:
                available_keys.append(api_key)
        return available_keys

    def _project_rpd_exhausted_message(self, model_id):
        if not model_id:
            return ""
        rpd_limit = self._reader_request_limit_for_model(model_id)
        if not rpd_limit or rpd_limit <= 0:
            return ""

        candidate_keys = []
        invalid_states = self._load_invalid_key_states()
        for api_key in self.api_keys:
            if not (api_key or "").strip():
                continue
            if api_key in self.disabled_api_keys or api_key in invalid_states:
                continue
            if self.settings_manager is not None:
                key_info = self.settings_manager.get_key_info(api_key)
                if key_info and self.settings_manager.is_key_limit_active(key_info, model_id):
                    continue
            candidate_keys.append(api_key)

        if not candidate_keys:
            return ""

        exhausted_keys = [
            api_key
            for api_key in candidate_keys
            if self._reader_key_model_quota_exhausted(api_key, model_id)
        ]
        if len(exhausted_keys) < len(candidate_keys):
            return ""
        reset_text = _next_policy_reset_text(_gemini_reset_policy())
        return (
            f"Дневной лимит {rpd_limit} запросов для модели {model_id} исчерпан на всех доступных ключах. "
            f"Следующий сброс около {reset_text}."
        )

    def _collect_target_chapter_indices(self, prefer_selection=False, include_done=False):
        if not self.bm:
            return []

        checked_indices = []
        if prefer_selection:
            checked_indices = self._checked_chapter_indices()

        target_indices = checked_indices or list(range(len(self.bm.chapters)))
        status_snapshot = self.bm.chapter_status_snapshot()
        skipped_indices = status_snapshot.get("skipped", set())
        done_indices = status_snapshot.get("done", set())
        prepared = []
        for idx in target_indices:
            if idx < 0 or idx >= len(self.bm.chapters):
                continue
            if idx in skipped_indices:
                continue
            if not include_done and idx in done_indices:
                continue
            prepared.append(idx)
        return sorted(dict.fromkeys(prepared))

    def _checked_chapter_indices(self):
        if not hasattr(self, "list_chapters"):
            return sorted(self._checked_chapter_indices_state)
        checked_indices = []
        for row in range(self.list_chapters.count()):
            item = self.list_chapters.item(row)
            if item is None or item.checkState() != Qt.CheckState.Checked:
                continue
            idx = item.data(Qt.ItemDataRole.UserRole)
            if isinstance(idx, int):
                checked_indices.append(idx)
        prepared = sorted(dict.fromkeys(checked_indices))
        if prepared or not self._checked_chapter_indices_state:
            return prepared
        return sorted(self._checked_chapter_indices_state)

    def _on_chapter_item_pressed(self, item):
        idx = item.data(Qt.ItemDataRole.UserRole) if item is not None else None
        self._chapter_last_press_index = idx if isinstance(idx, int) else None
        self._chapter_last_press_modifiers = QApplication.keyboardModifiers()

    def _apply_checked_range(self, anchor_idx, current_idx, desired_state):
        if not hasattr(self, "list_chapters"):
            return
        start_idx = min(anchor_idx, current_idx)
        end_idx = max(anchor_idx, current_idx)
        self._chapter_check_state_refresh = True
        try:
            for row in range(start_idx, end_idx + 1):
                item = self.list_chapters.item(row)
                if item is None:
                    continue
                if item.checkState() != desired_state:
                    item.setCheckState(desired_state)
        finally:
            self._chapter_check_state_refresh = False

    def _set_checked_chapter_indices(self, indices):
        normalized = {
            idx for idx in (indices or [])
            if isinstance(idx, int) and idx >= 0 and (not self.bm or idx < len(self.bm.chapters))
        }
        self._checked_chapter_indices_state = set(normalized)
        self._chapter_check_anchor_index = max(normalized) if normalized else None
        if not hasattr(self, "list_chapters") or self.list_chapters.count() <= 0:
            return

        self._chapter_check_state_refresh = True
        try:
            for row in range(self.list_chapters.count()):
                item = self.list_chapters.item(row)
                if item is None:
                    continue
                idx = item.data(Qt.ItemDataRole.UserRole)
                desired_state = Qt.CheckState.Checked if idx in normalized else Qt.CheckState.Unchecked
                if item.checkState() != desired_state:
                    item.setCheckState(desired_state)
        finally:
            self._chapter_check_state_refresh = False

    def _on_chapter_item_changed(self, item):
        if self._chapter_check_state_refresh or self._loading_settings:
            return
        idx = item.data(Qt.ItemDataRole.UserRole) if item is not None else None
        if isinstance(idx, int):
            modifiers = QApplication.keyboardModifiers()
            if modifiers == Qt.KeyboardModifier.NoModifier:
                modifiers = self._chapter_last_press_modifiers
            shift_active = bool(modifiers & Qt.KeyboardModifier.ShiftModifier)
            anchor_idx = self._chapter_check_anchor_index
            if shift_active and anchor_idx is not None and anchor_idx != idx:
                self._apply_checked_range(anchor_idx, idx, item.checkState())
            self._chapter_check_anchor_index = idx
        self._chapter_last_press_index = None
        self._chapter_last_press_modifiers = Qt.KeyboardModifier.NoModifier
        self._checked_chapter_indices_state = set(self._checked_chapter_indices())
        self._on_chapter_selection_changed()
        self.save_settings()

    def _selected_chapter_indices(self):
        selected_indices = []
        for item in self.list_chapters.selectedItems():
            idx = item.data(Qt.ItemDataRole.UserRole)
            if isinstance(idx, int):
                selected_indices.append(idx)
        return sorted(dict.fromkeys(selected_indices))

    def _format_chapter_ranges(self, indices):
        if not indices:
            return ""
        chapter_numbers = sorted({idx + 1 for idx in indices if isinstance(idx, int) and idx >= 0})
        if not chapter_numbers:
            return ""

        ranges = []
        start = chapter_numbers[0]
        prev = chapter_numbers[0]
        for number in chapter_numbers[1:]:
            if number == prev + 1:
                prev = number
                continue
            ranges.append(f"{start}-{prev}" if start != prev else str(start))
            start = prev = number
        ranges.append(f"{start}-{prev}" if start != prev else str(start))
        return ", ".join(ranges)

    def _selected_scope_summary(self):
        checked_indices = self._checked_chapter_indices()
        if self.chk_selected_only.isChecked():
            return self._format_chapter_ranges(checked_indices) if checked_indices else "отметки пусты"
        if not checked_indices:
            return "все"
        selection_text = self._format_chapter_ranges(checked_indices)
        return f"все (отмечено {len(checked_indices)}: {selection_text})"

    def _on_chapter_selection_changed(self, *_args):
        if not hasattr(self, "lbl_chapter_scope"):
            return
        summary = self._selected_scope_summary()
        self.lbl_chapter_scope.setText(f"Главы: {summary}")
        self.lbl_chapter_scope.setToolTip(f"Область обработки: {summary}")

    def clear_chapter_selection(self):
        self._set_checked_chapter_indices([])
        self.chk_selected_only.setChecked(False)
        self._on_chapter_selection_changed()
        self.save_settings()

    def _update_export_folder_status(self):
        if not hasattr(self, "lbl_export_folder"):
            return
        if not self.bm or not self.bm.book_dir:
            self.lbl_export_folder.setText("Экспорт: откройте книгу")
            self.lbl_export_folder.setToolTip("Сначала откройте книгу. Экспорт сохраняется в папку текущей книги.")
            return

        folder_name = os.path.basename(self.bm.book_dir.rstrip("\\/")) or self.bm.book_dir
        self.lbl_export_folder.setText(f"Экспорт: {folder_name}")
        self.lbl_export_folder.setToolTip(self.bm.book_dir)

    def _update_video_cover_status(self):
        if not hasattr(self, "lbl_video_cover"):
            return
        if not self.bm:
            self.lbl_video_cover.setText("Видео: откройте книгу")
            self.lbl_video_cover.setToolTip("Сначала откройте книгу, затем выберите картинку для видео.")
            return

        video_cover_path = self.bm.get_video_cover_path()
        if video_cover_path:
            file_name = os.path.basename(video_cover_path)
            self.lbl_video_cover.setText(f"Видео: {file_name}")
            self.lbl_video_cover.setToolTip(video_cover_path)
        else:
            self.lbl_video_cover.setText("Видео: картинка не выбрана")
            self.lbl_video_cover.setToolTip("Выбранная картинка будет скопирована в папку книги и использована для экспорта видео.")

    def open_export_folder(self):
        if not self.bm or not self.bm.book_dir:
            QMessageBox.information(self, "Экспорт", "Сначала откройте книгу.")
            return

        export_dir = self.bm.book_dir
        try:
            if platform.system() == "Windows":
                os.startfile(export_dir)
            elif platform.system() == "Darwin":
                subprocess.Popen(["open", export_dir])
            else:
                subprocess.Popen(["xdg-open", export_dir])
        except Exception as exc:
            QMessageBox.warning(self, "Экспорт", f"Не удалось открыть папку экспорта: {exc}")

    def select_video_cover(self):
        if not self.bm:
            QMessageBox.information(self, "Видео", "Сначала откройте книгу.")
            return

        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Выбрать картинку для видео",
            "",
            "Images (*.png *.jpg *.jpeg *.webp *.bmp);;All files (*.*)",
        )
        if not file_path:
            return

        try:
            stored_path = self.bm.save_video_cover_image(file_path)
        except Exception as exc:
            QMessageBox.warning(self, "Видео", f"Не удалось сохранить картинку: {exc}")
            return

        self._update_video_cover_status()
        self._refresh_runtime_controls()
        self.statusBar().showMessage(f"Картинка для видео сохранена: {os.path.basename(stored_path)}")

    def pick_chapters_dialog(self):
        if not self.bm:
            QMessageBox.information(self, "Главы", "Сначала откройте книгу.")
            return

        current_text = self._format_chapter_ranges(self._checked_chapter_indices())
        prompt = (
            "Введите номера глав через запятую или диапазоны через дефис.\n"
            "Пример: 1,3,5-8"
        )
        raw_text, ok = QInputDialog.getText(self, "Выбор глав", prompt, text=current_text)
        if not ok:
            return

        try:
            selected_indices = self._parse_chapter_selection_text(raw_text)
        except ValueError as exc:
            QMessageBox.warning(self, "Выбор глав", str(exc))
            return

        self._set_checked_chapter_indices(selected_indices)

        if selected_indices:
            self.chk_selected_only.setChecked(True)
            first_item = self.list_chapters.item(selected_indices[0])
            if first_item is not None:
                self.list_chapters.scrollToItem(first_item)
            self.statusBar().showMessage(f"Отмечены главы: {self._format_chapter_ranges(selected_indices)}")
        else:
            self.chk_selected_only.setChecked(False)
            self.statusBar().showMessage("Отметки глав очищены.")

        self._on_chapter_selection_changed()
        self.save_settings()

    def _parse_chapter_selection_text(self, raw_text):
        text = re.sub(r"\s*-\s*", "-", (raw_text or "").strip())
        if not text:
            return []
        if not self.bm:
            raise ValueError("Сначала откройте книгу.")

        parsed_indices = []
        total_chapters = len(self.bm.chapters)
        tokens = [token for token in re.split(r"[\s,;]+", text) if token]
        if not tokens:
            return []

        for token in tokens:
            range_match = re.fullmatch(r"(\d+)-(\d+)", token)
            if range_match:
                start_num = int(range_match.group(1))
                end_num = int(range_match.group(2))
                if start_num > end_num:
                    start_num, end_num = end_num, start_num
                if start_num < 1 or end_num > total_chapters:
                    raise ValueError(f"Диапазон {token} выходит за пределы книги (1-{total_chapters}).")
                parsed_indices.extend(range(start_num - 1, end_num))
                continue

            if token.isdigit():
                chapter_number = int(token)
                if chapter_number < 1 or chapter_number > total_chapters:
                    raise ValueError(f"Глава {chapter_number} вне диапазона 1-{total_chapters}.")
                parsed_indices.append(chapter_number - 1)
                continue

            raise ValueError(
                f"Не удалось разобрать '{token}'. Используйте формат вроде 1,3,5-8."
            )

        return sorted(dict.fromkeys(parsed_indices))

    def _collect_run_scope_indices(self, include_done=False, action_label="обработки"):
        prefer_selection = bool(self.chk_selected_only.isChecked())
        if prefer_selection and not self._checked_chapter_indices():
            QMessageBox.information(
                self,
                "Главы",
                f"Включён режим 'Только отмеченные главы', но главы не отмечены для {action_label}.",
            )
            return None
        return self._collect_target_chapter_indices(prefer_selection=prefer_selection, include_done=include_done)

    def _author_gender_script_missing_indices(self, chapter_indices):
        if not self.bm:
            return list(chapter_indices or [])
        missing = []
        for chapter_index in chapter_indices or []:
            script_text = self.bm.load_tts_script(chapter_index)
            raw_text = ""
            try:
                raw_text = self.bm.chapters[chapter_index].raw_text
            except Exception:
                raw_text = ""
            if (
                not _script_matches_voice_mode(script_text, "author_gender")
                or _find_author_gender_script_issues(script_text, raw_text)
            ):
                missing.append(chapter_index)
        return missing

    def _ensure_author_gender_scripts_ready(self, chapter_indices, action_label="озвучки"):
        missing_indices = self._author_gender_script_missing_indices(chapter_indices)
        if not missing_indices:
            return True

        missing_text = self._format_chapter_ranges(missing_indices)
        QMessageBox.warning(
            self,
            "AI-сценарий",
            (
                "Для режима 'Автор + Муж./Жен. роли' нужен корректный AI-сценарий без догадок и без текста, "
                "который не подтверждается исходной главой.\n\n"
                f"Подготовьте или пересоберите AI-сценарий для глав: {missing_text}.\n"
                f"После этого повторите запуск {action_label}."
            ),
        )
        return False

    def _load_script_for_chapter(self, chapter_index):
        if not self.bm or chapter_index is None or chapter_index < 0 or chapter_index >= len(self.bm.chapters):
            self._current_chapter_index = None
            self.script_view.setPlainText("")
            return

        self._current_chapter_index = chapter_index
        script_text = self.bm.load_tts_script(chapter_index)
        if script_text.strip():
            self.script_view.setPlainText(script_text)
        else:
            self.script_view.setPlainText(self.bm.chapters[chapter_index].raw_text)

    def save_current_script(self):
        if not self.bm or self._current_chapter_index is None:
            QMessageBox.information(self, "Сценарий", "Сначала выберите главу.")
            return
        script_text = self.script_view.toPlainText().strip()
        if not script_text:
            QMessageBox.warning(self, "Сценарий", "Нельзя сохранить пустой сценарий.")
            return
        self.bm.save_tts_script(self._current_chapter_index, script_text)
        self.refresh_chapters_list()
        self.statusBar().showMessage(f"Сценарий главы {self._current_chapter_index + 1} сохранён.")

    def _sync_prompt_views(self):
        if not hasattr(self, "preprocess_prompt_view"):
            return

        preprocess_text = self.preprocess_directive or DEFAULT_PREPROCESS_DIRECTIVE
        tts_text = self.tts_directive or DEFAULT_TTS_DIRECTIVE

        if self.preprocess_prompt_view.toPlainText() != preprocess_text:
            self.preprocess_prompt_view.setPlainText(preprocess_text)
        if self.tts_prompt_view.toPlainText() != tts_text:
            self.tts_prompt_view.setPlainText(tts_text)

        self._update_prompt_hint()

    def _update_prompt_hint(self):
        if not hasattr(self, "lbl_prompts_hint"):
            return
        if self._is_flash_tts_mode():
            self.lbl_prompts_hint.setText(
                "Промпты Flash TTS. Эти указания сейчас используются для AI-предобработки и TTS-исполнения."
            )
        elif self._selected_voice_mode() == "author_gender":
            self.lbl_prompts_hint.setText(
                "Промпты AI-сценария. В Live API для режима 'Автор + Муж./Жен. роли' отсюда берётся только AI-предобработка; именно она размечает Author/Male/Female без эвристик."
            )
        else:
            self.lbl_prompts_hint.setText(
                "Промпты AI-сценария. Они видны и редактируются здесь, но сейчас применяются только в Flash TTS и в Live API для режима 'Автор + Муж./Жен. роли'."
            )

    def save_prompt_settings(self):
        if not hasattr(self, "preprocess_prompt_view"):
            return
        self.preprocess_directive = self.preprocess_prompt_view.toPlainText().strip() or DEFAULT_PREPROCESS_DIRECTIVE
        self.tts_directive = self.tts_prompt_view.toPlainText().strip() or DEFAULT_TTS_DIRECTIVE
        self._sync_prompt_views()
        self.save_settings()
        self.statusBar().showMessage("Промпты AI-сценария и TTS сохранены.")

    def reset_prompt_settings(self):
        self.preprocess_directive = DEFAULT_PREPROCESS_DIRECTIVE
        self.tts_directive = DEFAULT_TTS_DIRECTIVE
        self._sync_prompt_views()
        self.save_settings()
        self.statusBar().showMessage("Промпты AI-сценария и TTS сброшены к стандартным.")

    def edit_flash_tts_prompts(self):
        dialog = PromptTuningDialog(self.preprocess_directive, self.tts_directive, self)
        if dialog.exec():
            preprocess_directive, tts_directive = dialog.get_values()
            self.preprocess_directive = preprocess_directive or DEFAULT_PREPROCESS_DIRECTIVE
            self.tts_directive = tts_directive or DEFAULT_TTS_DIRECTIVE
            self._sync_prompt_views()
            self.save_settings()
            self.statusBar().showMessage("Промпты AI-сценария и TTS обновлены.")

    def prepare_selected_scripts(self):
        allow_live_author_gender = (not self._is_flash_tts_mode()) and self._selected_voice_mode() == "author_gender"
        if not self._is_flash_tts_mode() and not allow_live_author_gender:
            QMessageBox.information(
                self,
                "Режим",
                "AI-подготовка сценария доступна в режиме Flash TTS и в Live API для режима 'Автор + Муж./Жен. роли'."
            )
            return
        if genai is None or genai_types is None:
            QMessageBox.warning(self, "Зависимости", "Для AI-подготовки сценария требуется пакет google-genai.")
            return
        if not self.api_keys:
            self.set_api_keys()
            if not self.api_keys:
                return
        if not self.bm:
            QMessageBox.warning(self, "Внимание", "Сначала откройте книгу!")
            return
        if self.workers:
            QMessageBox.warning(self, "Внимание", "Сначала дождитесь завершения текущей задачи или остановите её.")
            return

        target_indices = self._collect_run_scope_indices(include_done=True, action_label="AI-подготовки сценария")
        if target_indices is None:
            return
        if not target_indices:
            if self.chk_selected_only.isChecked():
                QMessageBox.information(self, "Сценарии", "Среди отмеченных глав нет доступных для AI-подготовки.")
                return
            QMessageBox.information(self, "Сценарии", "Нет глав для AI-подготовки.")
            return

        self._launch_flash_workers(target_indices, run_mode="prepare")

    def clear_selected_scripts(self, selected_items):
        if not self.bm:
            return
        removed_count = 0
        for item in selected_items:
            idx = item.data(Qt.ItemDataRole.UserRole)
            if not isinstance(idx, int):
                continue
            path = self.bm.get_tts_script_path(idx)
            if os.path.exists(path):
                try:
                    os.remove(path)
                    removed_count += 1
                except Exception as exc:
                    logger.warning(f"Не удалось удалить TTS-сценарий главы {idx + 1}: {exc}")
            self.bm.clear_tts_progress(idx)

        self.refresh_chapters_list()
        if self._current_chapter_index is not None:
            self._load_script_for_chapter(self._current_chapter_index)
        if removed_count:
            self.statusBar().showMessage(f"Удалено сценариев: {removed_count}.")

    def set_return_to_menu_handler(self, handler):
        self._return_to_menu_handler = handler

    def _return_to_menu(self):
        self.save_settings()
        self._returning_to_main_menu = True
        self.close()

    def apply_runtime_capabilities(self):
        runtime_notes = []

        if genai is None or genai_types is None:
            runtime_notes.append("нет google-genai: озвучка Gemini недоступна")

        if pyaudio is None:
            runtime_notes.append("нет PyAudio: live-воспроизведение отключено")

        edge_fallback_available = edge_tts is not None and AudioSegment is not None
        if not edge_fallback_available:
            self.chk_edge_fallback.setChecked(False)
            self.chk_edge_fallback.setEnabled(False)
            runtime_notes.append("нет edge-tts/pydub: Edge fallback отключён")
        else:
            self.chk_edge_fallback.setEnabled(True)

        if AudioSegment is None:
            self.chk_mp3.setChecked(False)
            self.chk_mp3.setEnabled(False)
            self.btn_combine.setEnabled(False)
            runtime_notes.append("нет pydub: MP3-экспорт отключён")

        if _resolve_tool_path("ffmpeg") is None or _resolve_tool_path("ffprobe") is None:
            self.btn_combine.setEnabled(False)
            self.btn_export_video.setEnabled(False)
            runtime_notes.append("нет ffmpeg/ffprobe: склейка MP3 недоступна")

        self._refresh_runtime_controls()

        if runtime_notes:
            self.statusBar().showMessage(" | ".join(runtime_notes))

    def _reader_ui_settings_payload(self):
        return {
            "engine": self._current_engine_id(),
            "model": self.combo_model.currentText(),
            "voice": self.combo_voices.currentData(),
            "voice_mode": self._selected_voice_mode(),
            "voice_secondary": self.combo_voice_secondary.currentData(),
            "voice_tertiary": self.combo_voice_tertiary.currentData(),
            "speed": self.combo_speed.currentText(),
            "live_segment_mode": self._selected_live_segment_mode(),
            "chunk": self.spin_chunk.value(),
            "worker_count": self.spin_workers.value(),
            "num_instances": self.spin_workers.value(),
            "selected_only": self.chk_selected_only.isChecked(),
            "checked_chapters": self._checked_chapter_indices(),
            "parallel_single_chapter": self.chk_parallel_single_chapter.isChecked(),
            "record": self.chk_mp3.isChecked(),
            "fast": self.chk_fast.isChecked(),
            "edge_fallback": self.chk_edge_fallback.isChecked(),
            "preprocess_model": self.combo_preprocess_model.currentText(),
            "preprocess_profile": self.combo_preprocess_profile.currentText(),
            "pipeline_mode": self._selected_pipeline_mode(),
            "preprocess_directive": self.preprocess_directive,
            "tts_directive": self.tts_directive,
        }

    def closeEvent(self, event):
        if self._running_tasks_exist():
            QMessageBox.warning(self, "Подождите", "Сначала остановите генерацию или дождитесь завершения фоновых задач.")
            event.ignore()
            return

        self.save_settings()

        if self._returning_to_main_menu:
            if callable(self._return_to_menu_handler):
                self._return_to_menu_handler()
            event.accept()
            return

        try:
            from gemini_translator.ui.dialogs.menu_utils import prompt_return_to_menu, return_to_main_menu
        except Exception:
            event.accept()
            return

        action = prompt_return_to_menu(self)
        if action == "cancel":
            event.ignore()
            return
        if action == "menu":
            if callable(self._return_to_menu_handler):
                self._return_to_menu_handler()
            else:
                return_to_main_menu()
        event.accept()



    def test_voice_sample(self):
        if genai is None or genai_types is None:
            QMessageBox.warning(self, "Зависимости", "Для теста голоса требуется пакет google-genai.")
            return
        if pyaudio is None:
            QMessageBox.warning(self, "Зависимости", "Для теста голоса требуется PyAudio.")
            return
        if not self.api_keys:
            self.set_api_keys()
            if not self.api_keys: return
            
        self.btn_test_voice.setEnabled(False)
        self.btn_test_voice.setText("⏳ Загрузка...")
        
        api_key = self.api_keys[0]
        actual_model_id = self._selected_model_id()
        voice = self.combo_voices.currentData()
        engine_mode = self._current_engine_id()
        voice_mode = self._selected_voice_mode()
        secondary_voice = self.combo_voice_secondary.currentData()
        tertiary_voice = self.combo_voice_tertiary.currentData()
        
        # Создаем и запускаем независимый мини-воркер
        self.tester_worker = VoiceSampleWorker(
            api_key,
            actual_model_id,
            voice,
            engine_mode=engine_mode,
            voice_mode=voice_mode,
            secondary_voice=secondary_voice,
            tertiary_voice=tertiary_voice,
            tts_directive=self.tts_directive,
            daily_request_limiter=self.daily_request_limiter,
        )
        
        def on_test_finish():
            self.btn_test_voice.setEnabled(True)
            self.btn_test_voice.setText("🔊 Плей")
            
        def on_test_error(err):
            QMessageBox.warning(self, "Ошибка теста", f"Не удалось получить голос от серверов Google:\n{err}")
            self.btn_test_voice.setEnabled(True)
            self.btn_test_voice.setText("🔊 Плей")
            
        self.tester_worker.finished_signal.connect(on_test_finish)
        self.tester_worker.error_signal.connect(on_test_error)
        self.tester_worker.start()



    def revoice_selected_chapters(self, selected_items):
        if not self.bm:
            return
            
        for item in selected_items:
            idx = item.data(Qt.ItemDataRole.UserRole)
            
            # Удаляем маркер "Готово"
            done_path = os.path.join(self.bm.book_dir, f"Ch{idx + 1}.done")
            if os.path.exists(done_path):
                os.remove(done_path)
                
            # Удаляем маркер "Пропуск"
            skip_path = os.path.join(self.bm.book_dir, f"Ch{idx + 1}.skip")
            if os.path.exists(skip_path):
                os.remove(skip_path)
                
            # Ставим системную метку утилизации/переозвучки
            revoice_path = os.path.join(self.bm.book_dir, f"Ch{idx + 1}.revoice")
            with open(revoice_path, 'w') as f:
                f.write("revoice")
            self.bm.clear_tts_progress(idx)
                
        self.refresh_chapters_list()





    def clean_stuck_files(self):
        if not self.bm:
            QMessageBox.warning(self, "Внимание", "Сначала откройте книгу!")
            return
        if self.workers:
            QMessageBox.warning(self, "Внимание", "Остановите процесс генерации перед очисткой!")
            return

        deleted_count = 0
        for i in range(len(self.bm.chapters)):
            mp3_p = self.bm.get_mp3_path(i)
            done_p = os.path.join(self.bm.book_dir, f"Ch{i + 1}.done")
            
            # Если есть MP3, но нет .done -> файл недописан (завис)
            if os.path.exists(mp3_p) and not os.path.exists(done_p):
                try:
                    os.remove(mp3_p)
                    deleted_count += 1
                except Exception as e:
                    logger.error(f"Не удалось удалить {mp3_p}: {e}")
                    
        QMessageBox.information(self, "Очистка", f"Очистка завершена.\nУдалено незаконченных файлов: {deleted_count}")
        self.refresh_chapters_list()




    def _enqueue_worker_progress(self, worker_id, chapter_index, step_index, total_steps):
        self._pending_worker_progress[worker_id] = (worker_id, chapter_index, step_index, total_steps)
        if not self._progress_flush_timer.isActive():
            self._progress_flush_timer.start(READER_PROGRESS_FLUSH_INTERVAL_MS)

    def _flush_worker_progress(self):
        if not self._pending_worker_progress:
            return
        pending = list(self._pending_worker_progress.values())
        self._pending_worker_progress.clear()
        for worker_id, chapter_index, step_index, total_steps in pending:
            row = self.worker_widgets.get(worker_id)
            if row is not None:
                row.update_progress(worker_id, chapter_index, step_index, total_steps)

    def _flush_log_buffer(self):
        if not hasattr(self, 'log_view') or not self._log_buffer:
            return
        self.log_view.appendPlainText("\n".join(self._log_buffer))
        self._log_buffer.clear()
        self.log_view.moveCursor(QTextCursor.MoveOperation.End)

    def add_log_to_ui(self, msg, level):
        """Добавляет запись в текстовое поле лога."""
        if not hasattr(self, 'log_view'):
            return
        timestamp = time.strftime("%H:%M:%S")
        self._log_buffer.append(f"[{timestamp}] [{level}] {msg}")
        if not self._log_flush_timer.isActive():
            self._log_flush_timer.start(READER_LOG_FLUSH_INTERVAL_MS)

    def dragEnterEvent(self, e):
        if e.mimeData().hasUrls():
            for url in e.mimeData().urls():
                if _is_supported_reader_book_path(url.toLocalFile()):
                    e.acceptProposedAction()
                    return
        e.ignore()

    def dropEvent(self, e):
        for url in e.mimeData().urls():
            path = url.toLocalFile()
            if _is_supported_reader_book_path(path):
                self.load_book(path)
                break

    def load_book(self, path):
        try:
            self.bm = BookManager(path, base_dir=self.reader_books_dir)
            self._refresh_project_settings_manager()
            self.load_settings()
            self.lbl_info.setText(f"📘 {self.bm.title}")
            self._current_chapter_index = None
            self.script_view.setPlainText("")
            self.refresh_chapters_list()
            self._update_export_folder_status()
            self._update_video_cover_status()
            self._refresh_runtime_controls()
            self.tabs.setCurrentIndex(0)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))

    def _update_chapter_list_item(self, item, idx, checked=None, status_snapshot=None):
        if not self.bm or item is None or idx < 0 or idx >= len(self.bm.chapters):
            return

        chap = self.bm.chapters[idx]
        item.setData(Qt.ItemDataRole.UserRole, idx)
        item.setFlags(item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
        item.setText(chap.title)
        item.setForeground(QColor("#ececec"))
        font = item.font()
        font.setStrikeOut(False)
        item.setFont(font)

        if status_snapshot is None:
            is_skipped = self.bm.is_chapter_skipped(idx)
            is_done = self.bm.is_chapter_done(idx)
            has_script = self.bm.has_tts_script(idx)
        else:
            is_skipped = idx in status_snapshot.get("skipped", set())
            is_done = idx in status_snapshot.get("done", set())
            has_script = idx in status_snapshot.get("scripts", set())

        if is_skipped:
            item.setText(f"❌ {chap.title} (Пропуск)")
            item.setForeground(QColor("#B0BEC5"))
            font.setStrikeOut(True)
            item.setFont(font)
        elif is_done:
            item.setText(f"✅ {chap.title}")
            item.setForeground(QColor("#4CAF50"))
        elif has_script:
            item.setText(f"📝 {chap.title}")

        if checked is not None:
            item.setCheckState(Qt.CheckState.Checked if checked else Qt.CheckState.Unchecked)

    def refresh_chapters_list(self):
        checked_indices = self._checked_chapter_indices()
        status_snapshot = self.bm.chapter_status_snapshot()
        self._chapter_check_state_refresh = True
        try:
            self.list_chapters.setUpdatesEnabled(False)
            self.list_chapters.clear()
            for i, chap in enumerate(self.bm.chapters):
                item = QListWidgetItem(chap.title)
                self._update_chapter_list_item(item, i, checked=i in checked_indices, status_snapshot=status_snapshot)
                self.list_chapters.addItem(item)
        finally:
            self.list_chapters.setUpdatesEnabled(True)
            self._chapter_check_state_refresh = False
        self._checked_chapter_indices_state = {
            idx for idx in checked_indices
            if isinstance(idx, int) and 0 <= idx < len(self.bm.chapters)
        }
        if self._chapter_check_anchor_index not in self._checked_chapter_indices_state:
            self._chapter_check_anchor_index = max(self._checked_chapter_indices_state) if self._checked_chapter_indices_state else None
        self._on_chapter_selection_changed()

    def show_chapter_context_menu(self, pos):
        if not self.bm:
            return
            
        selected_items = self.list_chapters.selectedItems()
        item_at_pos = self.list_chapters.itemAt(pos)
        
        # Умная обработка клика (как в Windows Explorer)
        if not selected_items:
            if not item_at_pos:
                return
            selected_items = [item_at_pos]
            item_at_pos.setSelected(True)
        elif item_at_pos and item_at_pos not in selected_items:
            self.list_chapters.clearSelection()
            selected_items =[item_at_pos]
            item_at_pos.setSelected(True)
            
        menu = QMenu(self)
        
        # Действие: Переозвучить группу
        revoice_action = QAction("♻️ Переозвучить выбранные", self)
        revoice_action.triggered.connect(lambda: self.revoice_selected_chapters(selected_items))
        menu.addAction(revoice_action)
        
        menu.addSeparator()

        prepare_action = QAction("🪄 Подготовить AI-сценарий", self)
        prepare_action.triggered.connect(self.prepare_selected_scripts)
        prepare_action.setEnabled(self._is_flash_tts_mode() or self._selected_voice_mode() == "author_gender")
        menu.addAction(prepare_action)

        clear_script_action = QAction("🗑️ Удалить сценарий", self)
        clear_script_action.triggered.connect(lambda: self.clear_selected_scripts(selected_items))
        menu.addAction(clear_script_action)

        menu.addSeparator()
        
        # Действие: Пропустить группу
        skip_action = QAction("❌ Пропустить / Включить выбранные", self)
        skip_action.triggered.connect(lambda: self.toggle_skip_chapter(selected_items))
        menu.addAction(skip_action)
        
        menu.exec(self.list_chapters.mapToGlobal(pos))

    def toggle_skip_chapter(self, selected_items):
        if not self.bm:
            return
            
        for item in selected_items:
            idx = item.data(Qt.ItemDataRole.UserRole)
            self.bm.toggle_chapter_skip(idx)
            
            # Если главу пропустили, снимаем маркер переозвучки во избежание конфликтов интерфейса
            revoice_path = os.path.join(self.bm.book_dir, f"Ch{idx + 1}.revoice")
            if self.bm.is_chapter_skipped(idx) and os.path.exists(revoice_path):
                os.remove(revoice_path)
                
        self.refresh_chapters_list()




    def on_chapter_clicked(self, item):
        idx = item.data(Qt.ItemDataRole.UserRole)
        if self.bm:
            self.txt_view.setPlainText(self.bm.chapters[idx].raw_text)
            self._load_script_for_chapter(idx)

    def on_chapter_done_ui(self, idx):
        item = self.list_chapters.item(idx)
        if item:
            self._update_chapter_list_item(item, idx)

    def on_script_ready_ui(self, idx):
        item = self.list_chapters.item(idx)
        if item:
            self._update_chapter_list_item(item, idx)
        if self._current_chapter_index == idx:
            self._load_script_for_chapter(idx)

    def _reset_dashboard(self):
        self._pending_worker_progress.clear()
        for i in reversed(range(self.dash_layout.count())):
            widget = self.dash_layout.itemAt(i).widget()
            if widget:
                widget.setParent(None)
        self.worker_widgets = {}

    def _active_worker_target_count(self):
        try:
            return max(1, int(self.spin_workers.value()))
        except Exception:
            return 1

    def _next_replacement_worker_id(self):
        used_ids = set()
        for worker in self.workers:
            try:
                used_ids.add(int(getattr(worker, "worker_id", -1)))
            except Exception:
                continue
        used_ids.update(worker_id for worker_id in self.worker_widgets.keys() if isinstance(worker_id, int))
        worker_id = 0
        while worker_id in used_ids:
            worker_id += 1
        return worker_id

    def _active_worker_api_keys(self):
        return {
            key
            for key in (getattr(worker, "api_key", "") for worker in self.workers)
            if (key or "").strip()
        }

    def _active_required_model_ids(self):
        if self._active_reader_engine == "flash_tts":
            run_mode = self._active_flash_run_mode or self._selected_pipeline_mode()
            model_ids = []
            if run_mode in {"auto", "prepare"}:
                model_ids.append(self._selected_preprocess_model_id())
            if run_mode != "prepare":
                model_ids.append(self._selected_model_id())
            return [model_id for model_id in dict.fromkeys(model_ids) if model_id]
        if self._active_reader_engine == "live" or self._active_job_kind == "tts_parallel_live":
            model_id = self._selected_model_id()
            return [model_id] if model_id else []
        return []

    def _replacement_api_keys(self, required_model_ids=None):
        active_keys = self._active_worker_api_keys()
        return [
            api_key
            for api_key in self._get_available_api_keys(required_model_ids or [])
            if api_key not in active_keys
        ]

    def _connect_reader_worker_signals(self, worker, *, chapter_done=False, script_ready=False):
        worker.worker_progress.connect(self._enqueue_worker_progress)
        worker.finished_signal.connect(self._on_worker_finished)
        if chapter_done:
            worker.chapter_done_ui_signal.connect(self.on_chapter_done_ui)
        if script_ready:
            worker.script_ready_signal.connect(self.on_script_ready_ui)
        worker.invalid_key_signal.connect(self._on_invalid_worker_key)
        worker.quota_key_signal.connect(self._on_quota_worker_key)
        worker.project_quota_signal.connect(self._on_project_quota_worker)
        worker.error_signal.connect(lambda _wid, msg: self.statusBar().showMessage(msg))

    def _start_replacement_worker_if_possible(self):
        if getattr(self, "_stop_requested", False):
            return False
        if self._active_manager_queue is None or self._active_manager_queue.qsize() <= 0:
            return False
        if self._project_quota_message:
            return False
        if self._parallel_live_state is not None and self._parallel_live_state.get("cancelled"):
            return False
        if len(self.workers) >= self._active_worker_target_count():
            return False

        required_model_ids = self._active_required_model_ids()
        replacement_keys = self._replacement_api_keys(required_model_ids)
        if not replacement_keys:
            return False

        worker_id = self._next_replacement_worker_id()
        api_key = replacement_keys[0]
        row = DashboardRow(worker_id)
        self.dash_layout.addWidget(row)
        self.worker_widgets[worker_id] = row

        try:
            if self._active_job_kind == "tts_parallel_live":
                if self._parallel_live_state is None:
                    raise RuntimeError("Parallel live state is missing.")
                worker = GeminiParallelChapterWorker(
                    worker_id,
                    api_key,
                    self.bm,
                    self._selected_model_id(),
                    self.combo_voices.currentData(),
                    self.combo_voice_secondary.currentData(),
                    self.combo_voice_tertiary.currentData(),
                    self.combo_speed.currentText(),
                    self._active_manager_queue,
                    self._parallel_live_state,
                    daily_request_limiter=self.daily_request_limiter,
                    voice_mode=self._selected_voice_mode(),
                    allow_edge_fallback=self.chk_edge_fallback.isChecked(),
                )
                self._connect_reader_worker_signals(worker)
            elif self._active_reader_engine == "flash_tts":
                run_mode = self._active_flash_run_mode or self._selected_pipeline_mode()
                live_playback = run_mode != "prepare" and self.player is not None
                worker = FlashTtsWorker(
                    worker_id,
                    api_key,
                    self.bm,
                    self.audio_queue if live_playback else None,
                    self._selected_model_id(),
                    self.combo_voices.currentData(),
                    self.combo_voice_secondary.currentData(),
                    self.combo_speed.currentText(),
                    self.chk_mp3.isChecked() if run_mode != "prepare" else False,
                    self.chk_fast.isChecked() if run_mode != "prepare" else True,
                    self.spin_chunk.value(),
                    self._active_manager_queue,
                    self._selected_preprocess_model_id(),
                    self.combo_preprocess_profile.currentData(),
                    self._selected_voice_mode(),
                    run_mode,
                    self.preprocess_directive,
                    self.tts_directive,
                    self.daily_request_limiter,
                    allow_edge_fallback=self.chk_edge_fallback.isChecked(),
                )
                self._connect_reader_worker_signals(
                    worker,
                    chapter_done=True,
                    script_ready=True,
                )
            elif self._active_reader_engine == "live":
                live_playback = self.player is not None
                worker = GeminiWorker(
                    worker_id,
                    api_key,
                    self.bm,
                    self.audio_queue if live_playback else None,
                    self._selected_model_id(),
                    self.combo_voices.currentData(),
                    "Ты диктор.",
                    self.combo_speed.currentText(),
                    self.chk_mp3.isChecked(),
                    self.chk_fast.isChecked(),
                    self.spin_chunk.value(),
                    self._selected_live_segment_mode(),
                    self._active_manager_queue,
                    daily_request_limiter=self.daily_request_limiter,
                    voice_mode=self._selected_voice_mode(),
                    secondary_voice=self.combo_voice_secondary.currentData(),
                    tertiary_voice=self.combo_voice_tertiary.currentData(),
                    allow_edge_fallback=self.chk_edge_fallback.isChecked(),
                )
                self._connect_reader_worker_signals(worker, chapter_done=True)
            else:
                raise RuntimeError("Unknown active reader engine.")
        except Exception as exc:
            self.worker_widgets.pop(worker_id, None)
            row.setParent(None)
            logger.warning(f"Не удалось запустить replacement-воркер: {exc}")
            return False

        self.workers.append(worker)
        worker.start()
        self.statusBar().showMessage(
            f"Ключ {_mask_api_key(api_key)} взят как замена; оставшаяся очередь продолжена."
        )
        return True

    def _on_worker_finished(self, worker_id):
        self._flush_worker_progress()
        self._pending_worker_progress.pop(worker_id, None)
        row_widget = self.worker_widgets.pop(worker_id, None)
        if row_widget:
            row_widget.setParent(None)

        self.workers = [worker for worker in self.workers if getattr(worker, "worker_id", None) != worker_id]
        if self._start_replacement_worker_if_possible():
            return
        if not self.workers:
            if self.player:
                self.player.stop()
                self.player = None
            self._set_reading_controls_running(False)
            self.refresh_chapters_list()
            if self._current_chapter_index is not None:
                self._load_script_for_chapter(self._current_chapter_index)
            remaining_chapters = self._active_manager_queue.qsize() if self._active_manager_queue is not None else 0
            if getattr(self, "_stop_requested", False):
                final_message = "Процесс остановлен."
            elif self._active_job_kind == "tts_parallel_live" and self._parallel_live_state is not None:
                final_message = self._finalize_parallel_live_chapter()
            elif self._project_quota_message:
                final_message = self._project_quota_message
            elif remaining_chapters > 0:
                final_message = (
                    f"Задача остановлена: осталось {remaining_chapters} глав(ы). "
                    "Проверьте пул API-ключей."
                )
            else:
                final_message = "AI-сценарии подготовлены." if self._active_job_kind == "prepare" else "Озвучка завершена."
                if self._run_had_invalid_keys:
                    final_message += " Невалидные ключи были исключены из запуска."
            self._active_manager_queue = None
            self._active_reader_engine = None
            self._active_flash_run_mode = None
            self._run_had_invalid_keys = False
            self._project_quota_message = ""
            self._stop_requested = False
            self.statusBar().showMessage(final_message)

    def _on_invalid_worker_key(self, worker_id, api_key, error_text, chapter_index):
        self.disabled_api_keys.add(api_key)
        self._run_had_invalid_keys = True
        self._mark_key_invalid(api_key, error_text, self._selected_model_id())
        self._update_worker_spinbox_limit()
        self._update_key_state_ui()
        chapter_label = f"глава {chapter_index + 1}" if chapter_index >= 0 else "текущая глава"
        masked_key = _mask_api_key(api_key)
        self.statusBar().showMessage(
            f"Отключён невалидный API-ключ {masked_key}; {chapter_label} возвращена в очередь."
        )

    def _next_worker_id(self):
        used_ids = {
            int(worker_id)
            for worker_id in self.worker_widgets.keys()
            if isinstance(worker_id, int)
        }
        for worker in self.workers:
            try:
                used_ids.add(int(getattr(worker, "worker_id", -1)))
            except (TypeError, ValueError):
                pass
        worker_id = 0
        while worker_id in used_ids:
            worker_id += 1
        return worker_id

    def _available_replacement_key(self, required_model_ids):
        active_keys = {
            getattr(worker, "api_key", "")
            for worker in self.workers
            if getattr(worker, "api_key", "")
        }
        for api_key in self._get_available_api_keys(required_model_ids):
            if api_key not in active_keys:
                return api_key
        return ""

    def _queue_has_pending_work(self):
        if self._active_manager_queue is None:
            return False
        try:
            return self._active_manager_queue.qsize() > 0
        except Exception:
            return True

    def _attach_common_worker_signals(self, worker, row):
        worker.worker_progress.connect(row.update_progress)
        worker.finished_signal.connect(self._on_worker_finished)
        worker.invalid_key_signal.connect(self._on_invalid_worker_key)
        worker.quota_key_signal.connect(self._on_quota_worker_key)
        worker.project_quota_signal.connect(self._on_project_quota_worker)
        worker.error_signal.connect(lambda _wid, msg: self.statusBar().showMessage(msg))

    def _start_replacement_worker(self):
        if not self._queue_has_pending_work():
            return ""

        worker_id = self._next_worker_id()
        worker = None

        if self._active_job_kind == "tts_parallel_live":
            if self._parallel_live_state is None:
                return ""
            replacement_key = self._available_replacement_key([self._selected_model_id()])
            if not replacement_key:
                return ""
            row = DashboardRow(worker_id)
            self.dash_layout.addWidget(row)
            self.worker_widgets[worker_id] = row
            worker = GeminiParallelChapterWorker(
                worker_id,
                replacement_key,
                self.bm,
                self._selected_model_id(),
                self.combo_voices.currentData(),
                self.combo_voice_secondary.currentData(),
                self.combo_voice_tertiary.currentData(),
                self.combo_speed.currentText(),
                self._active_manager_queue,
                self._parallel_live_state,
                daily_request_limiter=self.daily_request_limiter,
                voice_mode=self._selected_voice_mode(),
                allow_edge_fallback=self.chk_edge_fallback.isChecked(),
            )
            self._attach_common_worker_signals(worker, row)

        elif self._is_flash_tts_mode() or self._active_job_kind == "prepare":
            required_model_ids = (
                [self._selected_preprocess_model_id()]
                if self._active_job_kind == "prepare"
                else self._worker_models_for_limit()
            )
            replacement_key = self._available_replacement_key(required_model_ids)
            if not replacement_key:
                return ""
            row = DashboardRow(worker_id)
            self.dash_layout.addWidget(row)
            self.worker_widgets[worker_id] = row
            run_mode = "prepare" if self._active_job_kind == "prepare" else self._selected_pipeline_mode()
            live_playback = self.player is not None and run_mode != "prepare"
            worker = FlashTtsWorker(
                worker_id,
                replacement_key,
                self.bm,
                self.audio_queue if live_playback else None,
                self._selected_model_id(),
                self.combo_voices.currentData(),
                self.combo_voice_secondary.currentData(),
                self.combo_speed.currentText(),
                self.chk_mp3.isChecked() if run_mode != "prepare" else False,
                self.chk_fast.isChecked() if run_mode != "prepare" else True,
                self.spin_chunk.value(),
                self._active_manager_queue,
                self._selected_preprocess_model_id(),
                self.combo_preprocess_profile.currentData(),
                self._selected_voice_mode(),
                run_mode,
                self.preprocess_directive,
                self.tts_directive,
                self.daily_request_limiter,
                allow_edge_fallback=self.chk_edge_fallback.isChecked(),
            )
            self._attach_common_worker_signals(worker, row)
            worker.chapter_done_ui_signal.connect(self.on_chapter_done_ui)
            worker.script_ready_signal.connect(self.on_script_ready_ui)

        else:
            replacement_key = self._available_replacement_key([self._selected_model_id()])
            if not replacement_key:
                return ""
            row = DashboardRow(worker_id)
            self.dash_layout.addWidget(row)
            self.worker_widgets[worker_id] = row
            live_playback = self.player is not None
            worker = GeminiWorker(
                worker_id,
                replacement_key,
                self.bm,
                self.audio_queue if live_playback else None,
                self._selected_model_id(),
                self.combo_voices.currentData(),
                "Ты диктор.",
                self.combo_speed.currentText(),
                self.chk_mp3.isChecked(),
                self.chk_fast.isChecked(),
                self.spin_chunk.value(),
                self._selected_live_segment_mode(),
                self._active_manager_queue,
                daily_request_limiter=self.daily_request_limiter,
                voice_mode=self._selected_voice_mode(),
                secondary_voice=self.combo_voice_secondary.currentData(),
                tertiary_voice=self.combo_voice_tertiary.currentData(),
                allow_edge_fallback=self.chk_edge_fallback.isChecked(),
            )
            self._attach_common_worker_signals(worker, row)
            worker.chapter_done_ui_signal.connect(self.on_chapter_done_ui)

        worker.start_stagger_index = 0
        self.workers.append(worker)
        worker.start()
        return getattr(worker, "api_key", "")

    def _on_quota_worker_key(self, worker_id, api_key, model_id, error_text, chapter_index):
        self.disabled_api_keys.add(api_key)
        self._run_had_invalid_keys = True
        if self.settings_manager is not None and model_id:
            try:
                self.settings_manager.mark_key_as_exhausted(api_key, model_id)
            except Exception:
                pass
        self._update_worker_spinbox_limit()
        self._update_key_state_ui()
        chapter_label = f"глава {chapter_index + 1}" if chapter_index >= 0 else "текущая глава"
        masked_key = _mask_api_key(api_key)
        replacement_key = self._start_replacement_worker()
        if replacement_key:
            self.statusBar().showMessage(
                f"Ключ {masked_key} списан по лимиту {model_id}; "
                f"{chapter_label} возвращена в очередь, замена: {_mask_api_key(replacement_key)}."
            )
        else:
            self.statusBar().showMessage(
                f"Ключ {masked_key} списан по лимиту {model_id}; "
                f"{chapter_label} возвращена в очередь, свободной замены нет."
            )

    def _on_project_quota_worker(self, worker_id, model_id, error_text, chapter_index):
        chapter_label = f"глава {chapter_index + 1}" if chapter_index >= 0 else "текущая глава"
        self._project_quota_message = (
            f"Дневной лимит проекта для {model_id} исчерпан; {chapter_label} возвращена в очередь."
        )
        self.statusBar().showMessage(self._project_quota_message)
        for worker in list(self.workers):
            if getattr(worker, "worker_id", None) != worker_id:
                worker.stop()

    def _chapter_live_segments_for_parallel(self, chapter_index):
        chapter = self.bm.chapters[chapter_index]
        if self._selected_voice_mode() == "author_gender":
            saved_script = self.bm.load_tts_script(chapter_index) if self.bm is not None else ""
            if _script_matches_voice_mode(saved_script, "author_gender"):
                script_segments = _split_author_gender_script_segments(saved_script, self._selected_live_segment_mode())
                if script_segments:
                    return script_segments
            return []
        if self._selected_live_segment_mode() == "paragraphs":
            paragraphs = []
            for part in getattr(chapter, "paragraphs", []) or []:
                paragraphs.extend(_split_live_paragraph(part))
            paragraphs = [part.strip() for part in paragraphs if (part or "").strip()]
            if paragraphs:
                return paragraphs
            raw_text = (chapter.raw_text or "").strip()
            return [raw_text] if raw_text else []
        return [part.strip() for part in chapter.flat_sentences if (part or "").strip()]

    def _build_parallel_live_tasks(self, chapter_index):
        segments = self._chapter_live_segments_for_parallel(chapter_index)
        chunk_size = max(1, self.spin_chunk.value())
        tasks = []
        for start_idx in range(0, len(segments), chunk_size):
            task_segments = segments[start_idx:start_idx + chunk_size]
            if not task_segments:
                continue
            task_index = len(tasks)
            text_value = _join_live_request_segments(
                task_segments,
                self._selected_live_segment_mode(),
                self._selected_voice_mode(),
            )
            tasks.append(
                {
                    "chapter_index": chapter_index,
                    "task_index": task_index,
                    "text": text_value,
                }
            )
        total_tasks = len(tasks)
        for task in tasks:
            task["total_tasks"] = total_tasks
        return tasks

    def _parallel_live_temp_dir(self, chapter_index):
        return os.path.join(self.bm.book_dir, LIVE_PARALLEL_DIRNAME, f"Ch{chapter_index + 1}")

    def _cleanup_parallel_live_state(self, remove_files=True):
        state = self._parallel_live_state
        self._parallel_live_state = None
        if not state:
            return
        if remove_files:
            temp_dir = state.get("temp_dir")
            if temp_dir and os.path.isdir(temp_dir):
                try:
                    shutil.rmtree(temp_dir, ignore_errors=True)
                except Exception:
                    pass

    def _finalize_parallel_live_chapter(self):
        state = self._parallel_live_state
        if not state:
            return ""

        chapter_index = state["chapter_index"]
        total_tasks = int(state.get("total_tasks", 0) or 0)
        completed_count = int(state.get("completed_count", 0) or 0)
        remaining_tasks = state["task_queue"].qsize() if state.get("task_queue") is not None else 0

        if state.get("cancelled"):
            self._cleanup_parallel_live_state(remove_files=True)
            return "Параллельная озвучка главы остановлена."

        if remaining_tasks > 0 or completed_count < total_tasks:
            self._cleanup_parallel_live_state(remove_files=False)
            return (
                f"Параллельная озвучка главы {chapter_index + 1} остановлена: "
                f"готово {completed_count}/{total_tasks}, осталось {max(0, total_tasks - completed_count)} блок(ов)."
            )

        output_paths = [task["output_path"] for task in sorted(state.get("tasks", []), key=lambda item: item["task_index"])]
        missing_files = [path for path in output_paths if not os.path.exists(path)]
        if missing_files:
            self._cleanup_parallel_live_state(remove_files=False)
            return (
                f"Ошибка сборки главы {chapter_index + 1}: отсутствуют временные блоки "
                f"({len(missing_files)} шт.)."
            )

        chapter_path = self.bm.get_mp3_path(chapter_index)
        try:
            _combine_mp3_sequence(output_paths, chapter_path)
            self.bm.mark_chapter_done(chapter_index)
            self.on_chapter_done_ui(chapter_index)
            self._cleanup_parallel_live_state(remove_files=True)
            return (
                f"Глава {chapter_index + 1} озвучена параллельно: "
                f"{total_tasks} блок(ов), {state.get('worker_count', 1)} воркер(ов)."
            )
        except Exception as exc:
            self._cleanup_parallel_live_state(remove_files=False)
            return f"Ошибка сборки главы {chapter_index + 1}: {exc}"

    def _launch_parallel_live_workers(self, chapter_index, available_api_keys, requested_workers):
        if not self.chk_mp3.isChecked():
            QMessageBox.warning(
                self,
                "Режим вывода",
                "Для нескольких воркеров на одной главе нужна запись MP3. Включите 'Запись MP3'.",
            )
            return False

        tasks = self._build_parallel_live_tasks(chapter_index)
        if len(tasks) < 2:
            self.statusBar().showMessage("В главе слишком мало блоков для распараллеливания; используется обычный режим.")
            return False

        temp_dir = self._parallel_live_temp_dir(chapter_index)
        shutil.rmtree(temp_dir, ignore_errors=True)
        os.makedirs(temp_dir, exist_ok=True)

        task_queue = queue.Queue()
        for task in tasks:
            task["output_path"] = os.path.join(temp_dir, f"seg_{task['task_index']:05d}.mp3")
            task_queue.put(task)

        num_workers = min(requested_workers, len(available_api_keys), len(tasks))
        if num_workers <= 1:
            shutil.rmtree(temp_dir, ignore_errors=True)
            return False

        self._parallel_live_state = {
            "chapter_index": chapter_index,
            "total_tasks": len(tasks),
            "completed_count": 0,
            "completed_tasks": set(),
            "task_queue": task_queue,
            "temp_dir": temp_dir,
            "tasks": tasks,
            "segment_mode": self._selected_live_segment_mode(),
            "worker_count": num_workers,
            "cancelled": False,
            "lock": threading.Lock(),
        }

        self._active_job_kind = "tts_parallel_live"
        self._active_reader_engine = "live"
        self._active_flash_run_mode = None
        self._active_manager_queue = task_queue
        self._run_had_invalid_keys = False
        self._stop_requested = False
        self._reset_dashboard()

        self.statusBar().showMessage(
            f"Параллельная озвучка главы {chapter_index + 1}: {len(tasks)} блок(ов), {num_workers} воркер(ов). "
            "Live-воспроизведение для этого режима отключено."
        )

        for i in range(num_workers):
            row = DashboardRow(i)
            self.dash_layout.addWidget(row)
            self.worker_widgets[i] = row

            worker = GeminiParallelChapterWorker(
                i,
                available_api_keys[i],
                self.bm,
                self._selected_model_id(),
                self.combo_voices.currentData(),
                self.combo_voice_secondary.currentData(),
                self.combo_voice_tertiary.currentData(),
                self.combo_speed.currentText(),
                task_queue,
                self._parallel_live_state,
                daily_request_limiter=self.daily_request_limiter,
                voice_mode=self._selected_voice_mode(),
                allow_edge_fallback=self.chk_edge_fallback.isChecked(),
            )
            worker.worker_progress.connect(self._enqueue_worker_progress)
            worker.finished_signal.connect(self._on_worker_finished)
            worker.invalid_key_signal.connect(self._on_invalid_worker_key)
            worker.quota_key_signal.connect(self._on_quota_worker_key)
            worker.project_quota_signal.connect(self._on_project_quota_worker)
            worker.error_signal.connect(lambda _wid, msg: self.statusBar().showMessage(msg))
            self.workers.append(worker)
            worker.start()

        self._set_reading_controls_running(True)
        self.tabs.setCurrentIndex(2)
        return True

    def _launch_live_workers(self, target_indices):
        self._parallel_live_state = None
        q = queue.Queue()
        for idx in target_indices:
            q.put(idx)

        actual_model_id = self._selected_model_id()
        exhausted_message = self._project_rpd_exhausted_message(actual_model_id)
        if exhausted_message:
            QMessageBox.warning(self, "Лимит RPD", exhausted_message)
            return False
        available_api_keys = self._get_available_api_keys([actual_model_id])
        if not available_api_keys:
            QMessageBox.warning(
                self,
                "Ключи API",
                "Нет рабочих API-ключей для выбранной live-модели. "
                "Проверьте ключи или дождитесь сброса лимитов.",
            )
            return False

        requested_workers = max(1, self.spin_workers.value())
        num_workers = min(requested_workers, len(available_api_keys), q.qsize())
        if num_workers == 0:
            num_workers = 1

        parallel_single_chapter = (
            self.chk_parallel_single_chapter.isChecked()
            and len(target_indices) == 1
            and requested_workers > 1
        )
        if parallel_single_chapter:
            launched = self._launch_parallel_live_workers(target_indices[0], available_api_keys, requested_workers)
            if launched:
                return True

        multi = num_workers > 1
        live_playback = not self.chk_fast.isChecked() and not multi and pyaudio is not None
        if not live_playback and not self.chk_fast.isChecked() and not multi and pyaudio is None:
            self.statusBar().showMessage("PyAudio не найден: live-воспроизведение отключено, используется только экспорт.")
        if not live_playback and not self.chk_mp3.isChecked():
            QMessageBox.warning(self, "Режим вывода", "Отключены и live-воспроизведение, и запись MP3. Включите запись MP3 или установите PyAudio.")
            return False

        if num_workers < requested_workers:
            self.statusBar().showMessage(
                f"Запуск {num_workers} воркеров вместо {requested_workers}: ограничение по доступным ключам или числу глав."
            )

        if live_playback:
            self.player = AudioPlayer(self.audio_queue, 80)
            self.player.start()

        self._active_job_kind = "tts"
        self._active_reader_engine = "live"
        self._active_flash_run_mode = None
        self._active_manager_queue = q
        self._run_had_invalid_keys = False
        self._stop_requested = False
        self._reset_dashboard()

        for i in range(num_workers):
            row = DashboardRow(i)
            self.dash_layout.addWidget(row)
            self.worker_widgets[i] = row

            worker = GeminiWorker(
                i,
                available_api_keys[i],
                self.bm,
                self.audio_queue if live_playback else None,
                actual_model_id,
                self.combo_voices.currentData(),
                "Ты диктор.",
                self.combo_speed.currentText(),
                self.chk_mp3.isChecked(),
                self.chk_fast.isChecked(),
                self.spin_chunk.value(),
                self._selected_live_segment_mode(),
                q,
                daily_request_limiter=self.daily_request_limiter,
                voice_mode=self._selected_voice_mode(),
                secondary_voice=self.combo_voice_secondary.currentData(),
                tertiary_voice=self.combo_voice_tertiary.currentData(),
                allow_edge_fallback=self.chk_edge_fallback.isChecked(),
            )
            worker.worker_progress.connect(self._enqueue_worker_progress)
            worker.finished_signal.connect(self._on_worker_finished)
            worker.chapter_done_ui_signal.connect(self.on_chapter_done_ui)
            worker.invalid_key_signal.connect(self._on_invalid_worker_key)
            worker.quota_key_signal.connect(self._on_quota_worker_key)
            worker.project_quota_signal.connect(self._on_project_quota_worker)
            worker.error_signal.connect(lambda _wid, msg: self.statusBar().showMessage(msg))
            self.workers.append(worker)
            worker.start()

        self._set_reading_controls_running(True)
        self.tabs.setCurrentIndex(2)
        return True

    def _launch_flash_workers(self, target_indices, run_mode):
        self._parallel_live_state = None
        q = queue.Queue()
        for idx in target_indices:
            q.put(idx)

        required_model_ids = []
        if run_mode in {"auto", "prepare"}:
            required_model_ids.append(self._selected_preprocess_model_id())
        if run_mode != "prepare":
            required_model_ids.append(self._selected_model_id())

        for model_id in required_model_ids:
            exhausted_message = self._project_rpd_exhausted_message(model_id)
            if exhausted_message:
                QMessageBox.warning(self, "Лимит RPD", exhausted_message)
                return False

        available_api_keys = self._get_available_api_keys(required_model_ids)
        if not available_api_keys:
            pipeline_label = "AI-подготовки сценария" if run_mode == "prepare" else "Flash TTS pipeline"
            QMessageBox.warning(
                self,
                "Ключи API",
                f"Нет рабочих API-ключей для выбранного {pipeline_label}. "
                "Проверьте ключи или дождитесь сброса лимитов.",
            )
            return False

        requested_workers = max(1, self.spin_workers.value())
        num_workers = min(requested_workers, len(available_api_keys), q.qsize())
        if num_workers == 0:
            num_workers = 1

        if self._selected_voice_mode() == "author_gender" and run_mode != "prepare":
            QMessageBox.warning(self, "Режим", "Режим 'Автор + Муж./Жен. роли' поддерживается только в Live API.")
            return False

        if run_mode == "raw" and self._selected_voice_mode() == "duo":
            QMessageBox.warning(self, "Режим", "Двухголосый режим требует AI-сценарий. Выберите 'Авто' или 'По шагам'.")
            return False

        live_playback = run_mode != "prepare" and not self.chk_fast.isChecked() and num_workers == 1 and pyaudio is not None
        if live_playback:
            self.player = AudioPlayer(self.audio_queue, 80)
            self.player.start()
        elif run_mode != "prepare" and not self.chk_fast.isChecked() and num_workers == 1 and pyaudio is None:
            self.statusBar().showMessage("PyAudio не найден: live-воспроизведение для Flash TTS отключено, используется экспорт.")

        if num_workers < requested_workers:
            self.statusBar().showMessage(
                f"Запуск {num_workers} воркеров вместо {requested_workers}: ограничение по доступным ключам или числу глав."
            )

        self._active_job_kind = "prepare" if run_mode == "prepare" else "tts"
        self._active_reader_engine = "flash_tts"
        self._active_flash_run_mode = run_mode
        self._active_manager_queue = q
        self._run_had_invalid_keys = False
        self._stop_requested = False
        self._reset_dashboard()

        for i in range(num_workers):
            row = DashboardRow(i)
            self.dash_layout.addWidget(row)
            self.worker_widgets[i] = row

            worker = FlashTtsWorker(
                i,
                available_api_keys[i],
                self.bm,
                self.audio_queue if live_playback else None,
                self._selected_model_id(),
                self.combo_voices.currentData(),
                self.combo_voice_secondary.currentData(),
                self.combo_speed.currentText(),
                self.chk_mp3.isChecked() if run_mode != "prepare" else False,
                self.chk_fast.isChecked() if run_mode != "prepare" else True,
                self.spin_chunk.value(),
                q,
                self._selected_preprocess_model_id(),
                self.combo_preprocess_profile.currentData(),
                self._selected_voice_mode(),
                run_mode,
                self.preprocess_directive,
                self.tts_directive,
                self.daily_request_limiter,
                allow_edge_fallback=self.chk_edge_fallback.isChecked(),
            )
            worker.worker_progress.connect(self._enqueue_worker_progress)
            worker.finished_signal.connect(self._on_worker_finished)
            worker.chapter_done_ui_signal.connect(self.on_chapter_done_ui)
            worker.script_ready_signal.connect(self.on_script_ready_ui)
            worker.invalid_key_signal.connect(self._on_invalid_worker_key)
            worker.quota_key_signal.connect(self._on_quota_worker_key)
            worker.project_quota_signal.connect(self._on_project_quota_worker)
            worker.error_signal.connect(lambda _wid, msg: self.statusBar().showMessage(msg))
            self.workers.append(worker)
            worker.start()

        self._set_reading_controls_running(True)
        self.tabs.setCurrentIndex(2 if run_mode != "prepare" else 1)
        return True

    def toggle_play(self):
        if genai is None or genai_types is None:
            QMessageBox.warning(self, "Зависимости", "Для озвучки Gemini Reader требуется пакет google-genai.")
            return
        if not self.api_keys:
            self.set_api_keys()
            if not self.api_keys:
                return
        
        if not self.workers:
            if not self.bm:
                QMessageBox.warning(self, "Внимание", "Сначала откройте книгу!")
                return

            target_indices = self._collect_run_scope_indices(include_done=False, action_label="озвучки")
            if target_indices is None:
                return
            if not target_indices and self.chk_selected_only.isChecked():
                QMessageBox.information(self, "Главы", "Среди отмеченных глав нет доступных для озвучки.")
                return
            if not target_indices:
                QMessageBox.information(self, "Готово", "Все главы для озвучки уже завершены или пропущены!")
                return

            if self._is_flash_tts_mode():
                pipeline_mode = self._selected_pipeline_mode()
                self._launch_flash_workers(target_indices, run_mode=pipeline_mode)
            else:
                if self._selected_voice_mode() == "author_gender":
                    if not self._ensure_author_gender_scripts_ready(target_indices, action_label="озвучки"):
                        return
                self._launch_live_workers(target_indices)
        else:
            self.force_stop()




    def force_stop(self):
        if self._parallel_live_state is not None:
            self._parallel_live_state["cancelled"] = True
        self._stop_requested = bool(self.workers)
        for w in list(self.workers):
            w.stop()
        self._active_manager_queue = None
        self._active_reader_engine = None
        self._active_flash_run_mode = None
        self._run_had_invalid_keys = False
        self._project_quota_message = ""
        
        if self.player:
            self.player.stop()
            self.player = None
            
        self._set_reading_controls_running(False)
        self.statusBar().showMessage("Процесс останавливается..." if self._stop_requested else "Процесс остановлен.")

    def _start_audio_combiner(self, video_image_path=None):
        if not self.bm:
            return
        ffmpeg_path = _resolve_tool_path("ffmpeg")
        ffprobe_path = _resolve_tool_path("ffprobe")
        if ffmpeg_path is None or ffprobe_path is None:
            QMessageBox.warning(self, "Зависимости", "Для склейки MP3 и видео нужны ffmpeg и ffprobe рядом с приложением или в PATH.")
            return

        chapter_indices = None
        if self.chk_selected_only.isChecked():
            chapter_indices = self._collect_run_scope_indices(include_done=True, action_label="экспорта")
            if chapter_indices is None:
                return
            if not chapter_indices:
                QMessageBox.information(
                    self,
                    "Экспорт",
                    "Среди отмеченных глав нет доступных MP3 для экспорта.",
                )
                return

        self.combiner = AudioCombinerWorker(
            self.bm,
            self.combo_voices.currentData(),
            ffmpeg_path=ffmpeg_path,
            ffprobe_path=ffprobe_path,
            video_image_path=video_image_path,
            chapter_indices=chapter_indices,
        )
        self.combiner.progress_signal.connect(lambda m: self.statusBar().showMessage(m))

        def on_combine_finished(message):
            self.combiner = None
            self._refresh_runtime_controls()
            if str(message).lower().startswith("ошибка"):
                self.statusBar().showMessage(message)
            else:
                self.statusBar().showMessage("✅ Экспорт видео завершён!" if video_image_path else "✅ Склейка завершена!")
            QApplication.processEvents()
            QMessageBox.information(self, "Видео" if video_image_path else "Склейка", message)

        self.combiner.finished_signal.connect(on_combine_finished)
        self._refresh_runtime_controls()
        self.combiner.start()

    def run_combine(self):
        self._start_audio_combiner()

    def run_export_video(self):
        if not self.bm:
            return

        video_cover_path = self.bm.get_video_cover_path()
        if not video_cover_path:
            QMessageBox.information(self, "Видео", "Сначала выберите картинку для видео.")
            return
        self._start_audio_combiner(video_image_path=video_cover_path)

    def on_vol(self, v):
        if self.player: self.player.set_volume(v)

    def set_api_keys(self):
        dlg = ApiKeysDialog("\n".join(self.api_keys), self)
        if dlg.exec():
            new_keys = dlg.get_keys()
            self.disabled_api_keys.clear()
            self._save_api_keys_with_status_preservation(new_keys)
            self._update_worker_spinbox_limit()
            self._update_key_state_ui()



    def save_settings(self):
        if self._loading_settings:
            return
        data = self._reader_ui_settings_payload()
        if self.settings_manager is not None:
            self.settings_manager.save_ui_state({READER_SETTINGS_KEY: data})
        else:
            legacy_data = _load_legacy_settings()
            legacy_data.update(data)
            legacy_data["api_keys"] = self.api_keys
            with open(LEGACY_SETTINGS_FILE, "w", encoding="utf-8") as file_obj:
                json.dump(legacy_data, file_obj, ensure_ascii=False, indent=2)

        if self.project_settings_manager is not None:
            self.project_settings_manager.save_full_session_settings(data)
        return

        data = {
            "api_keys": self.api_keys,
            "model": self.combo_model.currentText(),
            "voice": self.combo_voices.currentData(), # Сохраняем чистый ID голоса
            "speed": self.combo_speed.currentText(),
            "chunk": self.spin_chunk.value(),
            "record": self.chk_mp3.isChecked(),
            "fast": self.chk_fast.isChecked()
        }
        with open("settings.json", "w") as f:
            json.dump(data, f)




    def _refresh_project_settings_manager(self):
        if SettingsManager is None or not self.bm or not self.bm.book_dir:
            self.project_settings_manager = None
            return
        project_settings_path = os.path.join(self.bm.book_dir, READER_PROJECT_SETTINGS_FILENAME)
        self.project_settings_manager = SettingsManager(config_file=project_settings_path)

    def _apply_reader_settings(self, data):
        self._loading_settings = True
        try:
            engine_id = data.get("engine", "live")
            engine_idx = self.combo_engine.findData(engine_id)
            if engine_idx >= 0:
                self.combo_engine.setCurrentIndex(engine_idx)
            else:
                self.combo_engine.setCurrentIndex(0)

            model_label = data.get("model")
            if model_label and model_label in [self.combo_model.itemText(i) for i in range(self.combo_model.count())]:
                self.combo_model.setCurrentText(model_label)

            voice_mode = data.get("voice_mode", "single")
            if engine_id == "flash_tts" and voice_mode == "author_gender":
                voice_mode = "duo"
            voice_mode_idx = self.combo_voice_mode.findData(voice_mode)
            if voice_mode_idx >= 0:
                self.combo_voice_mode.setCurrentIndex(voice_mode_idx)

            voice_id = data.get("voice")
            if voice_id and voice_id in VOICES_MAP:
                voice_idx = self.combo_voices.findData(voice_id)
                if voice_idx >= 0:
                    self.combo_voices.setCurrentIndex(voice_idx)

            voice_secondary = data.get("voice_secondary")
            if voice_secondary and voice_secondary in VOICES_MAP:
                secondary_idx = self.combo_voice_secondary.findData(voice_secondary)
                if secondary_idx >= 0:
                    self.combo_voice_secondary.setCurrentIndex(secondary_idx)

            voice_tertiary = data.get("voice_tertiary")
            if voice_tertiary and voice_tertiary in VOICES_MAP:
                tertiary_idx = self.combo_voice_tertiary.findData(voice_tertiary)
                if tertiary_idx >= 0:
                    self.combo_voice_tertiary.setCurrentIndex(tertiary_idx)

            speed_label = data.get("speed")
            if speed_label in SPEED_PROMPTS:
                self.combo_speed.setCurrentText(speed_label)

            live_segment_mode = data.get("live_segment_mode", "sentences")
            live_segment_idx = self.combo_live_segment_mode.findData(live_segment_mode)
            if live_segment_idx >= 0:
                self.combo_live_segment_mode.setCurrentIndex(live_segment_idx)

            preprocess_model = data.get("preprocess_model")
            if preprocess_model and preprocess_model in [self.combo_preprocess_model.itemText(i) for i in range(self.combo_preprocess_model.count())]:
                self.combo_preprocess_model.setCurrentText(preprocess_model)
            else:
                default_preprocess_label = _default_preprocess_model_label(self.preprocess_models_map)
                if default_preprocess_label:
                    self.combo_preprocess_model.setCurrentText(default_preprocess_label)

            preprocess_profile = data.get("preprocess_profile")
            if preprocess_profile and preprocess_profile in [self.combo_preprocess_profile.itemText(i) for i in range(self.combo_preprocess_profile.count())]:
                self.combo_preprocess_profile.setCurrentText(preprocess_profile)

            pipeline_mode = data.get("pipeline_mode", "auto")
            pipeline_idx = self.combo_pipeline_mode.findData(pipeline_mode if engine_id == "flash_tts" else "auto")
            if pipeline_idx >= 0:
                self.combo_pipeline_mode.setCurrentIndex(pipeline_idx)

            self.preprocess_directive = data.get("preprocess_directive") or DEFAULT_PREPROCESS_DIRECTIVE
            self.tts_directive = data.get("tts_directive") or DEFAULT_TTS_DIRECTIVE

            chunk_value = data.get("chunk", 2)
            try:
                chunk_value = int(chunk_value)
            except (TypeError, ValueError):
                chunk_value = 2
            self.spin_chunk.setValue(max(self.spin_chunk.minimum(), min(chunk_value, self.spin_chunk.maximum())))

            self.chk_mp3.setChecked(data.get("record", True))
            self.chk_fast.setChecked(data.get("fast", False))
            self.chk_edge_fallback.setChecked(data.get("edge_fallback", True))
            self.chk_selected_only.setChecked(bool(data.get("selected_only", False)))
            self._checked_chapter_indices_state = {
                idx for idx in (data.get("checked_chapters") or [])
                if isinstance(idx, int) and idx >= 0
            }
            self.chk_parallel_single_chapter.setChecked(bool(data.get("parallel_single_chapter", False)))

            self._update_worker_spinbox_limit()
            worker_count = data.get("worker_count", data.get("num_instances", 1))
            try:
                worker_count = int(worker_count)
            except (TypeError, ValueError):
                worker_count = 1
            self.spin_workers.setValue(max(1, min(worker_count, self.spin_workers.maximum())))
        finally:
            self._loading_settings = False
            self._set_checked_chapter_indices(self._checked_chapter_indices_state)
            self._on_chapter_selection_changed()
            self._refresh_live_segment_controls()
            self._refresh_runtime_controls()
            self._sync_prompt_views()

    def load_settings(self):
        legacy_data = _load_legacy_settings()
        data = dict(legacy_data)

        if self.settings_manager is not None:
            try:
                all_settings = self.settings_manager.load_settings()
            except Exception:
                all_settings = {}
            data.update(all_settings.get(READER_SETTINGS_KEY, {}))
            try:
                self.api_keys = self.settings_manager.get_api_keys() or legacy_data.get("api_keys", [])
            except Exception:
                self.api_keys = legacy_data.get("api_keys", [])
        else:
            self.api_keys = legacy_data.get("api_keys", [])

        self._prune_invalid_key_states(self.api_keys)

        if self.project_settings_manager is not None:
            try:
                project_settings = self.project_settings_manager.load_full_session_settings()
            except Exception:
                project_settings = {}
            if isinstance(project_settings, dict):
                data.update(project_settings)

        self._apply_reader_settings(data)
        self._update_key_state_ui()
        return

        engine_id = data.get("engine", "live")
        engine_idx = self.combo_engine.findData(engine_id)
        if engine_idx >= 0:
            self.combo_engine.setCurrentIndex(engine_idx)
        else:
            self.combo_engine.setCurrentIndex(0)

        m = data.get("model")
        if m and m in [self.combo_model.itemText(i) for i in range(self.combo_model.count())]:
            self.combo_model.setCurrentText(m)

        voice_mode = data.get("voice_mode", "single")
        if engine_id == "flash_tts" and voice_mode == "author_gender":
            voice_mode = "duo"
        voice_mode_idx = self.combo_voice_mode.findData(voice_mode)
        if voice_mode_idx >= 0:
            self.combo_voice_mode.setCurrentIndex(voice_mode_idx)

        v = data.get("voice")
        if v and v in VOICES_MAP:
            idx = self.combo_voices.findData(v)
            if idx >= 0:
                self.combo_voices.setCurrentIndex(idx)

        v_secondary = data.get("voice_secondary")
        if v_secondary and v_secondary in VOICES_MAP:
            idx = self.combo_voice_secondary.findData(v_secondary)
            if idx >= 0:
                self.combo_voice_secondary.setCurrentIndex(idx)

        v_tertiary = data.get("voice_tertiary")
        if v_tertiary and v_tertiary in VOICES_MAP:
            idx = self.combo_voice_tertiary.findData(v_tertiary)
            if idx >= 0:
                self.combo_voice_tertiary.setCurrentIndex(idx)

        s = data.get("speed")
        if s in SPEED_PROMPTS:
            self.combo_speed.setCurrentText(s)

        preprocess_model = data.get("preprocess_model")
        if preprocess_model and preprocess_model in [self.combo_preprocess_model.itemText(i) for i in range(self.combo_preprocess_model.count())]:
            self.combo_preprocess_model.setCurrentText(preprocess_model)
        else:
            default_preprocess_label = _default_preprocess_model_label(self.preprocess_models_map)
            if default_preprocess_label:
                self.combo_preprocess_model.setCurrentText(default_preprocess_label)

        preprocess_profile = data.get("preprocess_profile")
        if preprocess_profile and preprocess_profile in [self.combo_preprocess_profile.itemText(i) for i in range(self.combo_preprocess_profile.count())]:
            self.combo_preprocess_profile.setCurrentText(preprocess_profile)

        pipeline_mode = data.get("pipeline_mode", "auto")
        pipeline_idx = self.combo_pipeline_mode.findData(pipeline_mode if engine_id == "flash_tts" else "auto")
        if pipeline_idx >= 0:
            self.combo_pipeline_mode.setCurrentIndex(pipeline_idx)

        self.preprocess_directive = data.get("preprocess_directive") or DEFAULT_PREPROCESS_DIRECTIVE
        self.tts_directive = data.get("tts_directive") or DEFAULT_TTS_DIRECTIVE

        self.spin_chunk.setValue(data.get("chunk", 2))
        self.chk_mp3.setChecked(data.get("record", True))
        self.chk_fast.setChecked(data.get("fast", False))
        return

        if os.path.exists("settings.json"):
            try:
                with open("settings.json", "r") as f:
                    data = json.load(f)
                    self.api_keys = data.get("api_keys",[])
                    
                    # Восстанавливаем UI
                    m = data.get("model")
                    if m and m in[self.combo_model.itemText(i) for i in range(self.combo_model.count())]:
                        self.combo_model.setCurrentText(m)
                        
                    # Загрузка голоса по его скрытому ID
                    v = data.get("voice")
                    if v and v in VOICES_MAP:
                        idx = self.combo_voices.findData(v)
                        if idx >= 0:
                            self.combo_voices.setCurrentIndex(idx)
                    
                    s = data.get("speed")
                    if s in SPEED_PROMPTS: self.combo_speed.setCurrentText(s)
                    
                    self.spin_chunk.setValue(data.get("chunk", 2))
                    self.chk_mp3.setChecked(data.get("record", True))
                    self.chk_fast.setChecked(data.get("fast", False))
            except: 
                pass

    def open_book_dialog(self):
        p, _ = QFileDialog.getOpenFileName(self, "Открыть книгу", "", READER_BOOK_FILE_FILTER)
        if p: self.load_book(p)

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
