import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from gemini_translator.benchmark.evaluator import evaluate_translation
from gemini_translator.benchmark.runner import BenchmarkRunner, summarize_results


class PromptBenchmarkEvaluatorTests(unittest.TestCase):
    def test_evaluator_scores_preserved_glossary_and_placeholder(self):
        source = "<p>\u9752\u4e91\u5b97\u7684\u5f1f\u5b50</p><!-- MEDIA_0 -->"
        output = "<p>Sekta Lazurnogo Oblaka disciple</p><!-- MEDIA_0 -->"

        result = evaluate_translation(
            source,
            output,
            glossary_entries=[
                {
                    "original": "\u9752\u4e91\u5b97",
                    "rus": "Sekta Lazurnogo Oblaka",
                }
            ],
            checks={
                "required": ["disciple"],
                "forbidden": ["\u9752\u4e91\u5b97"],
                "placeholders": ["<!-- MEDIA_0 -->"],
                "max_length_ratio": 8.0,
            },
        )

        self.assertGreaterEqual(result.score, 90)
        self.assertEqual(result.metrics["placeholders"]["missing"], [])
        self.assertEqual(result.metrics["required_terms"]["missing"], [])

    def test_evaluator_penalizes_lost_structure_and_untranslated_text(self):
        source = "<p>\u9752\u4e91\u5b97\u7684\u5f1f\u5b50</p><!-- MEDIA_0 -->"
        output = "<div>\u9752\u4e91\u5b97</div>"

        result = evaluate_translation(
            source,
            output,
            glossary_entries=[
                {
                    "original": "\u9752\u4e91\u5b97",
                    "rus": "Sekta Lazurnogo Oblaka",
                }
            ],
            checks={"placeholders": ["<!-- MEDIA_0 -->"]},
        )

        self.assertLess(result.score, 70)
        self.assertTrue(any("missing placeholders" in issue for issue in result.issues))
        self.assertGreater(result.metrics["cjk_residue_chars"], 0)


class PromptBenchmarkRunnerTests(unittest.TestCase):
    def test_prompt_only_run_writes_reports_and_compiled_prompt(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            config_path = tmp_path / "benchmark.json"
            output_dir = tmp_path / "out"
            config = {
                "name": "unit",
                "prompts": [
                    {
                        "id": "raw",
                        "mode": "raw",
                        "template": "Translate:\n{text}\nGlossary:\n{glossary}",
                    }
                ],
                "models": [{"id": "dummy", "provider": "local", "model_id": "dummy"}],
                "cases": [
                    {
                        "id": "case1",
                        "source_html": "<p>Hello</p>",
                        "glossary": [{"original": "Hello", "rus": "Privet"}],
                    }
                ],
            }
            config_path.write_text(json.dumps(config), encoding="utf-8")

            runner = BenchmarkRunner(config_path, output_dir=output_dir, prompt_only=True)
            report = runner.run()

            self.assertEqual(len(report["results"]), 1)
            result = report["results"][0]
            self.assertEqual(result["status"], "prompt_only")
            self.assertGreater(result["prompt_tokens_estimate"], 0)
            self.assertTrue((output_dir / "results.json").exists())
            self.assertTrue((output_dir / "results.csv").exists())
            self.assertTrue((output_dir / "summary.md").exists())
            prompt_path = output_dir / result["prompt_path"]
            self.assertTrue(prompt_path.exists())
            self.assertIn("<p>Hello</p>", prompt_path.read_text(encoding="utf-8"))

    def test_prompt_only_accepts_docx_source_path(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_path = Path(tmp)
            docx_path = tmp_path / "source.docx"
            document = Document()
            document.add_heading("Chapter One", level=1)
            document.add_paragraph("Hello from DOCX.")
            document.save(docx_path)

            config_path = tmp_path / "benchmark.json"
            output_dir = tmp_path / "out"
            config = {
                "name": "unit-docx",
                "prompts": [
                    {
                        "id": "raw",
                        "mode": "raw",
                        "template": "Translate:\n{text}",
                    }
                ],
                "models": [{"id": "dummy", "provider": "local", "model_id": "dummy"}],
                "cases": [{"id": "case-docx", "source_path": str(docx_path)}],
            }
            config_path.write_text(json.dumps(config), encoding="utf-8")

            report = BenchmarkRunner(config_path, output_dir=output_dir, prompt_only=True).run()
            prompt_path = output_dir / report["results"][0]["prompt_path"]

            self.assertIn("Hello from DOCX", prompt_path.read_text(encoding="utf-8"))

    def test_summary_ranks_by_score_then_errors(self):
        summary = summarize_results(
            [
                {"prompt_id": "a", "model_id": "m", "status": "ok", "score": 80, "latency_ms": 10},
                {"prompt_id": "b", "model_id": "m", "status": "ok", "score": 95, "latency_ms": 20},
                {"prompt_id": "a", "model_id": "m", "status": "error", "latency_ms": 5},
            ]
        )

        self.assertEqual(summary[0]["prompt_id"], "b")
        self.assertEqual(summary[0]["avg_score"], 95)
        self.assertEqual(summary[1]["errors"], 1)


if __name__ == "__main__":
    unittest.main()
