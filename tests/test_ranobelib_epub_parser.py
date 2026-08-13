import os
import sys
import zipfile

from docx import Document


TESTS_DIR = os.path.dirname(__file__)
PROJECT_ROOT = os.path.dirname(TESTS_DIR)
RANOBELIB_DIR = os.path.join(PROJECT_ROOT, "ranobelib")

if RANOBELIB_DIR not in sys.path:
    sys.path.insert(0, RANOBELIB_DIR)

from parsers import FileParser  # noqa: E402
from models import ChapterData  # noqa: E402
from utils import is_safe_remote_http_url, parse_vol_and_chapter  # noqa: E402
from workers import (  # noqa: E402
    MEDIA_COVER_MODE_CODEX,
    MEDIA_COVER_MODE_NONE,
    MEDIA_COVER_MODE_RULATE,
    QIDIAN_RULATE_PROFILE_DIR,
    RANOBELIB_GENRES,
    RANOBELIB_TAGS,
    RulateToRanobeCreateWorker,
    RulateDownloadWorker,
    _clean_rulate_media_title,
    _find_cached_chromium_executable,
    _is_browser_missing_error,
    _launch_persistent_chromium_context,
    _normalize_allowed_catalog_items,
    _normalize_rulate_cover_url,
    _normalize_rulate_media_payload,
    _normalize_publisher_for_source,
    _parse_ranobelib_catalog_response,
    _prepare_ranobelib_author_payload,
    _ranobelib_title_status_value,
    _rulate_edit_info_url,
    _rulate_public_book_url,
    publisher_candidates_from_source_url,
    publisher_from_source_url,
)


def _write_epub(path, chapter_bodies):
    manifest_items = []
    spine_items = []
    files = {}
    for index, body in enumerate(chapter_bodies, start=1):
        item_id = f"chapter{index}"
        href = f"Text/ch{index}.xhtml"
        manifest_items.append(
            f'<item id="{item_id}" href="{href}" media-type="application/xhtml+xml"/>'
        )
        spine_items.append(f'<itemref idref="{item_id}"/>')
        files[f"OEBPS/{href}"] = body

    opf = (
        '<?xml version="1.0" encoding="UTF-8"?>'
        '<package xmlns="http://www.idpf.org/2007/opf" version="2.0">'
        "<manifest>"
        + "".join(manifest_items)
        + "</manifest><spine>"
        + "".join(spine_items)
        + "</spine></package>"
    )

    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("OEBPS/content.opf", opf)
        for name, content in files.items():
            archive.writestr(name, content)


def test_parse_epub_reads_body_text_split_by_br(tmp_path):
    long_para = " ".join(["regular paragraph text"] * 20)
    direct_text = "<br/>".join(
        [
            "direct body text one " * 10,
            "direct body text two " * 10,
            "direct body text three " * 10,
        ]
    )
    epub_path = tmp_path / "book.epub"
    _write_epub(
        epub_path,
        [
            (
                '<html><body><h1>Chapter 1. P tags</h1>'
                f"<p>{long_para}</p></body></html>"
            ),
            (
                '<html><body><h1>Chapter 2. BR tags</h1>'
                f"<br/>{direct_text}</body></html>"
            ),
        ],
    )

    chapters = FileParser.parse_epub(str(epub_path), "1")

    assert len(chapters) == 2
    assert chapters[0].number == 1.0
    assert chapters[1].number == 2.0
    assert "direct body text one" in chapters[1].content
    assert chapters[1].content.count("<p>") == 3


def test_parse_zip_docx_keeps_russian_chapter_parts(tmp_path):
    zip_path = tmp_path / "rulate.zip"
    titles = [
        "Глава 30 Король Валид становится королём зубрёжки. Часть 1",
        "Глава 30 Король Валид становится королём зубрёжки. Часть 2",
        "Глава 30 Король Валид становится королём зубрёжки. Часть 3",
    ]

    with zipfile.ZipFile(zip_path, "w") as archive:
        for index, title in enumerate(titles, start=1):
            docx_path = tmp_path / f"chapter_{index}.docx"
            doc = Document()
            doc.add_paragraph(f"Text for {title}")
            doc.save(docx_path)
            archive.write(docx_path, f"{title}.docx")

    chapters = FileParser.parse_zip_docx(str(zip_path), "1")

    assert [chapter.number for chapter in chapters] == [30.1, 30.2, 30.3]
    assert chapters[0].title == "Король Валид становится королём зубрёжки"


def test_parse_zip_docx_preserves_archive_order(tmp_path):
    zip_path = tmp_path / "rulate.zip"
    titles = [
        "z Глава 31 Кредитор",
        "a Глава 24 Глава XXIV Золотой и пурпурный летящий генерал",
        "m Глава 25 Глава XXV Благородство мужа и письмо ценою в золото",
    ]

    with zipfile.ZipFile(zip_path, "w") as archive:
        for index, title in enumerate(titles, start=1):
            docx_path = tmp_path / f"chapter_{index}.docx"
            doc = Document()
            doc.add_paragraph(f"Text for {title}")
            doc.save(docx_path)
            archive.write(docx_path, f"{title}.docx")

    chapters = FileParser.parse_zip_docx(str(zip_path), "1")

    assert [chapter.number for chapter in chapters] == [31.0, 24.0, 25.0]
    assert [chapter.title for chapter in chapters] == [
        "Кредитор",
        "Золотой и пурпурный летящий генерал",
        "Благородство мужа и письмо ценою в золото",
    ]


def test_parse_vol_and_chapter_cleans_nested_roman_heading_with_char_count():
    vol, number, title, num_found = parse_vol_and_chapter(
        "Т.1 Гл.24: Глава XXIV: «Золотой и пурпурный летящий генерал!» [11,886 зн.]",
        "1",
        1,
    )

    assert vol == "1"
    assert number == 24.0
    assert title == "Золотой и пурпурный летящий генерал!"
    assert num_found is True


def test_parse_vol_and_chapter_accepts_roman_and_russian_word_numbers():
    roman = parse_vol_and_chapter("Глава XXXVIII: «Башня Чэнфэн»", "1", 1)
    words = parse_vol_and_chapter(
        "Глава тридцать пятая: «Схватка с великим полководцем»",
        "1",
        1,
    )

    assert roman == ("1", 38.0, "Башня Чэнфэн", True)
    assert words == ("1", 35.0, "Схватка с великим полководцем", True)


def test_rulate_worker_applies_full_site_titles_to_downloaded_chapters():
    infos = [
        {
            "id": "101",
            "title": "Глава 31 Очень длинное необрезанное название главы с сайта Rulate",
            "number": 31.0,
        },
        {
            "id": "102",
            "title": "Глава 32 Второе длинное необрезанное название главы с сайта Rulate",
            "number": 32.0,
        },
    ]
    chapters = [
        ChapterData("1", 31.0, "Очень длинное необрезанное...", "text 31"),
        ChapterData("1", 32.0, "Второе длинное необрезанное...", "text 32"),
    ]
    worker = RulateDownloadWorker(
        "https://tl.rulate.ru/book/1",
        "1",
        chapter_ids=["101", "102"],
        chapter_infos=infos,
    )

    worker._apply_chapter_infos(chapters, infos)

    assert chapters[0].title == "Очень длинное необрезанное название главы с сайта Rulate"
    assert chapters[1].title == "Второе длинное необрезанное название главы с сайта Rulate"


def test_rulate_worker_applies_roman_site_title_cleanly():
    chapter = ChapterData("1", 24.0, "Глава XXIV", "text")
    worker = RulateDownloadWorker(
        "https://tl.rulate.ru/book/1",
        "1",
        chapter_ids=["101"],
        chapter_infos=[
            {
                "id": "101",
                "title": "Глава XXIV: «Золотой и пурпурный летящий генерал!»",
                "number": 24.0,
                "volume": "1",
            },
        ],
    )

    worker._apply_chapter_info(chapter, worker.chapter_infos[0])

    assert chapter.volume == "1"
    assert chapter.number == 24.0
    assert chapter.title == "Золотой и пурпурный летящий генерал!"


def test_rulate_worker_keeps_selected_chapter_order_after_matching():
    infos = [
        {"id": "131", "title": "Глава 31 Кредитор", "number": 31.0, "volume": "1"},
        {
            "id": "124",
            "title": "Глава XXIV: «Золотой и пурпурный летящий генерал!»",
            "number": 24.0,
            "volume": "1",
        },
        {
            "id": "125",
            "title": "Глава XXV: «Благородство мужа и письмо ценою в золото»",
            "number": 25.0,
            "volume": "1",
        },
    ]
    chapters = [
        ChapterData("1", 24.0, "old 24", "text 24"),
        ChapterData("1", 25.0, "old 25", "text 25"),
        ChapterData("1", 31.0, "old 31", "text 31"),
    ]
    worker = RulateDownloadWorker(
        "https://tl.rulate.ru/book/1",
        "1",
        chapter_ids=["131", "124", "125"],
        chapter_infos=infos,
    )

    worker._apply_chapter_infos(chapters, infos)

    assert [chapter.number for chapter in chapters] == [31.0, 24.0, 25.0]
    assert [chapter.title for chapter in chapters] == [
        "Кредитор",
        "Золотой и пурпурный летящий генерал!",
        "Благородство мужа и письмо ценою в золото",
    ]


def test_rulate_worker_infers_next_volume_when_numbers_restart():
    worker = RulateDownloadWorker("https://tl.rulate.ru/book/1", "1")
    infos = [
        {"id": "101", "title": "Chapter 1 Start", "number": 1.0, "downloadable": True},
        {"id": "102", "title": "Chapter 2 Middle", "number": 2.0, "downloadable": True},
        {"id": "201", "title": "Chapter 1 Reset", "number": 1.0, "downloadable": True},
        {"id": "202", "title": "Chapter 2 Continue", "number": 2.0, "downloadable": True},
    ]

    annotated = worker._annotate_chapter_infos(infos)

    assert [chapter["volume"] for chapter in annotated] == ["1", "1", "2", "2"]
    assert [chapter["number"] for chapter in annotated] == [1.0, 2.0, 1.0, 2.0]


def test_rulate_worker_annotates_roman_and_russian_word_numbers():
    worker = RulateDownloadWorker("https://tl.rulate.ru/book/1", "1")
    infos = [
        {"id": "124", "title": "Глава XXIV: «Золотой и пурпурный летящий генерал!»", "number": 0},
        {"id": "135", "title": "Глава тридцать пятая: «Схватка с великим полководцем»", "number": 0},
    ]

    annotated = worker._annotate_chapter_infos(infos)

    assert [chapter["volume"] for chapter in annotated] == ["1", "1"]
    assert [chapter["number"] for chapter in annotated] == [24.0, 35.0]


def test_rulate_worker_applies_inferred_volume_to_downloaded_chapter():
    worker = RulateDownloadWorker("https://tl.rulate.ru/book/1", "1")
    chapter = ChapterData("1", 1.0, "Short", "text")

    worker._apply_chapter_info(
        chapter,
        {"title": "Chapter 1 Site title", "number": 1.0, "volume": "2"},
    )

    assert chapter.volume == "2"
    assert chapter.number == 1.0
    assert chapter.title == "Site title"


def test_rulate_to_ranobelib_uses_qidian_rulate_cookie_profile():
    if "QIDIAN_RULATE_PROFILE_DIR" not in os.environ:
        assert ".qidian_rulate_creator" in str(QIDIAN_RULATE_PROFILE_DIR)
        assert "rulate_profile" in str(QIDIAN_RULATE_PROFILE_DIR)


def test_clean_rulate_media_title_removes_site_suffix():
    assert _clean_rulate_media_title("Моя новелла | Rulate") == "Моя новелла"
    assert _clean_rulate_media_title("Книга Моя новелла / читать онлайн") == "Моя новелла"


def test_normalize_rulate_media_payload_for_ranobelib_create():
    payload = _normalize_rulate_media_payload(
        {
            "title": "Моя новелла | Rulate",
            "description": "Описание: первая строка\n\n\nвторая строка",
            "cover_url": "/uploads/cover.webp",
            "author": " Автор ",
            "alt_names": ["My Novel", "My Novel"],
            "original_source_url": "https://www.qidian.com/book/1041604040/",
            "status": "Завершён",
            "year": "2021 год",
        },
        "https://tl.rulate.ru/book/123",
    )

    assert payload["title_ru"] == "Моя новелла"
    assert payload["original_title"] == "My Novel"
    assert payload["title_en"] == "My Novel"
    assert payload["alt_names"] == "My Novel"
    assert payload["alt_hieroglyph_title"] == ""
    assert payload["description"] == "первая строка\n\nвторая строка"
    assert payload["cover_url"] == "https://tl.rulate.ru/uploads/cover.webp"
    assert payload["source_url"] == "https://www.qidian.com/book/1041604040/"
    assert payload["publisher"] == "Qidian"
    assert payload["rulate_url"] == "https://tl.rulate.ru/book/123"
    assert payload["rulate_edit_url"] == "https://tl.rulate.ru/book/123/edit/info"
    assert payload["author"] == "Автор"
    assert payload["status_value"] == "2"
    assert payload["year"] == "2021"
    assert payload["rulate_genres"] == []
    assert payload["rulate_tags"] == []


def test_source_url_sets_known_publisher_candidates():
    assert publisher_from_source_url("https://www.qidian.com/book/1041604040/") == "Qidian"
    assert publisher_from_source_url("https://fanqienovel.com/page/7229603492648717324") == "Fanqie Manhua"
    assert publisher_candidates_from_source_url(
        "https://www.fanqienovel.com/page/7229603492648717324?enter_from=search"
    ) == ["Fanqie Manhua"]
    assert publisher_from_source_url("https://example.com/book/1") == ""


def test_legacy_fanqie_publisher_is_normalized_to_current_name():
    assert (
        _normalize_publisher_for_source("FanqNovel", "https://fanqienovel.com/page/7229603492648717324")
        == "Fanqie Manhua"
    )
    assert _normalize_publisher_for_source("FanqNovel", "https://www.qidian.com/book/1041604040/") == "FanqNovel"


def test_normalize_rulate_media_payload_filters_noise_and_logo_cover():
    payload = _normalize_rulate_media_payload(
        {
            "title": "Ночной страж Дафэна",
            "original_title": "a: --- Продолжается Завершён Брошен",
            "alt_names": ["大奉打更人"],
            "cover_url": "https://tl.rulate.ru/i/logo/rulate-24.png",
        },
        "https://tl.rulate.ru/book/204281/edit/info",
    )

    assert payload["alt_hieroglyph_title"] == "大奉打更人"
    assert "Продолжается" not in payload["original_title"]
    assert payload["alt_names"] == "大奉打更人"
    assert payload["cover_url"] == ""
    assert _normalize_rulate_cover_url("/uploads/book-cover.webp", payload["source_url"]).endswith(
        "/uploads/book-cover.webp"
    )


def test_rulate_cover_urls_reject_local_files_and_private_networks():
    source_url = "https://tl.rulate.ru/book/204281/edit/info"

    assert is_safe_remote_http_url("https://cdn.example.com/cover.webp") is True
    assert is_safe_remote_http_url("file:///etc/passwd") is False
    assert is_safe_remote_http_url("http://localhost:8080/cover.jpg") is False
    assert is_safe_remote_http_url("http://127.0.0.1/cover.jpg") is False
    assert is_safe_remote_http_url("http://192.168.1.10/cover.jpg") is False
    assert _normalize_rulate_cover_url("file:///etc/passwd", source_url) == ""
    assert _normalize_rulate_cover_url("http://127.0.0.1/private.jpg", source_url) == ""


def test_normalize_rulate_media_payload_splits_concatenated_rulate_catalog(monkeypatch):
    monkeypatch.setattr(
        "workers._load_rulate_allowed_tags",
        lambda: ["умный гг", "система", "магия"],
    )
    payload = _normalize_rulate_media_payload(
        {
            "title": "Каталог",
            "genres": ["комедияфэнтезиприключениябоевые искусстваповседневность"],
            "tags": ["умный ггсистемамагия"],
        },
        "https://tl.rulate.ru/book/123/edit/info",
    )

    assert payload["rulate_genres"] == [
        "комедия",
        "фэнтези",
        "приключения",
        "боевые искусства",
        "повседневность",
    ]
    assert payload["rulate_tags"] == ["умный гг", "система", "магия"]


def test_rulate_and_ranobelib_catalog_fields_are_kept_separate():
    worker = RulateToRanobeCreateWorker(
        "https://tl.rulate.ru/book/123",
        options={
            "rulate_genres": ["Фэнтези"],
            "rulate_tags": ["Магия"],
            "translator_team": "Test Team",
        },
    )
    data = worker._apply_options({"genres": ["Фэнтези"], "tags": ["Магия"]})

    assert data["rulate_genres"] == ["Фэнтези"]
    assert data["rulate_tags"] == ["Магия"]
    assert data["genres"] == []
    assert data["tags"] == []
    assert data["translator_team"] == "Test Team"

    worker = RulateToRanobeCreateWorker(
        "https://tl.rulate.ru/book/123",
        options={
            "rulate_genres": ["Фэнтези"],
            "rulate_tags": ["Магия"],
            "genres": ["Драма"],
            "tags": ["Умный ГГ"],
        },
    )
    data = worker._apply_options({})

    assert data["genres"] == ["Драма"]
    assert data["tags"] == ["Умный ГГ"]


def test_ranobelib_catalog_normalizer_splits_csv_and_glued_values():
    assert _normalize_allowed_catalog_items(
        ["Комедия, Повседневность", "Романтика"],
        RANOBELIB_GENRES,
        5,
    ) == ["Комедия", "Повседневность", "Романтика"]

    assert _normalize_allowed_catalog_items(
        ["СистемаСовременностьРеинкарнация"],
        RANOBELIB_TAGS,
        8,
    ) == ["Система", "Современность", "Реинкарнация"]


def test_rulate_to_ranobelib_uses_prefetched_metadata_without_reopening_rulate():
    worker = RulateToRanobeCreateWorker(
        "https://tl.rulate.ru/book/123",
        options={
            "title_ru": "Уже загружено",
            "source_url": "https://www.qidian.com/book/1041604040/",
            "rulate_edit_url": "https://tl.rulate.ru/book/123/edit/info",
        },
    )

    metadata = worker._read_rulate_metadata(playwright=None)

    assert metadata["title_ru"] == "Уже загружено"
    assert metadata["source_url"] == "https://www.qidian.com/book/1041604040/"
    assert metadata["rulate_url"] == "https://tl.rulate.ru/book/123"


def test_rulate_to_ranobelib_uses_translated_codex_cover_instead_of_rulate(tmp_path, monkeypatch):
    translated_cover = tmp_path / "translated.png"
    translated_cover.write_bytes(b"translated-cover" * 128)
    worker = RulateToRanobeCreateWorker(
        "https://tl.rulate.ru/book/123",
        options={
            "cover_mode": MEDIA_COVER_MODE_CODEX,
            "cover_path": str(translated_cover),
            "cover_url": "https://tl.rulate.ru/uploads/rulate-cover.webp",
        },
    )
    monkeypatch.setattr(
        "workers.urllib.request.urlopen",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(
            AssertionError("Rulate cover must not be downloaded in Codex mode")
        ),
    )

    metadata = worker._apply_options({})

    assert metadata["cover_mode"] == MEDIA_COVER_MODE_CODEX
    assert worker._download_cover(metadata) == str(translated_cover.resolve())


def test_rulate_to_ranobelib_can_skip_cover_entirely(monkeypatch):
    worker = RulateToRanobeCreateWorker(
        "https://tl.rulate.ru/book/123",
        options={
            "cover_mode": MEDIA_COVER_MODE_NONE,
            "cover_url": "https://tl.rulate.ru/uploads/rulate-cover.webp",
        },
    )
    monkeypatch.setattr(
        "workers.urllib.request.urlopen",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(
            AssertionError("No cover should be downloaded in none mode")
        ),
    )

    metadata = worker._apply_options({})

    assert worker._download_cover(metadata) is None


def test_unknown_cover_mode_falls_back_to_rulate():
    worker = RulateToRanobeCreateWorker(
        "https://tl.rulate.ru/book/123",
        options={"cover_mode": "unexpected"},
    )

    assert worker._apply_options({})["cover_mode"] == MEDIA_COVER_MODE_RULATE


def test_prepare_ranobelib_author_payload_uses_romanized_name(monkeypatch):
    def fake_translate(value, target_lang, source_lang="auto", timeout=20):
        return {"en": "Far Pupil", "ru": "Далёкий зрачок"}.get(target_lang, "")

    monkeypatch.setattr("workers._google_translate_or_empty", fake_translate)
    monkeypatch.setattr("workers._google_romanize_or_empty", lambda *args, **kwargs: "Yuan Tong")

    payload = _prepare_ranobelib_author_payload("远瞳")

    assert payload["name_en"] == "Yuan Tong"
    assert payload["name_ru"] == "Далёкий зрачок"
    assert "远瞳" in payload["aliases"]
    assert "Far Pupil" in payload["aliases"]


def test_ranobelib_author_autocomplete_searches_original_name_once():
    worker = RulateToRanobeCreateWorker("https://tl.rulate.ru/book/123")

    candidates = worker._author_autocomplete_candidates({"author": "远瞳"})

    assert candidates == ["远瞳"]


def test_ranobelib_autocomplete_accepts_resolved_display_name(monkeypatch):
    worker = RulateToRanobeCreateWorker("https://tl.rulate.ru/book/123")
    state = {"clicked": False}

    class FakeKeyboard:
        def type(self, _value, delay=0):
            return None

        def press(self, key):
            if key in {"ArrowDown", "Enter"}:
                raise AssertionError("fallback keyboard selection should not run after a successful click")

    class FakePage:
        keyboard = FakeKeyboard()

        def wait_for_timeout(self, _timeout):
            return None

    page = FakePage()
    monkeypatch.setattr(worker, "_group_contains_value", lambda *_args: False)
    monkeypatch.setattr(
        worker,
        "_group_selected_text",
        lambda *_args: "Chang Shi" if state["clicked"] else "",
    )
    monkeypatch.setattr(worker, "_ensure_group_autocomplete_open", lambda *_args: True)
    monkeypatch.setattr(worker, "_focus_group_autocomplete_input", lambda *_args: True)
    monkeypatch.setattr(worker, "_clear_active_autocomplete_input", lambda *_args: True)
    monkeypatch.setattr(worker, "_wait_for_autocomplete_suggestion", lambda *_args: True)

    def fake_click(*_args, **_kwargs):
        state["clicked"] = True
        return True

    monkeypatch.setattr(worker, "_click_autocomplete_suggestion", fake_click)

    assert worker._add_autocomplete_item(page, "Автор", "常世", fast_missing=True) is True


def test_ranobelib_description_is_filled_after_dynamic_fields(monkeypatch):
    worker = RulateToRanobeCreateWorker("https://tl.rulate.ru/book/123")
    events = []

    class FakePage:
        def evaluate(self, _script, _metadata):
            events.append("main fields")
            return {
                "original": True,
                "titleRu": True,
                "titleEn": True,
                "altNames": True,
                "type": True,
                "status": True,
                "age": True,
                "translationStatus": True,
                "chaptersUpload": True,
            }

    monkeypatch.setattr(worker, "log", lambda *_args: None)
    monkeypatch.setattr(worker, "_upload_cover", lambda *_args: events.append("cover"))
    monkeypatch.setattr(worker, "_ensure_author", lambda *_args, **_kwargs: events.append("author"))
    monkeypatch.setattr(worker, "_ensure_publisher", lambda *_args: events.append("publisher"))
    monkeypatch.setattr(worker, "_ensure_translator_team", lambda *_args: events.append("team"))
    monkeypatch.setattr(
        worker,
        "_add_autocomplete_items",
        lambda _page, group_label, _values, _limit: events.append(group_label),
    )
    monkeypatch.setattr(worker, "_try_fill_source_link", lambda *_args: events.append("source"))
    monkeypatch.setattr(worker, "_fill_description", lambda *_args: events.append("description"))

    worker._fill_ranobelib_form(
        FakePage(),
        {
            "description": "Описание",
            "genres": ["Фэнтези"],
            "tags": ["Магия"],
        },
        None,
    )

    assert events == [
        "cover",
        "main fields",
        "author",
        "publisher",
        "team",
        "Жанры",
        "Теги",
        "source",
        "description",
    ]


def test_ranobelib_description_retries_when_editor_is_replaced(monkeypatch):
    worker = RulateToRanobeCreateWorker("https://tl.rulate.ru/book/123")
    inserted = []
    matches = iter((False, True))

    class FakePage:
        def evaluate(self, _script, value):
            inserted.append(value)
            return {"ok": True, "characters": len(value)}

        def wait_for_timeout(self, _timeout):
            return None

    monkeypatch.setattr(worker, "log", lambda *_args: None)
    monkeypatch.setattr(worker, "_description_matches", lambda *_args: next(matches))

    assert worker._fill_description(FakePage(), "Описание") is True
    assert inserted == ["Описание", "Описание"]


def test_ranobelib_description_is_not_truncated_before_editor_fill(monkeypatch):
    worker = RulateToRanobeCreateWorker("https://tl.rulate.ru/book/123")
    description = ("Длинное описание без искусственного ограничения. " * 30).strip()
    inserted = []

    class FakePage:
        def evaluate(self, _script, value):
            inserted.append(value)
            return {"ok": True, "characters": len(value)}

        def wait_for_timeout(self, _timeout):
            return None

    monkeypatch.setattr(worker, "log", lambda *_args: None)
    monkeypatch.setattr(worker, "_description_matches", lambda _page, value: value == description)

    assert len(description) > 800
    assert worker._fill_description(FakePage(), description) is True
    assert inserted == [description]


def test_ranobelib_translator_team_search_uses_teams_group_even_when_other_team_exists(monkeypatch):
    worker = RulateToRanobeCreateWorker("https://tl.rulate.ru/book/123")
    calls = []

    monkeypatch.setattr(worker, "_group_has_any_value", lambda page, group_label: True)

    def fake_add(page, group_label, value, **kwargs):
        calls.append((group_label, value))
        return group_label == "Команды"

    monkeypatch.setattr(worker, "_add_autocomplete_item", fake_add)

    worker._ensure_translator_team(object(), {"translator_team": "Required Team"})

    assert calls == [("Команды", "Required Team")]


def test_ranobelib_title_status_value_defaults_to_ongoing():
    assert _ranobelib_title_status_value("продолжается") == "1"
    assert _ranobelib_title_status_value("выпуск прекращён") == "5"


def test_parse_ranobelib_catalog_response_strictly_uses_allowed_items():
    payload = _parse_ranobelib_catalog_response(
        """
        {
          "genres": ["Фэнтези", "Мистика", "Несуществующий жанр", "Приключения"],
          "tags": ["Магия", "Умный ГГ", "Чужой тег", "Фэнтези мир"],
          "age_rating": "18+",
          "title_status": "completed",
          "translation_status": "frozen",
          "release_year": "2024"
        }
        """
    )

    assert payload["genres"] == ["Фэнтези", "Мистика", "Приключения"]
    assert payload["tags"] == ["Магия", "Умный ГГ", "Фэнтези мир"]
    assert all(genre in RANOBELIB_GENRES for genre in payload["genres"])
    assert all(tag in RANOBELIB_TAGS for tag in payload["tags"])
    assert payload["age_value"] == "4"
    assert payload["status_value"] == "2"
    assert payload["translation_status_value"] == "3"
    assert payload["year"] == "2024"


def test_browser_missing_error_is_detected():
    error = RuntimeError("Executable doesn't exist at C:/ms-playwright/chromium/chrome.exe\nplaywright install")

    assert _is_browser_missing_error(error)


def test_cached_chromium_prefers_newest_revision(monkeypatch, tmp_path):
    older = tmp_path / "chromium-1000" / "chrome-win64" / "chrome.exe"
    newer = tmp_path / "chromium-1223" / "chrome-win64" / "chrome.exe"
    older.parent.mkdir(parents=True)
    newer.parent.mkdir(parents=True)
    older.write_text("", encoding="utf-8")
    newer.write_text("", encoding="utf-8")
    monkeypatch.setattr("workers._candidate_browser_cache_roots", lambda: [tmp_path])

    assert _find_cached_chromium_executable() == newer


def test_persistent_browser_launch_falls_back_from_stale_playwright_revision(monkeypatch, tmp_path):
    cached = tmp_path / "chromium-1228" / "chrome-win64" / "chrome.exe"
    cached.parent.mkdir(parents=True)
    cached.write_text("", encoding="utf-8")
    calls = []
    expected_context = object()

    class FakeChromium:
        def launch_persistent_context(self, **kwargs):
            calls.append(kwargs)
            if "executable_path" not in kwargs:
                raise RuntimeError(
                    "Executable doesn't exist at C:/ms-playwright/chromium-1234/chrome.exe\n"
                    "playwright install"
                )
            return expected_context

    class FakePlaywright:
        chromium = FakeChromium()

    monkeypatch.setattr("workers._find_cached_chromium_executable", lambda: cached)

    context = _launch_persistent_chromium_context(
        FakePlaywright(),
        user_data_dir=str(tmp_path / "profile"),
    )

    assert context is expected_context
    assert calls[1]["executable_path"] == str(cached)


def test_embedded_login_patch_does_not_bypass_resilient_browser_launcher():
    with open(os.path.join(PROJECT_ROOT, "main.py"), encoding="utf-8") as source_file:
        source = source_file.read()
    patch_body = source.split("def patch_ranobelib_login_worker():", 1)[1].split(
        "def iter_ranobelib_source_locations():",
        1,
    )[0]

    assert patch_body.count("workers_module._launch_persistent_chromium_context(") == 2
    assert "playwright.chromium.launch_persistent_context(" not in patch_body


def test_rulate_edit_info_url_is_used_for_metadata_source():
    assert (
        _rulate_edit_info_url("https://tl.rulate.ru/book/204281")
        == "https://tl.rulate.ru/book/204281/edit/info"
    )
    assert (
        _rulate_edit_info_url("https://tl.rulate.ru/book/204281/edit/info")
        == "https://tl.rulate.ru/book/204281/edit/info"
    )
    assert (
        _rulate_public_book_url("https://tl.rulate.ru/book/204281/edit/info")
        == "https://tl.rulate.ru/book/204281"
    )
