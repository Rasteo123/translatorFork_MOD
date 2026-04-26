from __future__ import annotations

import zipfile

from docx import Document

from gemini_translator.utils.document_importer import (
    create_epub_from_import,
    extract_document_chapters,
)


def _read_epub_text(epub_path, member_name):
    with zipfile.ZipFile(epub_path, "r") as epub_zip:
        return epub_zip.read(member_name).decode("utf-8")


def test_docx_import_splits_headings_and_creates_epub(tmp_path):
    docx_path = tmp_path / "source.docx"
    document = Document()
    document.add_heading("Chapter & One", level=1)
    paragraph = document.add_paragraph()
    paragraph.add_run("Bold").bold = True
    paragraph.add_run(" text")
    document.add_heading("Chapter Two", level=1)
    document.add_paragraph("Second chapter text.")
    document.save(docx_path)

    result = extract_document_chapters(docx_path)

    assert result.source_format == "docx"
    assert [chapter.title for chapter in result.chapters] == ["Chapter & One", "Chapter Two"]
    assert "<strong>Bold</strong>" in result.chapters[0].html

    epub_path = create_epub_from_import(result, docx_path, tmp_path, title="A & B")

    assert epub_path.exists()
    opf = _read_epub_text(epub_path, "OEBPS/content.opf")
    first_chapter = _read_epub_text(epub_path, "OEBPS/chapter_0001.xhtml")
    assert "<dc:title>A &amp; B</dc:title>" in opf
    assert "Chapter &amp; One" in first_chapter


def test_markdown_import_uses_headings_as_chapters(tmp_path):
    md_path = tmp_path / "book.md"
    md_path.write_text("# One\n\nText **bold**.\n\n## Two\n\nMore text.", encoding="utf-8")

    result = extract_document_chapters(md_path)

    assert result.source_format == "markdown"
    assert [chapter.title for chapter in result.chapters] == ["One", "Two"]
    assert "<strong>bold</strong>" in result.chapters[0].html


def test_html_import_splits_h1_h2_sections(tmp_path):
    html_path = tmp_path / "book.html"
    html_path.write_text(
        "<html><head><title>Book</title></head><body>"
        "<h1>Alpha</h1><p>First.</p><h2>Beta</h2><p>Second.</p>"
        "</body></html>",
        encoding="utf-8",
    )

    result = extract_document_chapters(html_path)

    assert result.source_format == "html"
    assert [chapter.title for chapter in result.chapters] == ["Alpha", "Beta"]
    assert "First." in result.chapters[0].html


def test_txt_import_detects_chapter_markers(tmp_path):
    txt_path = tmp_path / "plain.txt"
    txt_path.write_text("Chapter 1\nHello.\n\nChapter 2\nWorld.", encoding="utf-8")

    result = extract_document_chapters(txt_path)

    assert result.source_format == "txt"
    assert [chapter.title for chapter in result.chapters] == ["Chapter 1", "Chapter 2"]
