import logging
import json
import os
import re
import sys
import tempfile
import time
import traceback
import unicodedata
import urllib.parse
import urllib.request
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from playwright.sync_api import sync_playwright
from PyQt6.QtCore import QThread, pyqtSignal

from constants import (
    BROWSER_ARGS,
    BROWSER_PROFILE_DIR,
    BROWSER_RULATE_DIR,
    MAX_RETRIES,
    RETRY_DELAY_SEC,
    RUS_MONTHS,
    SELECTORS,
)
from models import ChapterData
from parsers import FileParser
from utils import format_num, format_timedelta, parse_vol_and_chapter

_CJK_RE = re.compile(r"[\u3400-\u9fff\u3040-\u30ff\uac00-\ud7af]")
QIDIAN_RULATE_PROFILE_DIR = Path(
    os.environ.get(
        "QIDIAN_RULATE_PROFILE_DIR",
        str(Path.home() / ".qidian_rulate_creator" / "rulate_profile"),
    )
)
QIDIAN_RULATE_PROFILE_DIR.mkdir(parents=True, exist_ok=True)

RANOBELIB_GENRES = (
    "Арт", "Безумие", "Боевик", "Боевые искусства", "Вампиры", "Военное", "Гарем",
    "Гендерная интрига", "Героическое фэнтези", "Демоны", "Детектив", "Дзёсэй",
    "Драма", "Игра", "Исекай", "История", "Киберпанк", "Кодомо", "Комедия",
    "Космос", "Магия", "Махо-сёдзё", "Машины", "Меха", "Мистика", "Музыка",
    "Научная фантастика", "Омегаверс", "Пародия", "Повседневность", "Полиция",
    "Постапокалиптика", "Приключения", "Психология", "Романтика",
    "Самурайский боевик", "Сверхъестественное", "Сёдзё", "Сёдзё-ай", "Сёнэн",
    "Сёнэн-ай", "Спорт", "Супер сила", "Сэйнэн", "Трагедия", "Триллер",
    "Ужасы", "Фантастика", "Фэнтези", "Школа", "Эротика", "Этти", "Юри", "Яой",
)

RANOBELIB_TAGS = (
    "Авантюристы", "Антигерой", "Бессмертные", "Боги", "Борьба за власть",
    "Брат и сестра", "Ведьма", "Видеоигры", "Викторианская эпоха",
    "Виртуальная реальность", "Владыка демонов", "Военные",
    "Воспоминания из другого мира", "Выживание", "ГГ - Мэри Сью", "ГГ женщина",
    "ГГ имба", "ГГ мужчина", "ГГ не ояш", "ГГ не человек", "ГГ ояш",
    "Главный герой бог", "Глупый ГГ", "Горничные", "Градостроение", "Гуро",
    "Гяру", "Демоны", "Драконы", "Древний мир", "Запугивание", "Зверолюди",
    "Зомби", "Исторические", "Исторические фигуры", "Космос", "Кулинария",
    "Культивирование", "ЛитРПГ", "Лоли", "Магия", "Мастурбация",
    "Машинный перевод", "Медицина", "Межгалактическая война", "Монстродевушки",
    "Монстры", "Мрачный мир", "Мурим", "Нетораре", "Ниндзя", "Обратный Гарем",
    "Офисные Работники", "Пираты", "Подземелья", "Политика", "Полиция",
    "Полностью CGI", "Полноцветный", "Преступники / Криминал",
    "Призраки / Духи", "Призыватели", "Прыжки между мирами",
    "Путешествие в другой мир", "Путешествие во времени", "Рабы", "Ранги силы",
    "Регрессия", "Реинкарнация", "Самураи", "Сёдзё-ай", "Сёнен-ай", "Система",
    "Скрытие личности", "Современность", "Спортивное тело", "Средневековье",
    "Традиционные игры", "Умный ГГ", "Фермерство", "Фэнтези мир",
    "Характерный рост", "Хикикомори", "Шоу-бизнес", "Эволюция", "Элементы РПГ",
    "Эльфы", "Юри", "Якудза", "Яндере", "Яой",
)


def _playwright_browser_install_hint() -> str:
    python_executable = sys.executable or "python"
    return (
        "Playwright не нашел совместимый Chromium. "
        f"Установите браузер командой: \"{python_executable}\" -m playwright install chromium"
    )


def _is_browser_missing_error(error: Exception) -> bool:
    text = str(error).lower()
    return (
        "executable doesn't exist" in text
        or "playwright install" in text
        or ("browsertype.launch" in text and "executable" in text)
        or ("chromium distribution" in text and "not found" in text)
    )


def _candidate_browser_cache_roots() -> list[Path]:
    roots: list[Path] = []
    env_value = os.environ.get("PLAYWRIGHT_BROWSERS_PATH")
    if env_value:
        roots.append(Path(env_value))

    module_root = Path(__file__).resolve().parents[1]
    for base in (module_root, Path.cwd()):
        roots.append(Path(base) / "playwright_runtime" / "ms-playwright")

    localappdata = os.environ.get("LOCALAPPDATA")
    if localappdata:
        roots.append(Path(localappdata) / "ms-playwright")

    unique = []
    seen = set()
    for root in roots:
        try:
            resolved = root.resolve()
        except Exception:
            resolved = root
        key = str(resolved).lower()
        if key not in seen and resolved.exists() and resolved.is_dir():
            seen.add(key)
            unique.append(resolved)
    return unique


def _revision_from_path(path: Path) -> int:
    match = re.search(r"chromium-(\d+)", str(path))
    if not match:
        return -1
    return int(match.group(1))


def _find_cached_chromium_executable() -> Path | None:
    candidates: list[Path] = []
    for root in _candidate_browser_cache_roots():
        candidates.extend(root.glob("chromium-*/chrome-win*/chrome.exe"))
        candidates.extend(root.glob("chromium_headless_shell-*/chrome-headless-shell-win*/chrome-headless-shell.exe"))
    existing = [candidate for candidate in candidates if candidate.exists() and candidate.is_file()]
    if not existing:
        return None
    return max(existing, key=_revision_from_path)


def _launch_persistent_chromium_context(
    playwright,
    *,
    user_data_dir: str,
    viewport: dict | None = None,
    headless: bool = False,
    log_callback=None,
):
    kwargs = {
        "user_data_dir": user_data_dir,
        "headless": headless,
        "args": BROWSER_ARGS,
    }
    if viewport:
        kwargs["viewport"] = viewport
    try:
        return playwright.chromium.launch_persistent_context(**kwargs)
    except Exception as error:
        if not _is_browser_missing_error(error):
            raise
        if log_callback:
            log_callback("WARNING", "Playwright Chromium не найден, пробую fallback-браузер.")

    cached_executable = _find_cached_chromium_executable()
    if cached_executable:
        try:
            if log_callback:
                log_callback("INFO", f"Playwright: запускаю Chromium из {cached_executable}.")
            return playwright.chromium.launch_persistent_context(
                **kwargs,
                executable_path=str(cached_executable),
            )
        except Exception as error:
            if log_callback:
                log_callback("WARNING", f"Кэшированный Chromium не запустился: {error}")

    for channel in ("chrome", "msedge"):
        try:
            if log_callback:
                log_callback("INFO", f"Playwright: пробую системный браузер {channel}.")
            return playwright.chromium.launch_persistent_context(**kwargs, channel=channel)
        except Exception as error:
            if log_callback:
                log_callback("WARNING", f"Системный браузер {channel} не запустился: {error}")

    raise RuntimeError(_playwright_browser_install_hint())


def _has_saved_ranobelib_auth(profile_dir) -> tuple[bool, str | None]:
    try:
        with sync_playwright() as p:
            context = _launch_persistent_chromium_context(
                p,
                user_data_dir=str(profile_dir),
                headless=True,
                viewport={"width": 1280, "height": 900},
            )
            try:
                page = context.pages[0] if context.pages else context.new_page()
                page.goto("https://ranobelib.me", wait_until="domcontentloaded", timeout=30000)
                page.wait_for_timeout(1200)
                auth_detected = page.evaluate(
                    """() => {
                        try {
                            const raw = localStorage.getItem("auth");
                            if (!raw) {
                                return false;
                            }
                            const parsed = JSON.parse(raw);
                            return !!(parsed && parsed.token && parsed.token.access_token);
                        } catch (error) {
                            return false;
                        }
                    }"""
                )
                return bool(auth_detected), None
            finally:
                context.close()
    except Exception as error:
        return False, str(error)


def _collapse_rulate_spaces(value: str | None) -> str:
    return re.sub(r"[ \t\r\f\v]+", " ", str(value or "")).strip()


def _clean_rulate_media_title(value: str | None) -> str:
    text = _collapse_rulate_spaces(value)
    text = re.sub(
        r"\s*(?:[|/]|[-–—])\s*(?:tl\.)?rulate(?:\.ru)?\b.*$",
        "",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\s*/\s*читать\s+онлайн.*$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^\s*(?:книга|ранобэ|новелла)\s+", "", text, flags=re.IGNORECASE)
    return text.strip(" -–—|/\t\r\n")


def _clean_rulate_description(value: str | None) -> str:
    text = str(value or "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t\f\v]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(
        r"^(?:описание|аннотация|синопсис)\s*:?\s*",
        "",
        text.strip(),
        flags=re.IGNORECASE,
    )
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


_BAD_RULATE_COVER_RE = re.compile(
    r"(?:/logo/|rulate-24|favicon|sprite|avatar|userpic|uploads/users|/users/|placeholder|blank|/icon|icon-)",
    re.IGNORECASE,
)
_BAD_RULATE_TITLE_RE = re.compile(
    r"(?:^a\s*:|---|продолжается.*заверш|заверш.*брошен|статус|выпуск оригинала)",
    re.IGNORECASE,
)


try:
    from pypinyin import lazy_pinyin as _pypinyin_lazy_pinyin
except Exception:
    _pypinyin_lazy_pinyin = None


def _is_bad_rulate_cover_url(value: str | None) -> bool:
    if not value:
        return False
    text = urllib.parse.unquote(str(value)).replace("\\", "/")
    return bool(_BAD_RULATE_COVER_RE.search(text))


def _normalize_rulate_cover_url(value: str | None, source_url: str) -> str:
    raw = _collapse_rulate_spaces(value)
    if not raw:
        return ""
    url = urllib.parse.urljoin(source_url, raw)
    if _is_bad_rulate_cover_url(url):
        return ""
    return url


def _is_bad_rulate_title_candidate(value: str | None) -> bool:
    text = _collapse_rulate_spaces(value)
    if not text:
        return True
    if len(text) > 140:
        return True
    return bool(_BAD_RULATE_TITLE_RE.search(text))


def _clean_rulate_title_candidate(value: str | None) -> str:
    cleaned = _clean_rulate_media_title(value)
    return "" if _is_bad_rulate_title_candidate(cleaned) else cleaned


def _as_clean_title_list(value) -> list[str]:
    result = []
    seen = set()
    for item in _as_clean_list(value):
        text = _clean_rulate_title_candidate(item)
        if not text:
            continue
        key = text.casefold()
        if key in seen:
            continue
        seen.add(key)
        result.append(text)
    return result


def _first_cjk_title(*values) -> str:
    for value in values:
        for text in _as_clean_title_list(value):
            for part in re.split(r"[/;\n]+", text):
                cleaned = _clean_rulate_title_candidate(part)
                if cleaned and _CJK_RE.search(cleaned):
                    return cleaned
    return ""


def _romanize_cjk_title(value: str | None) -> str:
    text = _clean_rulate_title_candidate(value)
    if not text:
        return ""
    if not _CJK_RE.search(text):
        return text
    if not _pypinyin_lazy_pinyin:
        return text
    parts = [part.strip() for part in _pypinyin_lazy_pinyin(text) if part and part.strip()]
    return " ".join(part.capitalize() for part in parts)


def _rulate_book_id_from_url(url: str | None) -> str:
    match = re.search(r"tl\.rulate\.ru/book/(\d+)", str(url or ""), re.IGNORECASE)
    return match.group(1) if match else ""


def _rulate_public_book_url(url: str | None) -> str:
    book_id = _rulate_book_id_from_url(url)
    return f"https://tl.rulate.ru/book/{book_id}" if book_id else str(url or "").strip()


def _rulate_edit_info_url(url: str | None) -> str:
    book_id = _rulate_book_id_from_url(url)
    return f"https://tl.rulate.ru/book/{book_id}/edit/info" if book_id else str(url or "").strip()


def _as_clean_list(value) -> list[str]:
    if value is None:
        return []
    if not isinstance(value, (list, tuple)):
        value = [value]

    result = []
    seen = set()
    for item in value:
        text = _collapse_rulate_spaces(item)
        if not text:
            continue
        key = text.casefold()
        if key in seen:
            continue
        seen.add(key)
        result.append(text)
    return result


def _format_ranobelib_alt_names(*values) -> str:
    result = []
    seen = set()
    for raw in values:
        for value in _as_clean_list(raw):
            for part in re.split(r"[/;\n]+", value):
                text = _clean_rulate_title_candidate(part)
                if not text:
                    continue
                key = text.casefold()
                if key in seen:
                    continue
                seen.add(key)
                result.append(text)
    return " / ".join(result)


def _load_rulate_allowed_genres() -> list[str]:
    try:
        from qidian_rulate.workers import RULATE_GENRES as allowed

        return list(allowed)
    except Exception:
        return []


def _load_rulate_allowed_tags() -> list[str]:
    try:
        from qidian_rulate.workers import load_rulate_tags

        return list(load_rulate_tags())
    except Exception:
        return []


def _compact_catalog_text(value: str) -> str:
    return re.sub(r"[\s,;，、/|]+", "", value or "").casefold()


def _split_known_catalog_blob(value: str, allowed: list[str]) -> list[str]:
    text = _collapse_rulate_spaces(value)
    if not text:
        return []

    parts = [part.strip() for part in re.split(r"[,;，、/|\n]+", text) if part.strip()]
    if len(parts) > 1:
        result = []
        for part in parts:
            result.extend(_split_known_catalog_blob(part, allowed))
        return result

    if not allowed:
        return [text]

    allowed_pairs = sorted(
        [(item, _compact_catalog_text(item)) for item in allowed if _compact_catalog_text(item)],
        key=lambda pair: len(pair[1]),
        reverse=True,
    )
    compact = _compact_catalog_text(text)
    if not compact:
        return []

    result = []
    pos = 0
    while pos < len(compact):
        match = None
        for item, compact_item in allowed_pairs:
            if compact.startswith(compact_item, pos):
                match = (item, compact_item)
                break
        if not match:
            if result:
                return result
            return [text]
        result.append(match[0])
        pos += len(match[1])
    return result or [text]


def _normalize_rulate_catalog_items(value, allowed: list[str]) -> list[str]:
    allowed_by_casefold = {item.casefold(): item for item in allowed}
    result = []
    seen = set()
    for item in _as_clean_list(value):
        for part in _split_known_catalog_blob(item, allowed):
            text = _collapse_rulate_spaces(part)
            if not text:
                continue
            canonical = allowed_by_casefold.get(text.casefold(), text)
            key = canonical.casefold()
            if key in seen:
                continue
            seen.add(key)
            result.append(canonical)
    return result


def _google_translate_or_empty(value: str | None, target_lang: str, source_lang: str = "auto", timeout: int = 20) -> str:
    text = _collapse_rulate_spaces(value)
    if not text:
        return ""
    endpoint = (
        "https://translate.googleapis.com/translate_a/single"
        f"?client=gtx&sl={urllib.parse.quote(source_lang)}"
        f"&tl={urllib.parse.quote(target_lang)}&dt=t&q={urllib.parse.quote(text)}"
    )
    try:
        with urllib.request.urlopen(endpoint, timeout=timeout) as response:
            payload = json.loads(response.read().decode("utf-8"))
        parts = payload[0] if payload and isinstance(payload[0], list) else []
        translated = "".join(str(part[0]) for part in parts if part and part[0])
        return _collapse_rulate_spaces(translated)
    except Exception:
        return ""


def _strip_transliteration_marks(value: str | None) -> str:
    text = unicodedata.normalize("NFKD", str(value or ""))
    text = "".join(char for char in text if not unicodedata.combining(char))
    return _collapse_rulate_spaces(text)


def _title_case_author_name(value: str | None) -> str:
    text = _collapse_rulate_spaces(value)
    if not text:
        return ""
    return " ".join(part[:1].upper() + part[1:] for part in text.split(" ") if part)


def _capitalize_first(value: str | None) -> str:
    text = _collapse_rulate_spaces(value)
    return text[:1].upper() + text[1:] if text else ""


def _google_romanize_or_empty(value: str | None, source_lang: str = "auto", timeout: int = 20) -> str:
    text = _collapse_rulate_spaces(value)
    if not text:
        return ""
    endpoint = (
        "https://translate.googleapis.com/translate_a/single"
        f"?client=gtx&sl={urllib.parse.quote(source_lang)}"
        f"&tl=en&dt=t&dt=rm&q={urllib.parse.quote(text)}"
    )
    try:
        with urllib.request.urlopen(endpoint, timeout=timeout) as response:
            payload = json.loads(response.read().decode("utf-8"))
        parts = payload[0] if payload and isinstance(payload[0], list) else []
        for part in parts:
            if isinstance(part, list) and len(part) > 3 and part[3]:
                return _title_case_author_name(_strip_transliteration_marks(part[3]))
    except Exception:
        return ""
    return ""


def _prepare_ranobelib_author_payload(author: str | None) -> dict:
    original = _collapse_rulate_spaces(author)
    if not original:
        return {"name_en": "", "name_ru": "", "aliases": "", "original": ""}

    romanized = _google_romanize_or_empty(original) or _romanize_cjk_title(original)
    translated_en = _google_translate_or_empty(original, "en")
    translated_ru = _google_translate_or_empty(original, "ru")

    name_en = romanized if _CJK_RE.search(original) and romanized else translated_en or original
    name_ru = translated_ru or name_en or original
    name_en = _title_case_author_name(name_en)
    name_ru = _capitalize_first(name_ru)
    aliases = _format_ranobelib_alt_names(original, romanized, translated_en)
    return {
        "name_en": name_en,
        "name_ru": name_ru,
        "aliases": aliases,
        "original": original,
    }


def _extract_release_year(value: str | None) -> str:
    match = re.search(r"(?:19|20)\d{2}", str(value or ""))
    return match.group(0) if match else ""


def _ranobelib_title_status_value(value: str | None) -> str:
    text = str(value or "").casefold()
    if re.search(r"заверш|оконч|complete|finished|完结", text):
        return "2"
    if re.search(r"анонс|announce", text):
        return "3"
    if re.search(r"приостанов|заморож|pause|hiatus|暂停", text):
        return "4"
    if re.search(r"прекращ|drop|dropped", text):
        return "5"
    return "1"


def _first_non_cjk_title(*values) -> str:
    for value in values:
        for text in _as_clean_title_list(value):
            for part in re.split(r"[/;\n]+", text):
                cleaned = _clean_rulate_title_candidate(part)
                if cleaned and not _CJK_RE.search(cleaned):
                    return cleaned
    return ""


def _normalize_rulate_media_payload(raw: dict | None, source_url: str) -> dict:
    raw = raw or {}
    title_ru = _clean_rulate_media_title(raw.get("title") or raw.get("title_ru"))
    alt_name_items = _as_clean_title_list(raw.get("alt_names"))
    alt_hieroglyph_title = _first_cjk_title(raw.get("original_title"), alt_name_items)
    original_title = (
        _romanize_cjk_title(alt_hieroglyph_title)
        or _first_non_cjk_title(raw.get("original_title"), alt_name_items)
        or title_ru
    )
    title_en = _clean_rulate_title_candidate(raw.get("title_en")) or _first_non_cjk_title(
        raw.get("english_title"),
        alt_name_items,
        title_ru,
    )
    alt_names = _format_ranobelib_alt_names(alt_name_items, alt_hieroglyph_title)
    status_text = " ".join(
        _as_clean_list(
            [
                raw.get("status"),
                raw.get("status_text"),
                raw.get("release_status"),
            ]
        )
    )
    year = _extract_release_year(raw.get("year") or status_text)
    original_source_url = _normalize_rulate_cover_url(raw.get("original_source_url"), source_url)

    genres = _normalize_rulate_catalog_items(raw.get("genres"), _load_rulate_allowed_genres())
    tags = _normalize_rulate_catalog_items(raw.get("tags"), _load_rulate_allowed_tags())
    return {
        "source_url": original_source_url or _rulate_public_book_url(source_url),
        "rulate_url": _rulate_public_book_url(source_url),
        "rulate_edit_url": _rulate_edit_info_url(source_url),
        "title_ru": title_ru or "Без названия",
        "original_title": original_title or title_ru or "Без названия",
        "title_en": title_en or title_ru or "Untitled",
        "alt_names": alt_names,
        "alt_hieroglyph_title": alt_hieroglyph_title,
        "description": _clean_rulate_description(raw.get("description")),
        "cover_url": _normalize_rulate_cover_url(raw.get("cover_url"), source_url),
        "author": _collapse_rulate_spaces(raw.get("author")),
        "genres": genres,
        "tags": tags,
        "rulate_genres": genres,
        "rulate_tags": tags,
        "year": year,
        "status_value": _ranobelib_title_status_value(status_text),
    }


def _normalize_allowed_catalog_items(value, allowed: tuple[str, ...], limit: int) -> list[str]:
    allowed_by_casefold = {item.casefold(): item for item in allowed}
    result = []
    seen = set()
    for item in _as_clean_list(value):
        for part in re.split(r"[,;\n]+", item):
            for piece in _split_known_catalog_blob(part, list(allowed)):
                text = _collapse_rulate_spaces(piece)
                canonical = allowed_by_casefold.get(text.casefold())
                if not canonical:
                    continue
                key = canonical.casefold()
                if key in seen:
                    continue
                seen.add(key)
                result.append(canonical)
                if len(result) >= limit:
                    break
            if len(result) >= limit:
                break
        if len(result) >= limit:
            break
    return result


def _extract_json_object(raw_response: str) -> dict:
    text = str(raw_response or "").strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)
    try:
        parsed = json.loads(text)
        return parsed if isinstance(parsed, dict) else {}
    except json.JSONDecodeError:
        pass
    match = re.search(r"\{.*\}", text, flags=re.DOTALL)
    if not match:
        return {}
    parsed = json.loads(match.group(0))
    return parsed if isinstance(parsed, dict) else {}


def _build_ranobelib_catalog_prompt(metadata: dict) -> str:
    genres = ", ".join(RANOBELIB_GENRES)
    tags = ", ".join(RANOBELIB_TAGS)
    rulate_genres = ", ".join(metadata.get("rulate_genres") or metadata.get("genres") or []) or "не указаны"
    rulate_tags = ", ".join(metadata.get("rulate_tags") or metadata.get("tags") or []) or "не указаны"
    return f"""Ты подбираешь метаданные для RanobeLib строго из разрешённых списков.

Верни только JSON без markdown:
{{
  "genres": ["..."],
  "tags": ["..."],
  "age_rating": "6+|12+|16+|18+",
  "title_status": "ongoing|completed|announced|paused|dropped",
  "translation_status": "ongoing|completed|frozen|abandoned",
  "release_year": "YYYY или пусто"
}}

Правила:
- genres: 3-5 пунктов строго из списка RanobeLib Genres.
- tags: 3-8 пунктов строго из списка RanobeLib Tags.
- Нельзя придумывать новые жанры/теги.
- Учитывай русское название, описание, жанры и теги Rulate.
- Возрастное ограничение подбирай осторожно: 18+ только при явной эротике/жести/гуро.

Название: {metadata.get("title_ru") or ""}
Автор: {metadata.get("author") or ""}
Год/статус Rulate: {metadata.get("year") or ""} {metadata.get("status_value") or ""}
Жанры Rulate: {rulate_genres}
Теги Rulate: {rulate_tags}
Описание:
{metadata.get("description") or ""}

RanobeLib Genres:
{genres}

RanobeLib Tags:
{tags}
"""


def _parse_ranobelib_catalog_response(raw_response: str) -> dict:
    payload = _extract_json_object(raw_response)
    age = str(payload.get("age_rating") or "").strip()
    if age not in {"6+", "12+", "16+", "18+"}:
        age = "16+"
    title_status_map = {
        "ongoing": "1",
        "completed": "2",
        "announced": "3",
        "paused": "4",
        "dropped": "5",
    }
    translation_status_map = {
        "ongoing": "1",
        "completed": "2",
        "frozen": "3",
        "abandoned": "4",
    }
    year = _extract_release_year(payload.get("release_year"))
    return {
        "genres": _normalize_allowed_catalog_items(payload.get("genres"), RANOBELIB_GENRES, 5),
        "tags": _normalize_allowed_catalog_items(payload.get("tags"), RANOBELIB_TAGS, 8),
        "age_rating": age,
        "age_value": {"6+": "1", "12+": "2", "16+": "3", "18+": "4"}[age],
        "status_value": title_status_map.get(str(payload.get("title_status") or "").strip(), "1"),
        "translation_status_value": translation_status_map.get(
            str(payload.get("translation_status") or "").strip(),
            "1",
        ),
        "year": year,
    }


_RULATE_MEDIA_EXTRACT_SCRIPT = r"""() => {
    const clean = (value) => (value || "").replace(/\s+/g, " ").trim();
    const cleanBlock = (value) => (value || "")
        .replace(/\r\n?/g, "\n")
        .replace(/[ \t\f\v]+/g, " ")
        .replace(/ *\n */g, "\n")
        .replace(/\n{3,}/g, "\n\n")
        .trim();
    const meta = (selector) => clean(document.querySelector(selector)?.getAttribute("content"));
    const abs = (url) => {
        if (!url) return "";
        try { return new URL(url, location.href).href; } catch (_error) { return ""; }
    };
    const badCoverCandidate = (url, attrs = "") => {
        const text = `${url || ""} ${attrs || ""}`.toLowerCase();
        return !url
            || text.includes("/logo/")
            || text.includes("rulate-24")
            || text.includes("favicon")
            || text.includes("sprite")
            || text.includes("avatar")
            || text.includes("userpic")
            || text.includes("placeholder")
            || text.includes("blank")
            || text.includes("/icon")
            || text.includes("icon-");
    };
    const addCoverCandidate = (candidates, url, attrs = "", width = 0, height = 0, allowPortrait = false) => {
        const href = abs(url);
        if (badCoverCandidate(href, attrs)) return;
        const portraitEnough = height >= 120 && width >= 70 && height >= width;
        const namedCover = /(cover|book|облож|poster|image)/i.test(attrs || "");
        if ((namedCover || (allowPortrait && portraitEnough)) && !candidates.includes(href)) {
            candidates.push(href);
        }
    };
    const fieldValue = (selector) => {
        const field = document.querySelector(selector);
        if (!field) return "";
        if (field.tagName === "SELECT") {
            const selected = field.options[field.selectedIndex];
            return clean(selected?.textContent || field.value || "");
        }
        return cleanBlock(field.value || field.getAttribute("value") || field.textContent || "");
    };
    const ckeditorValue = (instanceName, fallbackSelector) => {
        try {
            if (window.CKEDITOR && CKEDITOR.instances && CKEDITOR.instances[instanceName]) {
                const tmp = document.createElement("div");
                tmp.innerHTML = CKEDITOR.instances[instanceName].getData() || "";
                return cleanBlock(tmp.innerText || tmp.textContent || "");
            }
        } catch (_error) {}
        const value = fieldValue(fallbackSelector);
        if (!value) return "";
        const tmp = document.createElement("div");
        tmp.innerHTML = value;
        return cleanBlock(tmp.innerText || tmp.textContent || value);
    };
    const selectedMagicValues = (containerSelector) => {
        const root = document.querySelector(containerSelector);
        if (!root) return [];
        const selectors = [
            ".ms-sel-item",
            ".ms-sel-ctn .ms-sel-item",
            ".magic-suggest-selected",
            ".search-choice",
            ".search-choice span",
            ".select2-selection__choice",
            ".select2-selection__choice__display",
            ".tag",
            ".label",
            "li",
        ];
        const values = [];
        const nodeText = (node) => {
            const clone = node.cloneNode(true);
            for (const trash of Array.from(clone.querySelectorAll("button, input, svg, .ms-close-btn, .search-choice-close, .select2-selection__choice__remove"))) {
                trash.remove();
            }
            return clean(
                node.getAttribute("title")
                || node.getAttribute("data-name")
                || node.getAttribute("data-title")
                || clone.textContent
                || ""
            ).replace(/^×\s*/, "").replace(/\s*×$/, "").trim();
        };
        for (const selector of selectors) {
            for (const node of Array.from(root.querySelectorAll(selector))) {
                let text = nodeText(node);
                if (text && !values.includes(text)) values.push(text);
            }
            if (values.length) break;
        }
        if (!values.length) {
            const text = clean(root.innerText || root.textContent || "");
            const lower = text.toLowerCase();
            if (text && !lower.includes("type or click here") && !lower.includes("click here")) values.push(text);
        }
        return values;
    };
    const editInfoData = () => {
        const hasEditForm = !!(
            document.querySelector("#Book_t_title")
            || document.querySelector("#Book_s_title")
            || document.querySelector("#Book_descr")
            || document.querySelector('select[name="Book[status]"]')
        );
        if (!hasEditForm) return {};
        const coverCandidates = [];
        for (const selector of [
            "#Book_image",
            "#Book_img",
            "#Book_picture",
            "#Book_cover",
            'input[type="hidden"][name*="image"]',
            'input[type="hidden"][name*="cover"]',
        ]) {
            const value = fieldValue(selector);
            if (value) addCoverCandidate(coverCandidates, value, selector, 120, 90);
        }
        for (const img of Array.from(document.images)) {
            const src = img.currentSrc || img.src || img.getAttribute("data-src") || img.getAttribute("data-original") || "";
            const attrs = clean(`${img.className || ""} ${img.id || ""} ${img.alt || ""} ${img.title || ""} ${src || ""}`).toLowerCase();
            const width = img.naturalWidth || img.width || 0;
            const height = img.naturalHeight || img.height || 0;
            const context = clean(img.closest(".form-group, .control-group, .controls, .book-cover, .cover, [id*='cover'], [class*='cover']")?.innerText || "").toLowerCase();
            if (src && (attrs.includes("cover") || attrs.includes("облож") || context.includes("облож"))) {
                addCoverCandidate(coverCandidates, src, attrs, width, height);
            }
        }
        for (const node of Array.from(document.querySelectorAll("*"))) {
            const style = window.getComputedStyle(node);
            const match = (style.backgroundImage || "").match(/url\(["']?([^"')]+)["']?\)/);
            if (!match) continue;
            const rect = node.getBoundingClientRect();
            const attrs = clean(`${node.className || ""} ${node.id || ""} ${node.getAttribute("title") || ""}`).toLowerCase();
            if (attrs.includes("cover") || attrs.includes("облож")) {
                addCoverCandidate(coverCandidates, match[1], attrs, rect.width, rect.height);
            }
        }
        return {
            title: fieldValue("#Book_t_title") || fieldValue('input[name="Book[t_title]"]'),
            title_en: fieldValue("#Book_s_title") || fieldValue('input[name="Book[s_title]"]'),
            original_title: fieldValue("#Book_a_title_1") || fieldValue('input[name="Book[a_title][]"]'),
            alt_names: [
                fieldValue("#Book_a_title_1"),
                fieldValue("#Book_a_title_2"),
                fieldValue("#Book_a_title_3"),
            ].filter(Boolean),
            description: ckeditorValue("Book_descr", "#Book_descr"),
            cover_url: coverCandidates.find(Boolean) || "",
            author: fieldValue("#Book_author") || fieldValue('input[name="Book[author]"]'),
            original_source_url: fieldValue("#Book_source_url") || fieldValue('input[name="Book[source_url]"]'),
            genres: selectedMagicValues("#Book_genres"),
            tags: selectedMagicValues("#Book_tags"),
            status: fieldValue('select[name="Book[status]"]') || fieldValue("#Book_status"),
            year: fieldValue("#Book_year") || fieldValue('input[name="Book[year]"]'),
        };
    };
    const editPayload = editInfoData();
    const visibleText = (el) => {
        if (!el) return "";
        const style = window.getComputedStyle(el);
        if (style && (style.display === "none" || style.visibility === "hidden")) return "";
        return cleanBlock(el.innerText || el.textContent || "");
    };
    const bestText = (selectors) => {
        for (const selector of selectors) {
            const nodes = Array.from(document.querySelectorAll(selector));
            for (const node of nodes) {
                const text = visibleText(node);
                if (text && text.length >= 20) return text;
            }
        }
        return "";
    };
    const textNearLabel = (labels) => {
        const normalizedLabels = labels.map((label) => label.toLowerCase());
        const nodes = Array.from(document.querySelectorAll("li, tr, dl, p, div, span"));
        for (const node of nodes) {
            const text = clean(node.innerText || node.textContent || "");
            if (!text || text.length > 220) continue;
            const lower = text.toLowerCase();
            for (const label of normalizedLabels) {
                if (!lower.includes(label)) continue;
                let value = text.replace(new RegExp(`^.*?${label}\\s*[:：-]?\\s*`, "i"), "").trim();
                value = value.replace(/^[:：-]\s*/, "").trim();
                if (value && value !== text && value.length <= 120) return value;
            }
        }
        return "";
    };
    const listNearLabel = (labels) => {
        const value = textNearLabel(labels);
        if (!value) return [];
        return value.split(/[,，;、/]+/).map(clean).filter(Boolean);
    };

    const h1Candidates = Array.from(document.querySelectorAll("h1, .book-title, .book_name, .name"))
        .map(visibleText)
        .filter(Boolean)
        .sort((a, b) => b.length - a.length);
    const title = h1Candidates[0] || meta('meta[property="og:title"]') || meta('meta[name="twitter:title"]') || document.title;

    const coverCandidates = [];
    addCoverCandidate(coverCandidates, meta('meta[property="og:image"]'), "cover book", 120, 90);
    addCoverCandidate(coverCandidates, meta('meta[name="twitter:image"]'), "cover book", 120, 90);
    for (const img of Array.from(document.images)) {
        const src = img.currentSrc || img.src || img.getAttribute("data-src") || img.getAttribute("data-original") || "";
        if (!src) continue;
        const attrs = clean(`${img.className || ""} ${img.id || ""} ${img.alt || ""} ${img.title || ""} ${src || ""}`).toLowerCase();
        const width = img.naturalWidth || img.width || 0;
        const height = img.naturalHeight || img.height || 0;
        if (attrs.includes("cover") || attrs.includes("book") || (height >= 140 && width >= 90 && height >= width)) {
            addCoverCandidate(coverCandidates, src, attrs, width, height, true);
        }
    }
    for (const node of Array.from(document.querySelectorAll("*"))) {
        const style = window.getComputedStyle(node);
        const match = (style.backgroundImage || "").match(/url\(["']?([^"')]+)["']?\)/);
        if (!match) continue;
        const rect = node.getBoundingClientRect();
        const attrs = clean(`${node.className || ""} ${node.id || ""} ${node.getAttribute("title") || ""}`).toLowerCase();
        if (attrs.includes("cover") || attrs.includes("book") || attrs.includes("облож") || (rect.height >= 140 && rect.width >= 90 && rect.height >= rect.width)) {
            addCoverCandidate(coverCandidates, match[1], attrs, rect.width, rect.height, true);
        }
    }

    let description = bestText([
        "[itemprop='description']",
        ".book-description",
        ".book_desc",
        ".description",
        ".desc",
        ".summary",
        ".annotation",
        ".book-intro",
        "#book-description",
        "#description",
    ]);
    if (!description) {
        description = meta('meta[name="description"]') || meta('meta[property="og:description"]');
    }
    if (!description) {
        const labels = Array.from(document.querySelectorAll("h2, h3, h4, dt, b, strong, span, div"))
            .filter((node) => /^(описание|аннотация|синопсис)$/i.test(clean(node.textContent || "")));
        for (const label of labels) {
            const candidates = [
                label.nextElementSibling,
                label.parentElement?.nextElementSibling,
                label.parentElement?.querySelector("p"),
            ];
            for (const candidate of candidates) {
                const text = visibleText(candidate);
                if (text && text.length >= 20) {
                    description = text;
                    break;
                }
            }
            if (description) break;
        }
    }

    const fallbackPayload = {
        title,
        original_title: textNearLabel(["оригинальное название", "оригинал", "название оригинала"]),
        english_title: textNearLabel(["английское название", "english"]),
        alt_names: listNearLabel(["альтернативные названия", "другие названия", "альт. названия"]),
        description,
        cover_url: coverCandidates.find(Boolean) || "",
        author: textNearLabel(["автор", "author"]),
        genres: listNearLabel(["жанры", "жанр"]),
        tags: listNearLabel(["теги", "тег"]),
        status: textNearLabel(["статус", "выпуск"]),
        year: textNearLabel(["год релиза", "год выпуска", "выпуск"]),
    };
    return Object.assign({}, fallbackPayload, Object.fromEntries(
        Object.entries(editPayload).filter(([_key, value]) => {
            if (Array.isArray(value)) return value.length > 0;
            return !!value;
        })
    ));
}"""


def _merge_public_rulate_cover(page, raw_payload: dict, public_url: str, log_callback=None) -> dict:
    payload = dict(raw_payload or {})
    try:
        page.goto(public_url, wait_until="domcontentloaded", timeout=60000)
        page.wait_for_timeout(1500)
        public_payload = page.evaluate(_RULATE_MEDIA_EXTRACT_SCRIPT) or {}
        public_cover = _normalize_rulate_cover_url(public_payload.get("cover_url"), public_url)
        if public_cover:
            payload["cover_url"] = public_cover
            if log_callback:
                log_callback("INFO", "Rulate: обложка уточнена с публичной страницы книги.")
    except Exception as error:
        if log_callback:
            log_callback("WARNING", f"Rulate: не удалось уточнить обложку с публичной страницы: {error}")
    return payload


_RANOBELIB_MEDIA_FORM_FILL_SCRIPT = r"""(data) => {
    const norm = (value) => (value || "").replace(/\s+/g, " ").trim();
    const groups = () => Array.from(document.querySelectorAll(".form-group"));
    const findGroup = (label) => {
        const needle = label.toLowerCase();
        return groups().find((group) => {
            const text = norm(group.querySelector(".form-label")?.innerText || "");
            return text.toLowerCase().includes(needle);
        });
    };
    const setNativeValue = (el, value) => {
        const prototype = Object.getPrototypeOf(el);
        const descriptor = Object.getOwnPropertyDescriptor(prototype, "value")
            || Object.getOwnPropertyDescriptor(HTMLInputElement.prototype, "value")
            || Object.getOwnPropertyDescriptor(HTMLTextAreaElement.prototype, "value");
        if (descriptor && descriptor.set) descriptor.set.call(el, value);
        else el.value = value;
        el.dispatchEvent(new Event("input", { bubbles: true }));
        el.dispatchEvent(new Event("change", { bubbles: true }));
        el.dispatchEvent(new Event("blur", { bubbles: true }));
    };
    const fillText = (label, value, selector = "input, textarea") => {
        if (!value) return false;
        const group = findGroup(label);
        const field = group?.querySelector(selector);
        if (!field || field.disabled) return false;
        setNativeValue(field, value);
        return true;
    };
    const setSelect = (label, value) => {
        const group = findGroup(label);
        const field = group?.querySelector("select");
        if (!field || field.disabled) return false;
        setNativeValue(field, value);
        return true;
    };

    const result = {
        original: fillText("Оригинальное название", data.original_title || data.title_ru),
        titleRu: fillText("Название на русском", data.title_ru),
        titleEn: fillText("Название на английском", data.title_en || data.title_ru),
        altNames: data.alt_names ? fillText("Альтернативные названия", data.alt_names, "textarea, input") : true,
        type: setSelect("Тип", data.type_value || "12"),
        status: setSelect("Статус тайтла", data.status_value || "1"),
        age: setSelect("Возрастное ограничение", data.age_value || "3"),
        translationStatus: setSelect("Статус перевода", data.translation_status_value || "1"),
        chaptersUpload: setSelect("Загрузка глав", data.chapter_upload_value || "2"),
        year: data.year ? fillText("Год релиза", data.year) : false,
    };
    return result;
}"""

class RulateDownloadWorker(QThread):
    """
    Открывает Playwright-браузер (с persistent-профилем),
    переходит на страницу книги на rulate,
    выбирает нужные главы (чекбоксы), жмёт «Скачать .docx»,
    обрабатывает скачанный zip и возвращает список ChapterData.
    """
    log_signal = pyqtSignal(str, str)
    progress_signal = pyqtSignal(int)
    chapters_ready = pyqtSignal(list)       # список ChapterData
    chapter_list_ready = pyqtSignal(list)    # список dict для выбора: [{id, title, number}]
    finished_signal = pyqtSignal()

    def __init__(self, rulate_url: str, default_vol: str,
                 skip_after: int = 0, chapter_ids: list = None,
                 chapter_infos: list = None):
        """
        rulate_url: ссылка вида https://tl.rulate.ru/book/123870
        default_vol: том по умолчанию
        skip_after: пропустить главы с номером <= skip_after (0 = не пропускать)
        chapter_ids: если задан — скачиваем только эти главы (id из data-id).
                     Если None — сначала эмитим chapter_list_ready для выбора.
        """
        super().__init__()
        self.rulate_url = rulate_url.rstrip("/")
        self.default_vol = default_vol
        self.skip_after = skip_after
        self.chapter_ids = chapter_ids
        self.chapter_infos = chapter_infos or []
        self._chapter_info_by_id = {
            str(ch.get("id")): ch
            for ch in self.chapter_infos
            if ch and ch.get("id") is not None
        }
        self.is_running = True
        self._mode = "download" if chapter_ids else "list"

    def log(self, level: str, msg: str):
        self.log_signal.emit(level, msg)

    def stop(self):
        self.is_running = False

    def run(self):
        try:
            if self._mode == "list":
                self._fetch_chapter_list()
            else:
                self._download_chapters()
        except Exception as e:
            self.log("ERROR", f"Rulate: {e}")
            logging.error(traceback.format_exc())
        finally:
            self.finished_signal.emit()

    def _fetch_chapter_list(self):
        """Получить список глав со страницы книги."""
        self.log("INFO", "Rulate: открываю страницу книги…")
        try:
            with sync_playwright() as p:
                browser = _launch_persistent_chromium_context(
                    p,
                    user_data_dir=str(BROWSER_RULATE_DIR),
                    viewport={"width": 1280, "height": 900},
                    log_callback=self.log,
                )
                page = browser.pages[0]
                page.goto(self.rulate_url, timeout=60000)
                page.wait_for_timeout(3000)

                # Собираем информацию о главах
                chapters_info = page.evaluate("""() => {
                    const rows = document.querySelectorAll('tr.chapter_row');
                    const result = [];
                    for (const row of rows) {
                        const id = row.getAttribute('data-id');
                        const link = row.querySelector('td.t a');
                        const checkbox = row.querySelector('input.download_chapter');
                        if (!link) continue;
                        const titleCandidates = [
                            link.getAttribute('title'),
                            link.getAttribute('data-original-title'),
                            link.getAttribute('aria-label'),
                            row.getAttribute('data-title'),
                            link.textContent,
                        ].map(value => (value || '').trim()).filter(Boolean);
                        const title = titleCandidates.reduce(
                            (best, value) => value.length > best.length ? value : best,
                            ''
                        );
                        const hasCheckbox = !!checkbox;
                        // Пробуем извлечь номер из названия:
                        // 1) "Глава/Chapter/Ch N"
                        // 2) просто первое число в заголовке
                        // 3) "Часть/Part N" превращаем в дробную главу: 30 + Часть 2 -> 30.2
                        let match = title.match(/(?:Глава|Chapter|Ch|Гл)\\s*\\.?\\s*(\\d+(?:\\.\\d+)?)/i);
                        if (!match) {
                            match = title.match(/(\\d+(?:\\.\\d+)?)/);
                        }
                        let num = match ? parseFloat(match[1]) : 0;
                        const partMatch = title.match(/(?:Часть|Part)\\s*(\\d+)\\s*$/i);
                        if (num > 0 && partMatch) {
                            const base = String(num);
                            const part = partMatch[1];
                            num = parseFloat(base.includes(".") ? `${base}${part}` : `${base}.${part}`);
                        }
                        result.push({
                            id: id,
                            title: title,
                            number: num,
                            downloadable: hasCheckbox,
                        });
                    }
                    return result;
                }""")

                browser.close()

                if not chapters_info:
                    self.log("ERROR", "Rulate: главы не найдены на странице.")
                    return

                self.log("INFO", f"Rulate: найдено {len(chapters_info)} глав "
                         f"({sum(1 for c in chapters_info if c['downloadable'])} доступно для скачивания)")
                self.chapter_list_ready.emit(chapters_info)

        except Exception as e:
            self.log("ERROR", f"Rulate: ошибка при получении списка глав: {e}")
            logging.error(traceback.format_exc())

    @staticmethod
    def _chapter_has_part(title: str) -> bool:
        return bool(re.search(r"(?:Часть|Part)\s*\d+\s*$", title or "", re.IGNORECASE))

    @staticmethod
    def _chapter_base_number(number) -> int:
        try:
            value = float(number)
        except (TypeError, ValueError):
            return 0
        return int(value) if value > 0 else 0

    def _should_download_individually(self) -> bool:
        if not self.chapter_infos or len(self.chapter_ids or []) <= 1:
            return False

        seen_bases = set()
        for info in self.chapter_infos:
            title = info.get("title", "")
            number = info.get("number", 0)
            try:
                numeric = float(number)
            except (TypeError, ValueError):
                numeric = 0.0
            if self._chapter_has_part(title) or (numeric and numeric != int(numeric)):
                return True

            base = self._chapter_base_number(number)
            if base and base in seen_bases:
                return True
            if base:
                seen_bases.add(base)
        return False

    def _apply_chapter_info(self, chapter: ChapterData, info: dict | None):
        if not info:
            return

        title = info.get("title", "")
        fallback = int(float(info.get("number") or 1))
        vol, number, clean_title, num_found = parse_vol_and_chapter(
            title, self.default_vol, fallback
        )
        if info.get("number"):
            try:
                number = float(info["number"])
                num_found = True
            except (TypeError, ValueError):
                pass

        chapter.volume = vol
        chapter.number = number
        if clean_title:
            chapter.title = clean_title
        chapter._num_found = num_found

    def _apply_chapter_infos(self, chapters: list[ChapterData], infos: list[dict] | None):
        if not chapters or not infos:
            return

        unused_infos = list(infos)
        for chapter in chapters:
            matched = None
            for info in unused_infos:
                try:
                    info_number = float(info.get("number") or 0)
                except (TypeError, ValueError):
                    info_number = 0.0
                if info_number and abs(float(chapter.number) - info_number) < 0.0001:
                    matched = info
                    break

            if matched is None and len(chapters) == len(infos):
                matched = unused_infos[0]

            if matched:
                self._apply_chapter_info(chapter, matched)
                unused_infos.remove(matched)

    def _parse_downloaded_file(self, file_path: str, chapter_info: dict | None = None) -> list[ChapterData]:
        if file_path.lower().endswith(".zip"):
            chapters = FileParser.parse_zip_docx(file_path, self.default_vol, self.log)
        elif file_path.lower().endswith(".docx"):
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n".join(paragraphs)
            chapters = [ChapterData(self.default_vol, 1, "", content)]
        else:
            self.log("ERROR", f"Rulate: unknown downloaded file format: {file_path}")
            return []

        if chapter_info and len(chapters) == 1:
            self._apply_chapter_info(chapters[0], chapter_info)
        return chapters

    def _select_download_checkboxes(self, page, chapter_ids: list[str]) -> int:
        page.evaluate("""() => {
            document.querySelectorAll('input.download_chapter').forEach(cb => {
                cb.checked = false;
            });
        }""")

        return page.evaluate("""(ids) => {
            let count = 0;
            for (const id of ids) {
                const row = document.querySelector(`tr[data-id="${id}"]`);
                if (row) {
                    const cb = row.querySelector('input.download_chapter');
                    if (cb) {
                        cb.checked = true;
                        count++;
                    }
                }
            }
            return count;
        }""", [str(cid) for cid in chapter_ids])

    def _download_selected_file(self, page, chapter_ids: list[str], chapter_info: dict | None = None) -> list[ChapterData]:
        checked_count = self._select_download_checkboxes(page, chapter_ids)
        self.log("INFO", f"Rulate: отмечено {checked_count} из {len(chapter_ids)} глав")

        if checked_count == 0:
            self.log("ERROR", "Rulate: ни одна глава не была отмечена для скачивания.")
            return []

        self.log("INFO", "Rulate: запускаю скачивание .docx…")
        with page.expect_download(timeout=120000) as download_info:
            page.click('input[name="download_d"]')

        download = download_info.value
        self.log("INFO", f"Rulate: файл скачивается: {download.suggested_filename}")

        tmp_dir = tempfile.mkdtemp(prefix="rulate_")
        try:
            zip_path = os.path.join(tmp_dir, download.suggested_filename or "rulate_chapters.zip")
            download.save_as(zip_path)
            self.log("SUCCESS", f"Rulate: файл сохранён ({os.path.getsize(zip_path):,} байт)")
            self.log("INFO", "Rulate: разбираю скачанный архив…")
            return self._parse_downloaded_file(zip_path, chapter_info)
        finally:
            try:
                for name in os.listdir(tmp_dir):
                    os.remove(os.path.join(tmp_dir, name))
                os.rmdir(tmp_dir)
            except OSError:
                pass

    def _download_chapters(self):
        """Скачать выбранные главы через форму «Скачать .docx»."""
        if not self.chapter_ids:
            self.log("ERROR", "Rulate: не выбраны главы для скачивания.")
            return

        self.log("INFO", f"Rulate: скачиваю {len(self.chapter_ids)} глав…")

        try:
            with sync_playwright() as p:
                browser = _launch_persistent_chromium_context(
                    p,
                    user_data_dir=str(BROWSER_RULATE_DIR),
                    viewport={"width": 1280, "height": 900},
                    log_callback=self.log,
                )
                page = browser.pages[0]
                page.goto(self.rulate_url, timeout=60000)

                # Ждём загрузки таблицы глав
                page.wait_for_selector('tr.chapter_row', timeout=30000)
                page.wait_for_timeout(2000)

                self.log("INFO", "Rulate: выбираю главы…")
                self.progress_signal.emit(10)

                if self._should_download_individually():
                    total = len(self.chapter_ids)
                    chapters = []
                    self.log(
                        "INFO",
                        "Rulate: обнаружены главы с частями; скачиваю по одной, чтобы архив не потерял одноимённые части",
                    )
                    for index, chapter_id in enumerate(self.chapter_ids, start=1):
                        if not self.is_running:
                            break
                        chapter_info = self._chapter_info_by_id.get(str(chapter_id))
                        title = (chapter_info or {}).get("title") or str(chapter_id)
                        self.log("INFO", f"Rulate: [{index}/{total}] {title}")
                        try:
                            chapters.extend(
                                self._download_selected_file(
                                    page, [str(chapter_id)], chapter_info
                                )
                            )
                        except Exception as item_error:
                            self.log(
                                "ERROR",
                                f"Rulate: ошибка скачивания главы {title}: {item_error}",
                            )
                            logging.error(traceback.format_exc())
                        self.progress_signal.emit(10 + int(index * 90 / max(1, total)))

                    browser.close()
                    self.progress_signal.emit(100)

                    if chapters:
                        self.log("SUCCESS", f"Rulate: получено {len(chapters)} глав")
                        self.chapters_ready.emit(chapters)
                    else:
                        self.log("WARNING", "Rulate: из выбранных глав не удалось извлечь ни одной главы")
                    return

                # Снимаем все чекбоксы, затем ставим нужные
                page.evaluate("""() => {
                    document.querySelectorAll('input.download_chapter').forEach(cb => {
                        cb.checked = false;
                    });
                }""")

                ids_set = set(str(cid) for cid in self.chapter_ids)
                checked_count = page.evaluate("""(ids) => {
                    let count = 0;
                    for (const id of ids) {
                        const row = document.querySelector(`tr[data-id="${id}"]`);
                        if (row) {
                            const cb = row.querySelector('input.download_chapter');
                            if (cb) {
                                cb.checked = true;
                                count++;
                            }
                        }
                    }
                    return count;
                }""", list(ids_set))

                self.log("INFO", f"Rulate: отмечено {checked_count} из {len(ids_set)} глав")
                self.progress_signal.emit(30)

                if checked_count == 0:
                    self.log("ERROR", "Rulate: ни одна глава не была отмечена для скачивания.")
                    browser.close()
                    return

                # Ждём скачивания файла
                self.log("INFO", "Rulate: запускаю скачивание .docx…")

                with page.expect_download(timeout=120000) as download_info:
                    page.click('input[name="download_d"]')

                download = download_info.value
                self.progress_signal.emit(60)
                self.log("INFO", f"Rulate: файл скачивается: {download.suggested_filename}")

                # Сохраняем во временную папку
                tmp_dir = tempfile.mkdtemp(prefix="rulate_")
                zip_path = os.path.join(tmp_dir, download.suggested_filename or "rulate_chapters.zip")
                download.save_as(zip_path)

                self.progress_signal.emit(80)
                self.log("SUCCESS", f"Rulate: файл сохранён ({os.path.getsize(zip_path):,} байт)")

                browser.close()

                # Парсим скачанный zip
                self.log("INFO", "Rulate: разбираю скачанный архив…")

                if zip_path.lower().endswith(".zip"):
                    chapters = FileParser.parse_zip_docx(
                        zip_path, self.default_vol, self.log
                    )
                elif zip_path.lower().endswith(".docx"):
                    # Если скачан одним файлом (без zip)
                    doc = Document(zip_path)
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    content = "\n".join(paragraphs)
                    chapters = [ChapterData(self.default_vol, 1, "", content)]
                else:
                    self.log("ERROR", f"Rulate: неизвестный формат файла: {zip_path}")
                    return

                self._apply_chapter_infos(chapters, self.chapter_infos)
                self.progress_signal.emit(100)

                if chapters:
                    self.log("SUCCESS", f"Rulate: получено {len(chapters)} глав")
                    self.chapters_ready.emit(chapters)
                else:
                    self.log("WARNING", "Rulate: из архива не удалось извлечь ни одной главы")

                # Очистка
                try:
                    os.remove(zip_path)
                    os.rmdir(tmp_dir)
                except OSError:
                    pass

        except Exception as e:
            self.log("ERROR", f"Rulate: ошибка скачивания: {e}")
            logging.error(traceback.format_exc())


# ─── Рабочий поток: создание карточки RanobeLib из Rulate ───────────────────

class RulateToRanobeMetadataWorker(QThread):
    log_signal = pyqtSignal(str, str)
    metadata_ready = pyqtSignal(dict)
    finished_signal = pyqtSignal()

    def __init__(self, rulate_url: str):
        super().__init__()
        self.rulate_url = _rulate_public_book_url(rulate_url.strip().rstrip("/"))
        self.rulate_edit_url = _rulate_edit_info_url(rulate_url.strip().rstrip("/"))
        self._browser = None

    def log(self, level: str, msg: str):
        self.log_signal.emit(level, msg)

    def run(self):
        try:
            with sync_playwright() as p:
                self.log("INFO", "Rulate: открываю edit/info через профиль Qidian/Fanqie -> Rulate...")
                self.log("INFO", f"Rulate: страница данных: {self.rulate_edit_url}")
                self.log("INFO", f"Rulate: профиль куки: {QIDIAN_RULATE_PROFILE_DIR}")
                self._browser = _launch_persistent_chromium_context(
                    p,
                    user_data_dir=str(QIDIAN_RULATE_PROFILE_DIR),
                    viewport={"width": 1280, "height": 900},
                    log_callback=self.log,
                )
                page = self._browser.pages[0] if self._browser.pages else self._browser.new_page()
                page.goto(self.rulate_edit_url, wait_until="domcontentloaded", timeout=60000)
                page.wait_for_timeout(2500)
                raw_payload = page.evaluate(_RULATE_MEDIA_EXTRACT_SCRIPT)
                raw_payload = _merge_public_rulate_cover(page, raw_payload, self.rulate_url, self.log)
                metadata = _normalize_rulate_media_payload(raw_payload, self.rulate_edit_url)
                self.log("SUCCESS", f"Rulate: данные получены: {metadata['title_ru']}")
                self.metadata_ready.emit(metadata)
        except Exception as error:
            self.log("ERROR", f"Rulate: {error}")
            logging.error(traceback.format_exc())
        finally:
            if self._browser:
                try:
                    self._browser.close()
                except Exception:
                    pass
            self.finished_signal.emit()

    def stop(self):
        if self._browser:
            try:
                self._browser.close()
            except Exception:
                pass


class RanobeLibCatalogMatchWorker(QThread):
    log_signal = pyqtSignal(str, str)
    catalog_ready = pyqtSignal(dict)
    finished_signal = pyqtSignal()

    def __init__(
        self,
        metadata: dict,
        provider_id: str,
        model_settings: dict,
        active_keys: list[str],
        settings_manager,
    ):
        super().__init__()
        self.metadata = metadata or {}
        self.provider_id = provider_id
        self.model_settings = model_settings or {}
        self.active_keys = active_keys or []
        self.settings_manager = settings_manager

    def log(self, level: str, msg: str):
        self.log_signal.emit(level, msg)

    def run(self):
        try:
            if not self.provider_id:
                raise ValueError("Не выбран AI-сервис.")
            if not self.active_keys:
                raise ValueError("Не выбран активный ключ или сессия для AI-сервиса.")

            from qidian_rulate.workers import _run_ai_request

            self.log("INFO", "AI: подбираю жанры, теги и параметры RanobeLib...")
            raw_response = _run_ai_request(
                provider_id=self.provider_id,
                model_settings=self.model_settings,
                active_keys=self.active_keys,
                settings_manager=self.settings_manager,
                prompt=_build_ranobelib_catalog_prompt(self.metadata),
                log_callback=self.log,
                log_prefix="Rulate -> RanobeLib catalog",
                max_output_tokens=2048,
            )
            catalog = _parse_ranobelib_catalog_response(raw_response)
            if len(catalog["genres"]) < 3:
                raise ValueError("AI не вернул минимум 3 допустимых жанра RanobeLib.")
            if len(catalog["tags"]) < 3:
                raise ValueError("AI не вернул минимум 3 допустимых тега RanobeLib.")
            self.catalog_ready.emit(catalog)
            self.log("SUCCESS", "AI: жанры, теги и параметры подобраны.")
        except Exception as error:
            self.log("ERROR", f"AI: {error}")
            logging.error(traceback.format_exc())
        finally:
            self.finished_signal.emit()


class RulateToRanobeCreateWorker(QThread):
    log_signal = pyqtSignal(str, str)
    progress_signal = pyqtSignal(int)
    finished_signal = pyqtSignal()

    def __init__(self, rulate_url: str, options: dict | None = None):
        super().__init__()
        self.rulate_url = _rulate_public_book_url(rulate_url.strip().rstrip("/"))
        self.rulate_edit_url = _rulate_edit_info_url(rulate_url.strip().rstrip("/"))
        self.options = options or {}
        self.is_running = True
        self._rulate_browser = None
        self._ranobelib_browser = None
        self._tmp_files: list[str] = []

    def log(self, level: str, msg: str):
        self.log_signal.emit(level, msg)

    def stop(self):
        self.is_running = False
        for browser in (self._rulate_browser, self._ranobelib_browser):
            if browser:
                try:
                    browser.close()
                except Exception:
                    pass

    def _prefetched_metadata_from_options(self) -> dict | None:
        if not _collapse_rulate_spaces(self.options.get("title_ru")):
            return None
        metadata = dict(self.options)
        metadata["rulate_url"] = metadata.get("rulate_url") or self.rulate_url
        metadata["rulate_edit_url"] = metadata.get("rulate_edit_url") or self.rulate_edit_url
        metadata["source_url"] = metadata.get("source_url") or self.rulate_url
        return metadata

    def _read_rulate_metadata(self, playwright) -> dict:
        prefetched = self._prefetched_metadata_from_options()
        if prefetched:
            self.log("INFO", "Rulate: использую данные, уже загруженные в форме. Повторно Rulate не открываю.")
            return prefetched

        self.log("INFO", "Rulate: открываю edit/info через профиль Qidian/Fanqie -> Rulate...")
        self.log("INFO", f"Rulate: страница данных: {self.rulate_edit_url}")
        self.log("INFO", f"Rulate: профиль куки: {QIDIAN_RULATE_PROFILE_DIR}")
        self._rulate_browser = _launch_persistent_chromium_context(
            playwright,
            user_data_dir=str(QIDIAN_RULATE_PROFILE_DIR),
            viewport={"width": 1280, "height": 900},
            log_callback=self.log,
        )
        page = self._rulate_browser.pages[0] if self._rulate_browser.pages else self._rulate_browser.new_page()
        page.goto(self.rulate_edit_url, wait_until="domcontentloaded", timeout=60000)
        page.wait_for_timeout(2500)
        raw_payload = page.evaluate(_RULATE_MEDIA_EXTRACT_SCRIPT)
        raw_payload = _merge_public_rulate_cover(page, raw_payload, self.rulate_url, self.log)
        metadata = _normalize_rulate_media_payload(raw_payload, self.rulate_edit_url)
        self._rulate_browser.close()
        self._rulate_browser = None

        self.log("SUCCESS", f"Rulate: название найдено: {metadata['title_ru']}")
        if metadata.get("author"):
            self.log("INFO", f"Rulate: автор найден: {metadata['author']}")
        if metadata.get("description"):
            self.log("SUCCESS", f"Rulate: описание найдено ({len(metadata['description'])} символов)")
        else:
            self.log("WARNING", "Rulate: описание не найдено, поле описания останется пустым.")
        if metadata.get("cover_url"):
            self.log("INFO", "Rulate: ссылка на обложку найдена.")
        else:
            self.log("WARNING", "Rulate: обложка не найдена.")
        return metadata

    def _apply_options(self, metadata: dict) -> dict:
        data = dict(metadata or {})
        text_fields = {
            "title_ru": "title_ru",
            "original_title": "original_title",
            "title_en": "title_en",
            "alt_names": "alt_names",
            "description": "description",
            "author": "author",
            "year": "year",
            "cover_url": "cover_url",
        }
        for option_key, data_key in text_fields.items():
            value = self.options.get(option_key)
            if isinstance(value, str) and value.strip():
                data[data_key] = value.strip()
        alt_hieroglyph = self.options.get("alt_hieroglyph_title")
        if isinstance(alt_hieroglyph, str) and alt_hieroglyph.strip():
            data["alt_hieroglyph_title"] = alt_hieroglyph.strip()
            data["alt_names"] = _format_ranobelib_alt_names(
                data.get("alt_names"),
                alt_hieroglyph.strip(),
            )

        for key in (
            "type_value",
            "status_value",
            "age_value",
            "translation_status_value",
            "chapter_upload_value",
        ):
            value = self.options.get(key)
            if value not in (None, ""):
                data[key] = str(value)

        data["rulate_genres"] = _as_clean_list(self.options.get("rulate_genres") or data.get("genres"))
        data["rulate_tags"] = _as_clean_list(self.options.get("rulate_tags") or data.get("tags"))
        genres = _normalize_allowed_catalog_items(
            self.options.get("genres"),
            RANOBELIB_GENRES,
            5,
        )
        tags = _normalize_allowed_catalog_items(
            self.options.get("tags"),
            RANOBELIB_TAGS,
            8,
        )
        if genres:
            data["genres"] = genres
        else:
            data["genres"] = []
        if tags:
            data["tags"] = tags
        else:
            data["tags"] = []

        data["create_author"] = bool(self.options.get("create_author"))
        data["source_url"] = data.get("source_url") or self.rulate_url
        data.setdefault("type_value", "12")
        data.setdefault("status_value", "1")
        data.setdefault("age_value", "3")
        data.setdefault("translation_status_value", "1")
        data.setdefault("chapter_upload_value", "2")
        if not data.get("year"):
            data["year"] = "2026"
        return data

    def _download_cover(self, metadata: dict) -> str | None:
        cover_url = metadata.get("cover_url")
        if not cover_url:
            return None
        if _is_bad_rulate_cover_url(cover_url):
            self.log("WARNING", "Rulate: ссылка на обложку похожа на логотип/иконку, пропускаю загрузку.")
            return None

        self.log("INFO", "Rulate: скачиваю обложку для загрузки на RanobeLib...")
        parsed = urllib.parse.urlparse(cover_url)
        suffix = os.path.splitext(parsed.path)[1].lower()
        if suffix not in {".jpg", ".jpeg", ".png", ".webp", ".gif"}:
            suffix = ".jpg"

        request = urllib.request.Request(
            cover_url,
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/120.0 Safari/537.36"
                ),
                "Referer": self.rulate_edit_url,
            },
        )
        tmp = tempfile.NamedTemporaryFile(
            prefix="rulate_cover_",
            suffix=suffix,
            delete=False,
        )
        tmp_path = tmp.name
        try:
            with tmp:
                with urllib.request.urlopen(request, timeout=45) as response:
                    tmp.write(response.read())
            if os.path.getsize(tmp_path) < 1024:
                raise RuntimeError("скачанный файл обложки слишком маленький")
            self._tmp_files.append(tmp_path)
            self.log("SUCCESS", f"Rulate: обложка скачана ({os.path.getsize(tmp_path):,} байт)")
            return tmp_path
        except Exception:
            try:
                os.remove(tmp_path)
            except OSError:
                pass
            raise

    def _open_ranobelib_create_page(self, playwright):
        self.log("INFO", "RanobeLib: открываю создание тайтла...")
        self._ranobelib_browser = _launch_persistent_chromium_context(
            playwright,
            user_data_dir=str(BROWSER_PROFILE_DIR),
            viewport={"width": 1280, "height": 900},
            log_callback=self.log,
        )
        page = (
            self._ranobelib_browser.pages[0]
            if self._ranobelib_browser.pages
            else self._ranobelib_browser.new_page()
        )
        page.goto("https://ranobelib.me/ru/media/create", wait_until="domcontentloaded", timeout=60000)
        try:
            page.wait_for_selector(".media-edit-form, .ProseMirror", timeout=45000)
        except Exception as error:
            raise RuntimeError(
                "не открылась форма создания RanobeLib. "
                "Проверьте, что вы вошли через кнопку «Войти в RanobeLib»."
            ) from error
        return page

    def _upload_cover(self, page, cover_path: str | None):
        if not cover_path:
            return
        try:
            page.locator('input[type="file"][accept*="image"]').first.set_input_files(cover_path)
            page.wait_for_timeout(1500)
            self.log("SUCCESS", "RanobeLib: обложка подставлена в форму.")
        except Exception as error:
            self.log("WARNING", f"RanobeLib: не удалось подставить обложку автоматически: {error}")

    def _fill_description(self, page, description: str):
        if not description:
            return
        try:
            editor = page.locator(".ProseMirror").first
            editor.click(timeout=10000)
            page.keyboard.press("Control+A")
            page.keyboard.press("Backspace")
            page.keyboard.insert_text(description)
            self.log("SUCCESS", "RanobeLib: описание вставлено.")
        except Exception as error:
            self.log("WARNING", f"RanobeLib: не удалось вставить описание автоматически: {error}")

    def _try_fill_source_link(self, page, source_url: str):
        if not source_url:
            return
        try:
            clicked = page.evaluate(
                r"""() => {
                    const norm = (value) => (value || "").replace(/\s+/g, " ").trim();
                    const isVisible = (el) => {
                        const style = window.getComputedStyle(el);
                        const rect = el.getBoundingClientRect();
                        return style.display !== "none" && style.visibility !== "hidden"
                            && rect.width > 0 && rect.height > 0 && !el.disabled;
                    };
                    const isSourceLabel = (text) => {
                        const lower = norm(text).toLowerCase();
                        return lower.includes("ссыл") && lower.includes("оригинал");
                    };
                    const group = Array.from(document.querySelectorAll(".form-group")).find((item) => {
                        const label = norm(item.querySelector(".form-label")?.innerText || "");
                        return isSourceLabel(label) || isSourceLabel(item.innerText || "");
                    }) || Array.from(document.querySelectorAll("section, .section, .paper, div")).find((item) => {
                        const text = norm(item.innerText || "");
                        return text.length < 800 && isSourceLabel(text);
                    });
                    if (!group) return false;
                    const buttons = Array.from(group.querySelectorAll("button")).filter(isVisible);
                    const button = buttons.find((btn) => {
                        const text = norm(btn.innerText).toLowerCase();
                        return text.includes("добав") && text.includes("ссыл");
                    }) || buttons.find((btn) => norm(btn.innerText).toLowerCase().includes("добав")) || buttons[0];
                    if (!button) return false;
                    button.click();
                    return true;
                }"""
            )
            if not clicked:
                self.log("WARNING", "RanobeLib: не нашёл кнопку добавления ссылки на оригинал, пробую найти готовое поле.")
            page.wait_for_timeout(800)
            filled = page.evaluate(
                r"""(url) => {
                    const norm = (value) => (value || "").replace(/\s+/g, " ").trim();
                    const setNativeValue = (el, value) => {
                        const prototype = Object.getPrototypeOf(el);
                        const descriptor = Object.getOwnPropertyDescriptor(prototype, "value")
                            || Object.getOwnPropertyDescriptor(HTMLInputElement.prototype, "value")
                            || Object.getOwnPropertyDescriptor(HTMLTextAreaElement.prototype, "value");
                        if (descriptor && descriptor.set) descriptor.set.call(el, value);
                        else el.value = value;
                        el.dispatchEvent(new Event("input", { bubbles: true }));
                        el.dispatchEvent(new Event("change", { bubbles: true }));
                        el.dispatchEvent(new Event("blur", { bubbles: true }));
                    };
                    const visible = (el) => {
                        const style = window.getComputedStyle(el);
                        const rect = el.getBoundingClientRect();
                        return style.display !== "none" && style.visibility !== "hidden"
                            && rect.width > 0 && rect.height > 0 && !el.disabled;
                    };
                    const isSourceText = (text) => {
                        const lower = norm(text).toLowerCase();
                        return lower.includes("ссыл") && lower.includes("оригинал");
                    };
                    const group = Array.from(document.querySelectorAll(".form-group")).find((item) => {
                        const label = norm(item.querySelector(".form-label")?.innerText || "");
                        return isSourceText(label) || isSourceText(item.innerText || "");
                    }) || Array.from(document.querySelectorAll("section, .section, .paper, div")).find((item) => {
                        const text = norm(item.innerText || "");
                        return text.length < 1200 && isSourceText(text);
                    });
                    const scopes = [group, group?.nextElementSibling, group?.parentElement, document].filter(Boolean);
                    const fields = [];
                    for (const scope of scopes) {
                        for (const field of Array.from(scope.querySelectorAll("input, textarea"))) {
                            if (field.type === "file" || !visible(field)) continue;
                            if (!fields.includes(field)) fields.push(field);
                        }
                    }
                    const empty = fields.filter((field) => !norm(field.value));
                    const urlField = empty.find((field) => {
                        const haystack = norm(`${field.type || ""} ${field.placeholder || ""} ${field.closest(".form-group")?.innerText || ""}`).toLowerCase();
                        return field.type === "url"
                            || haystack.includes("url")
                            || haystack.includes("http")
                            || haystack.includes("ссыл")
                            || haystack.includes("оригинал");
                    });
                    const target = urlField || empty[empty.length - 1] || null;
                    if (!target) return false;
                    setNativeValue(target, url);
                    const targetGroup = target.closest(".form-group") || target.parentElement;
                    if (targetGroup) {
                        const siblingFields = Array.from(targetGroup.querySelectorAll("input, textarea")).filter((field) => {
                            return field !== target && field.type !== "file" && visible(field) && !norm(field.value);
                        });
                        const titleField = siblingFields.find((field) => {
                            const haystack = norm(`${field.placeholder || ""} ${field.closest(".form-group")?.innerText || ""}`).toLowerCase();
                            return haystack.includes("назв") || haystack.includes("name");
                        });
                        if (titleField) setNativeValue(titleField, "Оригинал");
                    }
                    return true;
                }""",
                source_url,
            )
            if filled:
                self.log("SUCCESS", "RanobeLib: ссылка на оригинал вставлена.")
            else:
                self.log("WARNING", "RanobeLib: ссылку на оригинал нужно вставить вручную.")
        except Exception as error:
            self.log("WARNING", f"RanobeLib: не удалось вставить ссылку на Rulate: {error}")

    def _click_group_button(self, page, group_label: str, button_text: str = "Добавить") -> bool:
        return bool(
            page.evaluate(
                r"""([groupLabel, buttonText]) => {
                    const norm = (value) => (value || "").replace(/\s+/g, " ").trim();
                    const visible = (el) => {
                        if (!el || el.disabled) return false;
                        const style = window.getComputedStyle(el);
                        const rect = el.getBoundingClientRect();
                        return style.display !== "none" && style.visibility !== "hidden"
                            && rect.width > 0 && rect.height > 0;
                    };
                    for (const old of Array.from(document.querySelectorAll("[data-codex-before-autocomplete]"))) {
                        old.removeAttribute("data-codex-before-autocomplete");
                    }
                    for (const input of Array.from(document.querySelectorAll("input"))) {
                        if (visible(input)) input.setAttribute("data-codex-before-autocomplete", "1");
                    }
                    for (const old of Array.from(document.querySelectorAll("[data-codex-autocomplete-group]"))) {
                        old.removeAttribute("data-codex-autocomplete-group");
                    }
                    const group = Array.from(document.querySelectorAll(".form-group")).find((item) => {
                        const label = norm(item.querySelector(".form-label")?.innerText || "");
                        return label.toLowerCase().includes(groupLabel.toLowerCase());
                    });
                    if (!group) return false;
                    group.setAttribute("data-codex-autocomplete-group", groupLabel);
                    const looksLikeRemoveButton = (btn) => {
                        const text = norm(
                            `${btn.innerText || ""} ${btn.getAttribute("aria-label") || ""} ${btn.getAttribute("title") || ""}`
                        ).toLowerCase();
                        const classes = String(btn.className || "").toLowerCase();
                        const chip = btn.closest(
                            ".tag, .label, .badge, .chip, .search-choice, .select2-selection__choice, .ms-sel-item"
                        );
                        return Boolean(chip)
                            || text === "x"
                            || text === "×"
                            || text.includes("удал")
                            || text.includes("remove")
                            || text.includes("delete")
                            || text.includes("close")
                            || classes.includes("remove")
                            || classes.includes("delete")
                            || classes.includes("close");
                    };
                    const buttons = Array.from(group.querySelectorAll("button"))
                        .filter((btn) => visible(btn) && !looksLikeRemoveButton(btn));
                    const buttonNeedle = buttonText.toLowerCase();
                    const button = buttons.find((btn) => {
                        const text = norm(
                            `${btn.innerText || ""} ${btn.getAttribute("aria-label") || ""} ${btn.getAttribute("title") || ""}`
                        ).toLowerCase();
                        return text.includes(buttonNeedle);
                    })
                        || buttons.find((btn) => {
                            const text = norm(
                                `${btn.innerText || ""} ${btn.getAttribute("aria-label") || ""} ${btn.getAttribute("title") || ""}`
                            ).toLowerCase();
                            const classes = String(btn.className || "").toLowerCase();
                            return text.includes("add") || text.includes("добав") || classes.includes("add");
                        })
                        || buttons[buttons.length - 1];
                    if (!button) return false;
                    button.scrollIntoView({ block: "center", inline: "center" });
                    button.click();
                    return true;
                }""",
                [group_label, button_text],
            )
        )

    def _focus_group_autocomplete_input(self, page, group_label: str) -> bool:
        return bool(
            page.evaluate(
                r"""(groupLabel) => {
                    const norm = (value) => (value || "").replace(/\s+/g, " ").trim();
                    const visible = (el) => {
                        if (!el || el.disabled) return false;
                        const style = window.getComputedStyle(el);
                        const rect = el.getBoundingClientRect();
                        return style.display !== "none" && style.visibility !== "hidden"
                            && rect.width > 0 && rect.height > 0;
                    };
                    const labelFor = (el) => norm(el?.querySelector(".form-label")?.innerText || "");
                    const wanted = groupLabel.toLowerCase();
                    const group = Array.from(document.querySelectorAll(".form-group")).find((item) => {
                        return labelFor(item).toLowerCase().includes(wanted);
                    });
                    if (!group) return false;

                    for (const old of Array.from(document.querySelectorAll("[data-codex-autocomplete-target]"))) {
                        old.removeAttribute("data-codex-autocomplete-target");
                    }

                    const scopes = [
                        group,
                        group.nextElementSibling,
                    ].filter(Boolean);
                    const fields = [];
                    const collectInput = (field) => {
                        const type = (field.type || "text").toLowerCase();
                        if (!["text", "search", ""].includes(type) || !visible(field)) return;
                        if (fields.includes(field)) return;
                        fields.push(field);
                    };
                    for (const scope of scopes) {
                        for (const field of Array.from(scope.querySelectorAll("input"))) {
                            collectInput(field);
                        }
                    }
                    if (document.activeElement && document.activeElement.tagName === "INPUT" && visible(document.activeElement)) {
                        const active = document.activeElement;
                        const activeGroup = active.closest(".form-group");
                        const activeLabel = labelFor(activeGroup).toLowerCase();
                        if ((group.contains(active) || activeLabel.includes(wanted) || !active.hasAttribute("data-codex-before-autocomplete")) && !fields.includes(active)) {
                            fields.unshift(active);
                        }
                    }
                    for (const field of Array.from(document.querySelectorAll("input"))) {
                        if (field.hasAttribute("data-codex-before-autocomplete")) continue;
                        collectInput(field);
                    }
                    const scored = fields.map((field) => {
                        const closestGroup = field.closest(".form-group");
                        const closestLabel = labelFor(closestGroup).toLowerCase();
                        const haystack = norm(`${field.className || ""} ${field.placeholder || ""} ${closestLabel}`).toLowerCase();
                        let score = 0;
                        if (group.contains(field)) score += 100;
                        if (closestLabel.includes(wanted)) score += 100;
                        if (!field.hasAttribute("data-codex-before-autocomplete")) score += 90;
                        if (document.activeElement === field) score += 40;
                        if (haystack.includes("autocomplete") || haystack.includes("поиск") || haystack.includes("введите")) score += 20;
                        if (haystack.includes("form-input__field") && field.hasAttribute("data-codex-before-autocomplete")) score -= 200;
                        if (norm(field.value)) score -= 20;
                        if (closestLabel.includes("оригинал") || closestLabel.includes("ссыл")) score -= 500;
                        if (closestLabel.includes("название")) score -= 500;
                        if (haystack.includes("url") || haystack.includes("http")) score -= 500;
                        return { field, score };
                    }).sort((a, b) => b.score - a.score);
                    const target = scored.find((item) => item.score >= 80)?.field || null;
                    if (!target) return false;
                    target.setAttribute("data-codex-autocomplete-target", "1");
                    target.focus();
                    target.click();
                    return true;
                }""",
                group_label,
            )
        )

    def _click_autocomplete_suggestion(self, page, value: str, allow_fallback: bool = True) -> bool:
        exact_text = re.compile(rf"^\s*{re.escape(value)}\s*$", re.IGNORECASE)
        locators = [page.locator(".autocomplete-suggestion__item:visible").filter(has_text=exact_text).first]
        if allow_fallback:
            locators.append(page.locator(".autocomplete-suggestion__item:visible").filter(has_text=value).first)
        for locator in locators:
            try:
                if locator.count() == 0:
                    continue
                locator.scroll_into_view_if_needed(timeout=2500)
                locator.click(timeout=3000, force=True)
                return True
            except Exception:
                pass

        box = page.evaluate(
            r"""([value, allowFallback]) => {
                const norm = (raw) => (raw || "").replace(/\s+/g, " ").trim();
                const wanted = norm(value).toLowerCase();
                const visible = (el) => {
                    if (!el) return false;
                    const style = window.getComputedStyle(el);
                    const rect = el.getBoundingClientRect();
                    return style.display !== "none" && style.visibility !== "hidden"
                        && rect.width > 0 && rect.height > 0;
                };
                const suggestions = Array.from(document.querySelectorAll(".autocomplete-suggestion__item"))
                    .filter(visible);
                const exact = suggestions.find((item) => norm(item.innerText).toLowerCase() === wanted);
                let target = exact || null;
                if (!target && allowFallback) {
                    const partial = suggestions.find((item) => norm(item.innerText).toLowerCase().includes(wanted));
                    target = partial || suggestions[0] || null;
                }
                if (!target) return null;
                target.scrollIntoView({ block: "center", inline: "center" });
                const rect = target.getBoundingClientRect();
                return {
                    x: rect.left + rect.width / 2,
                    y: rect.top + rect.height / 2,
                };
            }""",
            [value, allow_fallback],
        )
        if not box:
            return False
        try:
            page.mouse.move(box["x"], box["y"])
            page.wait_for_timeout(150)
            page.mouse.down()
            page.wait_for_timeout(120)
            page.mouse.up()
            return True
        except Exception:
            return False

    def _group_autocomplete_is_open(self, page, group_label: str) -> bool:
        try:
            return bool(
                page.evaluate(
                    r"""(groupLabel) => {
                        const norm = (value) => (value || "").replace(/\s+/g, " ").trim();
                        const visible = (el) => {
                            if (!el) return false;
                            const style = window.getComputedStyle(el);
                            const rect = el.getBoundingClientRect();
                            return style.display !== "none" && style.visibility !== "hidden"
                                && rect.width > 0 && rect.height > 0;
                        };
                        const wanted = groupLabel.toLowerCase();
                        const group = Array.from(document.querySelectorAll(".form-group")).find((item) => {
                            const label = norm(item.querySelector(".form-label")?.innerText || "");
                            return label.toLowerCase().includes(wanted);
                        });
                        if (!group) return false;
                        const button = Array.from(group.querySelectorAll("button")).find(visible);
                        if (!button || button.getAttribute("aria-expanded") !== "true") return false;
                        return Array.from(document.querySelectorAll(".autocomplete-suggestion__item")).some(visible)
                            || Array.from(document.querySelectorAll("input")).some((field) => {
                                return visible(field) && !field.hasAttribute("data-codex-before-autocomplete");
                            });
                    }""",
                    group_label,
                )
            )
        except Exception:
            return False

    def _close_autocomplete_popovers(self, page, except_group_label: str | None = None):
        try:
            page.evaluate(
                r"""(exceptGroupLabel) => {
                    const norm = (value) => (value || "").replace(/\s+/g, " ").trim();
                    const wanted = (exceptGroupLabel || "").toLowerCase();
                    const visible = (el) => {
                        if (!el) return false;
                        const style = window.getComputedStyle(el);
                        const rect = el.getBoundingClientRect();
                        return style.display !== "none" && style.visibility !== "hidden"
                            && rect.width > 0 && rect.height > 0;
                    };
                    for (const group of Array.from(document.querySelectorAll(".form-group"))) {
                        const label = norm(group.querySelector(".form-label")?.innerText || "").toLowerCase();
                        if (wanted && label.includes(wanted)) continue;
                        for (const button of Array.from(group.querySelectorAll("button"))) {
                            if (visible(button) && button.getAttribute("aria-expanded") === "true") {
                                button.click();
                            }
                        }
                    }
                }""",
                except_group_label,
            )
        except Exception:
            pass

    def _ensure_group_autocomplete_open(self, page, group_label: str) -> bool:
        self._close_autocomplete_popovers(page, except_group_label=group_label)
        if self._group_autocomplete_is_open(page, group_label):
            return True
        for _ in range(3):
            if not self._click_group_button(page, group_label):
                return False
            page.wait_for_timeout(900)
            if self._group_autocomplete_is_open(page, group_label):
                return True
        return False

    def _is_static_catalog_group(self, group_label: str) -> bool:
        label = str(group_label or "").casefold()
        return "жанр" in label or "тег" in label

    def _add_autocomplete_item(self, page, group_label: str, value: str) -> bool:
        value = _collapse_rulate_spaces(value)
        if not value:
            return False
        if self._group_contains_value(page, group_label, value):
            return True
        try:
            if not self._ensure_group_autocomplete_open(page, group_label):
                return False
            if self._is_static_catalog_group(group_label):
                for _ in range(2):
                    if self._click_autocomplete_suggestion(page, value, allow_fallback=False):
                        page.wait_for_timeout(900)
                        if self._group_contains_value(page, group_label, value):
                            return True
                    if not self._group_autocomplete_is_open(page, group_label):
                        self._ensure_group_autocomplete_open(page, group_label)
                    page.wait_for_timeout(500)

            focused = False
            for _ in range(5):
                if self._focus_group_autocomplete_input(page, group_label):
                    focused = True
                    break
                page.wait_for_timeout(350)
            if not focused:
                return False
            if not self._clear_active_autocomplete_input(page):
                return False
            page.keyboard.type(value, delay=60)
            page.wait_for_timeout(2400)
            success = False
            for _ in range(2):
                if self._click_autocomplete_suggestion(page, value):
                    page.wait_for_timeout(2600)
                    if self._group_contains_value(page, group_label, value):
                        success = True
                        break
                page.wait_for_timeout(1800)
            if not success:
                try:
                    page.keyboard.press("ArrowDown")
                    page.wait_for_timeout(400)
                    page.keyboard.press("Enter")
                    page.wait_for_timeout(2600)
                    success = self._group_contains_value(page, group_label, value)
                except Exception:
                    success = False
            try:
                page.keyboard.press("Escape")
            except Exception:
                pass
            page.wait_for_timeout(900)
            return success
        except Exception:
            try:
                page.keyboard.press("Escape")
            except Exception:
                pass
            page.wait_for_timeout(1200)
            return False

    def _clear_active_autocomplete_input(self, page) -> bool:
        try:
            return bool(
                page.evaluate(
                    r"""() => {
                        const visible = (el) => {
                            if (!el || el.disabled) return false;
                            const style = window.getComputedStyle(el);
                            const rect = el.getBoundingClientRect();
                            return style.display !== "none" && style.visibility !== "hidden"
                                && rect.width > 0 && rect.height > 0;
                        };
                        const setNativeValue = (el, value) => {
                            const prototype = Object.getPrototypeOf(el);
                            const descriptor = Object.getOwnPropertyDescriptor(prototype, "value")
                                || Object.getOwnPropertyDescriptor(HTMLInputElement.prototype, "value");
                            if (descriptor && descriptor.set) descriptor.set.call(el, value);
                            else el.value = value;
                            el.dispatchEvent(new Event("input", { bubbles: true }));
                            el.dispatchEvent(new Event("change", { bubbles: true }));
                        };
                        const target = document.querySelector('input[data-codex-autocomplete-target="1"]');
                        if (!visible(target)) return false;
                        target.focus();
                        target.click();
                        setNativeValue(target, "");
                        const active = document.activeElement;
                        return active && active.tagName === "INPUT" && visible(active);
                    }"""
                )
            )
        except Exception:
            return False

    def _group_contains_value(self, page, group_label: str, value: str) -> bool:
        value = _collapse_rulate_spaces(value)
        if not value:
            return False
        try:
            return bool(
                page.evaluate(
                    r"""([groupLabel, value]) => {
                        const norm = (raw) => (raw || "").replace(/\s+/g, " ").trim();
                        const group = Array.from(document.querySelectorAll(".form-group")).find((item) => {
                            const label = norm(item.querySelector(".form-label")?.innerText || "");
                            return label.toLowerCase().includes(groupLabel.toLowerCase());
                        });
                        if (!group) return false;
                        const clone = group.cloneNode(true);
                        for (const item of Array.from(clone.querySelectorAll(
                            ".form-label, .autocomplete-suggestion, .autocomplete-suggestion__list, input, textarea"
                        ))) {
                            item.remove();
                        }
                        for (const button of Array.from(clone.querySelectorAll("button"))) {
                            const text = norm(button.innerText || "").toLowerCase();
                            if (!text || text.includes("добавить")) button.remove();
                        }
                        const text = norm(clone.innerText || "");
                        return text.toLowerCase().includes(norm(value).toLowerCase());
                    }""",
                    [group_label, value],
                )
            )
        except Exception:
            return False

    def _wait_between_autocomplete_items(self, page):
        self._close_autocomplete_popovers(page)
        try:
            page.evaluate("document.activeElement && document.activeElement.blur && document.activeElement.blur()")
        except Exception:
            pass
        page.wait_for_timeout(900)

    def _add_autocomplete_items(self, page, group_label: str, values, limit: int):
        added = []
        for value in _as_clean_list(values)[:limit]:
            if self._add_autocomplete_item(page, group_label, value):
                added.append(value)
            else:
                self.log("WARNING", f"RanobeLib: не удалось добавить {group_label.lower()}: {value}")
            page.wait_for_timeout(600)
        self._wait_between_autocomplete_items(page)
        if added:
            self.log("SUCCESS", f"RanobeLib: добавлено в «{group_label}»: {', '.join(added)}")

    def _open_author_create_page(self, page):
        author_page = page.context.new_page()
        author_page.goto("https://ranobelib.me/ru/people/create", wait_until="domcontentloaded", timeout=30000)
        author_page.wait_for_timeout(1500)
        return author_page

    def _fill_author_create_form(self, author_page, author_payload: dict) -> dict:
        return author_page.evaluate(
            r"""(data) => {
                const norm = (value) => (value || "").replace(/\s+/g, " ").trim();
                const setNativeValue = (el, value) => {
                    const prototype = Object.getPrototypeOf(el);
                    const descriptor = Object.getOwnPropertyDescriptor(prototype, "value")
                        || Object.getOwnPropertyDescriptor(HTMLInputElement.prototype, "value")
                        || Object.getOwnPropertyDescriptor(HTMLTextAreaElement.prototype, "value");
                    if (descriptor && descriptor.set) descriptor.set.call(el, value);
                    else el.value = value;
                    el.dispatchEvent(new Event("input", { bubbles: true }));
                    el.dispatchEvent(new Event("change", { bubbles: true }));
                    el.dispatchEvent(new Event("blur", { bubbles: true }));
                };
                const groups = () => Array.from(document.querySelectorAll(".form-group"));
                const findGroup = (needles) => {
                    const normalized = needles.map((item) => item.toLowerCase());
                    return groups().find((group) => {
                        const label = norm(group.querySelector(".form-label")?.innerText || group.innerText || "").toLowerCase();
                        return normalized.every((needle) => label.includes(needle));
                    });
                };
                const fillGroup = (needles, value, selector = "input, textarea") => {
                    if (!value) return false;
                    const group = findGroup(needles);
                    const field = group?.querySelector(selector);
                    if (!field || field.disabled) return false;
                    setNativeValue(field, value);
                    return true;
                };
                const result = {
                    nameEn: fillGroup(["имя", "английском"], data.name_en)
                        || fillGroup(["имя", "ромадзи"], data.name_en),
                    nameRu: fillGroup(["имя", "русском"], data.name_ru),
                    aliases: fillGroup(["известен"], data.aliases, "textarea, input"),
                };
                if (data.description) {
                    const editor = document.querySelector(".ProseMirror");
                    if (editor) {
                        editor.focus();
                        editor.innerHTML = "";
                        const paragraph = document.createElement("p");
                        paragraph.textContent = data.description;
                        editor.appendChild(paragraph);
                        editor.dispatchEvent(new InputEvent("input", { bubbles: true, inputType: "insertText", data: data.description }));
                        result.description = true;
                    }
                }
                return result;
            }""",
            author_payload,
        )

    def _try_create_author(self, page, metadata: dict) -> bool:
        author_payload = _prepare_ranobelib_author_payload(metadata.get("author"))
        if not author_payload.get("original"):
            return False
        try:
            self.log("INFO", "RanobeLib: автор не найден в подсказках, открываю форму создания автора.")
            author_page = self._open_author_create_page(page)
            filled = self._fill_author_create_form(author_page, author_payload)
            missed = []
            if not filled.get("nameEn"):
                missed.append("Имя (на Английском или Ромадзи)")
            if not filled.get("nameRu"):
                missed.append("Имя на русском")
            if missed:
                self.log("WARNING", "RanobeLib: в форме автора не заполнены автоматически: " + ", ".join(missed))
            self.log(
                "SUCCESS",
                "RanobeLib: форма нового автора заполнена. Кнопку «Создать» нужно нажать вручную после проверки.",
            )
            return True
        except Exception as error:
            self.log("WARNING", f"RanobeLib: не удалось открыть/заполнить форму автора: {error}")
            return False

    def _author_autocomplete_candidates(self, metadata: dict) -> list[str]:
        author = _collapse_rulate_spaces(metadata.get("author"))
        if not author:
            return []
        payload = _prepare_ranobelib_author_payload(author)
        primary = _collapse_rulate_spaces(payload.get("name_en")) or _collapse_rulate_spaces(payload.get("original"))
        return [primary] if primary else []

    def _ensure_author(self, page, metadata: dict, create_if_missing: bool):
        author = _collapse_rulate_spaces(metadata.get("author"))
        if not author:
            return
        for candidate in self._author_autocomplete_candidates(metadata):
            if self._add_autocomplete_item(page, "Автор", candidate):
                self.log("SUCCESS", f"RanobeLib: автор добавлен: {candidate}")
                return
        if create_if_missing and self._try_create_author(page, metadata):
            self.log("WARNING", "RanobeLib: после ручного создания автора его нужно добавить в карточку.")
            return
        self.log("WARNING", f"RanobeLib: автора нужно проверить вручную: {author}")

    def _fill_ranobelib_form(self, page, metadata: dict, cover_path: str | None):
        self._upload_cover(page, cover_path)
        filled = page.evaluate(_RANOBELIB_MEDIA_FORM_FILL_SCRIPT, metadata)
        required = {
            "original": "Оригинальное название",
            "titleRu": "Название на русском",
            "titleEn": "Название на английском",
            "altNames": "Альтернативные названия",
            "type": "Тип",
            "status": "Статус тайтла",
            "age": "Возрастное ограничение",
            "translationStatus": "Статус перевода",
            "chaptersUpload": "Загрузка глав",
        }
        missed = [label for key, label in required.items() if not filled.get(key)]
        if missed:
            self.log("WARNING", "RanobeLib: не заполнены автоматически: " + ", ".join(missed))
        else:
            self.log("SUCCESS", "RanobeLib: основные поля заполнены.")
        self._fill_description(page, metadata.get("description", ""))
        self._ensure_author(
            page,
            metadata,
            create_if_missing=bool(metadata.get("create_author")),
        )
        self._add_autocomplete_items(page, "Жанры", metadata.get("genres"), 5)
        self._add_autocomplete_items(page, "Теги", metadata.get("tags"), 8)
        self._try_fill_source_link(page, metadata.get("source_url", ""))

        if metadata.get("author"):
            self.log(
                "INFO",
                f"RanobeLib: автор из Rulate: {metadata['author']}",
            )
        if metadata.get("rulate_genres"):
            self.log("INFO", "Rulate: жанры для ручной проверки: " + ", ".join(metadata["rulate_genres"]))
        if metadata.get("rulate_tags"):
            self.log("INFO", "Rulate: теги для ручной проверки: " + ", ".join(metadata["rulate_tags"]))

    def _wait_until_browser_closed(self):
        self.log(
            "WARNING",
            "Проверьте форму RanobeLib вручную. Автоматически «Создать» не нажимаю; закройте браузер после проверки.",
        )
        while self.is_running and self._ranobelib_browser:
            try:
                pages = [page for page in self._ranobelib_browser.pages if not page.is_closed()]
                if not pages:
                    break
                pages[0].wait_for_timeout(1000)
            except Exception:
                break

    def _cleanup_tmp_files(self):
        for path in self._tmp_files:
            try:
                os.remove(path)
            except OSError:
                pass
        self._tmp_files.clear()

    def run(self):
        try:
            with sync_playwright() as p:
                self.progress_signal.emit(10)
                metadata = self._read_rulate_metadata(p)
                metadata = self._apply_options(metadata)
                self.progress_signal.emit(35)
                cover_path = self._download_cover(metadata)
                self.progress_signal.emit(50)
                page = self._open_ranobelib_create_page(p)
                self.progress_signal.emit(70)
                self._fill_ranobelib_form(page, metadata, cover_path)
                self.progress_signal.emit(100)
                self._wait_until_browser_closed()
        except Exception as e:
            self.log("ERROR", f"Rulate -> RanobeLib: {e}")
            logging.error(traceback.format_exc())
        finally:
            for browser in (self._rulate_browser, self._ranobelib_browser):
                if browser:
                    try:
                        browser.close()
                    except Exception:
                        pass
            self._rulate_browser = None
            self._ranobelib_browser = None
            self._cleanup_tmp_files()
            self.finished_signal.emit()


# ─── Рабочий поток: определение последней главы на RanobeLib ────────────────

class LastChapterDetector(QThread):
    """
    Определяет номер последней залитой главы на RanobeLib.

    Способ 1 (основной): открыть страницу add-chapter — поле «Глава»
                         содержит предложенный номер (последняя + 1).
    Способ 2 (фолбэк):  открыть ?section=chapters — первая глава в списке
                         (отсортирован новые→старые) и есть последняя.
    """
    log_signal = pyqtSignal(str, str)
    result_signal = pyqtSignal(float, str)  # (номер_главы, описание)
    finished_signal = pyqtSignal()

    def __init__(self, ranobelib_url: str):
        super().__init__()
        self.raw_url = ranobelib_url.strip()
        # URL add-chapter для способа 1
        self.add_chapter_url = self.raw_url
        if not self.add_chapter_url.rstrip("/").endswith("/add-chapter"):
            self.add_chapter_url = self.add_chapter_url.rstrip("/") + "/add-chapter"
        # URL ?section=chapters для способа 2
        self.book_url = re.sub(r"/add-chapter\s*$", "", self.raw_url)
        self.chapters_url = self.book_url.rstrip("/") + "?section=chapters"
        self.is_running = True

    def log(self, level: str, msg: str):
        self.log_signal.emit(level, msg)

    def stop(self):
        self.is_running = False

    def run(self):
        try:
            self._detect()
        except Exception as e:
            self.log("ERROR", f"Детектор: {e}")
            logging.error(traceback.format_exc())
        finally:
            self.finished_signal.emit()

    def _detect(self):
        self.log("INFO", "Определяю последнюю залитую главу на RanobeLib…")
        last_chapter_num = 0.0
        last_chapter_desc = ""

        try:
            with sync_playwright() as p:
                browser = _launch_persistent_chromium_context(
                    p,
                    user_data_dir=str(BROWSER_PROFILE_DIR),
                    headless=True,
                    viewport={"width": 1280, "height": 900},
                    log_callback=self.log,
                )
                page = browser.pages[0]

                # ── Способ 1: страница add-chapter ──
                # Поле «Глава» содержит предложенный номер (последняя + 1)
                try:
                    page.goto(self.add_chapter_url, timeout=30000)
                    page.wait_for_selector(
                        SELECTORS["chapter_input"], state="visible", timeout=10000
                    )
                    suggested = page.input_value(SELECTORS["chapter_input"]).strip()
                    if suggested:
                        suggested_num = float(suggested)
                        if suggested_num > 1:
                            last_chapter_num = suggested_num - 1
                            last_chapter_desc = f"на add-chapter предложена {format_num(suggested_num)}"
                        elif suggested_num == 1:
                            # Предложена глава 1 — значит глав ещё нет
                            last_chapter_num = 0
                            last_chapter_desc = "глав ещё нет"
                        self.log("INFO",
                                 f"Способ 1: поле «Глава» = {suggested} → "
                                 f"последняя = {format_num(last_chapter_num)}")
                except Exception as e:
                    self.log("WARNING", f"Способ 1 (add-chapter) не сработал: {e}")

                # ── Способ 2 (фолбэк): ?section=chapters ──
                if last_chapter_num == 0 and last_chapter_desc != "глав ещё нет":
                    try:
                        self.log("INFO", "Пробую способ 2: ?section=chapters…")
                        page.goto(self.chapters_url, timeout=30000)
                        page.wait_for_timeout(4000)

                        result = page.evaluate("""() => {
                            const body = document.body.innerText;

                            // �?щем все упоминания "Глава X" на странице
                            const matches = [...body.matchAll(
                                /(?:Том\\s*\\d+\\s+)?Глава\\s+(\\d+(?:\\.\\d+)?)/gi
                            )];
                            let maxNum = 0;
                            let maxMatch = '';
                            for (const m of matches) {
                                const num = parseFloat(m[1]);
                                if (num > maxNum) {
                                    maxNum = num;
                                    maxMatch = m[0];
                                }
                            }
                            return { number: maxNum, description: maxMatch };
                        }""")

                        num = result.get("number", 0)
                        if num > 0:
                            last_chapter_num = num
                            last_chapter_desc = result.get("description", "")
                            self.log("INFO",
                                     f"Способ 2: найдена последняя глава {format_num(num)} ({last_chapter_desc})")
                    except Exception as e:
                        self.log("WARNING", f"Способ 2 (?section=chapters) не сработал: {e}")

                browser.close()

                if last_chapter_num > 0:
                    self.log("SUCCESS",
                             f"Последняя глава на RanobeLib: {format_num(last_chapter_num)} ({last_chapter_desc})")
                elif last_chapter_desc == "глав ещё нет":
                    self.log("INFO", "На RanobeLib ещё нет глав. Пропуск не требуется.")
                else:
                    self.log("WARNING", "Не удалось определить последнюю главу. "
                             "Возможно, требуется авторизация на RanobeLib.")

                self.result_signal.emit(last_chapter_num, last_chapter_desc)

        except Exception as e:
            self.log("ERROR", f"Детектор: ошибка: {e}")
            self.result_signal.emit(0.0, "Ошибка")


# ─── Рабочий поток: загрузка глав ───────────────────────────────────────────

class UploadWorker(QThread):
    log_signal = pyqtSignal(str, str)
    progress_signal = pyqtSignal(int)
    stats_signal = pyqtSignal(int, int, int)  # ok, errors, skipped
    eta_signal = pyqtSignal(str)               # "~12мин 30сек"
    finished_signal = pyqtSignal()
    chapter_done_signal = pyqtSignal(int)      # Feature 2: индекс завершённой главы

    def __init__(
        self,
        url: str,
        chapters_list: list[ChapterData],
        schedule_enabled: bool,
        start_time: datetime,
        interval_minutes: int,
        paid_enabled: bool,
        price: int,
        force_num: bool,
    ):
        super().__init__()
        self.url = url
        self.chapters_list = chapters_list
        self.schedule_enabled = schedule_enabled
        self.current_publish_time = start_time
        self.interval_minutes = interval_minutes
        self.paid_enabled = paid_enabled
        self.price = price
        self.force_num = force_num
        self.is_running = True
        self.limit_date = datetime.now() + timedelta(days=60)

        # Статистика
        self._ok = 0
        self._errors = 0
        self._skipped = 0
        self._times: list[float] = []  # время загрузки каждой главы

    def log(self, level: str, msg: str):
        self.log_signal.emit(level, msg)

    # ── Настройка времени в поповере ──

    def _adjust_time_column(self, page, popover, col_index: int, target_val: str):
        try:
            btn_up = popover.locator(SELECTORS["arrow_up"]).nth(col_index)
            btn_down = popover.locator(SELECTORS["arrow_down"]).nth(col_index)
            container = btn_up.locator("xpath=../..")
            target_int = int(target_val)

            for _ in range(60):
                if not self.is_running:
                    return
                raw = container.inner_text()
                m = re.search(r"\d+", raw)
                if not m:
                    page.wait_for_timeout(100)
                    continue
                curr = int(m.group(0))
                if curr == target_int:
                    return

                diff = target_int - curr
                # Выбираем кратчайший путь (для часов 0–23)
                go_up = diff > 0 if abs(diff) < 30 else diff <= 0
                (btn_up if go_up else btn_down).click()
                page.wait_for_timeout(70)

        except Exception as e:
            self.log("WARNING", f"Сбой настройки времени (колонка {col_index}): {e}")

    def _set_calendar_date(self, page, popover, target_dt: datetime):
        try:
            target_year = target_dt.year
            target_month = target_dt.month

            for _ in range(24):
                page.wait_for_timeout(200)
                header = popover.text_content().lower()

                found_year = -1
                found_month = -1
                ym = re.search(r"20\d{2}", header)
                if ym:
                    found_year = int(ym.group(0))
                for name, num in RUS_MONTHS.items():
                    if name in header:
                        found_month = num
                        break

                if found_month == -1 or found_year == -1:
                    popover.locator(SELECTORS["month_arrow_right"]).click()
                    page.wait_for_timeout(500)
                    continue

                if found_year == target_year and found_month == target_month:
                    break

                curr_val = found_year * 12 + found_month
                tgt_val = target_year * 12 + target_month
                arrow = (
                    SELECTORS["month_arrow_right"]
                    if curr_val < tgt_val
                    else SELECTORS["month_arrow_left"]
                )
                popover.locator(arrow).click()
                page.wait_for_timeout(600)

            # Выбор дня
            day_int = target_dt.day
            day_str = str(day_int)
            page.wait_for_timeout(300)

            candidates = popover.get_by_text(day_str, exact=True).all()
            real_days = [
                el
                for el in candidates
                if el.locator("xpath=..").locator("svg").count() == 0
            ]

            day_el = None
            if not real_days:
                self.log("WARNING", f"День {day_str} не найден, пробуем первый кандидат")
                if candidates:
                    day_el = candidates[0]
            elif len(real_days) > 1:
                day_el = real_days[0] if day_int <= 15 else real_days[-1]
            else:
                day_el = real_days[0]

            if day_el:
                for attempt in range(3):
                    try:
                        day_el.click(force=True, timeout=1000)
                        break
                    except Exception:
                        page.wait_for_timeout(200)

            page.wait_for_timeout(500)

            # Часы и минуты
            self._adjust_time_column(page, popover, 0, f"{target_dt.hour:02d}")
            self._adjust_time_column(page, popover, 1, f"{target_dt.minute:02d}")

        except Exception as e:
            self.log("ERROR", f"Ошибка календаря: {e}")

    def _normalize_num_text(self, raw: str) -> str:
        text = (raw or "").strip().replace(",", ".")
        if not text:
            return ""
        try:
            return format_num(float(text))
        except Exception:
            return text

    def _read_suggested_number(self, page) -> str:
        try:
            page.wait_for_selector(SELECTORS["chapter_input"], state="visible", timeout=10000)
            raw = page.input_value(SELECTORS["chapter_input"])
            return self._normalize_num_text(raw)
        except Exception:
            return ""

    def _assert_next_number_changed(self, page, before_num: str):
        """
        Верификация успешного сохранения:
        если после сохранения предложенный номер на add-chapter не изменился,
        считаем, что главу нужно перезалить.
        """
        if not before_num:
            return

        page.goto(self.url, timeout=30000)
        page.wait_for_selector(SELECTORS["volume_input"], state="visible", timeout=15000)
        after_num = self._read_suggested_number(page)
        if after_num and after_num == before_num:
            raise RuntimeError(
                f"После сохранения номер следующей главы не изменился ({after_num})"
            )

    def _wait_before_retry(self, page):
        try:
            page.wait_for_timeout(RETRY_DELAY_SEC * 1000)
        except Exception:
            time.sleep(RETRY_DELAY_SEC)

    # ── Загрузка одной главы (с retry) ──

    def _upload_chapter(self, page, chapter: ChapterData) -> bool:
        """
        Загрузить одну главу. Возвращает True при успехе.
        """
        for attempt in range(1, MAX_RETRIES + 1):
            if not self.is_running:
                return False
            attempt_publish_time = self.current_publish_time
            try:
                # Навигация
                try:
                    page.goto(self.url, timeout=30000)
                    page.wait_for_selector(
                        SELECTORS["volume_input"], state="visible", timeout=15000
                    )
                except Exception:
                    self.log("WARNING", f"Релоад страницы (попытка {attempt})...")
                    page.reload()
                    page.wait_for_selector(SELECTORS["volume_input"], timeout=20000)

                suggested_before = self._read_suggested_number(page)

                # Заполнение полей
                page.fill(SELECTORS["volume_input"], str(chapter.volume))

                if self.force_num:
                    try:
                        page.fill(
                            SELECTORS["chapter_input"], format_num(chapter.number)
                        )
                    except Exception as e:
                        self.log("WARNING", f"Не удалось вписать номер главы: {e}")

                if chapter.title:
                    page.fill(SELECTORS["title_input"], chapter.title)

                # Вставка контента в редактор
                page.click(SELECTORS["editor_area"])
                content = chapter.content
                is_html = content.strip().startswith("<")

                if is_html and len(content) > 5000:
                    # HTML-контент (из epub) — быстрая вставка через innerHTML
                    page.evaluate(
                        """(text) => {
                            const editor = document.querySelector('.ProseMirror');
                            if (editor) { editor.innerHTML = text; }
                        }""",
                        content,
                    )
                    page.wait_for_timeout(300)
                elif not is_html and len(content) > 5000:
                    # Plain-text (из docx/txt) — оборачиваем абзацы в <p>
                    html_content = "".join(
                        f"<p>{line}</p>" for line in content.split("\n") if line.strip()
                    )
                    page.evaluate(
                        """(text) => {
                            const editor = document.querySelector('.ProseMirror');
                            if (editor) { editor.innerHTML = text; }
                        }""",
                        html_content,
                    )
                    page.wait_for_timeout(300)
                else:
                    # Короткий контент — insert_text сохраняет переносы
                    page.keyboard.insert_text(content)

                # Платный доступ + расписание
                if self.paid_enabled:
                    self._configure_paid(page, chapter)
                elif self.schedule_enabled:
                    self._configure_schedule(page)

                # Сохранение
                save_btn = page.locator(SELECTORS["submit_btn"])
                if save_btn.is_visible():
                    save_btn.click()
                else:
                    page.get_by_text("Создать", exact=True).click()

                # Ждём подтверждения (проверяем, что страница сменилась или появилось сообщение)
                page.wait_for_timeout(2000)
                self._assert_next_number_changed(page, suggested_before)

                self.log(
                    "SUCCESS",
                    f"Глава {format_num(chapter.number)} сохранена"
                    + (f" (попытка {attempt})" if attempt > 1 else ""),
                )
                return True

            except Exception as e:
                if self.schedule_enabled:
                    # При ретрае не сдвигаем расписание для этой же главы
                    self.current_publish_time = attempt_publish_time
                self.log(
                    "WARNING",
                    f"Попытка {attempt}/{MAX_RETRIES} для Гл.{format_num(chapter.number)}: {e}",
                )
                if attempt < MAX_RETRIES:
                    self._wait_before_retry(page)
                else:
                    self.log(
                        "ERROR",
                        f"Глава {format_num(chapter.number)} — все попытки исчерпаны.",
                    )
                    return False

        return False

    def _configure_paid(self, page, chapter: ChapterData):
        try:
            page.click(SELECTORS["gear_btn"])
            popover = page.locator(SELECTORS["popover"]).first
            popover.wait_for(state="visible", timeout=3000)

            paid_chk = popover.locator('input[type="checkbox"]').first
            if not paid_chk.is_checked():
                paid_chk.click(force=True)
            page.wait_for_timeout(300)

            price_field = popover.locator(SELECTORS["price_input"])
            if price_field.is_visible():
                price_field.fill(str(self.price))

            if self.schedule_enabled:
                self._set_calendar_date(page, popover, self.current_publish_time)
                self.current_publish_time += timedelta(minutes=self.interval_minutes)

            page.mouse.click(0, 0)
            page.wait_for_timeout(300)
        except Exception as e:
            self.log("WARNING", f"Сбой платных настроек: {e}")

    def _configure_schedule(self, page):
        try:
            page.click(SELECTORS["clock_btn"])
            popover = page.locator(SELECTORS["popover"]).first
            popover.wait_for(state="visible", timeout=3000)

            sched_chk = popover.locator('input[type="checkbox"]').first
            if not sched_chk.is_checked():
                sched_chk.click(force=True)
            page.wait_for_timeout(300)

            self._set_calendar_date(page, popover, self.current_publish_time)
            page.mouse.click(0, 0)
            self.current_publish_time += timedelta(minutes=self.interval_minutes)
        except Exception as e:
            self.log("WARNING", f"Сбой отложки: {e}")

    # ── Основной цикл ──

    def run(self):
        self.log("INFO", "Запуск браузера Chrome...")
        try:
            with sync_playwright() as p:
                browser = _launch_persistent_chromium_context(
                    p,
                    user_data_dir=str(BROWSER_PROFILE_DIR),
                    viewport={"width": 1280, "height": 900},
                    log_callback=self.log,
                )
                page = browser.pages[0]
                total = len(self.chapters_list)

                for index, chapter in enumerate(self.chapters_list):
                    if not self.is_running:
                        self._skipped += total - index
                        break

                    if (
                        self.schedule_enabled
                        and self.current_publish_time > self.limit_date
                    ):
                        self.log("ERROR", "ПРЕДЕЛ В 60 ДНЕЙ. Остановка.")
                        self._skipped += total - index
                        break

                    # Лог статуса
                    parts = []
                    if self.schedule_enabled:
                        parts.append(
                            self.current_publish_time.strftime("%d.%m %H:%M")
                        )
                    else:
                        parts.append("Сразу")
                    if self.paid_enabled:
                        parts.append(f"Плат: {self.price}₽")
                    self.log(
                        "INFO",
                        f"[{index + 1}/{total}] Т.{chapter.volume} "
                        f"Гл.{format_num(chapter.number)} ({', '.join(parts)})",
                    )

                    t0 = time.monotonic()
                    ok = self._upload_chapter(page, chapter)
                    elapsed = time.monotonic() - t0
                    self._times.append(elapsed)

                    if ok:
                        self._ok += 1
                        self.chapter_done_signal.emit(index)
                    else:
                        self._errors += 1

                    progress = int(((index + 1) / total) * 100)
                    self.progress_signal.emit(progress)
                    self.stats_signal.emit(self._ok, self._errors, self._skipped)

                    # ETA
                    remaining = total - (index + 1)
                    if remaining > 0 and self._times:
                        avg = sum(self._times) / len(self._times)
                        eta = timedelta(seconds=avg * remaining)
                        self.eta_signal.emit(f"~{format_timedelta(eta)}")
                    else:
                        self.eta_signal.emit("—")

                    page.wait_for_timeout(1000)

                browser.close()

        except Exception as e:
            self.log("ERROR", f"Критическая ошибка браузера: {e}")
            logging.error(traceback.format_exc())

        self.stats_signal.emit(self._ok, self._errors, self._skipped)
        self.finished_signal.emit()

    def stop(self):
        self.is_running = False


# ─── Рабочий поток: авторизация ─────────────────────────────────────────────

class LoginWorker(QThread):
    log_signal = pyqtSignal(str, str)
    finished_signal = pyqtSignal()

    def __init__(self, site="ranobelib"):
        super().__init__()
        self._browser = None
        self._site = site  # "ranobelib" или "rulate"

    def run(self):
        try:
            profile_dir = BROWSER_PROFILE_DIR if self._site == "ranobelib" else BROWSER_RULATE_DIR
            start_url = "https://ranobelib.me" if self._site == "ranobelib" else "https://tl.rulate.ru"
            site_label = "RanobeLib" if self._site == "ranobelib" else "Rulate"

            with sync_playwright() as p:
                self._browser = _launch_persistent_chromium_context(
                    p,
                    user_data_dir=str(profile_dir),
                    log_callback=self.log_signal.emit,
                )
                page = self._browser.pages[0]
                try:
                    page.goto(start_url, timeout=60000)
                except Exception:
                    pass
                self.log_signal.emit(
                    "WARNING", f">>> ВОЙДИТЕ В АККАУНТ {site_label} И ЗАКРОЙТЕ БРАУЗЕР <<<"
                )
                try:
                    while True:
                        page.wait_for_timeout(1000)
                except Exception:
                    pass
                self.log_signal.emit("SUCCESS", f"Браузер {site_label} закрыт. Куки сохранены.")
        except Exception as e:
            self.log_signal.emit("ERROR", f"Ошибка авторизации: {e}")
        finally:
            self.finished_signal.emit()

    def stop(self):
        if self._browser:
            try:
                self._browser.close()
            except Exception:
                pass


